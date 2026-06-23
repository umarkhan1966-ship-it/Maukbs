"""
BusinessVault — Clean Foundation (Module 1)
==========================================
Database schema + navigation shell + login stub.
All other modules will be added on top of this file.
"""

import sqlite3
import hashlib
import os
import secrets
from datetime import datetime
from fastapi import FastAPI, Request, Form, Cookie
from fastapi.responses import HTMLResponse, RedirectResponse

app     = FastAPI()
DB_FILE = "business_vault.db"

STORE_GPS = {
    "Uxbridge": (51.5462, -0.4791),
    "Newbury":  (51.4014, -1.3231)
}
GEOFENCE_RADIUS_M = 200

# ══════════════════════════════════════════════════════════════════════════════
# DATABASE
# ══════════════════════════════════════════════════════════════════════════════

def db():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

def q(sql, params=(), fetch=False):
    conn = db()
    cur  = conn.cursor()
    cur.execute(sql, params)
    result = cur.fetchall() if fetch else None
    conn.commit()
    conn.close()
    return result


# ══════════════════════════════════════════════════════════════════════════════
# PASSWORD HASHING (salted PBKDF2 — no external dependency)
# ══════════════════════════════════════════════════════════════════════════════

PBKDF2_ITERATIONS = 200_000

def hash_password(pw: str) -> str:
    """Return a salted PBKDF2 hash string: pbkdf2_sha256$iters$salt$hash."""
    salt = secrets.token_bytes(16)
    dk   = hashlib.pbkdf2_hmac("sha256", pw.encode(), salt, PBKDF2_ITERATIONS)
    return f"pbkdf2_sha256${PBKDF2_ITERATIONS}${salt.hex()}${dk.hex()}"

def verify_password(pw: str, stored: str) -> bool:
    """Verify a password against a stored hash.

    Supports the new salted PBKDF2 format and falls back to the legacy
    unsalted SHA-256 hashes so existing accounts keep working until they
    log in once (login upgrades them automatically)."""
    if not stored:
        return False
    if stored.startswith("pbkdf2_sha256$"):
        try:
            _, iters, salt_hex, hash_hex = stored.split("$")
            dk = hashlib.pbkdf2_hmac("sha256", pw.encode(),
                                     bytes.fromhex(salt_hex), int(iters))
            return secrets.compare_digest(dk.hex(), hash_hex)
        except Exception:
            return False
    # Legacy unsalted SHA-256
    legacy = hashlib.sha256(pw.encode()).hexdigest()
    return secrets.compare_digest(legacy, stored)


def init_db():
    conn = db()
    c    = conn.cursor()

    # ── Users ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            user_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            username   TEXT UNIQUE NOT NULL,
            password   TEXT NOT NULL,
            full_name  TEXT,
            role       TEXT NOT NULL DEFAULT 'staff',
            store_name TEXT,
            is_active  INTEGER DEFAULT 1,
            created_at TEXT DEFAULT (date('now'))
        )
    """)

    # ── Staff Profiles ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_profiles (
            staff_id       INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_number   INTEGER,
            first_name     TEXT NOT NULL,
            last_name      TEXT NOT NULL,
            store_name     TEXT,
            sex            TEXT,
            phone          TEXT,
            email          TEXT,
            address_1      TEXT,
            address_2      TEXT,
            address_3      TEXT,
            address_4      TEXT,
            postcode       TEXT,
            date_joined    TEXT,
            date_left      TEXT,
            leaving_reason TEXT,
            date_of_birth  TEXT,
            contracted_hrs REAL,
            hourly_rate    REAL,
            is_salaried    TEXT DEFAULT 'N',
            salary_amount  REAL,
            is_active      INTEGER DEFAULT 1
        )
    """)

    # ── Supplier Invoices (Retail) ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS supplier_invoices (
            invoice_id     INTEGER PRIMARY KEY AUTOINCREMENT,
            seq_no         INTEGER,
            supplier_name  TEXT NOT NULL,
            store_name     TEXT NOT NULL,
            invoice_number TEXT,
            invoice_date   TEXT,
            gross_amount   REAL DEFAULT 0,
            vat_amount     REAL DEFAULT 0,
            net_amount     REAL DEFAULT 0,
            payment_terms  INTEGER,
            due_date       TEXT,
            paid_date      TEXT,
            amount_paid    REAL DEFAULT 0,
            credit_note    REAL DEFAULT 0,
            is_paid        TEXT DEFAULT 'No',
            payment_method TEXT,
            comments       TEXT,
            pdf_path          TEXT,
            approval_status   TEXT DEFAULT 'approved',
            submitted_by      TEXT,
            created_at        TEXT DEFAULT (date('now')),
            UNIQUE(invoice_number, store_name)
        )
    """)

    # ── Property Invoices / Expenses ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS property_invoices (
            invoice_id     INTEGER PRIMARY KEY AUTOINCREMENT,
            property_name  TEXT NOT NULL,
            supplier_name  TEXT NOT NULL,
            invoice_number TEXT,
            invoice_date   TEXT,
            expense_type   TEXT,
            gross_amount   REAL DEFAULT 0,
            vat_amount     REAL DEFAULT 0,
            net_amount     REAL DEFAULT 0,
            due_date       TEXT,
            paid_date      TEXT,
            amount_paid    REAL DEFAULT 0,
            credit_note    REAL DEFAULT 0,
            is_paid        TEXT DEFAULT 'No',
            payment_method TEXT,
            comments       TEXT,
            pdf_path          TEXT,
            approval_status   TEXT DEFAULT 'approved',
            submitted_by      TEXT,
            created_at        TEXT DEFAULT (date('now'))
        )
    """)

    # ── Rental Income ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS rental_income (
            record_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            property_name    TEXT NOT NULL,
            tenant_name      TEXT,
            rent_from        TEXT,
            rent_to          TEXT,
            agreed_rent      REAL DEFAULT 0,
            agency_comm      REAL DEFAULT 0,
            agency_vat       REAL DEFAULT 0,
            tds_fee          REAL DEFAULT 0,
            gas_elec_cert    REAL DEFAULT 0,
            inventory_fee    REAL DEFAULT 0,
            deposit_fee      REAL DEFAULT 0,
            tenancy_setup    REAL DEFAULT 0,
            repairs          REAL DEFAULT 0,
            repairs_vat      REAL DEFAULT 0,
            mortgage         REAL DEFAULT 0,
            net_rent         REAL DEFAULT 0,
            date_received    TEXT,
            notes            TEXT,
            created_at       TEXT DEFAULT (date('now'))
        )
    """)

    # ── Properties ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS properties (
            property_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            short_name     TEXT UNIQUE NOT NULL,
            full_address   TEXT NOT NULL,
            purchase_price REAL,
            mortgage       REAL,
            monthly_mortgage REAL,
            purchase_date  TEXT,
            first_rented   TEXT,
            is_active      INTEGER DEFAULT 1,
            notes          TEXT
        )
    """)

    # ── Rotas ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS store_rotas (
            rota_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_name TEXT NOT NULL,
            store_name TEXT NOT NULL,
            work_day   TEXT NOT NULL,
            shift_time TEXT DEFAULT 'OFF',
            UNIQUE(staff_name, store_name, work_day)
        )
    """)

    # ── Timesheets ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS timesheets (
            timesheet_id   INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_name     TEXT NOT NULL,
            store_name     TEXT NOT NULL,
            work_date      TEXT NOT NULL,
            clock_in_time  TEXT,
            clock_out_time TEXT,
            status_flag    TEXT,
            absence_type   TEXT,
            comments       TEXT,
            UNIQUE(staff_name, store_name, work_date)
        )
    """)

    # ── Daily Sales ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS daily_sales (
            sale_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name     TEXT NOT NULL,
            sale_date      TEXT NOT NULL,
            week_ending    TEXT,
            category       TEXT NOT NULL,
            amount         REAL DEFAULT 0,
            entered_by     TEXT,
            created_at     TEXT DEFAULT (datetime('now')),
            UNIQUE(store_name, sale_date, category)
        )
    """)

    # ── Sales Targets ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS sales_targets (
            target_id      INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name     TEXT NOT NULL,
            year           INTEGER NOT NULL,
            month          INTEGER NOT NULL,
            target_amount  REAL DEFAULT 0,
            ly_actual      REAL DEFAULT 0,
            target_pct     REAL DEFAULT 1.05,
            UNIQUE(store_name, year, month)
        )
    """)

    # ── Sessions (server-side login tokens) ──
    c.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            token      TEXT PRIMARY KEY,
            username   TEXT NOT NULL,
            created_at TEXT DEFAULT (datetime('now')),
            expires_at TEXT NOT NULL
        )
    """)

    # ── Seed default owner account (password: changeme) ──
    pw_hash = hash_password("changeme")
    c.execute("""
        INSERT OR IGNORE INTO users (username, password, full_name, role)
        VALUES ('owner', ?, 'Business Owner', 'owner')
    """, (pw_hash,))

    # ── Seed properties from known data ──
    for short, full, price, mort, m_mort, pdate in [
        ("104 Dane",  "104 Dane Road",      550000, 412482, 0,      "2021-05-07"),
        ("53 Ampth",  "53 Ampthill Way",    160000, 114906, 549.54, "2024-10-31"),
        ("26 Ampth",  "26 Ampthill Way",    165000, 123750, 0,      "2025-07-26"),
    ]:
        c.execute("""
            INSERT OR IGNORE INTO properties
                (short_name, full_address, purchase_price, mortgage, monthly_mortgage, purchase_date)
            VALUES (?,?,?,?,?,?)
        """, (short, full, price, mort, m_mort, pdate))

    conn.commit()
    conn.close()
    print("✅ Database initialised.")

init_db()


# ══════════════════════════════════════════════════════════════════════════════
# AUTH HELPERS (stub — full login added in later module)
# ══════════════════════════════════════════════════════════════════════════════

def get_session(token: str | None) -> dict | None:
    """Return the logged-in user dict for a valid, unexpired session token."""
    if not token:
        return None
    rows = q("""SELECT u.* FROM sessions s
                JOIN users u ON u.username = s.username
                WHERE s.token = ?
                  AND s.expires_at > datetime('now')
                  AND u.is_active = 1""",
             (token,), fetch=True)
    return dict(rows[0]) if rows else None

def require_login(token: str | None):
    """Redirect to login if not authenticated."""
    user = get_session(token)
    if not user:
        return RedirectResponse("/login", status_code=303), None
    return None, user


# ══════════════════════════════════════════════════════════════════════════════
# HTML SHELL
# ══════════════════════════════════════════════════════════════════════════════

SALES_CATS = [
    "Digital Printing", "Other D&P", "Instant Prints", "Reprint/Enlarge",
    "Internet Orders", "Passport", "Film Media", "Graphic Design",
    "Large Format", "Toner/ Laser Output", "Batteries", "Frames & Albums",
    "Photogifts", "Backup to Media", "DVD Transfer", "Studio",
    "Sundry", "Promotions", "RCS (STD VAT)", "RCS (ZERO)",
    "Photobooks", "TYPE B Sales"
]

def page(title: str, content: str, user: dict, active: str = "") -> str:
    role     = user.get("role", "staff")
    name     = user.get("full_name") or user.get("username", "")
    store    = user.get("store_name") or ""
    is_owner = role == "owner"
    is_mgr   = role in ("owner", "manager")

    # Nav items: (label, href, icon, min_role)
    nav = [
        ("Dashboard",   "/",              "&#11035;", "staff"),
        ("My Profile",  "/my-profile",    "&#128100;","staff"),
        ("Sales",       "/sales",         "&#128200;","staff"),
        ("Invoices",    "/invoices",      "&#129534;","staff"),
        ("Staff",       "/staff",         "&#128100;","manager"),
        ("Rota",        "/rota",          "&#128197;","manager"),
        ("Timesheets",  "/timesheets",    "&#9200;",  "manager"),
        ("Property",    "/property",      "&#127968;","owner"),
        ("Settings",    "/settings",      "&#9881;",  "owner"),
    ]

    nav_html = ""
    for label, href, icon, min_role in nav:
        if min_role == "owner"   and not is_owner: continue
        if min_role == "manager" and not is_mgr:   continue
        active_cls = "bg-white/15 font-black" if active == label.lower() else "hover:bg-white/10"
        nav_html += f"<a href='{href}' class='flex items-center gap-2 px-3 py-2 rounded-lg text-sm font-semibold transition {active_cls}'>{icon} {label}</a>"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title} — BusinessVault</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700;900&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
  <script src="https://cdn.jsdelivr.net/npm/@tailwindcss/browser@4"></script>
  <style>
    body {{ font-family: 'DM Sans', sans-serif; }}
    .mono {{ font-family: 'DM Mono', monospace; }}
    ::-webkit-scrollbar {{ width: 6px; height: 6px; }}
    ::-webkit-scrollbar-track {{ background: #f1f5f9; }}
    ::-webkit-scrollbar-thumb {{ background: #cbd5e1; border-radius: 3px; }}
    .card {{ background: white; border-radius: 16px; border: 1px solid #e2e8f0; padding: 24px; }}
    .btn-primary {{ background:#1e3a5f; color:white; font-weight:700; padding:8px 20px; border-radius:10px; font-size:14px; transition:all .15s; display:inline-block; }}
    .btn-primary:hover {{ background:#16304f; }}
    .btn-secondary {{ background:#f1f5f9; color:#334155; font-weight:700; padding:8px 20px; border-radius:10px; font-size:14px; transition:all .15s; display:inline-block; }}
    .btn-secondary:hover {{ background:#e2e8f0; }}
    .btn-danger {{ background:#fee2e2; color:#dc2626; font-weight:700; padding:8px 20px; border-radius:10px; font-size:14px; transition:all .15s; display:inline-block; }}
    .btn-danger:hover {{ background:#fecaca; }}
    .btn-success {{ background:#dcfce7; color:#16a34a; font-weight:700; padding:8px 20px; border-radius:10px; font-size:14px; transition:all .15s; display:inline-block; }}
    .btn-success:hover {{ background:#bbf7d0; }}
    .badge-paid {{ background:#dcfce7; color:#16a34a; font-size:11px; font-weight:700; padding:2px 8px; border-radius:6px; }}
    .badge-overdue {{ background:#fee2e2; color:#dc2626; font-size:11px; font-weight:700; padding:2px 8px; border-radius:6px; }}
    .badge-partial {{ background:#fef3c7; color:#d97706; font-size:11px; font-weight:700; padding:2px 8px; border-radius:6px; }}
    .badge-unpaid {{ background:#f1f5f9; color:#64748b; font-size:11px; font-weight:700; padding:2px 8px; border-radius:6px; }}
    .tbl {{ width:100%; border-collapse:collapse; font-size:13px; }}
    .tbl th {{ background:#0f2942; color:white; padding:10px 12px; text-align:left; font-size:11px; font-weight:700; text-transform:uppercase; letter-spacing:.05em; white-space:nowrap; }}
    .tbl td {{ padding:10px 12px; border-bottom:1px solid #f1f5f9; vertical-align:middle; }}
    .tbl tr:hover td {{ background:#f8fafc; }}
    .tbl tr:last-child td {{ border-bottom:none; }}
    input, select, textarea {{
      width:100%; border:1px solid #e2e8f0; border-radius:8px;
      padding:8px 12px; font-size:14px; font-family:'DM Sans',sans-serif;
      outline:none; transition:border .15s; background:white;
    }}
    input:focus, select:focus, textarea:focus {{ border-color:#1e3a5f; }}
    input[type=number]::-webkit-outer-spin-button,
    input[type=number]::-webkit-inner-spin-button {{ -webkit-appearance:none; margin:0; }}
    input[type=number] {{ -moz-appearance:textfield; }}
    label {{ font-size:12px; font-weight:700; color:#64748b; text-transform:uppercase; letter-spacing:.05em; display:block; margin-bottom:4px; }}
    .flash-success {{ background:#dcfce7; border:1px solid #86efac; color:#15803d; padding:12px 16px; border-radius:10px; font-size:14px; font-weight:600; }}
    .flash-error   {{ background:#fee2e2; border:1px solid #fca5a5; color:#dc2626; padding:12px 16px; border-radius:10px; font-size:14px; font-weight:600; }}
  </style>
</head>
<body class="bg-slate-100 min-h-screen">

  <!-- Sidebar -->
  <div class="fixed top-0 left-0 h-full w-52 z-40"
       style="background:linear-gradient(180deg,#0f2942 0%,#1e3a5f 100%);">
    <div class="p-5 border-b border-white/10">
      <div class="text-white font-black text-lg tracking-tight">BusinessVault</div>
      <div class="text-blue-300 text-xs font-semibold mt-0.5">Maukbs Ltd</div>
    </div>
    <nav class="p-3 space-y-1 text-white">
      {nav_html}
    </nav>
    <div class="absolute bottom-0 left-0 right-0 p-4 border-t border-white/10">
      <div class="text-white text-xs font-bold truncate">{name}</div>
      <div class="text-blue-300 text-xs capitalize">{role}{' · ' + store if store else ''}</div>
      <a href="/logout" class="text-blue-300 hover:text-white text-xs mt-1 inline-block transition">Sign out →</a>
    </div>
  </div>

  <!-- Main content -->
  <div class="ml-52 min-h-screen">
    <div class="max-w-7xl mx-auto p-6 space-y-6">
      {content}
    </div>
  </div>

</body>
</html>"""


# ══════════════════════════════════════════════════════════════════════════════
# LOGIN
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/login", response_class=HTMLResponse)
def login_page(error: str = ""):
    err_html = f"<p class='flash-error'>{error}</p>" if error else ""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>BusinessVault — Sign In</title>
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;600;700;900&display=swap" rel="stylesheet">
  <script src="https://cdn.jsdelivr.net/npm/@tailwindcss/browser@4"></script>
  <style>body{{font-family:'DM Sans',sans-serif;}}</style>
</head>
<body class="bg-slate-100 min-h-screen flex items-center justify-center p-4">
  <div class="w-full max-w-sm">
    <div class="text-center mb-8">
      <div class="text-3xl font-black text-slate-800 tracking-tight">BusinessVault</div>
      <div class="text-slate-500 text-sm mt-1">Maukbs Ltd · Management System</div>
    </div>
    <div style="background:white;border-radius:20px;padding:32px;border:1px solid #e2e8f0;box-shadow:0 4px 24px rgba(0,0,0,.06)">
      {err_html}
      <form action="/login" method="POST" class="space-y-4 {'mt-4' if error else ''}">
        <div>
          <label style="font-size:12px;font-weight:700;color:#64748b;text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px">Username</label>
          <input name="username" type="text" required autofocus
            style="width:100%;border:1px solid #e2e8f0;border-radius:8px;padding:10px 14px;font-size:15px;outline:none;font-family:'DM Sans',sans-serif;">
        </div>
        <div>
          <label style="font-size:12px;font-weight:700;color:#64748b;text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px">Password</label>
          <input name="password" type="password" required
            style="width:100%;border:1px solid #e2e8f0;border-radius:8px;padding:10px 14px;font-size:15px;outline:none;font-family:'DM Sans',sans-serif;">
        </div>
        <button type="submit"
          style="width:100%;background:#0f2942;color:white;font-weight:700;padding:12px;border-radius:10px;font-size:15px;border:none;cursor:pointer;font-family:'DM Sans',sans-serif;margin-top:4px;">
          Sign In →
        </button>
      </form>
    </div>
    <p class="text-center text-xs text-slate-400 mt-6">Maukbs Ltd · Authorised users only</p>
  </div>
</body>
</html>"""


@app.post("/login")
def do_login(username: str = Form(...), password: str = Form(...)):
    rows = q("SELECT * FROM users WHERE username=? AND is_active=1",
             (username,), fetch=True)
    user = dict(rows[0]) if rows else None
    if not user or not verify_password(password, user["password"]):
        return RedirectResponse("/login?error=Invalid+username+or+password", status_code=303)

    # Transparently upgrade legacy unsalted hashes on successful login.
    if not user["password"].startswith("pbkdf2_sha256$"):
        q("UPDATE users SET password=? WHERE username=?",
          (hash_password(password), username))

    # Clear out expired sessions, then issue a fresh random token.
    q("DELETE FROM sessions WHERE expires_at <= datetime('now')")
    token = secrets.token_urlsafe(32)
    q("INSERT INTO sessions (token, username, expires_at) "
      "VALUES (?, ?, datetime('now', '+7 days'))",
      (token, username))

    resp = RedirectResponse("/", status_code=303)
    resp.set_cookie("session", token, httponly=True, samesite="lax",
                    max_age=86400 * 7)
    return resp


@app.get("/logout")
def do_logout(session: str | None = Cookie(default=None)):
    if session:
        q("DELETE FROM sessions WHERE token=?", (session,))
    resp = RedirectResponse("/login", status_code=303)
    resp.delete_cookie("session")
    return resp


# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD (placeholder — will be filled in later modules)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/", response_class=HTMLResponse)
def dashboard(session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    today    = datetime.now().strftime("%A, %d %B %Y")
    is_owner = user["role"] == "owner"

    # Quick summary counts
    overdue  = q("""SELECT COUNT(*) as n FROM supplier_invoices
                    WHERE is_paid!='Yes' AND due_date < date('now')""", fetch=True)
    overdue_n = overdue[0]["n"] if overdue else 0

    overdue_val = q("""SELECT COALESCE(SUM(gross_amount-amount_paid-credit_note),0) as v
                       FROM supplier_invoices
                       WHERE is_paid!='Yes' AND due_date < date('now')""", fetch=True)
    overdue_v = overdue_val[0]["v"] if overdue_val else 0

    active_staff = q("SELECT COUNT(*) as n FROM staff_profiles WHERE is_active=1", fetch=True)
    staff_n = active_staff[0]["n"] if active_staff else 0

    # This week's sales (both stores)
    week_sales = q("""SELECT COALESCE(SUM(amount),0) as v FROM daily_sales
                      WHERE sale_date >= date('now','-7 days')""", fetch=True)
    week_v = week_sales[0]["v"] if week_sales else 0

    def stat_card(icon, label, value, sub, colour):
        return f"""
        <div class='card flex items-start gap-4'>
          <div class='text-3xl'>{icon}</div>
          <div>
            <div class='text-xs font-bold text-slate-400 uppercase tracking-wide'>{label}</div>
            <div class='text-2xl font-black' style='color:{colour}'>{value}</div>
            <div class='text-xs text-slate-400 mt-0.5'>{sub}</div>
          </div>
        </div>"""

    cards = f"""
    <div>
      <div class='text-2xl font-black text-slate-800'>Good {'morning' if datetime.now().hour < 12 else 'afternoon'}, {user['full_name'] or user['username'].title()} 👋</div>
      <div class='text-slate-400 text-sm mt-1'>{today}</div>
    </div>
    <div class='grid grid-cols-2 gap-4' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
      {stat_card('🚨', 'Overdue Invoices', overdue_n, f'£{overdue_v:,.2f} outstanding', '#dc2626')}
      {stat_card('📈', 'Sales This Week', f'£{week_v:,.2f}', 'Both stores combined', '#16a34a')}
      {stat_card('👤', 'Active Staff', staff_n, 'Across both stores', '#1e3a5f')}
      {stat_card('🏠', 'Properties', '3', '104 Dane · 53 Ampth · 26 Ampth', '#7c3aed') if is_owner else ''}
    </div>

    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(300px,1fr))'>
      <div class='card'>
        <div class='font-black text-slate-700 mb-3'>⚡ Quick Actions</div>
        <div class='space-y-2'>
          <a href='/invoices' class='btn-primary block text-center'>🧾 Manage Invoices</a>
          <a href='/sales' class='btn-primary block text-center'>📈 Enter Today's Sales</a>
          <a href='/rota' class='btn-secondary block text-center'>📅 View Rota</a>
          {'<a href="/property" class="btn-secondary block text-center">🏠 Property Portfolio</a>' if is_owner else ''}
        </div>
      </div>
      <div class='card'>
        <div class='font-black text-slate-700 mb-3'>📋 Modules</div>
        <div class='text-sm text-slate-500 space-y-2'>
          <div class='flex justify-between items-center py-1 border-b border-slate-100'>
            <span>🧾 Invoice Management</span><span class='badge-paid'>Ready</span>
          </div>
          <div class='flex justify-between items-center py-1 border-b border-slate-100'>
            <span>📈 Sales & Franchise</span><span class='badge-unpaid'>Coming next</span>
          </div>
          <div class='flex justify-between items-center py-1 border-b border-slate-100'>
            <span>👤 Staff & Rota</span><span class='badge-unpaid'>Coming soon</span>
          </div>
          <div class='flex justify-between items-center py-1'>
            <span>🏠 Property Portfolio</span><span class='badge-unpaid'>Coming soon</span>
          </div>
        </div>
      </div>
    </div>"""

    return page("Dashboard", cards, user, "dashboard")


# ══════════════════════════════════════════════════════════════════════════════
# PLACEHOLDER ROUTES (so nav links don't 404)
# ══════════════════════════════════════════════════════════════════════════════

def placeholder(title, icon, session):
    redir, user = require_login(session)
    if redir: return redir
    content = f"""
    <div class='text-2xl font-black text-slate-800'>{icon} {title}</div>
    <div class='card text-center py-16 text-slate-400'>
      <div class='text-4xl mb-3'>🚧</div>
      <div class='font-bold text-lg'>Coming in the next build</div>
      <div class='text-sm mt-1'>This module is being built now</div>
    </div>"""
    return page(title, content, user, title.lower())

# sales routes added below

# invoices route replaced below

# staff routes replaced below

# rota routes added below

# timesheets route added below

@app.get("/property",   response_class=HTMLResponse)
def property_page(session: str | None = Cookie(default=None)):
    return placeholder("Property", "🏠", session)

@app.get("/settings",   response_class=HTMLResponse)
def settings_page(session: str | None = Cookie(default=None)):
    return placeholder("Settings", "⚙️", session)

# mobile-clock route added below


# ══════════════════════════════════════════════════════════════════════════════
# MODULE 2 — INVOICE MANAGER
# ══════════════════════════════════════════════════════════════════════════════
#
# Single-page VBA-style interface:
#   • Search form at top (supplier name, invoice no, serial no, store, status)
#   • Results list below — click any row to populate the edit form
#   • Add / Edit / Delete / Mark Paid (full, part, credit note)
#   • Works for both Retail stores and Property portfolio
#

import os, uuid, re
from typing import Optional

def extract_pdf_data(pdf_bytes: bytes) -> dict:
    """Try to extract invoice fields from a PDF using pdfplumber.
    Returns a dict with whatever fields we can find — caller fills the rest manually."""
    result = {}
    try:
        import pdfplumber, io
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)

        lines = text.split("\n")
        full  = text.lower()

        # Supplier name — usually first non-empty line or line before "Invoice"
        for line in lines[:8]:
            line = line.strip()
            if len(line) > 3 and not any(w in line.lower() for w in
               ["invoice", "tax", "vat", "date", "ltd", "limited", "plc"]):
                result["supplier_name"] = line
                break
        # Also try lines containing the company name near "from" or at very top
        if not result.get("supplier_name") and lines:
            result["supplier_name"] = lines[0].strip()

        # Invoice number — look for "invoice no", "inv no", "invoice #", "invoice number"
        inv_patterns = [
            r"invoice\s*(?:no\.?|number|#)[:\s]+([A-Z0-9\-\/]+)",
            r"inv\.?\s*(?:no\.?|#)[:\s]+([A-Z0-9\-\/]+)",
            r"(?:^|\s)(INV[-\s]?[0-9]+)",
        ]
        for pat in inv_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                result["invoice_number"] = m.group(1).strip()
                break

        # Invoice date
        date_patterns = [
            r"(?:invoice\s*date|date\s*of\s*invoice|date)[:\s]+([0-9]{1,2}[\s/\-][A-Za-z0-9]{1,3}[\s/\-][0-9]{2,4})",
            r"(?:dated?)[:\s]+([0-9]{1,2}[\s/\-][A-Za-z0-9]{2,3}[\s/\-][0-9]{2,4})",
            r"([0-9]{2}/[0-9]{2}/[0-9]{4})",
            r"([0-9]{2}-[0-9]{2}-[0-9]{4})",
            r"([0-9]{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+[0-9]{4})",
        ]
        for pat in date_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                raw = m.group(1).strip()
                # Try to parse and normalise to YYYY-MM-DD
                from datetime import datetime as dt
                for fmt in ("%d/%m/%Y","%d-%m-%Y","%d %B %Y","%d %b %Y",
                            "%d/%m/%y","%B %d, %Y","%b %d, %Y"):
                    try:
                        result["invoice_date"] = dt.strptime(raw, fmt).strftime("%Y-%m-%d")
                        break
                    except: pass
                if not result.get("invoice_date"):
                    result["invoice_date_raw"] = raw
                break

        # Gross / total amount
        amount_patterns = [
            r"(?:total|gross|amount\s*due|total\s*due|balance\s*due|total\s*payable)[:\s£]+([0-9,]+\.?[0-9]*)",
            r"(?:total\s*inc\.?\s*vat|total\s*including\s*vat)[:\s£]+([0-9,]+\.?[0-9]*)",
            r"£\s*([0-9,]+\.[0-9]{2})\s*$",
        ]
        for pat in amount_patterns:
            m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
            if m:
                try:
                    result["gross_amount"] = float(m.group(1).replace(",",""))
                    break
                except: pass

        # VAT amount
        vat_patterns = [
            r"(?:vat|tax)[:\s£]+([0-9,]+\.?[0-9]*)",
            r"(?:vat\s*@\s*20%)[:\s£]+([0-9,]+\.?[0-9]*)",
        ]
        for pat in vat_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    result["vat_amount"] = float(m.group(1).replace(",",""))
                    break
                except: pass

        # Net amount
        net_patterns = [
            r"(?:net|subtotal|sub\s*total|amount\s*ex\.?\s*vat)[:\s£]+([0-9,]+\.?[0-9]*)",
        ]
        for pat in net_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    result["net_amount"] = float(m.group(1).replace(",",""))
                    break
                except: pass

        # If we have gross and vat but no net, calculate it
        if result.get("gross_amount") and result.get("vat_amount") and not result.get("net_amount"):
            result["net_amount"] = round(result["gross_amount"] - result["vat_amount"], 2)

        # Payment terms
        terms_m = re.search(r"(?:payment\s*terms?|net)[:\s]+(\d+)\s*days?", text, re.IGNORECASE)
        if terms_m:
            result["payment_terms"] = int(terms_m.group(1))

        result["_raw_text"] = text[:500]  # first 500 chars for debugging

    except Exception as e:
        result["_error"] = str(e)

    return result

UPLOAD_DIR = "invoice_pdfs"
os.makedirs(UPLOAD_DIR, exist_ok=True)

PAYMENT_METHODS = ["", "Direct Debit", "Card", "Cash", "Cheque", "Online", "Amex"]
EXPENSE_TYPES   = [
    "Mortgage", "Insurance", "Legal Fees", "Management Fees",
    "Repairs & Maintenance", "Gas/Electric Certificate", "Inventory Fee",
    "Deposit Fee", "Tenancy Setup", "Rates", "Utilities", "Other"
]

# ── Helper: which ledger does this user see? ──────────────────────────────────

def ledger_options(user: dict) -> list[tuple]:
    """Return (value, label) pairs for store/ledger selector.
    Staff see only their store. Managers see both stores.
    Owner sees both stores + all properties."""
    opts = []
    role  = user.get("role", "staff")
    store = user.get("store_name", "")
    if role == "owner":
        opts += [("Uxbridge", "🏪 Uxbridge (Retail)"),
                 ("Newbury",  "🏪 Newbury (Retail)")]
        props = q("SELECT short_name, full_address FROM properties ORDER BY short_name",
                  fetch=True) or []
        for p in props:
            opts.append((f"PROP:{p['short_name']}", f"🏠 {p['full_address']}"))
    elif role == "manager":
        opts += [("Uxbridge", "🏪 Uxbridge (Retail)"),
                 ("Newbury",  "🏪 Newbury (Retail)")]
    else:
        # Store staff — only their assigned store
        s = store or "Uxbridge"
        opts += [(s, f"🏪 {s} (Retail)")]
    return opts


def is_property_ledger(store_val: str) -> bool:
    return store_val.startswith("PROP:")


def prop_name(store_val: str) -> str:
    return store_val.replace("PROP:", "")


# ── Fetch invoices with optional filters ─────────────────────────────────────

def fetch_invoices(ledger: str, search: str, status: str,
                   pg: int, page_size: int = 30):
    is_prop = is_property_ledger(ledger)
    table   = "property_invoices" if is_prop else "supplier_invoices"
    loc_col = "property_name"     if is_prop else "store_name"
    loc_val = prop_name(ledger)   if is_prop else ledger

    conds  = [f"{loc_col} = ?"]
    params = [loc_val]

    if search.strip():
        conds.append("(supplier_name LIKE ? OR invoice_number LIKE ? OR CAST(seq_no AS TEXT) LIKE ?)")
        params += [f"%{search}%", f"%{search}%", f"%{search}%"]

    today = datetime.now().strftime("%Y-%m-%d")
    if status == "overdue":
        conds.append(f"is_paid != 'Yes' AND due_date < '{today}' AND COALESCE(approval_status,'approved')='approved'")
    elif status == "unpaid":
        conds.append("is_paid != 'Yes' AND COALESCE(approval_status,'approved')='approved'")
    elif status == "paid":
        conds.append("is_paid = 'Yes'")
    elif status == "partial":
        conds.append("is_paid != 'Yes' AND amount_paid > 0")
    elif status == "pending":
        conds.append("approval_status = 'pending'")
    else:
        # Default: exclude pending from main view unless owner/manager reviewing
        pass

    where  = "WHERE " + " AND ".join(conds)
    total  = q(f"SELECT COUNT(*) as n FROM {table} {where}", params, fetch=True)
    total_n = total[0]["n"] if total else 0

    balance_expr = "COALESCE(gross_amount,0)-COALESCE(amount_paid,0)-COALESCE(credit_note,0)"
    rows = q(f"""
        SELECT *, {balance_expr} AS balance
        FROM {table} {where}
        ORDER BY due_date ASC, invoice_id DESC
        LIMIT ? OFFSET ?
    """, params + [page_size, (pg-1)*page_size], fetch=True) or []

    # Convert to dicts so .get() works safely throughout
    return [dict(r) for r in rows], total_n


# ── Main invoice page ─────────────────────────────────────────────────────────

@app.get("/invoices", response_class=HTMLResponse)
def invoices_page(
    session:  str | None = Cookie(default=None),
    ledger:   str = "Uxbridge",
    search:   str = "",
    status:   str = "",
    pg:       int = 1,
    edit_id:  int = 0,
    msg:      str = "",
    msg_type: str = "success"
):
    redir, user = require_login(session)
    if redir: return redir

    today      = datetime.now().strftime("%Y-%m-%d")
    is_prop    = is_property_ledger(ledger)
    table      = "property_invoices" if is_prop else "supplier_invoices"
    loc_col    = "property_name"     if is_prop else "store_name"
    loc_val    = prop_name(ledger)   if is_prop else ledger
    ledgers    = ledger_options(user)
    PAGE_SIZE  = 30

    # If edit_id given, load that invoice into the form
    edit_inv = None
    if edit_id:
        rows = q(f"SELECT * FROM {table} WHERE invoice_id=?", (edit_id,), fetch=True)
        if rows:
            edit_inv = dict(rows[0])

    invoices, total_n = fetch_invoices(ledger, search, status, pg, PAGE_SIZE)
    total_pages = max(1, (total_n + PAGE_SIZE - 1) // PAGE_SIZE)

    # Pending approvals count (managers/owners only)
    pending_count = 0
    if user["role"] in ("owner", "manager"):
        p1 = q(f"SELECT COUNT(*) as n FROM {table} WHERE {loc_col}=? AND approval_status='pending'",
               (loc_val,), fetch=True)
        pending_count = p1[0]["n"] if p1 else 0

    # Summary totals for this ledger
    tots = q(f"""
        SELECT
          COUNT(*) as total_count,
          COALESCE(SUM(CASE WHEN is_paid!='Yes' AND due_date < '{today}' THEN gross_amount-amount_paid-credit_note ELSE 0 END),0) as overdue_val,
          COUNT(CASE WHEN is_paid!='Yes' AND due_date < '{today}' THEN 1 END) as overdue_count,
          COALESCE(SUM(CASE WHEN is_paid='Yes' THEN amount_paid ELSE 0 END),0) as paid_val
        FROM {table} WHERE {loc_col}=?
    """, (loc_val,), fetch=True)
    t = dict(tots[0]) if tots else {}

    # ── Flash message ──
    flash = ""
    if msg:
        cls = "flash-success" if msg_type == "success" else "flash-error"
        flash = f"<div class='{cls}'>{msg}</div>"

    # ── Ledger selector ──
    ledger_opts = ""
    for val, label in ledgers:
        sel = "selected" if val == ledger else ""
        ledger_opts += f"<option value='{val}' {sel}>{label}</option>"

    # ── Summary bar ──
    summary = f"""
    <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(160px,1fr))'>
      <div class='card py-3 text-center'>
        <div class='text-xs font-bold text-slate-400 uppercase'>Total Invoices</div>
        <div class='text-2xl font-black text-slate-800'>{t.get('total_count',0)}</div>
      </div>
      <div class='card py-3 text-center'>
        <div class='text-xs font-bold text-slate-400 uppercase'>Overdue</div>
        <div class='text-2xl font-black text-rose-600'>{t.get('overdue_count',0)}</div>
        <div class='text-xs text-rose-400 mono'>£{t.get('overdue_val',0):,.2f}</div>
      </div>
      <div class='card py-3 text-center'>
        <div class='text-xs font-bold text-slate-400 uppercase'>Total Paid (YTD)</div>
        <div class='text-2xl font-black text-emerald-600'>£{t.get('paid_val',0):,.2f}</div>
      </div>
    </div>"""

    # ── Search & filter bar ──
    status_opts = ""
    for val, label in [("","All"),("overdue","Overdue"),("unpaid","Unpaid"),
                        ("partial","Partial"),("paid","Paid")]:
        sel = "selected" if val == status else ""
        status_opts += f"<option value='{val}' {sel}>{label}</option>"

    search_bar = f"""
    <div class='card'>
      <form method='GET' action='/invoices' class='flex flex-wrap gap-3 items-end'>
        <input type='hidden' name='ledger' value='{ledger}'>
        <div style='flex:2;min-width:200px'>
          <label>Search supplier, invoice no. or serial no.</label>
          <input type='text' name='search' value='{search}'
            placeholder='e.g. Bestway, INV-001, 42...'>
        </div>
        <div style='min-width:130px'>
          <label>Status</label>
          <select name='status'>{status_opts}</select>
        </div>
        <div style='display:flex;gap:8px;align-items:flex-end'>
          <button type='submit' class='btn-primary'>🔍 Search</button>
          <a href='/invoices?ledger={ledger}' class='btn-secondary'>✕ Clear</a>
        </div>
      </form>
    </div>"""

    # ── Add / Edit form ──
    inv    = edit_inv or {}
    is_edit = bool(edit_inv)
    form_action = f"/invoices/save/{edit_id}" if is_edit else "/invoices/save/0"
    form_title  = f"✏️ Edit Invoice — {inv.get('supplier_name','')} {inv.get('invoice_number','')}" if is_edit else "➕ New Invoice"
    cancel_url  = f"/invoices?ledger={ledger}"

    def fi(name, label, ftype="text", val=None, req=False, opts=None, placeholder=""):
        """Render a form field."""
        safe_val = val if val is not None else ""
        req_attr = "required" if req else ""
        step     = "step='0.01'" if ftype == "number" else ""
        ph       = f"placeholder='{placeholder}'" if placeholder else ""
        if opts is not None:
            o_html = ""
            for ov, ol in opts:
                sel = "selected" if str(safe_val) == str(ov) else ""
                o_html += f"<option value='{ov}' {sel}>{ol}</option>"
            return f"<div><label>{label}</label><select name='{name}' {req_attr}>{o_html}</select></div>"
        return f"<div><label>{label}</label><input type='{ftype}' name='{name}' value='{safe_val}' {req_attr} {step} {ph}></div>"

    # Payment status fields (only show if editing)
    payment_fields = ""
    if is_edit:
        paid_opts  = [("No","Unpaid"),("Yes","Paid")]
        meth_opts  = [(m, m or "-- Select --") for m in PAYMENT_METHODS]
        balance    = (inv.get("gross_amount") or 0) - (inv.get("amount_paid") or 0) - (inv.get("credit_note") or 0)
        payment_fields = f"""
        <div class='col-span-2' style='border-top:1px solid #e2e8f0;padding-top:12px;margin-top:4px'>
          <div class='text-xs font-bold text-slate-500 uppercase tracking-wide mb-3'>Payment Details</div>
          <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(150px,1fr))'>
            {fi('is_paid',        'Status',          opts=paid_opts,  val=inv.get('is_paid','No'))}
            {fi('paid_date',      'Paid Date',        'date',          inv.get('paid_date',''))}
            {fi('payment_method', 'Payment Method',   opts=meth_opts,  val=inv.get('payment_method',''))}
            {fi('amount_paid',    'Amount Paid (£)',  'number',        inv.get('amount_paid',0))}
            {fi('credit_note',    'Credit Note (£)',  'number',        inv.get('credit_note',0))}
          </div>
          <div class='text-xs text-slate-400 mt-2 mono'>
            Balance outstanding: <strong class='{'text-rose-600' if balance > 0 else 'text-emerald-600'}'>£{balance:,.2f}</strong>
          </div>
        </div>"""

    # Property-specific field
    prop_or_store_field = ""
    if is_prop:
        prop_or_store_field = fi('expense_type', 'Expense Type',
            opts=[(e, e or "-- Select --") for e in [""] + EXPENSE_TYPES],
            val=inv.get('expense_type',''))
    
    # Seq no (retail only)
    seq_field = "" if is_prop else fi('seq_no','Serial No.','number', inv.get('seq_no',''))

    form_html = f"""
    <div class='card' id='invoice-form'>
      <!-- PDF Upload — one file does both: auto-fills fields AND saves with invoice -->
      <div style='background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;padding:12px 16px;margin-bottom:16px'>
        <div style='font-size:13px;font-weight:700;color:#0369a1;margin-bottom:8px'>
          📎 Attach Invoice PDF
          <span style='font-weight:400;color:#64748b;font-size:12px;margin-left:8px'>
            — uploads once, auto-fills fields AND saves the PDF with the record
          </span>
        </div>
        <div style='display:flex;gap:10px;align-items:center;flex-wrap:wrap'>
          <input type='file' name='pdf_file' id='pdf_prefill' accept='.pdf'
            onchange='extractPdf()'
            style='flex:1;min-width:200px;border:1px solid #bae6fd;background:white;padding:5px 10px;border-radius:8px;font-size:13px'>
          <span id='pdf_status' style='font-size:12px;color:#0369a1'></span>
        </div>
        <div style='font-size:11px;color:#94a3b8;margin-top:6px'>
          Fields auto-fill from the PDF where possible. Check and adjust anything that looks wrong before saving.
        </div>
      </div>
      <div class='flex justify-between items-center mb-4'>
        <div class='font-black text-slate-800'>{form_title}</div>
        {'<a href="' + cancel_url + '" class="btn-secondary text-xs">✕ Cancel Edit</a>' if is_edit else ''}
      </div>
      <form action='{form_action}' method='POST' enctype='multipart/form-data'>
        <input type='hidden' name='ledger' value='{ledger}'>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(180px,1fr))'>
          {seq_field}
          {fi('supplier_name',  'Supplier Name',    val=inv.get('supplier_name',''),  req=True)}
          {fi('invoice_number', 'Invoice Number',   val=inv.get('invoice_number',''))}
          {fi('invoice_date',   'Invoice Date',     'date', inv.get('invoice_date',''))}
          {fi('due_date',       'Due Date',         'date', inv.get('due_date',''))}
          {fi('gross_amount',   'Gross Amount (£)', 'number', inv.get('gross_amount',0))}
          {fi('vat_amount',     'VAT Amount (£)',   'number', inv.get('vat_amount',0))}
          {fi('net_amount',     'Net Amount (£)',   'number', inv.get('net_amount',0))}
          {fi('payment_terms',  'Terms (days)',     'number', inv.get('payment_terms',''))}
          {prop_or_store_field}
          <!-- PDF attached via the strip above -->
          <div style='grid-column:1/-1'>
            {fi('comments','Comments', val=inv.get('comments',''))}
          </div>
          {payment_fields}
        </div>
        <div class='flex gap-3 mt-4'>
          <button type='submit' class='btn-primary'>{'💾 Update Invoice' if is_edit else '➕ Save Invoice'}</button>
          {'<a href="/invoices/delete/' + str(edit_id) + '?ledger=' + ledger + '" class="btn-danger" onclick=\"return confirm(\'Delete this invoice?\');\">🗑️ Delete</a>' if is_edit else ''}
          <a href='{cancel_url}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>"""

    # ── Invoice list ──
    rows_html = ""
    for row in invoices:
        paid    = row["amount_paid"]  or 0
        credit  = row["credit_note"]  or 0 if not is_prop else 0
        balance = row["balance"]      or 0
        today_s = datetime.now().strftime("%Y-%m-%d")

        approval = row.get("approval_status", "approved")
        if approval == "pending":
            badge = "<span style='background:#fef3c7;color:#92400e;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>⏳ PENDING</span>"
            row_cls = "style='background:#fffbeb'"
        elif row["is_paid"] == "Yes":
            badge = "<span class='badge-paid'>PAID</span>"
            row_cls = ""
        elif row["due_date"] and row["due_date"] < today_s:
            badge = "<span class='badge-overdue'>OVERDUE</span>"
            row_cls = "style='background:#fff5f5'"
        elif paid > 0:
            badge = "<span class='badge-partial'>PARTIAL</span>"
            row_cls = "style='background:#fffbeb'"
        else:
            badge = "<span class='badge-unpaid'>UNPAID</span>"
            row_cls = ""

        seq_td = f"<td class='mono' style='color:#94a3b8;font-size:11px'>{row['seq_no'] or ''}</td>" if not is_prop else ""
        pdf_td = ""
        if row.get("pdf_path"):
            inv_pdf_url = f"/invoices/pdf/{row['invoice_id']}?ledger={ledger}"
        if row.get("pdf_path"):
            inv_id  = row['invoice_id']
            pdf_url = f'/invoices/pdf/{inv_id}?ledger={ledger}'
            pdf_td  = '<a href="#" onclick="event.stopPropagation();showPdf('' + pdf_url + '');return false;" style="color:#1e3a5f;font-size:11px;font-weight:700">&#128206; View</a>'

        # Approve/reject buttons for pending invoices (managers/owners only)
        approval_td = ""
        row_approval = row.get("approval_status", "approved")
        if row_approval == "pending" and user["role"] in ("owner","manager"):
            approval_td = f"""
            <a href='/invoices/approve/{row['invoice_id']}?ledger={ledger}'
               style='background:#dcfce7;color:#16a34a;font-size:11px;font-weight:700;
                      padding:3px 8px;border-radius:6px;text-decoration:none;margin-right:4px'
               onclick='event.stopPropagation()'>✅ Approve</a>
            <a href='/invoices/reject/{row['invoice_id']}?ledger={ledger}'
               style='background:#fee2e2;color:#dc2626;font-size:11px;font-weight:700;
                      padding:3px 8px;border-radius:6px;text-decoration:none'
               onclick='event.stopPropagation()'
               onclick="return confirm('Reject this invoice?')">❌ Reject</a>"""

        rows_html += f"""
        <tr {row_cls} onclick="selectInvoice({row['invoice_id']}, '{ledger}')"
            style='cursor:pointer' id='row-{row['invoice_id']}'>
          {seq_td}
          <td style='font-weight:700;color:#0f172a'>{row['supplier_name']}</td>
          <td class='mono' style='font-size:12px'>{row['invoice_number'] or '—'}</td>
          <td class='mono' style='font-size:12px;color:#64748b'>{row['invoice_date'] or '—'}</td>
          <td class='mono' style='font-size:12px;color:#64748b'>{row['due_date'] or '—'}</td>
          <td class='mono' style='font-weight:700'>£{row['gross_amount']:,.2f}</td>
          <td class='mono' style='color:#16a34a'>{'£'+f'{paid:,.2f}' if paid else '—'}</td>
          <td class='mono' style='font-weight:700;color:{"#dc2626" if balance > 0 else "#16a34a"}'>£{balance:,.2f}</td>
          <td>{badge}</td>
          <td style='font-size:12px;color:#64748b'>{row['payment_method'] or '—'}</td>
          <td>{pdf_td}</td>
          <td>{approval_td}</td>
        </tr>"""

    # Seq header only for retail
    seq_th = "<th>Serial</th>" if not is_prop else ""

    # Pagination
    pag_html = ""
    if total_pages > 1:
        base = f"/invoices?ledger={ledger}&search={search}&status={status}&page="
        pag_html = "<div class='flex gap-2 flex-wrap justify-center'>"
        for p in range(1, total_pages + 1):
            cls = "btn-primary" if p == pg else "btn-secondary"
            pag_html += f"<a href='{base}{p}' class='{cls}' style='padding:6px 14px'>{p}</a>"
        pag_html += "</div>"

    list_html = f"""
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:16px 20px;background:#0f2942;display:flex;justify-content:space-between;align-items:center'>
        <div style='color:white;font-weight:700;font-size:14px'>
          {total_n} invoices
          {'· <span style="color:#fbbf24">'+str(t.get('overdue_count',0))+' overdue</span>' if t.get('overdue_count',0) > 0 else ''}
        </div>
        <div style='color:#93c5fd;font-size:12px'>Click any row to edit</div>
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead>
            <tr>
              {seq_th}
              <th>Supplier</th><th>Invoice No.</th><th>Inv. Date</th>
              <th>Due Date</th><th>Gross</th><th>Paid</th>
              <th>Balance</th><th>Status</th><th>Method</th><th>PDF</th>
            </tr>
          </thead>
          <tbody>{rows_html if rows_html else "<tr><td colspan='10' style='text-align:center;padding:32px;color:#94a3b8'>No invoices found</td></tr>"}</tbody>
        </table>
      </div>
    </div>
    {pag_html}"""

    # ── JS: click row to scroll to form and load edit ──
    js = """
    <script>
    function selectInvoice(id, ledger) {
      document.querySelectorAll('.tbl tbody tr').forEach(r => r.style.outline = '');
      const row = document.getElementById('row-' + id);
      if (row) row.style.outline = '2px solid #1e3a5f';
      window.location.href = '/invoices?ledger=' + ledger + '&edit_id=' + id + '#invoice-form';
    }

    // ── Smart field calculations ──
    document.addEventListener('DOMContentLoaded', function() {
      const gross = document.querySelector('[name="gross_amount"]');
      const vat   = document.querySelector('[name="vat_amount"]');
      const net   = document.querySelector('[name="net_amount"]');
      const idate = document.querySelector('[name="invoice_date"]');
      const ddate = document.querySelector('[name="due_date"]');
      const terms = document.querySelector('[name="payment_terms"]');

      // Auto-calc net = gross - vat
      function recalcNet() {
        if (gross && vat && net) {
          const g = parseFloat(gross.value) || 0;
          const v = parseFloat(vat.value)   || 0;
          // Only auto-fill net if it's empty or was previously auto-filled
          if (!net.dataset.manual) net.value = (g - v).toFixed(2);
        }
      }
      // Auto-calc VAT = gross * 1/6 (standard 20% VAT on gross)
      function recalcVat() {
        if (gross && vat && !vat.dataset.manual) {
          const g = parseFloat(gross.value) || 0;
          vat.value = (g / 6).toFixed(2);
          recalcNet();
        }
      }
      // Auto-calc due date from invoice date + terms
      function recalcDueDate() {
        if (idate && ddate && terms && idate.value && terms.value && !ddate.dataset.manual) {
          const d = new Date(idate.value);
          d.setDate(d.getDate() + parseInt(terms.value));
          ddate.value = d.toISOString().split('T')[0];
        }
      }

      if (gross) gross.addEventListener('input', recalcVat);
      if (vat)   { vat.addEventListener('input',   () => { vat.dataset.manual='1'; recalcNet(); }); }
      if (net)   { net.addEventListener('input',   () => { net.dataset.manual='1'; }); }
      if (idate) idate.addEventListener('change', recalcDueDate);
      if (terms) terms.addEventListener('input',  recalcDueDate);
      if (ddate) ddate.addEventListener('input',  () => { ddate.dataset.manual='1'; });
    });

    // ── PDF auto-fill ──
    async function extractPdf() {
      const fileInput = document.getElementById('pdf_prefill');
      const status    = document.getElementById('pdf_status');
      if (!fileInput.files.length) return;
      // File is already attached to the form — just extract the data
      status.textContent = '⏳ Reading PDF...';
      const formData = new FormData();
      formData.append('pdf_file', fileInput.files[0]);
      try {
        const resp = await fetch('/invoices/extract-pdf', { method:'POST', body:formData });
        const data = await resp.json();
        if (data.error) { status.textContent = '❌ ' + data.error; return; }

        // Fill fields if found in PDF
        const fill = (name, val) => {
          const el = document.querySelector('[name="' + name + '"]');
          if (el && val !== undefined && val !== null && val !== '') {
            el.value = val;
            el.style.background = '#f0fdf4';  // green tint = auto-filled
          }
        };
        fill('supplier_name',  data.supplier_name);
        fill('invoice_number', data.invoice_number);
        fill('invoice_date',   data.invoice_date);
        fill('gross_amount',   data.gross_amount);
        fill('vat_amount',     data.vat_amount);
        fill('net_amount',     data.net_amount);
        fill('payment_terms',  data.payment_terms);

        // Trigger calculations for any fields NOT found in PDF
        const gross = document.querySelector('[name="gross_amount"]');
        if (gross) gross.dispatchEvent(new Event('input'));
        const idate = document.querySelector('[name="invoice_date"]');
        if (idate) idate.dispatchEvent(new Event('change'));

        let found = Object.keys(data).filter(k => !k.startsWith('_') && data[k]).length;
        status.textContent = '✅ ' + found + ' fields found — please check and adjust as needed';
        status.style.color = '#16a34a';

      } catch(e) {
        status.textContent = '❌ Could not read PDF — please fill manually';
        status.style.color = '#dc2626';
      }
    }
    </script>

    <!-- PDF preview panel (slides in from right) -->
    <div id="pdfPanel" style="display:none;position:fixed;top:0;right:0;width:45%;height:100vh;
         background:white;box-shadow:-4px 0 24px rgba(0,0,0,.15);z-index:1000;flex-direction:column">
      <div style="background:#0f2942;color:white;padding:12px 16px;display:flex;justify-content:space-between;align-items:center">
        <span style="font-weight:700;font-size:14px">📎 Invoice PDF</span>
        <button onclick="closePdf()"
          style="background:rgba(255,255,255,.15);color:white;border:none;border-radius:6px;
                 padding:4px 12px;cursor:pointer;font-weight:700">✕ Close</button>
      </div>
      <iframe id="pdfFrame" src="" style="flex:1;width:100%;height:calc(100vh - 48px);border:none"></iframe>
    </div>
    <script>
    function showPdf(url) {
      document.getElementById('pdfFrame').src = url;
      const panel = document.getElementById('pdfPanel');
      panel.style.display = 'flex';
      // Shrink main content to make room
      document.querySelector('.ml-52').style.marginRight = '45%';
    }
    function closePdf() {
      document.getElementById('pdfPanel').style.display = 'none';
      document.querySelector('.ml-52').style.marginRight = '0';
      document.getElementById('pdfFrame').src = '';
    }
    </script>"""

    # ── Pending approvals banner ──
    if pending_count > 0:
        flash += f"""<div style='background:#fef3c7;border:1px solid #fbbf24;border-radius:10px;
            padding:12px 16px;display:flex;justify-content:space-between;align-items:center'>
          <span style='font-weight:700;color:#92400e'>
            ⏳ {pending_count} invoice{'s' if pending_count>1 else ''} awaiting your approval
          </span>
          <a href='/invoices?ledger={ledger}&status=pending'
             style='background:#d97706;color:white;font-weight:700;padding:6px 14px;
                    border-radius:8px;font-size:13px;text-decoration:none'>
            Review Now →
          </a>
        </div>"""

    # ── Ledger switcher ──
    ledger_switcher = f"""
    <div class='flex flex-wrap gap-2 items-center'>
      <div class='text-xl font-black text-slate-800'>🧾 Invoice Manager</div>
      <select onchange="window.location='/invoices?ledger='+this.value"
        style='border:1px solid #e2e8f0;border-radius:8px;padding:6px 12px;font-size:14px;font-weight:600;max-width:260px'>
        {ledger_opts}
      </select>
      <a href='/invoices/recent-payments' class='btn-secondary' style='margin-left:auto'>📋 Recent Payments</a>
    </div>"""

    content = "\n".join([flash, ledger_switcher, summary, search_bar, form_html, list_html, js])
    return page("Invoices", content, user, "invoices")


# ── Save invoice (add or update) ──────────────────────────────────────────────

@app.post("/invoices/save/{invoice_id}")
async def save_invoice(
    request:    Request,
    invoice_id: int,
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    form   = await request.form()
    ledger = form.get("ledger", "Uxbridge")
    is_prop = is_property_ledger(ledger)
    table   = "property_invoices" if is_prop else "supplier_invoices"
    loc_col = "property_name"     if is_prop else "store_name"
    loc_val = prop_name(ledger)   if is_prop else ledger

    def fv(key, default=""):
        v = form.get(key, default)
        return v.strip() if isinstance(v, str) else v

    def fnum(key):
        try: return float(form.get(key, 0) or 0)
        except: return 0.0

    def fint(key):
        try: return int(form.get(key, 0) or 0)
        except: return None

    # Handle PDF upload
    pdf_path = None
    pdf_file = form.get("pdf_file")
    if pdf_file and hasattr(pdf_file, "filename") and pdf_file.filename:
        ext      = os.path.splitext(pdf_file.filename)[1].lower()
        filename = f"{uuid.uuid4().hex}{ext}"
        full_path = os.path.join(UPLOAD_DIR, filename)
        with open(full_path, "wb") as f:
            f.write(await pdf_file.read())
        pdf_path = full_path

    supplier   = fv("supplier_name")
    inv_no     = fv("invoice_number")
    inv_date   = fv("invoice_date") or None
    due_date   = fv("due_date")     or None
    gross      = fnum("gross_amount")
    vat        = fnum("vat_amount")
    net        = fnum("net_amount")
    terms      = fint("payment_terms")
    comments   = fv("comments")     or None
    is_paid    = fv("is_paid", "No")
    paid_date  = fv("paid_date")    or None
    pay_method = fv("payment_method") or None
    amt_paid   = fnum("amount_paid")
    credit     = fnum("credit_note")
    seq_no     = fint("seq_no")
    exp_type   = fv("expense_type") or None

    if not supplier:
        return RedirectResponse(f"/invoices?ledger={ledger}&msg=Supplier+name+is+required&msg_type=error",
                                status_code=303)

    from urllib.parse import quote as urlquote

    # ── Approval status based on role ──
    role = user.get("role", "staff")
    approval_status = "approved" if role in ("owner", "manager") else "pending"
    submitted_by    = user.get("username", "")

    # ── Duplicate check (supplier + invoice_number + store, warn only) ──
    force = fv("force_save")
    if invoice_id == 0 and inv_no and not force:
        dup = q(f"SELECT invoice_id, supplier_name FROM {table} WHERE {loc_col}=? AND supplier_name=? AND invoice_number=?",
                (loc_val, supplier, inv_no), fetch=True)
        if dup:
            # Return duplicate warning page
            warn_url = f"/invoices?ledger={ledger}&edit_id={dup[0]['invoice_id']}"
            return HTMLResponse(f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;700;900&display=swap" rel="stylesheet">
<style>body{{font-family:'DM Sans',sans-serif;background:#f8fafc;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0}}</style>
</head><body>
<div style='background:white;border-radius:20px;padding:40px;max-width:480px;width:90%;border:2px solid #fbbf24;box-shadow:0 8px 32px rgba(0,0,0,.08)'>
  <div style='font-size:40px;text-align:center;margin-bottom:16px'>⚠️</div>
  <h2 style='font-weight:900;color:#92400e;text-align:center;margin:0 0 8px'>Possible Duplicate Invoice</h2>
  <p style='color:#64748b;font-size:14px;text-align:center;margin:0 0 20px'>
    An invoice from <strong>{supplier}</strong> with number <strong>{inv_no}</strong>
    already exists in {loc_val}.
  </p>
  <div style='background:#fef3c7;border-radius:10px;padding:12px 16px;font-size:13px;color:#92400e;margin-bottom:24px'>
    This may be a genuine duplicate. Check the existing record before saving again.
  </div>
  <div style='display:flex;flex-direction:column;gap:10px'>
    <a href='{warn_url}' style='background:#1e3a5f;color:white;font-weight:700;padding:12px;border-radius:10px;text-align:center;text-decoration:none;font-size:14px'>
      👁️ View Existing Invoice
    </a>
    <form method='POST' action='/invoices/save/0'>
      <input type='hidden' name='ledger'          value='{ledger}'>
      <input type='hidden' name='supplier_name'   value='{supplier}'>
      <input type='hidden' name='invoice_number'  value='{inv_no}'>
      <input type='hidden' name='invoice_date'    value='{fv("invoice_date")}'>
      <input type='hidden' name='due_date'        value='{fv("due_date")}'>
      <input type='hidden' name='gross_amount'    value='{gross}'>
      <input type='hidden' name='vat_amount'      value='{vat}'>
      <input type='hidden' name='net_amount'      value='{net}'>
      <input type='hidden' name='payment_terms'   value='{terms or ""}'>
      <input type='hidden' name='comments'        value='{fv("comments")}'>
      <input type='hidden' name='seq_no'          value='{seq_no or ""}'>
      <input type='hidden' name='force_save'      value='1'>
      <button type='submit' style='width:100%;background:#dc2626;color:white;font-weight:700;padding:12px;border-radius:10px;font-size:14px;border:none;cursor:pointer'>
        ⚠️ Save Anyway (Different Supplier?)
      </button>
    </form>
    <a href='/invoices?ledger={ledger}' style='color:#64748b;text-align:center;font-size:13px;text-decoration:none'>← Cancel, go back</a>
  </div>
</div>
</body></html>""")

    # ── Validation warnings (non-blocking, stored as comment note) ──
    warnings = []
    if gross > 0 and vat > 0:
        expected_vat = round(gross / 6, 2)
        if abs(vat - expected_vat) > 1.0:
            warnings.append(f"VAT £{vat:.2f} doesn't match standard 20% (expected ~£{expected_vat:.2f})")
    if gross > 10000:
        warnings.append(f"Large invoice amount: £{gross:,.2f} — please double-check")
    if due_date and due_date < datetime.now().strftime("%Y-%m-%d"):
        warnings.append("Due date is in the past")
    warning_note = (" | WARNINGS: " + "; ".join(warnings)) if warnings else ""
    if warning_note and comments:
        comments = comments + warning_note
    elif warning_note:
        comments = warning_note.strip(" | ")

    if invoice_id == 0:
        # New invoice
        if is_prop:
            q(f"""INSERT OR IGNORE INTO {table}
                (property_name, supplier_name, invoice_number, invoice_date,
                 expense_type, gross_amount, vat_amount, net_amount, due_date,
                 payment_terms, comments, is_paid, pdf_path)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
              (loc_val, supplier, inv_no, inv_date, exp_type,
               gross, vat, net, due_date, terms, comments, is_paid, pdf_path))
        else:
            q(f"""INSERT OR IGNORE INTO {table}
                (store_name, seq_no, supplier_name, invoice_number, invoice_date,
                 gross_amount, vat_amount, net_amount, due_date, payment_terms,
                 comments, is_paid, pdf_path, approval_status, submitted_by)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
              (loc_val, seq_no, supplier, inv_no, inv_date,
               gross, vat, net, due_date, terms, comments, is_paid, pdf_path,
               approval_status, submitted_by))
        if approval_status == "pending":
            msg = f"Invoice submitted for approval — {supplier} {inv_no}"
        else:
            msg = f"Invoice added — {supplier} {inv_no}"
    else:
        # Update existing
        if is_prop:
            q(f"""UPDATE {table} SET
                supplier_name=?, invoice_number=?, invoice_date=?,
                expense_type=?, gross_amount=?, vat_amount=?, net_amount=?,
                due_date=?, payment_terms=?, comments=?, is_paid=?,
                paid_date=?, payment_method=?, amount_paid=?,
                {', pdf_path=?' if pdf_path else ''}
                WHERE invoice_id=?""",
              ([supplier, inv_no, inv_date, exp_type, gross, vat, net,
                due_date, terms, comments, is_paid, paid_date, pay_method, credit]
               + ([pdf_path] if pdf_path else []) + [invoice_id]))
        else:
            q(f"""UPDATE {table} SET
                seq_no=?, supplier_name=?, invoice_number=?, invoice_date=?,
                gross_amount=?, vat_amount=?, net_amount=?,
                due_date=?, payment_terms=?, comments=?, is_paid=?,
                paid_date=?, payment_method=?, amount_paid=?, credit_note=?
                {', pdf_path=?' if pdf_path else ''}
                WHERE invoice_id=?""",
              ([seq_no, supplier, inv_no, inv_date, gross, vat, net,
                due_date, terms, comments, is_paid, paid_date,
                pay_method, amt_paid, credit]
               + ([pdf_path] if pdf_path else []) + [invoice_id]))
        msg = f"Invoice updated — {supplier} {inv_no}"

    from urllib.parse import quote as urlquote
    return RedirectResponse(
        f"/invoices?ledger={ledger}&msg={urlquote(msg)}&msg_type=success#invoice-form",
        status_code=303)


# ── Delete invoice ────────────────────────────────────────────────────────────

@app.get("/invoices/delete/{invoice_id}")
def delete_invoice(
    invoice_id: int,
    ledger:     str = "Uxbridge",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse(f"/invoices?ledger={ledger}&msg=Not+authorised&msg_type=error",
                                status_code=303)
    table = "property_invoices" if is_property_ledger(ledger) else "supplier_invoices"
    q(f"DELETE FROM {table} WHERE invoice_id=?", (invoice_id,))
    from urllib.parse import quote as urlquote
    return RedirectResponse(
        f"/invoices?ledger={ledger}&msg={urlquote('Invoice deleted')}&msg_type=success",
        status_code=303)


# ── Serve PDF ─────────────────────────────────────────────────────────────────

@app.get("/invoices/pdf/{invoice_id}")
def serve_pdf(
    invoice_id: int,
    ledger:     str = "Uxbridge",
    session:    str | None = Cookie(default=None)
):
    from fastapi.responses import FileResponse
    redir, user = require_login(session)
    if redir: return redir
    table = "property_invoices" if is_property_ledger(ledger) else "supplier_invoices"
    rows  = q(f"SELECT pdf_path FROM {table} WHERE invoice_id=?", (invoice_id,), fetch=True)
    if rows and rows[0]["pdf_path"] and os.path.exists(rows[0]["pdf_path"]):
        return FileResponse(rows[0]["pdf_path"], media_type="application/pdf")
    return HTMLResponse("<p>PDF not found</p>", status_code=404)


# ── PDF extract endpoint (AJAX) ──────────────────────────────────────────────

from fastapi.responses import JSONResponse

@app.post("/invoices/extract-pdf")
async def extract_pdf_ajax(request: Request, session: str | None = Cookie(default=None)):
    """Receive a PDF upload, extract fields, return JSON for JS to fill the form."""
    redir, user = require_login(session)
    if redir: return JSONResponse({"error": "Not logged in"}, status_code=401)
    form = await request.form()
    pdf_file = form.get("pdf_file")
    if not pdf_file or not hasattr(pdf_file, "read"):
        return JSONResponse({"error": "No file"})
    data = extract_pdf_data(await pdf_file.read())
    return JSONResponse(data)


# ── Recent payments ───────────────────────────────────────────────────────────

# ── Approve / Reject invoice ─────────────────────────────────────────────────

@app.get("/invoices/approve/{invoice_id}")
def approve_invoice(
    invoice_id: int,
    ledger:     str = "Uxbridge",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse(f"/invoices?ledger={ledger}&msg=Not+authorised&msg_type=error",
                                status_code=303)
    table = "property_invoices" if is_property_ledger(ledger) else "supplier_invoices"
    q(f"UPDATE {table} SET approval_status='approved' WHERE invoice_id=?", (invoice_id,))
    from urllib.parse import quote as urlquote
    return RedirectResponse(
        f"/invoices?ledger={ledger}&msg={urlquote('Invoice approved ✅')}&msg_type=success",
        status_code=303)


@app.get("/invoices/reject/{invoice_id}")
def reject_invoice(
    invoice_id: int,
    ledger:     str = "Uxbridge",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse(f"/invoices?ledger={ledger}&msg=Not+authorised&msg_type=error",
                                status_code=303)
    table = "property_invoices" if is_property_ledger(ledger) else "supplier_invoices"
    q(f"UPDATE {table} SET approval_status='rejected' WHERE invoice_id=?", (invoice_id,))
    from urllib.parse import quote as urlquote
    return RedirectResponse(
        f"/invoices?ledger={ledger}&msg={urlquote('Invoice rejected and flagged')}&msg_type=error",
        status_code=303)


@app.get("/invoices/recent-payments", response_class=HTMLResponse)
def recent_payments(session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    from collections import defaultdict

    rows = q("""
        SELECT 'retail' as ledger_type, store_name as location,
               supplier_name, invoice_number, gross_amount,
               amount_paid, credit_note, paid_date, payment_method, is_paid,
               COALESCE(gross_amount,0)-COALESCE(amount_paid,0)-COALESCE(credit_note,0) as balance
        FROM supplier_invoices
        WHERE paid_date IS NOT NULL OR amount_paid > 0
        UNION ALL
        SELECT 'property', property_name,
               supplier_name, invoice_number, gross_amount,
               amount_paid, 0, paid_date, payment_method, is_paid,
               COALESCE(gross_amount,0)-COALESCE(amount_paid,0) as balance
        FROM property_invoices
        WHERE paid_date IS NOT NULL OR amount_paid > 0
        ORDER BY paid_date DESC
        LIMIT 200
    """, fetch=True) or []

    by_date = defaultdict(list)
    for r in rows:
        by_date[r["paid_date"] or "Unknown"].append(r)

    rows_html = ""
    for date_key in sorted(by_date.keys(), reverse=True):
        day_rows  = by_date[date_key]
        day_total = sum(r["amount_paid"] or 0 for r in day_rows)
        rows_html += f"""
        <tr style='background:#f8fafc'>
          <td colspan='8' style='font-weight:900;color:#0f2942;padding:10px 12px;font-size:13px'>
            📅 {date_key}
            <span style='float:right;color:#16a34a;font-weight:700'>Day total: £{day_total:,.2f}</span>
          </td>
        </tr>"""
        for r in day_rows:
            paid    = r["amount_paid"] or 0
            balance = r["balance"]     or 0
            status  = "PAID" if r["is_paid"] == "Yes" else f"Outstanding £{balance:,.2f}"
            status_cls = "badge-paid" if r["is_paid"] == "Yes" else "badge-partial"
            rows_html += f"""
            <tr>
              <td style='font-size:11px;color:#94a3b8'>{r['location']}</td>
              <td style='font-weight:700'>{r['supplier_name']}</td>
              <td class='mono' style='font-size:12px'>{r['invoice_number'] or '—'}</td>
              <td class='mono'>£{r['gross_amount']:,.2f}</td>
              <td class='mono' style='color:#16a34a;font-weight:700'>£{paid:,.2f}</td>
              <td style='font-size:12px;color:#64748b'>{r['payment_method'] or '—'}</td>
              <td><span class='{status_cls}'>{status}</span></td>
            </tr>"""

    content = f"""
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>📋 Recent Payments</div>
      <a href='/invoices' class='btn-secondary'>← Back to Invoices</a>
    </div>
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead>
            <tr>
              <th>Store/Property</th><th>Supplier</th><th>Invoice No.</th>
              <th>Gross</th><th>Paid</th><th>Method</th><th>Status</th>
            </tr>
          </thead>
          <tbody>
            {rows_html or '<tr><td colspan="7" style="text-align:center;padding:32px;color:#94a3b8">No payments recorded yet</td></tr>'}
          </tbody>
        </table>
      </div>
    </div>"""
    return page("Recent Payments", content, user, "invoices")


# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3 — STAFF MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

import math

# UK Bank Holidays 2026 (hardcoded, will be managed via settings later)
UK_BANK_HOLIDAYS_2026 = [
    "2026-01-01", "2026-04-03", "2026-04-06",
    "2026-05-04", "2026-05-25", "2026-08-31",
    "2026-12-25", "2026-12-28"
]

ABSENCE_TYPES = {
    "H":  "Holiday",
    "S":  "Sick",
    "B":  "Bank Holiday",
    "L":  "Lateness",
    "AL": "Authorised Leave",
    "UL": "Unauthorised Leave",
    "MA": "Maternity",
    "PA": "Paternity",
    "JP": "Jury Service",
    "TO": "TOIL",
    "WFH":"Working From Home",
}

def calc_entitlement(contracted_hrs: float) -> float:
    """5.6 weeks × contracted hours per week, including bank holidays."""
    if not contracted_hrs:
        return 0.0
    return round(5.6 * contracted_hrs, 1)  # entitlement in hours

def hrs_to_days(hrs: float, contracted_hrs: float) -> float:
    """Convert hours to days based on contracted daily hours."""
    if not contracted_hrs or contracted_hrs == 0:
        return 0.0
    daily = contracted_hrs / 5
    return round(hrs / daily, 1) if daily > 0 else 0.0

def is_full_time(contracted_hrs: float) -> bool:
    return contracted_hrs is not None and contracted_hrs >= 30.0

def fmt_entitlement(hrs: float, contracted_hrs: float) -> str:
    """Format entitlement as days for FT (>=30h), hours for PT (<30h)."""
    if not contracted_hrs:
        return "—"
    if is_full_time(contracted_hrs):
        days = hrs_to_days(hrs, contracted_hrs)
        return f"{days} days"
    else:
        return f"{hrs} hrs ({hrs_to_days(hrs, contracted_hrs)} days)"

def get_leave_summary(staff_id: int, year: int = None) -> dict:
    """Return entitlement, taken, balance for a staff member."""
    if year is None:
        year = datetime.now().year
    staff = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not staff:
        return {}
    s = dict(staff[0])
    contracted = s.get("contracted_hrs") or 0

    # Check for custom entitlement first
    custom = q("""SELECT effective_days, statutory_days FROM leave_entitlements
                  WHERE staff_id=? AND year=?""", (staff_id, year), fetch=True)
    if custom:
        c = dict(custom[0])
        effective_days   = c["effective_days"] or c["statutory_days"] or 0
        entitlement_hrs  = effective_days * (contracted/5) if contracted else effective_days
    else:
        entitlement_hrs = calc_entitlement(contracted)

    # Days taken this year by type
    daily = contracted / 5 if contracted else 7.5

    taken_h = q("""SELECT COUNT(*) as n FROM leave_requests
                   WHERE staff_id=? AND status='approved'
                   AND leave_type='H' AND strftime('%Y',date_from)=?""",
                (staff_id, str(year)), fetch=True)
    holiday_days = taken_h[0]["n"] if taken_h else 0
    taken_hrs    = holiday_days * daily

    bh_used = q("""SELECT COUNT(*) as n FROM leave_requests
                   WHERE staff_id=? AND status='approved'
                   AND leave_type='B' AND strftime('%Y',date_from)=?""",
                (staff_id, str(year)), fetch=True)
    bh_days = bh_used[0]["n"] if bh_used else 0
    bh_hrs  = bh_days * daily

    sick_used = q("""SELECT COUNT(*) as n FROM leave_requests
                     WHERE staff_id=? AND status='approved'
                     AND leave_type='S' AND strftime('%Y',date_from)=?""",
                  (staff_id, str(year)), fetch=True)
    sick_days = sick_used[0]["n"] if sick_used else 0

    balance_hrs = entitlement_hrs - taken_hrs - bh_hrs
    # Sick days tracked separately — don't affect holiday balance
    daily = contracted / 5 if contracted else 7.5

    ft = is_full_time(contracted)
    return {
        "entitlement_hrs":  entitlement_hrs,
        "entitlement_days": hrs_to_days(entitlement_hrs, contracted),
        "entitlement_fmt":  fmt_entitlement(entitlement_hrs, contracted),
        "taken_hrs":        round(taken_hrs, 1),
        "taken_days":       holiday_days,
        "taken_fmt":        f"{holiday_days} days",
        "bh_days":          bh_days,
        "bh_hrs":           round(bh_hrs, 1),
        "sick_days":        sick_days,
        "balance_hrs":      round(balance_hrs, 1),
        "balance_days":     hrs_to_days(balance_hrs, contracted),
        "balance_fmt":      fmt_entitlement(round(balance_hrs,1), contracted),
        "daily_hrs":        round(daily, 2),
        "contracted_hrs":   contracted,
        "is_full_time":     ft,
    }


# ── Add leave_requests table if not exists ────────────────────────────────────

def ensure_staff_tables():
    conn = db()
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS leave_requests (
            request_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            leave_type    TEXT NOT NULL DEFAULT 'H',
            date_from     TEXT NOT NULL,
            date_to       TEXT NOT NULL,
            days_taken    REAL DEFAULT 1,
            status        TEXT DEFAULT 'pending',
            requested_by  TEXT,
            approved_by   TEXT,
            approved_at   TEXT,
            notes         TEXT,
            created_at    TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_documents (
            doc_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            doc_type      TEXT,
            file_path     TEXT,
            uploaded_at   TEXT DEFAULT (datetime('now')),
            notes         TEXT,
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS leave_entitlements (
            entitlement_id   INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id         INTEGER NOT NULL,
            year             INTEGER NOT NULL,
            statutory_days   REAL,
            custom_days      REAL,
            effective_days   REAL,
            notes            TEXT,
            UNIQUE(staff_id, year),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS document_templates (
            template_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_type       TEXT NOT NULL,
            version        INTEGER DEFAULT 1,
            file_path      TEXT NOT NULL,
            file_name      TEXT,
            is_current     INTEGER DEFAULT 1,
            uploaded_by    TEXT,
            uploaded_at    TEXT DEFAULT (datetime('now')),
            notes          TEXT,
            UNIQUE(doc_type, version)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_documents (
            doc_id         INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id       INTEGER NOT NULL,
            doc_type       TEXT NOT NULL,
            version        INTEGER DEFAULT 1,
            file_path      TEXT NOT NULL,
            file_name      TEXT,
            is_current     INTEGER DEFAULT 1,
            generated      INTEGER DEFAULT 0,
            uploaded_by    TEXT,
            uploaded_at    TEXT DEFAULT (datetime('now')),
            notes          TEXT,
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS pay_history (
            pay_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            effective_date TEXT NOT NULL,
            hourly_rate   REAL NOT NULL,
            previous_rate REAL,
            change_reason TEXT,
            recorded_by   TEXT,
            created_at    TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS nmw_rates (
            nmw_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            effective_date TEXT NOT NULL,
            rate_21_plus  REAL,
            rate_18_20    REAL,
            rate_16_17    REAL,
            rate_apprentice REAL,
            UNIQUE(effective_date)
        )
    """)
    # Seed NMW rates from historical data
    nmw_data = [
        ("2026-04-01", 12.71, 10.85, 8.00,  7.55),
        ("2025-04-01", 12.21, 10.00, 7.55,  7.55),
        ("2024-04-01", 11.44,  8.60, 6.40,  6.40),
        ("2023-04-01", 10.42,  7.49, 5.28,  5.28),
        ("2022-04-01",  9.50,  6.83, 4.81,  4.81),
        ("2021-04-01",  8.91,  6.56, 4.62,  4.30),
        ("2020-04-01",  8.72,  6.45, 4.55,  4.15),
        ("2019-04-01",  8.21,  6.15, 4.35,  3.90),
        ("2018-04-01",  7.83,  5.90, 4.20,  3.70),
        ("2017-04-01",  7.50,  5.60, 4.05,  3.50),
    ]
    for row in nmw_data:
        try: c.execute("INSERT OR IGNORE INTO nmw_rates (effective_date,rate_21_plus,rate_18_20,rate_16_17,rate_apprentice) VALUES(?,?,?,?,?)", row)
        except: pass
    conn.commit()
    conn.close()

ensure_staff_tables()

os.makedirs("staff_docs", exist_ok=True)


# ── Staff list page ───────────────────────────────────────────────────────────

@app.get("/staff", response_class=HTMLResponse)
def staff_page(
    session:  str | None = Cookie(default=None),
    store:    str = "",
    show:     str = "active",
    msg:      str = "",
    msg_type: str = "success"
):
    redir, user = require_login(session)
    if redir: return redir
    is_owner = user["role"] == "owner"

    # Build filter
    conds  = []
    params = []
    if show == "active":
        conds.append("is_active = 1")
    elif show == "leavers":
        if not is_owner:
            return RedirectResponse("/staff", status_code=303)
        conds.append("is_active = 0")
    # Store filter
    if store:
        conds.append("store_name = ?")
        params.append(store)
    elif user["role"] == "manager" and user.get("store_name"):
        conds.append("store_name = ?")
        params.append(user["store_name"])

    where = ("WHERE " + " AND ".join(conds)) if conds else ""
    staff = q(f"SELECT * FROM staff_profiles {where} ORDER BY store_name, first_name",
              params, fetch=True) or []

    # Pending leave requests count
    pending_leave = q("""SELECT COUNT(*) as n FROM leave_requests lr
                         JOIN staff_profiles sp ON lr.staff_id=sp.staff_id
                         WHERE lr.status='pending'""", fetch=True)
    pending_n = pending_leave[0]["n"] if pending_leave else 0

    flash = f"<div class='flash-{'success' if msg_type=='success' else 'error'}'>{msg}</div>" if msg else ""

    # ── Tab bar ──
    tabs = ""
    for val, label in [("active","Active Staff"),("leavers","Former Staff (Leavers)")]:
        if val == "leavers" and not is_owner:
            continue
        active_cls = "border-b-2 border-blue-900 font-black text-blue-900" if show == val else "text-slate-500 hover:text-slate-700"
        tabs += f"<a href='/staff?show={val}' class='px-4 py-2 text-sm {active_cls} transition'>{label}</a>"

    # ── Store filter buttons ──
    store_btns = ""
    if is_owner or user["role"] == "manager":
        for sv, sl in [("","Both Stores"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]:
            cls = "btn-primary" if store == sv else "btn-secondary"
            store_btns += f"<a href='/staff?show={show}&store={sv}' class='{cls}' style='padding:6px 14px;font-size:13px'>{sl}</a>"

    # ── Staff cards ──
    cards_html = ""
    year = datetime.now().year
    for s in staff:
        s = dict(s)
        sid   = s["staff_id"]
        name  = f"{s['first_name']} {s['last_name']}"
        store_badge = f"<span style='background:#e0f2fe;color:#0369a1;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>{s.get('store_name','')}</span>"
        status_badge = "<span class='badge-paid'>Active</span>" if s["is_active"] else "<span class='badge-overdue'>Left</span>"

        # Quick leave summary
        leave = get_leave_summary(sid, year)
        bal      = leave.get("balance_days", 0)
        taken    = leave.get("taken_days", 0)
        entit    = leave.get("entitlement_days", 0)
        bal_fmt  = leave.get("balance_fmt", "—")
        tak_fmt  = leave.get("taken_fmt", "—")
        ent_fmt  = leave.get("entitlement_fmt", "—")
        bal_col  = "#16a34a" if bal > 5 else ("#d97706" if bal > 0 else "#dc2626")

        rate = f"£{s['hourly_rate']:.2f}/hr" if s.get("hourly_rate") else "—"
        hrs  = f"{s['contracted_hrs']}h/wk" if s.get("contracted_hrs") else "—"
        joined = s.get("date_joined") or "—"

        cards_html += f"""
        <div class='card' style='padding:0;overflow:hidden'>
          <div style='background:#f8fafc;padding:12px 16px;display:flex;justify-content:space-between;align-items:center;border-bottom:1px solid #e2e8f0'>
            <div>
              <div style='font-weight:900;font-size:15px;color:#0f172a'>{name}</div>
              <div style='display:flex;gap:6px;margin-top:4px'>{store_badge} {status_badge}</div>
            </div>
            <div style='display:flex;gap:8px'>
              <a href='/staff/{sid}' class='btn-primary' style='padding:5px 12px;font-size:12px'>👁 View</a>
              <a href='/staff/{sid}/edit' class='btn-secondary' style='padding:5px 12px;font-size:12px'>✏️ Edit</a>
            </div>
          </div>
          <div style='padding:12px 16px;display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:8px'>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Joined</div>
                 <div style='font-size:13px;font-weight:600;color:#334155'>{joined}</div></div>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Hours</div>
                 <div style='font-size:13px;font-weight:600;color:#334155'>{hrs}</div></div>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Rate</div>
                 <div style='font-size:13px;font-weight:600;color:#334155'>{rate}</div></div>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Leave Balance {year}</div>
                 <div style='font-size:13px;font-weight:700;color:{bal_col}'>{bal_fmt} left <span style='color:#94a3b8;font-weight:400'>({tak_fmt} of {ent_fmt})</span></div></div>
          </div>
        </div>"""

    if not cards_html:
        cards_html = "<div class='card text-center' style='padding:40px;color:#94a3b8'>No staff found</div>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>👤 Staff</div>
      <div style='display:flex;gap:8px;flex-wrap:wrap'>
        {'<a href="/staff/leave-requests" class="btn-secondary" style="position:relative">📋 Leave Requests' + (f'<span style="position:absolute;top:-6px;right:-6px;background:#dc2626;color:white;border-radius:50%;width:18px;height:18px;font-size:10px;font-weight:900;display:flex;align-items:center;justify-content:center">{pending_n}</span>' if pending_n > 0 else '') + '</a>' if is_owner or user["role"]=="manager" else ''}
        <a href='/staff/leave-planner' class='btn-secondary'>📅 Leave Planner</a>
        {'<a href="/staff/pay-overview" class="btn-secondary">💰 Pay Overview</a>' if is_owner else ''}
        {'<a href="/staff/document-templates" class="btn-secondary">📋 Doc Templates</a>' if is_owner else ''}
        {'<a href="/staff/new" class="btn-primary">➕ Add Staff Member</a>' if is_owner or user["role"]=="manager" else ''}
      </div>
    </div>
    <div style='display:flex;gap:0;border-bottom:1px solid #e2e8f0'>{tabs}</div>
    <div style='display:flex;gap:8px;flex-wrap:wrap'>{store_btns}</div>
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(420px,1fr))'>
      {cards_html}
    </div>"""

    return page("Staff", content, user, "staff")


# ── Individual staff profile ──────────────────────────────────────────────────

@app.get("/staff/document-templates", response_class=HTMLResponse)
def document_templates(session: str | None = Cookie(default=None), msg: str = ""):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    templates = q("SELECT * FROM document_templates ORDER BY doc_type, version DESC",
                  fetch=True) or []

    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    from collections import defaultdict
    by_type = defaultdict(list)
    for t in templates:
        by_type[dict(t)["doc_type"]].append(dict(t))

    tmpl_html = ""
    for dtype in DOC_TYPES:
        type_tmpls = by_type.get(dtype, [])
        current    = next((t for t in type_tmpls if t["is_current"]), None)
        older      = [t for t in type_tmpls if not t["is_current"]]

        current_html = ""
        if current:
            current_html = f"""
            <div style='background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:12px 14px'>
              <div style='font-size:13px;font-weight:700;color:#166534;margin-bottom:4px'>
                ✅ Current — v{current['version']} uploaded {current['uploaded_at'][:10]}
              </div>
              <div style='font-size:11px;color:#64748b;margin-bottom:10px'>{current.get('notes') or ''}</div>
              <div style='display:flex;gap:12px;align-items:center'>
                <a href='/staff/document-templates/{current["template_id"]}/download'
                   style='color:#64748b;font-size:12px;text-decoration:underline'>⬇️ download template</a>
                <a href='/staff/document-templates/{current["template_id"]}/delete'
                   onclick='return confirm("Delete this version? The previous version will become current.")'
                   class='btn-danger' style='padding:5px 14px;font-size:12px'>🗑️ Delete</a>
              </div>
            </div>"""

        older_html = "".join(
            f"<div style='font-size:12px;color:#94a3b8;padding:4px 10px'>v{t['version']} — {t['uploaded_at'][:10]} (superseded)</div>"
            for t in older
        )

        tmpl_html += f"""
        <div class='card'>
          <div style='font-weight:900;color:#0f2942;margin-bottom:8px'>{dtype}</div>
          {current_html or "<div style='color:#94a3b8;font-size:13px;padding:8px 0'>No template uploaded yet</div>"}
          {older_html}
          <form action='/staff/document-templates/upload' method='POST'
                enctype='multipart/form-data'
                style='margin-top:12px;padding-top:12px;border-top:1px solid #f1f5f9'
                onsubmit='showUploading(this)'>
            <input type='hidden' name='doc_type' value='{dtype}'>
            <div style='margin-bottom:8px'>
              <label style='font-size:11px;font-weight:700;color:#64748b;
                            text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px'>
                Select Template File (.docx or .dotx)
              </label>
              <input type='file' name='template_file' accept='.docx,.dotx' required
                     style='width:100%;border:1px solid #e2e8f0;border-radius:8px;
                            padding:8px 10px;font-size:13px;background:white;cursor:pointer'
                     onchange='previewFile(this)'>
              <div id='preview_{dtype.replace(" ","_")}' style='font-size:12px;color:#16a34a;
                   font-weight:700;margin-top:4px;display:none'>
                ✅ Selected: <span class='filename'></span>
              </div>
            </div>
            <div style='margin-bottom:8px'>
              <label style='font-size:11px;font-weight:700;color:#64748b;
                            text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px'>
                Version Notes
              </label>
              <input type='text' name='notes' id='notes_{dtype.replace(" ","_")}'
                     placeholder='e.g. Updated Nov 2024 — new holiday clause'
                     style='width:100%;border:1px solid #e2e8f0;border-radius:8px;
                            padding:8px 10px;font-size:13px'>
            </div>
            <button type='submit' class='btn-primary' style='width:100%;padding:8px;font-size:13px'>
              ⬆️ Upload New Version
            </button>
          </form>
        </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>📋 Document Templates</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>
    <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
      <div style='font-size:13px;font-weight:700;color:#92400e'>📝 How to set up templates</div>
      <div style='font-size:13px;color:#78350f;margin-top:4px'>
        Create a Word (.docx) document with your letter/contract content.
        Use these placeholders where you want staff details inserted:
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{FULL_NAME}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{STORE}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{DATE_JOINED}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{HOURLY_RATE}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{TODAY}}}}</code>
        and more. See the full list when generating a document.
      </div>
    </div>
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(380px,1fr))'>
      {tmpl_html}
    </div>
    <script>
    function previewFile(input) {{
      if (!input.files.length) return;
      const fname = input.files[0].name;
      const form  = input.closest('form');
      // Show selected filename
      const preview = form.querySelector('[id^="preview_"]');
      if (preview) {{
        preview.querySelector('.filename').textContent = fname;
        preview.style.display = 'block';
      }}
      // Auto-fill notes with filename + date if empty
      const notes = form.querySelector('[id^="notes_"]');
      if (notes && !notes.value) {{
        const today = new Date().toLocaleDateString('en-GB', {{day:'2-digit',month:'short',year:'numeric'}});
        notes.value = fname + ' — uploaded ' + today;
      }}
    }}
    function showUploading(form) {{
      const btn = form.querySelector('button[type="submit"]');
      if (btn) {{ btn.textContent = '⏳ Uploading...'; btn.disabled = true; }}
    }}
    </script>"""

    return page("Document Templates", content, user, "staff")


@app.post("/staff/document-templates/upload")
async def upload_template(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff/document-templates", status_code=303)

    form     = await request.form()
    doc_type = form.get("doc_type","")
    tmpl     = form.get("template_file")
    notes    = str(form.get("notes","") or "").strip()

    if not tmpl or not hasattr(tmpl,"filename") or not tmpl.filename:
        return RedirectResponse("/staff/document-templates?msg=No+file+selected", status_code=303)

    existing = q("SELECT MAX(version) as v FROM document_templates WHERE doc_type=?",
                 (doc_type,), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1

    q("UPDATE document_templates SET is_current=0 WHERE doc_type=?", (doc_type,))

    filename = f"template_{doc_type.replace(' ','_')}_v{next_ver}.docx"
    filepath = os.path.join(TEMPLATES_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(await tmpl.read())

    q("""INSERT INTO document_templates
            (doc_type, version, file_path, file_name, is_current, uploaded_by, notes)
         VALUES(?,?,?,?,1,?,?)""",
      (doc_type, next_ver, filepath, tmpl.filename, user.get("username"), notes or None))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/document-templates?msg={uq(doc_type + ' template uploaded (v' + str(next_ver) + ')')}",
        status_code=303)


@app.get("/staff/document-templates/{template_id}/delete")
def delete_template(template_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff/document-templates", status_code=303)
    # Get the doc_type before deleting
    rows = q("SELECT * FROM document_templates WHERE template_id=?", (template_id,), fetch=True)
    if rows:
        t = dict(rows[0])
        # Delete the file from disk
        if os.path.exists(t["file_path"]):
            os.remove(t["file_path"])
        # Delete from database
        q("DELETE FROM document_templates WHERE template_id=?", (template_id,))
        # If this was current, make the previous version current
        q("""UPDATE document_templates SET is_current=1
             WHERE doc_type=? AND template_id=(
                SELECT MAX(template_id) FROM document_templates WHERE doc_type=?
             )""", (t["doc_type"], t["doc_type"]))
    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/document-templates?msg={uq('Template version deleted')}",
        status_code=303)


@app.get("/staff/document-templates/{template_id}/download")
def download_template(template_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM document_templates WHERE template_id=?", (template_id,), fetch=True)
    if not rows: return HTMLResponse("<p>Not found</p>", status_code=404)
    t = dict(rows[0])
    if not os.path.exists(t["file_path"]):
        return HTMLResponse("<p>File not found</p>", status_code=404)
    return FileResponse(t["file_path"], filename=t["file_name"] or os.path.basename(t["file_path"]),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")



@app.get("/staff/new", response_class=HTMLResponse)
def new_staff_form(session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse("/staff", status_code=303)
    return render_staff_form(user, None)


@app.get("/staff/leave-requests", response_class=HTMLResponse)
def leave_requests(session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/staff", status_code=303)

    pending = q("""
        SELECT lr.*, sp.first_name, sp.last_name, sp.store_name, sp.contracted_hrs
        FROM leave_requests lr
        JOIN staff_profiles sp ON lr.staff_id=sp.staff_id
        WHERE lr.status='pending'
        ORDER BY lr.date_from ASC
    """, fetch=True) or []

    recent = q("""
        SELECT lr.*, sp.first_name, sp.last_name, sp.store_name
        FROM leave_requests lr
        JOIN staff_profiles sp ON lr.staff_id=sp.staff_id
        WHERE lr.status != 'pending'
        ORDER BY lr.created_at DESC LIMIT 20
    """, fetch=True) or []

    def req_row(lr, show_actions=True):
        lr    = dict(lr)
        name  = f"{lr['first_name']} {lr['last_name']}"
        ltype = ABSENCE_TYPES.get(lr['leave_type'], lr['leave_type'])
        badge = {"approved":"<span class='badge-paid'>Approved</span>",
                 "pending": "<span class='badge-partial'>Pending</span>",
                 "declined":"<span class='badge-overdue'>Declined</span>"}.get(lr["status"],"")
        actions = ""
        if show_actions:
            actions = f"""
            <a href='/staff/leave-requests/{lr['request_id']}/approve'
               class='btn-success' style='padding:4px 10px;font-size:11px'>✅ Approve</a>
            <a href='/staff/leave-requests/{lr['request_id']}/decline'
               class='btn-danger' style='padding:4px 10px;font-size:11px'>❌ Decline</a>"""
        return f"""<tr>
          <td style='font-weight:700'>{name}</td>
          <td style='font-size:12px;color:#64748b'>{lr.get('store_name','')}</td>
          <td>{ltype}</td>
          <td class='mono'>{lr['date_from']}</td>
          <td class='mono'>{lr['date_to']}</td>
          <td class='mono'>{lr['days_taken']}</td>
          <td>{badge}</td>
          <td style='font-size:12px;color:#64748b'>{lr.get('notes') or '—'}</td>
          <td><div style='display:flex;gap:4px'>{actions}</div></td>
        </tr>"""

    pending_rows = "".join(req_row(lr) for lr in pending)
    recent_rows  = "".join(req_row(lr, False) for lr in recent)

    content = f"""
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>📋 Leave Requests</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>

    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#d97706;color:white;font-weight:700;font-size:14px'>
        ⏳ Pending Approval ({len(pending)})
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Staff</th><th>Store</th><th>Type</th><th>From</th><th>To</th><th>Days</th><th>Status</th><th>Notes</th><th>Action</th></tr></thead>
          <tbody>{pending_rows or '<tr><td colspan="9" style="text-align:center;padding:24px;color:#94a3b8">No pending requests</td></tr>'}</tbody>
        </table>
      </div>
    </div>

    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>Recent Decisions</div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Staff</th><th>Store</th><th>Type</th><th>From</th><th>To</th><th>Days</th><th>Status</th><th>Notes</th><th></th></tr></thead>
          <tbody>{recent_rows or '<tr><td colspan="9" style="text-align:center;padding:24px;color:#94a3b8">No recent decisions</td></tr>'}</tbody>
        </table>
      </div>
    </div>"""
    return page("Leave Requests", content, user, "staff")


@app.get("/staff/leave-planner", response_class=HTMLResponse)
def leave_planner(
    session: str | None = Cookie(default=None),
    year:    int = 0,
    store:   str = ""
):
    redir, user = require_login(session)
    if redir: return redir
    if not year: year = datetime.now().year

    # Get store filter
    if user["role"] == "manager" and user.get("store_name"):
        store = store or user["store_name"]

    # Get active staff
    conds  = ["is_active=1"]
    params = []
    if store:
        conds.append("store_name=?")
        params.append(store)
    staff = q(f"SELECT * FROM staff_profiles WHERE {' AND '.join(conds)} ORDER BY first_name",
              params, fetch=True) or []

    # Get all approved leave for this year
    leave_data = q("""
        SELECT lr.staff_id, lr.date_from, lr.date_to, lr.leave_type
        FROM leave_requests lr
        WHERE lr.status='approved'
          AND strftime('%Y',lr.date_from)=?
    """, (str(year),), fetch=True) or []

    # Build a set of (staff_id, date) → leave_type
    leave_map = {}
    for lr in leave_data:
        lr = dict(lr)
        try:
            d1 = datetime.strptime(lr["date_from"], "%Y-%m-%d")
            d2 = datetime.strptime(lr["date_to"],   "%Y-%m-%d")
            cur = d1
            while cur <= d2:
                leave_map[(lr["staff_id"], cur.strftime("%Y-%m-%d"))] = lr["leave_type"]
                cur = cur.replace(day=cur.day+1) if cur.day < 28 else cur.replace(
                    month=cur.month+1 if cur.month<12 else 1,
                    year=cur.year+1 if cur.month==12 else cur.year, day=1)
        except: pass

    # BH set
    bh_set = set(UK_BANK_HOLIDAYS_2026)

    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    import calendar

    # Build the planner grid
    grid_html = ""
    for mi, mname in enumerate(months, 1):
        _, days_in_month = calendar.monthrange(year, mi)
        # Header row
        grid_html += f"<tr><td style='background:#0f2942;color:white;font-weight:900;font-size:12px;padding:6px 10px;white-space:nowrap'>{mname}</td>"
        for d in range(1, 32):
            if d > days_in_month:
                grid_html += "<td style='background:#f8fafc'></td>"
                continue
            date_str = f"{year}-{mi:02d}-{d:02d}"
            dow = datetime(year, mi, d).weekday()
            is_weekend = dow >= 5
            is_bh = date_str in bh_set
            if is_bh:
                bg = "#fef3c7"; txt = "<span style='font-size:9px;color:#92400e;font-weight:700'>BH</span>"
            elif is_weekend:
                bg = "#f1f5f9"; txt = ""
            else:
                bg = "white"; txt = ""
            grid_html += f"<td style='background:{bg};border:1px solid #e2e8f0;text-align:center;padding:2px;min-width:28px;font-size:10px'>{txt}</td>"
        grid_html += "</tr>"

        # Staff rows for this month
        for s in staff:
            s = dict(s)
            sid   = s["staff_id"]
            # First name or nickname (first 3 chars)
            initials = s["first_name"][:3]
            grid_html += f"<tr><td style='font-size:11px;font-weight:700;color:#334155;padding:3px 10px;white-space:nowrap;border-bottom:1px solid #f1f5f9'>{initials}</td>"
            for d in range(1, 32):
                if d > days_in_month:
                    grid_html += "<td style='background:#f8fafc'></td>"
                    continue
                date_str = f"{year}-{mi:02d}-{d:02d}"
                dow = datetime(year, mi, d).weekday()
                is_weekend = dow >= 5
                ltype = leave_map.get((sid, date_str))
                if ltype == "H":
                    bg = "#dcfce7"; cell = f"<span style='font-size:9px;font-weight:900;color:#166534'>{initials}</span>"
                elif ltype == "S":
                    bg = "#fee2e2"; cell = f"<span style='font-size:9px;font-weight:900;color:#991b1b'>S</span>"
                elif ltype == "B":
                    bg = "#fef3c7"; cell = ""
                elif is_weekend:
                    bg = "#f1f5f9"; cell = ""
                else:
                    bg = "white"; cell = ""
                grid_html += f"<td style='background:{bg};border:1px solid #f1f5f9;text-align:center;padding:1px;font-size:9px'>{cell}</td>"
            grid_html += "</tr>"
        # Spacer between months
        grid_html += f"<tr><td colspan='32' style='height:4px;background:#f8fafc'></td></tr>"

    # Store filter buttons
    store_btns = ""
    if user["role"] in ("owner","manager"):
        for sv,sl in [("","Both"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]:
            cls = "btn-primary" if store==sv else "btn-secondary"
            store_btns += f"<a href='/staff/leave-planner?year={year}&store={sv}' class='{cls}' style='padding:5px 12px;font-size:12px'>{sl}</a>"

    # Legend
    legend = """
    <div style='display:flex;gap:12px;flex-wrap:wrap;font-size:12px;font-weight:600'>
      <span><span style='display:inline-block;width:14px;height:14px;background:#dcfce7;border:1px solid #86efac;border-radius:3px;vertical-align:middle'></span> Holiday</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#fee2e2;border:1px solid #fca5a5;border-radius:3px;vertical-align:middle'></span> Sick</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#fef3c7;border:1px solid #fcd34d;border-radius:3px;vertical-align:middle'></span> Bank Holiday</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#f1f5f9;border:1px solid #e2e8f0;border-radius:3px;vertical-align:middle'></span> Weekend</span>
    </div>"""

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>📅 Leave Planner {year}</div>
      <div style='display:flex;gap:8px;align-items:center;flex-wrap:wrap'>
        {store_btns}
        <a href='/staff/leave-planner?year={year-1}&store={store}' class='btn-secondary' style='padding:5px 12px;font-size:12px'>← {year-1}</a>
        <a href='/staff/leave-planner?year={year+1}&store={store}' class='btn-secondary' style='padding:5px 12px;font-size:12px'>{year+1} →</a>
        <a href='/staff' class='btn-secondary' style='padding:5px 12px;font-size:12px'>← Staff</a>
      </div>
    </div>
    {legend}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table style='border-collapse:collapse;font-family:DM Mono,monospace;width:100%'>
          <!-- Day numbers header -->
          <tr>
            <td style='background:#0f2942;color:white;font-size:11px;font-weight:700;padding:6px 10px;white-space:nowrap'>Month / Day</td>
            {"".join(f"<td style='background:#0f2942;color:white;font-size:10px;font-weight:700;text-align:center;padding:3px;min-width:28px'>{d}</td>" for d in range(1,32))}
          </tr>
          {grid_html}
        </table>
      </div>
    </div>"""
    return page("Leave Planner", content, user, "staff")



# ── Pay History & NMW Routes ──────────────────────────────────────────────────

def get_nmw_for_person(dob_str: str, check_date: str = None) -> float:
    """Return current NMW rate for a person based on their age."""
    if not dob_str: return 0.0
    if not check_date: check_date = datetime.now().strftime("%Y-%m-%d")
    try:
        dob  = datetime.strptime(dob_str, "%Y-%m-%d")
        chk  = datetime.strptime(check_date, "%Y-%m-%d")
        age  = (chk - dob).days // 365
        # Get most recent NMW rates effective on or before check_date
        rates = q("""SELECT * FROM nmw_rates WHERE effective_date <= ?
                     ORDER BY effective_date DESC LIMIT 1""",
                  (check_date,), fetch=True)
        if not rates: return 0.0
        r = dict(rates[0])
        if age >= 21: return r["rate_21_plus"]
        elif age >= 18: return r["rate_18_20"]
        else: return r["rate_16_17"]
    except: return 0.0



@app.get("/staff/pay-overview", response_class=HTMLResponse)
def pay_overview(session: str | None = Cookie(default=None)):
    """Owner-only overview of all staff pay vs NMW."""
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    today = datetime.now().strftime("%Y-%m-%d")
    staff = q("SELECT * FROM staff_profiles WHERE is_active=1 ORDER BY store_name, first_name",
              fetch=True) or []

    rows_html = ""
    warnings  = 0
    for s in staff:
        s       = dict(s)
        current = s.get("hourly_rate") or 0
        nmw     = get_nmw_for_person(s.get("date_of_birth",""), today)
        diff    = round(current - nmw, 2) if nmw else 0
        annual  = current * (s.get("contracted_hrs") or 0) * 52

        if nmw == 0:
            status = "<span class='badge-unpaid'>No DOB</span>"
        elif diff < 0:
            status = f"<span class='badge-overdue'>⚠️ £{abs(diff):.2f} BELOW</span>"
            warnings += 1
        elif diff < 0.50:
            status = f"<span class='badge-partial'>Near NMW +£{diff:.2f}</span>"
        else:
            status = f"<span class='badge-paid'>✅ +£{diff:.2f}</span>"

        dob = s.get("date_of_birth","")
        age = ((datetime.now() - datetime.strptime(dob,"%Y-%m-%d")).days//365) if dob else "?"
        rows_html += f"""<tr>
          <td style='font-weight:700'>{s['first_name']} {s['last_name']}</td>
          <td style='font-size:12px;color:#64748b'>{s.get('store_name','')}</td>
          <td>{age}</td>
          <td class='mono' style='font-weight:700'>£{current:.2f}</td>
          <td class='mono' style='color:#64748b'>£{nmw:.2f}</td>
          <td>{status}</td>
          <td class='mono'>£{annual:,.0f}</td>
          <td><a href='/staff/{s["staff_id"]}/pay-history' class='btn-secondary' style='padding:3px 10px;font-size:11px'>History</a></td>
        </tr>"""

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>💰 Pay Overview — All Staff</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>
    {'<div class="flash-error">⚠️ ' + str(warnings) + ' staff member(s) may be below National Minimum Wage — please review</div>' if warnings else ''}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Name</th><th>Store</th><th>Age</th><th>Current Rate</th><th>NMW</th><th>Status</th><th>Annual Equiv.</th><th></th></tr></thead>
          <tbody>{rows_html}</tbody>
        </table>
      </div>
    </div>"""
    return page("Pay Overview", content, user, "staff")



# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3b — DOCUMENT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

import shutil
from fastapi.responses import FileResponse
from docx import Document as DocxDocument
from docx.shared import Pt
import io

DOCS_DIR      = "staff_docs"
TEMPLATES_DIR = "doc_templates"
os.makedirs(DOCS_DIR,      exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

DOC_TYPES = [
    "Offer Letter",
    "Employment Contract",
    "Right to Work",
    "P45/P46",
    "New Employee Notification",
    "DBS Check",
    "Other",
]

# Store addresses for mail merge
STORE_ADDRESSES = {
    "Uxbridge": "Sappy Properties (Uxbridge) Llp T/A Snappy Snaps, 178 High Street, Uxbridge, Middlesex, UB8 1LA",
    "Newbury":  "Maukbs Ltd T/A Snappy Snaps, 95 Northbrook Street, Newbury, Berkshire, RG14 1AA",
}

def get_merge_fields(staff: dict) -> dict:
    """Return all merge fields for Word template substitution.
    Supports both <<field>> (your existing format) and {{FIELD}} formats.
    """
    today    = datetime.now().strftime("%d %B %Y")
    name     = f"{staff.get('first_name','')} {staff.get('last_name','')}".strip()
    store    = staff.get('store_name','')
    store_addr = STORE_ADDRESSES.get(store, store)
    hrs      = staff.get('contracted_hrs') or 0
    rate     = staff.get('hourly_rate') or 0
    emp_type = staff.get('employment_type') or ('Full-time' if hrs >= 30 else 'Part-time')
    pay_type = 'salary' if staff.get('is_salaried') == 'Y' else 'hourly rate of'
    wages    = f"£{staff.get('salary_amount',0):,.2f} per annum" if staff.get('is_salaried') == 'Y' else f"£{rate:.2f} per hour"
    job_title  = staff.get('job_title') or 'Sales Assistant'
    reports_to = staff.get('reports_to') or 'Store Manager'

    # Support both << >> and {{ }} formats
    fields = {}

    # Your existing << >> format
    angle = {
        "<<employee name>>":           name,
        "<<employee first name>>":     staff.get('first_name',''),
        "<<address line 1>>":          staff.get('address_1','') or '',
        "<<address line 2>>":          staff.get('address_2','') or '',
        "<<address line 3>>":          staff.get('address_3','') or '',
        "<<address line 4>>":          staff.get('address_4','') or '',
        "<<post code>>":               staff.get('postcode','') or '',
        "<<today's date>>":            today,
        "<<position>>":                job_title,
        "<<FT or PT>>":                emp_type,
        "<<salary or hourly>>":        pay_type,
        "<<wages>>":                   wages,
        "<<employer>>":                "Maukbs Ltd T/A Snappy Snaps",
        "<<employer and store address>>": store_addr,
        "<<store address>>":           store_addr,
        "<<s tore address >>":         store_addr,
        "<<reporting to>>":            reports_to,
        "<<contracted hours>>":        f"{hrs} hours per week",
        "<<hourly rate>>":             f"£{rate:.2f}",
        "<<date of joining>>":         staff.get('date_joined','') or '',
        "<<date of birth>>":           staff.get('date_of_birth','') or '',
        "<<p osition>>":               job_title,
        "<<e mployer>>":               "Maukbs Ltd T/A Snappy Snaps",
    }
    fields.update(angle)

    # Also {{ }} format for new templates
    curly = {
        "{{FULL_NAME}}":        name,
        "{{FIRST_NAME}}":       staff.get('first_name',''),
        "{{LAST_NAME}}":        staff.get('last_name',''),
        "{{ADDRESS_1}}":        staff.get('address_1','') or '',
        "{{ADDRESS_2}}":        staff.get('address_2','') or '',
        "{{ADDRESS_3}}":        staff.get('address_3','') or '',
        "{{POSTCODE}}":         staff.get('postcode','') or '',
        "{{EMAIL}}":            staff.get('email','') or '',
        "{{PHONE}}":            staff.get('phone','') or '',
        "{{STORE}}":            store,
        "{{STORE_ADDRESS}}":    store_addr,
        "{{DATE_JOINED}}":      staff.get('date_joined','') or '',
        "{{DATE_OF_BIRTH}}":    staff.get('date_of_birth','') or '',
        "{{CONTRACTED_HOURS}}": str(hrs),
        "{{HOURLY_RATE}}":      f"£{rate:.2f}" if rate else '',
        "{{JOB_TITLE}}":        job_title,
        "{{REPORTS_TO}}":       reports_to,
        "{{EMP_TYPE}}":         emp_type,
        "{{TODAY}}":            today,
        "{{YEAR}}":             str(datetime.now().year),
        "{{EMPLOYER}}":         "Maukbs Ltd T/A Snappy Snaps",
    }
    fields.update(curly)
    return fields


def fill_word_template(template_path: str, fields: dict) -> bytes:
    """Fill a Word .docx/.dotx template with merge fields and return as bytes.
    Handles split runs where a merge field spans multiple runs in the same paragraph.
    """
    # For .dotx files, copy to a temp .docx first
    import tempfile, shutil
    if template_path.lower().endswith('.dotx'):
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        shutil.copy(template_path, tmp.name)
        doc = DocxDocument(tmp.name)
        os.unlink(tmp.name)
    else:
        doc = DocxDocument(template_path)

    def replace_para_text(para):
        """Replace merge fields even when split across runs."""
        # First try simple per-run replacement
        for run in para.runs:
            for key, val in fields.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

        # Then handle split runs by rebuilding full paragraph text
        full_text = para.text
        changed = False
        for key, val in fields.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
                changed = True

        if changed:
            # Put all text in first run, clear others
            if para.runs:
                para.runs[0].text = full_text
                for run in para.runs[1:]:
                    run.text = ''

    for para in doc.paragraphs:
        replace_para_text(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para_text(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Staff documents page ──────────────────────────────────────────────────────


@app.get("/staff/{staff_id}", response_class=HTMLResponse)
def staff_profile(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])

    # Access control — staff can only see their own profile
    if user["role"] == "staff":
        my_rows = q("SELECT staff_id FROM staff_profiles WHERE first_name||' '||last_name=?",
                    (user.get("full_name",""),), fetch=True)
        my_id = my_rows[0]["staff_id"] if my_rows else None
        if my_id != staff_id:
            return RedirectResponse(f"/staff/{my_id}" if my_id else "/", status_code=303)

    is_leaver = not s["is_active"]
    if is_leaver and user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    year  = datetime.now().year
    leave = get_leave_summary(staff_id, year)

    # Leave history
    leave_hist = q("""
        SELECT * FROM leave_requests
        WHERE staff_id=? ORDER BY date_from DESC LIMIT 20
    """, (staff_id,), fetch=True) or []

    # Recent attendance
    attendance = q("""
        SELECT * FROM timesheets WHERE staff_name=?
        ORDER BY work_date DESC LIMIT 10
    """, (f"{s['first_name']} {s['last_name']}",), fetch=True) or []

    name = f"{s['first_name']} {s['last_name']}"

    # ── Leave history table ──
    leave_rows = ""
    for lr in leave_hist:
        lr = dict(lr)
        at = ABSENCE_TYPES.get(lr["leave_type"], lr["leave_type"])
        status_badge = {
            "approved": "<span class='badge-paid'>Approved</span>",
            "pending":  "<span class='badge-partial'>Pending</span>",
            "declined": "<span class='badge-overdue'>Declined</span>",
        }.get(lr["status"], lr["status"])
        leave_rows += f"""
        <tr>
          <td>{lr['date_from']}</td>
          <td>{lr['date_to']}</td>
          <td>{at}</td>
          <td class='mono'>{lr['days_taken']}</td>
          <td>{status_badge}</td>
          <td style='font-size:12px;color:#64748b'>{lr.get('notes') or '—'}</td>
        </tr>"""

    # ── Attendance table ──
    att_rows = ""
    for a in attendance:
        a = dict(a)
        ci = a.get("clock_in_time") or "—"
        co = a.get("clock_out_time") or "⚠️ Open"
        att_rows += f"<tr><td class='mono'>{a['work_date']}</td><td class='mono' style='color:#16a34a'>{ci}</td><td class='mono' style='color:#dc2626'>{co}</td><td><span class='badge-{'paid' if a.get('status_flag')=='GPS_VERIFIED' else 'unpaid'}'>{a.get('status_flag') or '—'}</span></td></tr>"

    can_edit = user["role"] in ("owner", "manager")

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Staff</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>{name}</div>
        <div style='color:#64748b;font-size:13px'>{s.get('store_name') or ''} {'· Left ' + s['date_left'] if is_leaver and s.get('date_left') else ''}</div>
      </div>
      <div style='display:flex;gap:8px;flex-wrap:wrap'>
        {'<a href="/staff/' + str(staff_id) + '/edit" class="btn-primary">✏️ Edit Profile</a>' if can_edit else ''}
        <a href='/staff/{staff_id}/request-leave' class='btn-secondary'>📅 Request Leave</a>
        {'<a href="/staff/' + str(staff_id) + '/pay-history" class="btn-secondary">💰 Pay History</a>' if can_edit else ''}
        {'<a href="/staff/' + str(staff_id) + '/set-entitlement" class="btn-secondary">⚙️ Set Entitlement</a>' if can_edit else ''}
        <a href='/staff/{staff_id}/documents' class='btn-secondary'>&#128193; Documents</a>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>&#128203; Onboarding</a>
      </div>
    </div>

    <!-- Summary cards -->
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(150px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Entitlement {year}</div>
        <div style='font-size:20px;font-weight:900;color:#0f2942'>{leave.get("entitlement_fmt","—")}</div>
        <div style='font-size:10px;color:#94a3b8'>inc. bank holidays</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Holiday Taken</div>
        <div style='font-size:20px;font-weight:900;color:#d97706'>{leave.get("taken_days",0)} days</div>
        <div style='font-size:10px;color:#94a3b8'>{leave.get("bh_days",0)} bank hols used</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Holiday Balance</div>
        <div style='font-size:20px;font-weight:900;color:{"#16a34a" if leave.get("balance_days",0)>0 else "#dc2626"}'>{leave.get("balance_fmt","—")}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Sick Days {year}</div>
        <div style='font-size:20px;font-weight:900;color:{"#dc2626" if leave.get("sick_days",0)>0 else "#0f2942"}'>{leave.get("sick_days",0)}</div>
        <div style='font-size:10px;color:#94a3b8'>does not affect holiday</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Contract</div>
        <div style='font-size:24px;font-weight:900;color:#0f2942'>{s.get('contracted_hrs') or '—'}</div>
        <div style='font-size:11px;color:#94a3b8'>hrs/week</div>
      </div>
    </div>

    <!-- Personal details -->
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Personal Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr));font-size:13px'>
        <div><span style='color:#94a3b8;font-weight:700'>Date of Birth</span><br>{s.get('date_of_birth') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Phone</span><br>{s.get('phone') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Email</span><br>{s.get('email') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Address</span><br>{', '.join(filter(None,[s.get('address_1'),s.get('address_2'),s.get('address_3'),s.get('postcode')])) or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Date Joined</span><br>{s.get('date_joined') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Hourly Rate</span><br>{'£'+str(s['hourly_rate'])+'/hr' if s.get('hourly_rate') else '—'}</div>
      </div>
    </div>


        <!-- Leave history -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px;display:flex;justify-content:space-between;align-items:center'>
        <span>📅 Leave History {year}</span>
        <a href='/staff/{staff_id}/request-leave' style='background:rgba(255,255,255,.15);color:white;font-size:12px;font-weight:700;padding:4px 12px;border-radius:6px;text-decoration:none'>+ Request Leave</a>
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>From</th><th>To</th><th>Type</th><th>Days</th><th>Status</th><th>Notes</th></tr></thead>
          <tbody>{leave_rows or '<tr><td colspan="6" style="text-align:center;padding:24px;color:#94a3b8">No leave recorded this year</td></tr>'}</tbody>
        </table>
      </div>
    </div>

    <!-- Attendance -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>⏱ Recent Attendance</div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Date</th><th>Clock In</th><th>Clock Out</th><th>Status</th></tr></thead>
          <tbody>{att_rows or '<tr><td colspan="4" style="text-align:center;padding:24px;color:#94a3b8">No attendance records</td></tr>'}</tbody>
        </table>
      </div>
    </div>"""

    return page(name, content, user, "staff")


# ── Add / Edit staff ──────────────────────────────────────────────────────────


@app.get("/staff/{staff_id}/edit", response_class=HTMLResponse)
def edit_staff_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])
    # Staff can only edit their own profile (limited fields)
    if user["role"] == "staff":
        my_rows = q("SELECT staff_id FROM staff_profiles WHERE first_name||' '||last_name=?",
                    (user.get("full_name",""),), fetch=True)
        my_id = my_rows[0]["staff_id"] if my_rows else None
        if my_id != staff_id:
            return RedirectResponse("/", status_code=303)
    return render_staff_form(user, s)

def render_staff_form(user: dict, s: dict | None) -> HTMLResponse:
    is_edit   = s is not None
    is_owner  = user["role"] == "owner"
    is_mgr    = user["role"] in ("owner", "manager")
    is_self   = user["role"] == "staff"
    title     = f"✏️ Edit — {s['first_name']} {s['last_name']}" if is_edit else "➕ New Staff Member"
    action    = f"/staff/{s['staff_id']}/save" if is_edit else "/staff/save"
    back_url  = f"/staff/{s['staff_id']}" if is_edit else "/staff"
    sv        = s or {}

    def fi(name, label, ftype="text", val=None, req=False, opts=None, disabled=False, placeholder=""):
        safe = val if val is not None else ""
        req_a = "required" if req else ""
        dis_a = "disabled style='background:#f8fafc;color:#94a3b8'" if disabled else ""
        step  = "step='0.01'" if ftype=="number" else ""
        ph    = f"placeholder='{placeholder}'" if placeholder else ""
        if opts is not None:
            o = "".join(f"<option value='{ov}' {'selected' if str(safe)==str(ov) else ''}>{ol}</option>"
                        for ov,ol in opts)
            return f"<div><label>{label}</label><select name='{name}' {req_a} {dis_a}>{o}</select></div>"
        return f"<div><label>{label}</label><input type='{ftype}' name='{name}' value='{safe}' {req_a} {dis_a} {step} {ph}></div>"

    store_opts = [("","-- Select --"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]
    sex_opts   = [("","--"),("M","Male"),("F","Female"),("O","Other")]
    active_opts= [("1","Active"),("0","Left / Leaver")]

    # Personal details — staff can edit these themselves
    personal = f"""
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Personal Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
        {fi('first_name','First Name','text',sv.get('first_name'),req=True,disabled=is_self)}
        {fi('last_name', 'Last Name', 'text',sv.get('last_name'), req=True,disabled=is_self)}
        {fi('date_of_birth','Date of Birth','date',sv.get('date_of_birth'))}
        {fi('sex','Gender',opts=sex_opts,val=sv.get('sex'))}
        {fi('phone','Phone','text',sv.get('phone'),placeholder='07700 000000')}
        {fi('email','Email','email',sv.get('email'))}
        {fi('address_1','Address Line 1','text',sv.get('address_1'))}
        {fi('address_2','Address Line 2','text',sv.get('address_2'))}
        {fi('address_3','Town/City','text',sv.get('address_3'))}
        {fi('postcode','Postcode','text',sv.get('postcode'))}
      </div>
    </div>"""

    # Employment details — manager/owner only
    employment = ""
    if is_mgr:
        employment = f"""
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
        {fi('staff_number','Staff Number','number',sv.get('staff_number'))}
        {fi('store_name','Store',opts=store_opts,val=sv.get('store_name'),req=True)}
        {fi('job_title','Job Title','text',sv.get('job_title',''),placeholder='e.g. Sales Assistant')}
        {fi('employment_type','Employment Type',opts=[('Full-time','Full-time'),('Part-time','Part-time')],val=sv.get('employment_type','Part-time'))}
        {fi('reports_to','Reports To','text',sv.get('reports_to',''),placeholder='e.g. Store Manager')}
        {fi('date_joined','Date Joined','date',sv.get('date_joined'))}
        {fi('contracted_hrs','Contracted Hours/Week','number',sv.get('contracted_hrs'),placeholder='e.g. 37.5')}
        {fi('hourly_rate','Hourly Rate (£)','number',sv.get('hourly_rate'),placeholder='e.g. 11.44')}
        {fi('is_salaried','Salaried?',opts=[('N','No'),('Y','Yes')],val=sv.get('is_salaried','N'))}
        {fi('salary_amount','Salary Amount (£/yr)','number',sv.get('salary_amount'))}
        {fi('is_active','Status',opts=active_opts,val=str(sv.get('is_active',1)))}
        {fi('date_left','Date Left','date',sv.get('date_left')) if is_edit else ''}
        {fi('leaving_reason','Reason for Leaving','text',sv.get('leaving_reason')) if is_edit else ''}
      </div>
    </div>"""

    content = f"""
    <div class='flex justify-between items-center'>
      <div>
        <a href='{back_url}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>{title}</div>
      </div>
    </div>
    <form action='{action}' method='POST' enctype='multipart/form-data'>
      {personal}
      {employment}
      <div class='card'>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>{'💾 Save Changes' if is_edit else '➕ Add Staff Member'}</button>
          <a href='{back_url}' class='btn-secondary'>Cancel</a>
          {'<a href="/staff/' + str(s["staff_id"]) + '/delete" class="btn-danger" onclick="return confirm(\'Are you sure?\')">🗑️ Delete</a>' if is_edit and is_owner else ''}
        </div>
      </div>
    </form>"""

    return HTMLResponse(page(title, content, user, "staff"))

@app.post("/staff/save")
async def save_new_staff(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form = await request.form()
    fv = lambda k, d="": str(form.get(k, d) or d).strip()
    fn = lambda k: float(form.get(k, 0) or 0) if form.get(k) else None
    q("""INSERT INTO staff_profiles
        (staff_number,first_name,last_name,store_name,sex,phone,email,
         address_1,address_2,address_3,postcode,date_joined,date_of_birth,
         contracted_hrs,hourly_rate,is_salaried,salary_amount,is_active)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
      (form.get("staff_number") or None, fv("first_name"), fv("last_name"),
       fv("store_name"), fv("sex"), fv("phone"), fv("email"),
       fv("address_1"), fv("address_2"), fv("address_3"), fv("postcode"),
       fv("date_joined") or None, fv("date_of_birth") or None,
       fn("contracted_hrs"), fn("hourly_rate"),
       fv("is_salaried","N"), fn("salary_amount"),
       int(form.get("is_active", 1))))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff?msg={uq('Staff member added successfully')}", status_code=303)

@app.post("/staff/{staff_id}/save")
async def save_staff(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form = await request.form()
    fv = lambda k, d="": str(form.get(k, d) or d).strip()
    fn = lambda k: float(form.get(k, 0) or 0) if form.get(k) else None
    is_mgr = user["role"] in ("owner","manager")

    if is_mgr:
        q("""UPDATE staff_profiles SET
            staff_number=?,first_name=?,last_name=?,store_name=?,sex=?,
            phone=?,email=?,address_1=?,address_2=?,address_3=?,postcode=?,
            date_joined=?,date_of_birth=?,contracted_hrs=?,hourly_rate=?,
            is_salaried=?,salary_amount=?,is_active=?,date_left=?,leaving_reason=?
            WHERE staff_id=?""",
          (form.get("staff_number") or None, fv("first_name"), fv("last_name"),
           fv("store_name"), fv("sex"), fv("phone"), fv("email"),
           fv("address_1"), fv("address_2"), fv("address_3"), fv("postcode"),
           fv("date_joined") or None, fv("date_of_birth") or None,
           fn("contracted_hrs"), fn("hourly_rate"),
           fv("is_salaried","N"), fn("salary_amount"),
           int(form.get("is_active",1)),
           fv("date_left") or None, fv("leaving_reason") or None,
           staff_id))
    else:
        # Staff can only update personal contact details
        q("""UPDATE staff_profiles SET
            phone=?,email=?,address_1=?,address_2=?,address_3=?,postcode=?,date_of_birth=?
            WHERE staff_id=?""",
          (fv("phone"),fv("email"),fv("address_1"),fv("address_2"),
           fv("address_3"),fv("postcode"),fv("date_of_birth") or None, staff_id))

    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}?msg={uq('Profile updated')}", status_code=303)


# ── Leave request ─────────────────────────────────────────────────────────────

@app.get("/staff/{staff_id}/request-leave", response_class=HTMLResponse)
def request_leave_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"
    leave = get_leave_summary(staff_id)
    bal   = leave.get("balance_days", 0)

    type_opts = "".join(f"<option value='{k}'>{v}</option>" for k,v in ABSENCE_TYPES.items())

    content = f"""
    <div>
      <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>📅 Request Leave — {name}</div>
    </div>
    <div class='card' style='max-width:520px'>
      <div style='background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;padding:12px 16px;margin-bottom:16px;font-size:13px'>
        Current balance: <strong style='color:#0369a1'>{leave.get('balance_fmt','—')} remaining</strong>
      </div>
      <form action='/staff/{staff_id}/request-leave' method='POST' class='space-y-4'>
        <div><label>Leave Type</label>
          <select name='leave_type'>{type_opts}</select></div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          <div><label>From Date</label>
            <input type='date' name='date_from' required></div>
          <div><label>To Date</label>
            <input type='date' name='date_to' required></div>
        </div>
        <div><label>Notes (optional)</label>
          <textarea name='notes' rows='2' placeholder='Any additional details...'></textarea></div>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>📤 Submit Request</button>
          <a href='/staff/{staff_id}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>
    <script>
    // Auto-calculate working days when dates change
    document.addEventListener('DOMContentLoaded', function() {{
      const from = document.querySelector('[name="date_from"]');
      const to   = document.querySelector('[name="date_to"]');
      if (from) from.addEventListener('change', function() {{
        if (to && !to.value) to.value = from.value;
      }});
    }});
    </script>"""
    return page("Request Leave", content, user, "staff")

@app.post("/staff/{staff_id}/request-leave")
async def submit_leave(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form       = await request.form()
    leave_type = form.get("leave_type","H")
    date_from  = form.get("date_from","")
    date_to    = form.get("date_to","")
    notes      = str(form.get("notes","") or "").strip()

    # Calculate working days
    try:
        d1 = datetime.strptime(date_from, "%Y-%m-%d")
        d2 = datetime.strptime(date_to,   "%Y-%m-%d")
        days = 0
        cur  = d1
        while cur <= d2:
            if cur.weekday() < 5 and cur.strftime("%Y-%m-%d") not in UK_BANK_HOLIDAYS_2026:
                days += 1
            cur = cur.replace(day=cur.day+1) if cur.day < 28 else cur.replace(
                month=cur.month+1 if cur.month<12 else 1,
                year=cur.year+1 if cur.month==12 else cur.year, day=1)
    except:
        days = 1

    # Auto-approve for manager/owner, else pending
    status = "approved" if user["role"] in ("owner","manager") else "pending"

    q("""INSERT INTO leave_requests
        (staff_id, leave_type, date_from, date_to, days_taken,
         status, requested_by, notes)
        VALUES(?,?,?,?,?,?,?,?)""",
      (staff_id, leave_type, date_from, date_to, days,
       status, user.get("username"), notes or None))

    from urllib.parse import quote as uq
    msg = "Leave approved ✅" if status=="approved" else "Leave request submitted — awaiting approval ⏳"
    return RedirectResponse(f"/staff/{staff_id}?msg={uq(msg)}", status_code=303)


# ── Leave requests management (manager/owner) ─────────────────────────────────


@app.get("/staff/leave-requests/{req_id}/approve")
def approve_leave(req_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    q("UPDATE leave_requests SET status='approved', approved_by=?, approved_at=datetime('now') WHERE request_id=?",
      (user.get("username"), req_id))
    return RedirectResponse("/staff/leave-requests", status_code=303)

@app.get("/staff/leave-requests/{req_id}/decline")
def decline_leave(req_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    q("UPDATE leave_requests SET status='declined', approved_by=?, approved_at=datetime('now') WHERE request_id=?",
      (user.get("username"), req_id))
    return RedirectResponse("/staff/leave-requests", status_code=303)


# ── Annual Leave Planner (visual calendar) ────────────────────────────────────


@app.get("/staff/{staff_id}/pay-history", response_class=HTMLResponse)
def pay_history_page(staff_id: int, session: str | None = Cookie(default=None), msg: str = ""):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse("/staff", status_code=303)

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Current NMW
    today   = datetime.now().strftime("%Y-%m-%d")
    nmw     = get_nmw_for_person(s.get("date_of_birth",""), today)
    current = s.get("hourly_rate") or 0
    diff    = round(current - nmw, 2)
    dob_str = s.get("date_of_birth","")
    age     = ((datetime.now() - datetime.strptime(dob_str,"%Y-%m-%d")).days // 365) if dob_str else 0

    # NMW status badge
    if nmw == 0:
        nmw_badge = "<span class='badge-unpaid'>DOB not set</span>"
    elif diff < 0:
        nmw_badge = f"<span class='badge-overdue'>⚠️ £{abs(diff):.2f} BELOW NMW</span>"
    elif diff < 0.50:
        nmw_badge = f"<span class='badge-partial'>⚠️ Only £{diff:.2f} above NMW</span>"
    else:
        nmw_badge = f"<span class='badge-paid'>✅ £{diff:.2f} above NMW</span>"

    # Pay history
    history = q("""SELECT * FROM pay_history WHERE staff_id=?
                   ORDER BY effective_date DESC""", (staff_id,), fetch=True) or []

    hist_rows = ""
    for h in history:
        h = dict(h)
        prev = f"£{h['previous_rate']:.2f}" if h.get("previous_rate") else "—"
        change = ""
        if h.get("previous_rate") and h.get("hourly_rate"):
            pct = ((h["hourly_rate"] - h["previous_rate"]) / h["previous_rate"]) * 100
            change = f"<span style='color:{'#16a34a' if pct>=0 else '#dc2626'};font-weight:700'>{'▲' if pct>=0 else '▼'} {abs(pct):.1f}%</span>"
        hist_rows += f"""<tr>
          <td class='mono'>{h['effective_date']}</td>
          <td class='mono' style='font-weight:700'>£{h['hourly_rate']:.2f}/hr</td>
          <td class='mono'>{prev}</td>
          <td>{change}</td>
          <td style='font-size:12px;color:#64748b'>{h.get('change_reason') or '—'}</td>
          <td style='font-size:12px;color:#94a3b8'>{h.get('recorded_by') or '—'}</td>
        </tr>"""

    # NMW history table
    nmw_rows = ""
    all_nmw = q("SELECT * FROM nmw_rates ORDER BY effective_date DESC", fetch=True) or []
    for n in all_nmw:
        n   = dict(n)
        age_rate = n["rate_21_plus"] if age >= 21 else (n["rate_18_20"] if age >= 18 else n["rate_16_17"])
        nmw_rows += f"""<tr>
          <td class='mono'>{n['effective_date']}</td>
          <td class='mono'>£{n['rate_21_plus']:.2f}</td>
          <td class='mono'>£{n['rate_18_20']:.2f}</td>
          <td class='mono'>£{n['rate_16_17']:.2f}</td>
          <td class='mono' style='font-weight:700;color:#0369a1'>£{age_rate:.2f}</td>
        </tr>"""

    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>💰 Pay History — {name}</div>
      </div>
    </div>

    <!-- Current pay status -->
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(160px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Current Rate</div>
        <div style='font-size:28px;font-weight:900;color:#0f2942'>£{current:.2f}</div>
        <div style='font-size:11px;color:#94a3b8'>per hour</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>NMW (Age {age})</div>
        <div style='font-size:28px;font-weight:900;color:#0f2942'>£{nmw:.2f}</div>
        <div style='font-size:11px;color:#94a3b8'>minimum wage</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Status</div>
        <div style='margin-top:8px'>{nmw_badge}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Annual Equiv.</div>
        <div style='font-size:22px;font-weight:900;color:#0f2942'>£{(current * (s.get('contracted_hrs') or 0) * 52):,.0f}</div>
        <div style='font-size:11px;color:#94a3b8'>based on {s.get('contracted_hrs') or 0}h/wk</div>
      </div>
    </div>

    <!-- Record new pay change -->
    <div class='card' style='max-width:500px'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>➕ Record Pay Change</div>
      <form action='/staff/{staff_id}/pay-history' method='POST' class='grid gap-3' style='grid-template-columns:1fr 1fr'>
        <div><label>Effective Date</label>
          <input type='date' name='effective_date' value='{today}' required></div>
        <div><label>New Hourly Rate (£)</label>
          <input type='number' step='0.01' name='hourly_rate' placeholder='e.g. 12.71' required></div>
        <div style='grid-column:1/-1'><label>Reason</label>
          <input type='text' name='change_reason' placeholder='e.g. Annual review, NMW increase, Promotion'></div>
        <div style='grid-column:1/-1'>
          <button type='submit' class='btn-primary'>💾 Save Pay Change</button>
        </div>
      </form>
    </div>

    <!-- Pay history -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>Pay History</div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Effective Date</th><th>Rate</th><th>Previous</th><th>Change</th><th>Reason</th><th>Recorded By</th></tr></thead>
          <tbody>{hist_rows or '<tr><td colspan="6" style="text-align:center;padding:24px;color:#94a3b8">No pay history recorded yet</td></tr>'}</tbody>
        </table>
      </div>
    </div>

    <!-- NMW reference table -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>
        National Minimum Wage Reference
        <span style='font-size:12px;font-weight:400;color:#93c5fd;margin-left:8px'>Highlighted column = this employee's applicable rate</span>
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Effective</th><th>21+</th><th>18-20</th><th>16-17</th><th style='background:#1e3a5f'>Applicable Rate</th></tr></thead>
          <tbody>{nmw_rows}</tbody>
        </table>
      </div>
    </div>"""
    return page(f"Pay History — {name}", content, user, "staff")


@app.post("/staff/{staff_id}/pay-history")
async def save_pay_change(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}", status_code=303)
    form = await request.form()
    new_rate  = float(form.get("hourly_rate", 0) or 0)
    eff_date  = form.get("effective_date","")
    reason    = str(form.get("change_reason","") or "").strip()

    # Get previous rate
    current = q("SELECT hourly_rate FROM staff_profiles WHERE staff_id=?",
                (staff_id,), fetch=True)
    prev_rate = dict(current[0])["hourly_rate"] if current else None

    q("""INSERT INTO pay_history (staff_id,effective_date,hourly_rate,previous_rate,change_reason,recorded_by)
         VALUES(?,?,?,?,?,?)""",
      (staff_id, eff_date, new_rate, prev_rate, reason or None, user.get("username")))

    # Update current rate on profile
    q("UPDATE staff_profiles SET hourly_rate=? WHERE staff_id=?", (new_rate, staff_id))

    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/pay-history?msg={uq('Pay change recorded')}", status_code=303)


@app.get("/staff/{staff_id}/set-entitlement", response_class=HTMLResponse)
def set_entitlement_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}", status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"
    year = datetime.now().year
    contracted = s.get("contracted_hrs") or 0
    # Statutory minimum is always 5.6 weeks x 5 days = 28 days
    # (regardless of hours — part-timers get pro-rated days but
    #  the day count is the same; the difference is day length)
    statutory  = 28.0
    # Pro-rata = 5.6 weeks × contracted hours ÷ 5 (actual days entitlement)
    pro_rata   = round(5.6 * contracted / 5, 1) if contracted else 28.0
    existing   = q("SELECT * FROM leave_entitlements WHERE staff_id=? AND year=?",
                   (staff_id, year), fetch=True)
    current    = dict(existing[0]) if existing else {}
    cur_val    = current.get("custom_days") or pro_rata

    html  = f"""
    <div>
      <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>
        &larr; Back to {name}</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>
        &#9881;&#65039; Leave Entitlement &mdash; {name} ({year})</div>
    </div>
    <div class='card' style='max-width:520px'>
      <div style='background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;
                  padding:12px 16px;margin-bottom:16px;font-size:13px'>
        <div><strong>Contracted hours:</strong> {contracted}h/week</div>
        <div><strong>Statutory minimum:</strong> 28 days (5.6 weeks)</div>
        <div style='font-size:12px;color:#64748b;margin-top:4px'>
          For part-time staff you can set below 28 days — pro-rata
          is acceptable (e.g. 3 days/wk = 5.6 × 3 = 16.8 days minimum).
          Set whatever you have agreed contractually.
        </div>
      </div>
      <form action='/staff/{staff_id}/set-entitlement' method='POST' class='space-y-4'>
        <input type='hidden' name='year' value='{year}'>
        <div>
          <label>Custom Entitlement (days) for {year}</label>
          <div style='display:flex;gap:8px;align-items:center'>
            <input type='number' step='0.5' name='custom_days' id='custom_days_input'
                   value='{cur_val}' min='1' style='flex:1'>
            <button type='button'
              onclick="document.getElementById('custom_days_input').value='{pro_rata}';
                       document.getElementById('notes_input').value='Reset to statutory pro-rata';"
              class='btn-secondary' style='white-space:nowrap;padding:8px 14px;font-size:12px'>
              🔄 Reset to Statutory ({pro_rata} days)
            </button>
          </div>
          <div style='font-size:11px;color:#94a3b8;margin-top:4px'>
            Statutory pro-rata for {contracted}h/wk =
            5.6 weeks × {contracted}h ÷ 5 days/wk = {pro_rata} days
          </div>
        </div>
        <div>
          <label>Notes</label>
          <input type='text' name='notes' id='notes_input'
                 value='{current.get("notes") or ""}'
                 placeholder='e.g. Pro-rated, contractual agreement'>
        </div>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>💾 Save Entitlement</button>
          <a href='/staff/{staff_id}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>"""
    return page(f"Entitlement", html, user, "staff")


@app.post("/staff/{staff_id}/set-entitlement")
async def save_entitlement(staff_id: int, request: Request,
                           session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form        = await request.form()
    year        = int(form.get("year", datetime.now().year))
    custom_days = float(form.get("custom_days", 0) or 0)
    notes       = str(form.get("notes","") or "").strip()
    s           = q("SELECT contracted_hrs FROM staff_profiles WHERE staff_id=?",
                    (staff_id,), fetch=True)
    contracted  = dict(s[0])["contracted_hrs"] if s else 0
    statutory   = 28.0  # Always 28 days statutory minimum
    q("""INSERT INTO leave_entitlements
            (staff_id,year,statutory_days,custom_days,effective_days,notes)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,year) DO UPDATE SET
            custom_days=excluded.custom_days,
            effective_days=excluded.effective_days,
            notes=excluded.notes""",
      (staff_id, year, statutory, custom_days, custom_days, notes or None))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}?msg={uq('Entitlement updated')}",
                            status_code=303)



@app.get("/staff/{staff_id}/documents", response_class=HTMLResponse)
def staff_documents(
    staff_id: int,
    session:  str | None = Cookie(default=None),
    msg:      str = "",
    msg_type: str = "success"
):
    redir, user = require_login(session)
    if redir: return redir

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get all documents for this staff member
    docs = q("""SELECT * FROM staff_documents WHERE staff_id=?
                ORDER BY doc_type, version DESC""",
             (staff_id,), fetch=True) or []

    # Get available templates
    templates = q("SELECT * FROM document_templates WHERE is_current=1 ORDER BY doc_type",
                  fetch=True) or []
    template_types = {dict(t)["doc_type"] for t in templates}

    flash = f"<div class='flash-{'success' if msg_type=='success' else 'error'}'>{msg}</div>" if msg else ""

    # Group docs by type
    from collections import defaultdict
    by_type = defaultdict(list)
    for d in docs:
        by_type[dict(d)["doc_type"]].append(dict(d))

    # Build document cards
    doc_cards = ""
    for dtype in DOC_TYPES:
        type_docs = by_type.get(dtype, [])
        has_template = dtype in template_types

        # Current version
        current = next((d for d in type_docs if d["is_current"]), None)
        older   = [d for d in type_docs if not d["is_current"]]

        current_html = ""
        if current:
            gen_badge = "<span style='background:#dbeafe;color:#1d4ed8;font-size:10px;font-weight:700;padding:2px 6px;border-radius:4px'>AUTO-GENERATED</span>" if current["generated"] else ""
            current_html = f"""
            <div style='background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:10px 14px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
              <div>
                <div style='font-size:13px;font-weight:700;color:#166534'>
                  ✅ v{current['version']} — {current['uploaded_at'][:10]} {gen_badge}
                </div>
                <div style='font-size:11px;color:#64748b'>{current.get('notes') or ''}</div>
              </div>
              <div style='display:flex;gap:6px'>
                <a href='/staff/{staff_id}/documents/{current["doc_id"]}/download'
                   class='btn-secondary' style='padding:4px 10px;font-size:11px'>⬇️ Download</a>
                <a href='/staff/{staff_id}/documents/{current["doc_id"]}/view'
                   class='btn-secondary' style='padding:4px 10px;font-size:11px' target='_blank'>👁 View</a>
              </div>
            </div>"""

        older_html = ""
        if older:
            older_html = "<div style='margin-top:6px'>"
            for od in older:
                older_html += f"""
                <div style='display:flex;justify-content:space-between;align-items:center;padding:6px 10px;font-size:12px;color:#64748b;border-bottom:1px solid #f1f5f9'>
                  <span>v{od['version']} — {od['uploaded_at'][:10]}</span>
                  <a href='/staff/{staff_id}/documents/{od["doc_id"]}/download'
                     style='color:#1e3a5f;font-weight:700;font-size:11px'>⬇️ Download</a>
                </div>"""
            older_html += "</div>"

        # Upload / generate form
        action_html = f"""
        <div style='margin-top:10px;padding-top:10px;border-top:1px solid #f1f5f9'>
          <div style='display:flex;gap:8px;flex-wrap:wrap'>
            <form action='/staff/{staff_id}/documents/upload' method='POST'
                  enctype='multipart/form-data' style='display:flex;gap:6px;flex:1;min-width:200px'>
              <input type='hidden' name='doc_type' value='{dtype}'>
              <input type='file' name='doc_file' accept='.pdf,.doc,.docx,.dotx'
                     style='flex:1;font-size:12px;padding:4px 8px'>
              <button type='submit' class='btn-primary' style='padding:4px 12px;font-size:11px;white-space:nowrap'>
                ⬆️ Upload
              </button>
            </form>
            {"<a href='/staff/" + str(staff_id) + "/documents/generate?doc_type=" + dtype + "' class='btn-secondary' style='padding:4px 12px;font-size:11px;white-space:nowrap'>⚡ Auto-fill</a>" if has_template else ""}
          </div>
        </div>"""

        doc_cards += f"""
        <div class='card'>
          <div style='font-weight:900;color:#0f2942;margin-bottom:8px;font-size:14px'>{dtype}</div>
          {current_html or "<div style='color:#94a3b8;font-size:13px'>No document uploaded yet</div>"}
          {older_html}
          {action_html}
        </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>📁 Documents — {name}</div>
      </div>
      {'<a href="/staff/document-templates" class="btn-secondary">📋 Manage Templates</a>' if user["role"] == "owner" else ''}
    </div>
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(400px,1fr))'>
      {doc_cards}
    </div>"""

    return page(f"Documents — {name}", content, user, "staff")


# ── Upload document ───────────────────────────────────────────────────────────

@app.post("/staff/{staff_id}/documents/upload")
async def upload_staff_doc(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    form     = await request.form()
    doc_type = form.get("doc_type","Other")
    doc_file = form.get("doc_file")
    notes    = str(form.get("notes","") or "").strip()

    if not doc_file or not hasattr(doc_file, "filename") or not doc_file.filename:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/documents?msg={uq('No file selected')}&msg_type=error",
            status_code=303)

    # Get next version number
    existing = q("""SELECT MAX(version) as v FROM staff_documents
                    WHERE staff_id=? AND doc_type=?""",
                 (staff_id, doc_type), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1

    # Mark previous versions as not current
    q("UPDATE staff_documents SET is_current=0 WHERE staff_id=? AND doc_type=?",
      (staff_id, doc_type))

    # Save file
    ext      = os.path.splitext(doc_file.filename)[1].lower()
    filename = f"staff_{staff_id}_{doc_type.replace(' ','_')}_v{next_ver}{ext}"
    filepath = os.path.join(DOCS_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(await doc_file.read())

    q("""INSERT INTO staff_documents
            (staff_id, doc_type, version, file_path, file_name,
             is_current, generated, uploaded_by, notes)
         VALUES(?,?,?,?,?,1,0,?,?)""",
      (staff_id, doc_type, next_ver, filepath,
       doc_file.filename, user.get("username"), notes or None))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/documents?msg={uq(doc_type + ' uploaded successfully')}",
        status_code=303)


# ── Auto-generate document from template ─────────────────────────────────────

@app.get("/staff/{staff_id}/documents/generate", response_class=HTMLResponse)
def generate_doc_form(
    staff_id: int,
    doc_type: str = "",
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}/documents", status_code=303)

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get available templates
    templates = q("SELECT * FROM document_templates WHERE is_current=1 ORDER BY doc_type",
                  fetch=True) or []

    # Show merge fields preview
    fields     = get_merge_fields(s)
    fields_html = ""
    for k, v in fields.items():
        fields_html += f"""
        <div style='display:flex;gap:12px;padding:4px 0;border-bottom:1px solid #f1f5f9;font-size:12px'>
          <code style='color:#7c3aed;min-width:180px'>{k}</code>
          <span style='color:#334155'>{v or '—'}</span>
        </div>"""

    type_opts = ""
    for t in templates:
        td = dict(t)
        sel = "selected" if td["doc_type"] == doc_type else ""
        type_opts += f'<option value="{td["doc_type"]}" {sel}>{td["doc_type"]} (v{td["version"]})</option>'

    content = f"""
    <div>
      <a href='/staff/{staff_id}/documents' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Documents</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>⚡ Auto-Generate Document — {name}</div>
    </div>
    <div class='grid gap-6' style='grid-template-columns:1fr 1fr'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Generate Document</div>
        <form action='/staff/{staff_id}/documents/generate' method='POST' class='space-y-4'>
          <div>
            <label>Document Type</label>
            <select name='doc_type' required>
              <option value=''>-- Select template --</option>
              {type_opts}
            </select>
          </div>
          <div>
            <label>Notes (optional)</label>
            <input type='text' name='notes' placeholder='e.g. Initial offer, Updated contract'>
          </div>
          <button type='submit' class='btn-primary'>⚡ Generate & Download</button>
        </form>
        {'<div class="flash-error" style="margin-top:12px">No templates uploaded yet. <a href=\'  /staff/document-templates\' style=\'color:#1e3a5f;font-weight:700\'>Upload templates here →</a></div>' if not templates else ''}
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Available Merge Fields</div>
        <div style='font-size:11px;color:#64748b;margin-bottom:8px'>
          Use these placeholders in your Word template — they will be replaced with this staff member's details.
        </div>
        <div style='max-height:400px;overflow-y:auto'>
          {fields_html}
        </div>
      </div>
    </div>"""

    return page("Generate Document", content, user, "staff")


@app.post("/staff/{staff_id}/documents/generate")
async def generate_doc(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    form     = await request.form()
    doc_type = form.get("doc_type","")
    notes    = str(form.get("notes","") or "").strip()

    # Get staff details
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])

    # Get template
    tmpl = q("SELECT * FROM document_templates WHERE doc_type=? AND is_current=1",
             (doc_type,), fetch=True)
    if not tmpl:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/documents?msg={uq('No template found for ' + doc_type)}&msg_type=error",
            status_code=303)
    tmpl = dict(tmpl[0])

    # Fill template
    fields   = get_merge_fields(s)
    doc_bytes = fill_word_template(tmpl["file_path"], fields)

    # Save generated file
    existing = q("""SELECT MAX(version) as v FROM staff_documents
                    WHERE staff_id=? AND doc_type=?""",
                 (staff_id, doc_type), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1
    q("UPDATE staff_documents SET is_current=0 WHERE staff_id=? AND doc_type=?",
      (staff_id, doc_type))

    filename = f"staff_{staff_id}_{doc_type.replace(' ','_')}_v{next_ver}.docx"
    filepath = os.path.join(DOCS_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(doc_bytes)

    q("""INSERT INTO staff_documents
            (staff_id, doc_type, version, file_path, file_name,
             is_current, generated, uploaded_by, notes)
         VALUES(?,?,?,?,?,1,1,?,?)""",
      (staff_id, doc_type, next_ver, filepath, filename,
       user.get("username"), notes or None))

    # Return the file for download
    name = f"{s['first_name']} {s['last_name']}"
    download_name = f"{doc_type} - {name}.docx"
    return FileResponse(filepath, filename=download_name,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ── View/Download document ────────────────────────────────────────────────────

@app.get("/staff/{staff_id}/documents/{doc_id}/download")
def download_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=?",
             (doc_id, staff_id), fetch=True)
    if not rows: return HTMLResponse("<p>Document not found</p>", status_code=404)
    d = dict(rows[0])
    if not os.path.exists(d["file_path"]):
        return HTMLResponse("<p>File not found on disk</p>", status_code=404)
    ext = os.path.splitext(d["file_path"])[1].lower()
    media = "application/pdf" if ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(d["file_path"], filename=d["file_name"] or os.path.basename(d["file_path"]),
                        media_type=media)

@app.get("/staff/{staff_id}/documents/{doc_id}/view")
def view_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=?",
             (doc_id, staff_id), fetch=True)
    if not rows: return HTMLResponse("<p>Document not found</p>", status_code=404)
    d = dict(rows[0])
    if not os.path.exists(d["file_path"]):
        return HTMLResponse("<p>File not found on disk</p>", status_code=404)
    return FileResponse(d["file_path"], media_type="application/pdf")


# ── Document Templates Management (owner only) ────────────────────────────────



# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3c — ONBOARDING FORMS
# ══════════════════════════════════════════════════════════════════════════════

def ensure_onboarding_tables():
    conn = db()
    c    = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS onboarding_forms (
            form_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id       INTEGER NOT NULL,
            form_type      TEXT NOT NULL,
            status         TEXT DEFAULT 'not_started',
            started_at     TEXT,
            completed_at   TEXT,
            form_data      TEXT,
            pdf_path       TEXT,
            UNIQUE(staff_id, form_type),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    conn.commit()
    conn.close()

ensure_onboarding_tables()

ONBOARD_FORMS = [
    ("employment_application", "Employment Application",  "staff"),
    ("p46",                    "P46 Tax Form",            "staff"),
    ("new_employee_notify",    "New Employee Notification","owner"),
    ("offer_letter",           "Offer Letter",            "owner"),
    ("employment_contract",    "Employment Contract",     "owner"),
    ("right_to_work",          "Right to Work Checked",  "owner"),
]

# Form types that link to the digital form (others are document-based)
DIGITAL_FORMS = {"employment_application", "p46", "new_employee_notify"}

def get_onboarding_status(staff_id: int) -> dict:
    """Return completion status for each onboarding form and document."""
    rows = q("SELECT form_type, status FROM onboarding_forms WHERE staff_id=?",
             (staff_id,), fetch=True) or []
    status_map = {dict(r)["form_type"]: dict(r)["status"] for r in rows}

    # Check staff_documents for document-based items
    doc_rows = q("SELECT doc_type, is_current FROM staff_documents WHERE staff_id=? AND is_current=1",
                 (staff_id,), fetch=True) or []
    doc_types = {dict(d)["doc_type"] for d in doc_rows}

    # Map document types to onboarding form types
    doc_type_map = {
        "offer_letter":        "Offer Letter",
        "employment_contract": "Employment Contract",
        "right_to_work":       "Right to Work",
    }

    result = {}
    for ftype, flabel, fwho in ONBOARD_FORMS:
        # Check if it's a document-based item
        if ftype in doc_type_map:
            doc_label = doc_type_map[ftype]
            status = "completed" if doc_label in doc_types else status_map.get(ftype, "not_started")
        else:
            status = status_map.get(ftype, "not_started")
        result[ftype] = {
            "label":   flabel,
            "who":     fwho,
            "status":  status,
            "is_doc":  ftype not in DIGITAL_FORMS,
        }
    return result

def onboard_status_badge(status: str) -> str:
    return {
        "not_started": "<span style='background:#f1f5f9;color:#64748b;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>Not Started</span>",
        "in_progress": "<span style='background:#fef3c7;color:#d97706;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>In Progress</span>",
        "completed":   "<span style='background:#dcfce7;color:#16a34a;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>✅ Complete</span>",
    }.get(status, status)


# ── Onboarding overview on staff profile ─────────────────────────────────────

@app.get("/staff/{staff_id}/onboarding", response_class=HTMLResponse)
def onboarding_overview(
    staff_id: int,
    session:  str | None = Cookie(default=None),
    msg:      str = ""
):
    redir, user = require_login(session)
    if redir: return redir

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    ob_status = get_onboarding_status(staff_id)
    is_owner  = user["role"] == "owner"
    flash     = f"<div class='flash-success'>{msg}</div>" if msg else ""

    # Build checklist
    checklist = ""
    all_done  = all(v["status"] == "completed" for v in ob_status.values())

    for ftype, info in ob_status.items():
        # Staff only see their own forms
        if info["who"] == "owner" and not is_owner:
            continue

        status  = info["status"]
        badge   = onboard_status_badge(status)
        is_doc  = info.get("is_doc", False)

        if is_doc:
            # Document-based — link to documents page
            btn_url = f"/staff/{staff_id}/documents"
            btn_lbl = "Go to Documents →" if status != "completed" else "View Documents →"
            btn_cls = "btn-primary" if status != "completed" else "btn-secondary"
        else:
            btn_url = f"/staff/{staff_id}/onboarding/{ftype}"
            btn_lbl = "Start →" if status == "not_started" else ("Continue →" if status == "in_progress" else "View →")
            btn_cls = "btn-primary" if status != "completed" else "btn-secondary"

        # Upload signed copy option (for forms only, not doc-based items)
        upload_html = ""
        if not is_doc and status != "completed":
            upload_html = f"""
        <form action='/staff/{staff_id}/onboarding/{ftype}/upload-paper' method='POST'
              enctype='multipart/form-data'
              style='display:inline-flex;gap:6px;align-items:center;margin-left:8px;margin-top:6px'>
          <input type='file' name='paper_form' accept='.pdf,.jpg,.jpeg,.png'
                 style='font-size:11px;max-width:160px;padding:3px'>
          <button type='submit' class='btn-secondary' style='padding:3px 8px;font-size:11px;white-space:nowrap'>
            &#128196; Upload Signed Copy (PDF/Scan)
          </button>
        </form>"""

        # PDF download if completed
        pdf_link = ""
        if status == "completed":
            pdf_row = q("SELECT pdf_path FROM onboarding_forms WHERE staff_id=? AND form_type=?",
                        (staff_id, ftype), fetch=True)
            if pdf_row and dict(pdf_row[0]).get("pdf_path"):
                pdf_link = f"<a href='/staff/{staff_id}/onboarding/{ftype}/pdf' target='_blank' style='color:#1e3a5f;font-size:12px;font-weight:700;margin-left:8px'>📄 PDF</a>"

        owner_tag = "<span style='font-size:10px;color:#94a3b8;margin-left:6px'>(owner only)</span>" if info["who"] == "owner" else ""

        checklist += f"""
        <div style='padding:14px 16px;border-bottom:1px solid #f1f5f9'>
          <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
            <div>
              <div style='font-weight:700;color:#0f172a;font-size:14px'>{info['label']}{owner_tag}</div>
              <div style='margin-top:4px'>{badge}{pdf_link}</div>
            </div>
            <a href='{btn_url}' class='{btn_cls}' style='padding:6px 16px;font-size:13px'>{btn_lbl}</a>
          </div>
          {upload_html}
        </div>"""

    completion_bar = ""
    completed_n = sum(1 for v in ob_status.values() if v["status"] == "completed")
    total_n     = len(ob_status)
    pct         = int(completed_n / total_n * 100)
    bar_col     = "#16a34a" if pct == 100 else ("#d97706" if pct > 0 else "#e2e8f0")
    completion_bar = f"""
    <div class='card'>
      <div style='display:flex;justify-content:space-between;margin-bottom:6px'>
        <span style='font-size:13px;font-weight:700;color:#0f172a'>Onboarding Progress</span>
        <span style='font-size:13px;font-weight:700;color:{bar_col}'>{completed_n}/{total_n} complete</span>
      </div>
      <div style='background:#f1f5f9;border-radius:99px;height:8px'>
        <div style='background:{bar_col};border-radius:99px;height:8px;width:{pct}%;transition:width .3s'></div>
      </div>
      {'<div style="font-size:12px;color:#16a34a;font-weight:700;margin-top:6px">🎉 All onboarding forms complete!</div>' if pct==100 else ''}
    </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>📋 Onboarding — {name}</div>
      </div>
    </div>
    {completion_bar}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>
        Onboarding Checklist
      </div>
      {checklist}
    </div>"""

    return page("Onboarding", content, user, "staff")


# ── Employment Application Form ───────────────────────────────────────────────

@app.get("/staff/{staff_id}/onboarding/employment_application", response_class=HTMLResponse)
def employment_application_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get any saved data
    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='employment_application'",
              (staff_id,), fetch=True)
    import json
    data = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}

    # Only pre-fill from profile if form has been started before
    has_data = bool(data)
    def fv(k, default=""): return data.get(k, default) if has_data else ""
    def fi(name, label, ftype="text", val=None, req=False, placeholder=""):
        v    = val if val is not None else fv(name)
        req_a = "required" if req else ""
        ph    = f"placeholder='{placeholder}'" if placeholder else ""
        return f"<div><label>{label}</label><input type='{ftype}' name='{name}' value='{v}' {req_a} {ph}></div>"

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>Employment Application — {name}</div>
      <div style='font-size:13px;color:#64748b;margin-top:2px'>Snappy Snaps — Equal Opportunity Employer</div>
    </div>

    <form action='/staff/{staff_id}/onboarding/employment_application' method='POST' class='space-y-6'>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Position & Personal Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('position_applied', 'Position Applied For', req=True)}
          {fi('full_name', 'Full Name', val=fv('full_name') or f"{s.get('first_name','')} {s.get('last_name','')}", req=True)}
          {fi('address',         'Address',               val=fv('address') or ', '.join(filter(None,[s.get('address_1',''),s.get('address_2',''),s.get('address_3',''),s.get('postcode','')])))}
          <div><label>Telephone No.</label>
            <input type='text' name='phone' value='{fv("phone") or s.get("phone","")}'
              placeholder='01234 567890'>
          </div>
          <div><label>Mobile No. <span style="font-size:10px;color:#94a3b8;font-weight:400">(preferred format: 07700 123456)</span></label>
            <input type='text' name='mobile' value='{fv("mobile")}'
              placeholder='07700 123456'>
          </div>
          {fi('ni_number',       'National Insurance No.',placeholder='AB 12 34 56 C')}
          <div><label>Driving Licence</label>
            <select name='driving_licence'>
              <option value=''>-- Select --</option>
              <option {'selected' if fv('driving_licence')=='Full' else ''}>Full</option>
              <option {'selected' if fv('driving_licence')=='Provisional' else ''}>Provisional</option>
              <option {'selected' if fv('driving_licence')=='None' else ''}>None</option>
            </select></div>
          <div><label>Work Permit Required?</label>
            <select name='work_permit'>
              <option value='No' {'selected' if fv('work_permit','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('work_permit')=='Yes' else ''}>Yes</option>
            </select></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Health Information</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('health_state', 'Current State of Health', placeholder='e.g. Good')}
          <div><label>Respiratory Problems?</label>
            <select name='respiratory'><option value='No'>No</option><option value='Yes' {'selected' if fv('respiratory')=='Yes' else ''}>Yes</option></select></div>
          <div><label>Skin Irritation?</label>
            <select name='skin_irritation'><option value='No'>No</option><option value='Yes' {'selected' if fv('skin_irritation')=='Yes' else ''}>Yes</option></select></div>
          <div style='grid-column:1/-1'>
            <label>Absence from work through illness in past 12 months</label>
            <textarea name='illness_absence' rows='2' placeholder='Please give details if any'>{fv('illness_absence')}</textarea>
          </div>
          <div><label>Do you smoke?</label>
            <select name='smoking'>
              <option value='Never' {'selected' if fv('smoking','Never')=='Never' else ''}>Never</option>
              <option value='Socially' {'selected' if fv('smoking')=='Socially' else ''}>Socially</option>
              <option value='Sometimes' {'selected' if fv('smoking')=='Sometimes' else ''}>Sometimes</option>
              <option value='Over 20/day' {'selected' if fv('smoking')=='Over 20/day' else ''}>Over 20/day</option>
            </select></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Educational History</div>
        <div class='grid gap-3' style='grid-template-columns:1fr'>
          <div><label>O Levels / GCSEs (subjects and grades)</label>
            <textarea name='gcse' rows='2'>{fv('gcse')}</textarea></div>
          <div><label>A Levels</label>
            <textarea name='a_levels' rows='2'>{fv('a_levels')}</textarea></div>
          <div><label>University / Degree</label>
            <input type='text' name='university' value='{fv('university')}'></div>
          <div><label>Other Qualifications or Skills</label>
            <textarea name='other_quals' rows='2'>{fv('other_quals')}</textarea></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>General Information</div>
        <div class='grid gap-3' style='grid-template-columns:1fr'>
          <div><label>What do you seek most from this position?</label>
            <textarea name='seeks' rows='2'>{fv('seeks')}</textarea></div>
          <div><label>Where do you see yourself in 5 years?</label>
            <textarea name='five_years' rows='2'>{fv('five_years')}</textarea></div>
          <div><label>Interests and hobbies</label>
            <textarea name='hobbies' rows='2'>{fv('hobbies')}</textarea></div>
          <div><label>Greatest strengths</label>
            <textarea name='strengths' rows='2'>{fv('strengths')}</textarea></div>
          <div><label>Greatest weaknesses</label>
            <textarea name='weaknesses' rows='2'>{fv('weaknesses')}</textarea></div>
          <div><label>Any court convictions or outstanding hearings?</label>
            <textarea name='convictions' rows='2' placeholder='Please declare if any'>{fv('convictions')}</textarea></div>
          <div><label>Have you previously applied to or worked at Snappy Snaps?</label>
            <select name='prev_snappy'>
              <option value='No' {'selected' if fv('prev_snappy','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('prev_snappy')=='Yes' else ''}>Yes</option>
            </select></div>
          {fi('prev_snappy_details', 'If yes, please give details', placeholder='Store, position, dates')}
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment History (most recent first)</div>"""

    for i in range(1, 4):
        content += f"""
        <div style='border:1px solid #e2e8f0;border-radius:10px;padding:14px;margin-bottom:10px'>
          <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Employer {i}</div>
          <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
            <div><label>Employer Name</label><input type='text' name='emp{i}_name' value='{fv(f"emp{i}_name")}'></div>
            <div><label>Address</label><input type='text' name='emp{i}_address' value='{fv(f"emp{i}_address")}'></div>
            <div><label>Date Commenced</label><input type='date' name='emp{i}_start' value='{fv(f"emp{i}_start")}'></div>
            <div><label>Date Left</label><input type='date' name='emp{i}_end' value='{fv(f"emp{i}_end")}'></div>
            <div><label>Position</label><input type='text' name='emp{i}_position' value='{fv(f"emp{i}_position")}'></div>
            <div><label>Salary</label><input type='text' name='emp{i}_salary' value='{fv(f"emp{i}_salary")}'></div>
            <div style='grid-column:1/-1'><label>Reason for leaving</label>
              <input type='text' name='emp{i}_reason' value='{fv(f"emp{i}_reason")}'></div>
          </div>
        </div>"""

    content += f"""
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>References</div>
        <div class='grid gap-6' style='grid-template-columns:1fr 1fr'>"""

    for i in range(1, 3):
        content += f"""
          <div>
            <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Reference {i}</div>
            <div class='space-y-2'>
              <div><label>Name</label><input type='text' name='ref{i}_name' value='{fv(f"ref{i}_name")}'></div>
              <div><label>Address</label><textarea name='ref{i}_address' rows='2'>{fv(f"ref{i}_address")}</textarea></div>
            </div>
          </div>"""

    content += f"""
        </div>
      </div>

      <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
        <div style='font-size:13px;color:#92400e;font-weight:600;margin-bottom:12px'>
          Declaration: I warrant that the information given is complete, true and accurate.
          I understand that any false statement may disqualify me from employment.
        </div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          {fi('declaration_name', 'Printed Name', val=f"{s.get('first_name','')} {s.get('last_name','')}", req=True)}
          {fi('declaration_date', 'Date', 'date', req=True)}
        </div>
      </div>

      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Submit & Generate PDF</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("Employment Application", content, user, "staff")


@app.post("/staff/{staff_id}/onboarding/employment_application")
async def save_employment_application(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    q("""INSERT INTO onboarding_forms (staff_id, form_type, status, started_at, completed_at, form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status,
            completed_at=excluded.completed_at,
            form_data=excluded.form_data""",
      (staff_id, "employment_application", status, now, now if status=="completed" else None,
       json.dumps(data)))

    # Update staff profile with key fields
    q("""UPDATE staff_profiles SET phone=?, address_1=?
         WHERE staff_id=?""",
      (data.get("mobile") or data.get("phone",""),
       data.get("address",""), staff_id))

    from urllib.parse import quote as uq
    msg = "Application submitted ✅" if status=="completed" else "Progress saved"
    return RedirectResponse(
        f"/staff/{staff_id}/onboarding?msg={uq(msg)}", status_code=303)


# ── P46 Form ──────────────────────────────────────────────────────────────────

@app.get("/staff/{staff_id}/onboarding/p46", response_class=HTMLResponse)
def p46_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='p46'",
              (staff_id,), fetch=True)
    import json
    data  = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}
    has_data = bool(data)
    def fv(k, d=""): return data.get(k, d) if has_data else d

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>P46 — Employee without a P45</div>
      <div style='font-size:13px;color:#64748b;margin-top:2px'>Section one — to be completed by the employee</div>
    </div>
    <form action='/staff/{staff_id}/onboarding/p46' method='POST' class='space-y-6'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Your Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Title</label>
            <select name='title'>
              <option>Mr</option><option>Mrs</option><option>Miss</option><option>Ms</option><option>Dr</option>
            </select></div>
          <div><label>Surname</label><input type='text' name='surname' value='{fv("surname", s.get("last_name",""))}' required></div>
          <div><label>First Name(s)</label><input type='text' name='first_name' value='{fv("first_name", s.get("first_name",""))}' required></div>
          <div><label>Gender</label>
            <select name='gender'>
              <option value='Male' {'selected' if fv('gender','Male')=='Male' else ''}>Male</option>
              <option value='Female' {'selected' if fv('gender')=='Female' else ''}>Female</option>
            </select></div>
          <div><label>Date of Birth</label><input type='date' name='dob' value='{fv("dob", s.get("date_of_birth",""))}' required></div>
          <div><label>National Insurance Number</label><input type='text' name='nino' value='{fv("nino")}' placeholder='AB 12 34 56 C' required></div>
          <div style='grid-column:1/-1'><label>Address</label>
            <input type='text' name='address' value='{fv("address", ", ".join(filter(None,[s.get("address_1",""),s.get("address_2",""),s.get("address_3",""),s.get("postcode","")])))}'></div>
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Your Present Circumstances</div>
        <div style='font-size:13px;color:#64748b;margin-bottom:12px'>Please select the statement that applies to you:</div>
        <div class='space-y-3'>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='A' {'checked' if fv('circumstance')=='A' else ''} style='width:auto;margin-top:3px'>
            <span><strong>A</strong> — This is my first job since last 6 April and I have not been receiving taxable Jobseeker's Allowance, Employment and Support Allowance or a state/occupational pension.</span>
          </label>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='B' {'checked' if fv('circumstance')=='B' else ''} style='width:auto;margin-top:3px'>
            <span><strong>B</strong> — This is now my only job, but since last 6 April I have had another job or received taxable Jobseeker's Allowance or Employment Support Allowance.</span>
          </label>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='C' {'checked' if fv('circumstance')=='C' else ''} style='width:auto;margin-top:3px'>
            <span><strong>C</strong> — I have another job or receive a state or occupational pension.</span>
          </label>
        </div>
        <div style='margin-top:12px'>
          <label style='display:flex;gap:10px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='student_loan' value='D' {'checked' if fv('student_loan')=='D' else ''} style='width:auto'>
            <span><strong>D</strong> — I have a Student Loan to repay (box D on P46)</span>
          </label>
        </div>
      </div>
      <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
        <div style='font-size:13px;color:#92400e;font-weight:600;margin-bottom:10px'>
          Declaration: I confirm that this information is correct.
        </div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          <div><label>Date</label><input type='date' name='sign_date' value='{fv("sign_date")}' required></div>
        </div>
      </div>

      {'"""' if user["role"] != "owner" else f"""
      <div class='card' style='border:2px solid #0f2942'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:4px'>Section 2 — To be completed by the Employer</div>
        <div style='font-size:12px;color:#94a3b8;margin-bottom:12px'>Owner only — not visible to staff</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Employer Name &amp; Address</label>
            <select name='s2_employer'>
              <option value='Sappy Properties (Uxbridge) Llp T/A Snappy Snaps, 178 High Street, Uxbridge, Middlesex, UB8 1LA' {'selected' if fv('s2_employer','').startswith('Sappy') else ''}>
                Snappy Snaps Uxbridge — 178 High Street, Uxbridge UB8 1LA
              </option>
              <option value='Maukbs Ltd T/A Snappy Snaps, 95 Northbrook Street, Newbury, Berkshire, RG14 1AA' {'selected' if fv('s2_employer','').startswith('Maukbs') else ''}>
                Snappy Snaps Newbury — 95 Northbrook Street, Newbury RG14 1AA
              </option>
            </select>
          </div>
          <div><label>Date Employment Started</label>
            <input type='date' name='s2_start_date' value='{fv("s2_start_date") or s.get("date_joined","")}'></div>
          <div><label>Job Title</label>
            <input type='text' name='s2_job_title' value='{fv("s2_job_title") or s.get("job_title","")}'
              placeholder='e.g. Sales Assistant'></div>
          <div><label>Works/Payroll Number</label>
            <input type='text' name='s2_payroll_no' value='{fv("s2_payroll_no")}' placeholder='e.g. P001'></div>
          <div><label>Employer PAYE Reference</label>
            <input type='text' name='s2_paye_ref' value='{fv("s2_paye_ref")}' placeholder='e.g. 123/AB456'></div>
          <div><label>Tax Code Used</label>
            <input type='text' name='s2_tax_code' value='{fv("s2_tax_code")}' placeholder='e.g. 1257L'></div>
        </div>
        <div style='margin-top:12px'>
          <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Tax Code Basis</div>
          <div class='space-y-2'>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='A_cumulative' {'checked' if fv("s2_tax_basis")=="A_cumulative" else ''} style='width:auto'>
              Box A — Emergency code on a cumulative basis
            </label>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='B_week1' {'checked' if fv("s2_tax_basis")=="B_week1" else ''} style='width:auto'>
              Box B — Emergency code on a non-cumulative Week 1/Month 1 basis
            </label>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='C_BR' {'checked' if fv("s2_tax_basis")=="C_BR" else ''} style='width:auto'>
              Box C — Code BR (or 0T if employee fails to complete Section 1) Week 1/Month 1 basis
            </label>
          </div>
        </div>
      </div>""" if user["role"] == "owner" else ""}

      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Submit</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("P46", content, user, "staff")


@app.post("/staff/{staff_id}/onboarding/p46")
async def save_p46(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    q("""INSERT INTO onboarding_forms (staff_id,form_type,status,started_at,completed_at,form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status, completed_at=excluded.completed_at, form_data=excluded.form_data""",
      (staff_id,"p46",status,now,now if status=="completed" else None, json.dumps(data)))
    # Update DOB on profile
    if data.get("dob"):
        q("UPDATE staff_profiles SET date_of_birth=? WHERE staff_id=?", (data["dob"], staff_id))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/onboarding?msg={uq('P46 saved')}", status_code=303)


# ── New Employee Notification (owner only) ────────────────────────────────────

@app.get("/staff/{staff_id}/onboarding/new_employee_notify", response_class=HTMLResponse)
def new_employee_notify_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/onboarding", status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='new_employee_notify'",
              (staff_id,), fetch=True)
    import json
    data     = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}
    has_data = bool(data)
    def fv(k, d=""): return data.get(k, d) if has_data else d

    def fi(nm, lbl, ft="text", val=None, ph=""):
        v  = val if val is not None else fv(nm)
        ph = f"placeholder='{ph}'" if ph else ""
        return f"<div><label>{lbl}</label><input type='{ft}' name='{nm}' value='{v}' {ph}></div>"

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>New Employee Notification — {name}</div>
      <div style='font-size:12px;color:#94a3b8;margin-top:2px'>Owner only — not visible to staff</div>
    </div>
    <form action='/staff/{staff_id}/onboarding/new_employee_notify' method='POST' class='space-y-6'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employee Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('surname',      'Surname',       val=fv('surname') or s.get('last_name',''))}
          {fi('first_name',   'First Name(s)', val=fv('first_name') or s.get('first_name',''))}
          <div><label>Title</label><select name='title'><option>Mr</option><option>Mrs</option><option>Miss</option><option>Ms</option></select></div>
          <div><label>Gender</label><select name='gender'><option>Male</option><option>Female</option></select></div>
          <div><label>Married</label><select name='married'><option value='No'>No</option><option value='Yes' {'selected' if fv('married')=='Yes' else ''}>Yes</option></select></div>
          {fi('dob',          'Date of Birth', 'date', s.get('date_of_birth',''))}
          {fi('nino',         'NI Number',     ph='AB 12 34 56 C')}
          {fi('start_date',   'Start Date',    'date', s.get('date_joined',''))}
          {fi('address',      'Employee Address', val=fv('address') or ', '.join(filter(None,[s.get('address_1',''),s.get('address_2',''),s.get('address_3',''),s.get('postcode','')])))}
          {fi('postcode',     'Post Code',     val=fv('postcode') or s.get('postcode',''))}
          {fi('phone',        'Phone',         val=fv('phone') or s.get('phone',''))}
          {fi('emergency',    'Emergency Contact')}
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment Details (Employer)</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Employer Name & Address</label>
            <select name='employer_name'>
              <option value='Sappy Properties (Uxbridge) Llp T/A Snappy Snaps, 178 High Street, Uxbridge, Middlesex, UB8 1LA' {'selected' if "Uxbridge" in fv("employer_name","") else ""}>
                Snappy Snaps Uxbridge — 178 High Street, Uxbridge, Middlesex UB8 1LA
              </option>
              <option value='Maukbs Ltd T/A Snappy Snaps, 95 Northbrook Street, Newbury, Berkshire, RG14 1AA' {'selected' if "Newbury" in fv("employer_name","") else ""}>
                Snappy Snaps Newbury — 95 Northbrook Street, Newbury, Berkshire RG14 1AA
              </option>
            </select>
          </div>
          <div><label>Pay Frequency</label>
            <select name='pay_frequency'>
              <option {'selected' if fv('pay_frequency','Monthly')=='Monthly' else ''}>Monthly</option>
              <option {'selected' if fv('pay_frequency')=='Weekly' else ''}>Weekly</option>
              <option {'selected' if fv('pay_frequency')=='4 Weekly' else ''}>4 Weekly</option>
            </select></div>
          {fi('pay_day',       'Pay Day & Date', ph='e.g. Last Friday of month')}
          <div><label>Pay Method</label>
            <select name='pay_method'>
              <option {'selected' if fv('pay_method','BACS')=='BACS' else ''}>BACS</option>
              <option {'selected' if fv('pay_method')=='Cash' else ''}>Cash</option>
              <option {'selected' if fv('pay_method')=='Cheque' else ''}>Cheque</option>
            </select></div>
          {fi('payroll_no',    'Payroll No.',    ph='e.g. P001')}
          {fi('tax_code',      'Tax Code',       ph='e.g. 1257L')}
          {fi('nic_letter',    'NIC Letter',     ph='e.g. A')}
          {fi('contracted_hrs','Contracted Hours/Week', val=fv('contracted_hrs') or str(s.get('contracted_hrs','')))}
          {fi('wage',          'Wage/Salary',    ph='e.g. £12.71 per hour')}
          {fi('holiday_start', 'Holiday Year Start', 'date', fv('holiday_start','2026-01-01'))}
          {fi('holiday_end',   'Holiday Year End',   'date', fv('holiday_end','2026-12-31'))}
          {fi('holiday_days',  'Holiday Entitlement (days)', ph='e.g. 19')}
          <div><label>Employment Type</label>
            <select name='emp_type'>
              <option {'selected' if fv('emp_type','Permanent')=='Permanent' else ''}>Permanent</option>
              <option {'selected' if fv('emp_type')=='Temporary' else ''}>Temporary</option>
            </select></div>
          <div><label>Student?</label>
            <select name='is_student'>
              <option value='No' {'selected' if fv('is_student','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('is_student')=='Yes' else ''}>Yes</option>
            </select></div>
          <div><label>Only Employment?</label>
            <select name='only_employment'>
              <option value='Yes' {'selected' if fv('only_employment','Yes')=='Yes' else ''}>Yes</option>
              <option value='No' {'selected' if fv('only_employment')=='No' else ''}>No</option>
            </select></div>
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Right to Work Documents Checked</div>
        <div class='space-y-2'>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_passport' value='1' {'checked' if fv('rtw_passport') else ''} style='width:auto'>
            UK or EEA Passport
          </label>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_birth_cert' value='1' {'checked' if fv('rtw_birth_cert') else ''} style='width:auto'>
            Full British Birth Certificate
          </label>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_work_permit' value='1' {'checked' if fv('rtw_work_permit') else ''} style='width:auto'>
            Work Permit with Passport
          </label>
        </div>
      </div>
      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Mark Complete</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("New Employee Notification", content, user, "staff")


@app.post("/staff/{staff_id}/onboarding/new_employee_notify")
async def save_new_employee_notify(
    staff_id: int, request: Request, session: str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    q("""INSERT INTO onboarding_forms (staff_id,form_type,status,started_at,completed_at,form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status, completed_at=excluded.completed_at, form_data=excluded.form_data""",
      (staff_id,"new_employee_notify",status,now,now if status=="completed" else None,json.dumps(data)))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/onboarding?msg={uq('Notification saved')}", status_code=303)



# ── Upload paper onboarding form ──────────────────────────────────────────────

@app.post("/staff/{staff_id}/onboarding/{form_type}/upload-paper")
async def upload_paper_form(
    staff_id:  int,
    form_type: str,
    request:   Request,
    session:   str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    form      = await request.form()
    paper     = form.get("paper_form")

    if not paper or not hasattr(paper, "filename") or not paper.filename:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/onboarding?msg={uq('No file selected')}&msg_type=error",
            status_code=303)

    # Save the file
    ext      = os.path.splitext(paper.filename)[1].lower()
    filename = f"onboard_{staff_id}_{form_type}_paper{ext}"
    filepath = os.path.join(DOCS_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(await paper.read())

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Mark form as completed with paper upload
    q("""INSERT INTO onboarding_forms
            (staff_id, form_type, status, started_at, completed_at, form_data, pdf_path)
         VALUES(?,?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status='completed',
            completed_at=excluded.completed_at,
            pdf_path=excluded.pdf_path""",
      (staff_id, form_type, "completed", now, now,
       '{"source":"paper_upload"}', filepath))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/onboarding?msg={uq('Paper form uploaded and marked complete')}",
        status_code=303)



# ══════════════════════════════════════════════════════════════════════════════
# MODULE 4 — ROTA & CLOCKING
# ══════════════════════════════════════════════════════════════════════════════

import math
from datetime import timedelta

DAYS = ["Sun","Mon","Tue","Wed","Thu","Fri","Sat"]
FULL_DAYS = ["Sunday","Monday","Tuesday","Wednesday","Thursday","Friday","Saturday"]

def calc_paid_hours(raw_hrs: float) -> float:
    """Deduct 30 min unpaid break for shifts of 4 hours or more."""
    if not raw_hrs: return 0.0
    return round(raw_hrs - 0.5, 2) if raw_hrs >= 4.0 else round(raw_hrs, 2)

def parse_hours(start: str, end: str) -> float:
    """Calculate raw hours between two HH:MM time strings."""
    try:
        sh, sm = map(int, start.split(':'))
        eh, em = map(int, end.split(':'))
        return (eh*60+em - sh*60-sm) / 60
    except: return 0.0


def ensure_rota_tables():
    conn = db()
    c    = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS rota_templates (
            template_id  INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id     INTEGER NOT NULL,
            store_name   TEXT NOT NULL,
            day_of_week  INTEGER NOT NULL,
            shift_start  TEXT,
            shift_end    TEXT,
            hours        REAL,
            is_off       INTEGER DEFAULT 0,
            UNIQUE(staff_id, day_of_week),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS rotas (
            rota_id      INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name   TEXT NOT NULL,
            week_start   TEXT NOT NULL,
            status       TEXT DEFAULT 'draft',
            published_at TEXT,
            published_by TEXT,
            notes        TEXT,
            UNIQUE(store_name, week_start)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS rota_shifts (
            shift_id     INTEGER PRIMARY KEY AUTOINCREMENT,
            rota_id      INTEGER NOT NULL,
            staff_id     INTEGER NOT NULL,
            shift_date   TEXT NOT NULL,
            shift_start  TEXT,
            shift_end    TEXT,
            hours        REAL,
            is_off       INTEGER DEFAULT 0,
            absence_type TEXT,
            notes        TEXT,
            UNIQUE(rota_id, staff_id, shift_date),
            FOREIGN KEY (rota_id)  REFERENCES rotas(rota_id),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    conn.commit()
    conn.close()

ensure_rota_tables()


def get_week_start(date_str: str = None) -> str:
    """Return Sunday of the week containing the given date (or today)."""
    d = datetime.strptime(date_str, "%Y-%m-%d") if date_str else datetime.now()
    # Go back to Sunday
    days_since_sunday = d.weekday() + 1  # weekday() 0=Mon, so +1 for Sun
    if days_since_sunday == 7:
        days_since_sunday = 0
    return (d - timedelta(days=days_since_sunday)).strftime("%Y-%m-%d")


def get_week_dates(week_start: str) -> list:
    """Return list of 7 date strings for Sun-Sat week."""
    start = datetime.strptime(week_start, "%Y-%m-%d")
    return [(start + timedelta(days=i)).strftime("%Y-%m-%d") for i in range(7)]


def get_or_create_rota(store_name: str, week_start: str) -> dict:
    """Get existing rota or create from templates."""
    rows = q("SELECT * FROM rotas WHERE store_name=? AND week_start=?",
             (store_name, week_start), fetch=True)
    if rows:
        rota = dict(rows[0])
        # Load shifts
        shifts = q("""SELECT rs.*, sp.first_name, sp.last_name
                      FROM rota_shifts rs
                      JOIN staff_profiles sp ON rs.staff_id=sp.staff_id
                      WHERE rs.rota_id=?""",
                   (rota["rota_id"],), fetch=True) or []
        rota["shifts"] = [dict(s) for s in shifts]
        return rota

    # Create new rota from templates
    q("INSERT OR IGNORE INTO rotas (store_name, week_start, status) VALUES(?,?,'draft')",
      (store_name, week_start))
    rota_rows = q("SELECT * FROM rotas WHERE store_name=? AND week_start=?",
                  (store_name, week_start), fetch=True)
    rota     = dict(rota_rows[0])
    rota_id  = rota["rota_id"]
    week_dates = get_week_dates(week_start)

    # Get active staff for this store
    staff = q("SELECT * FROM staff_profiles WHERE store_name=? AND is_active=1",
              (store_name,), fetch=True) or []

    # Get approved leave for this week
    leave_map = {}
    for s in staff:
        leaves = q("""SELECT date_from, date_to, leave_type FROM leave_requests
                      WHERE staff_id=? AND status='approved'
                        AND date_from <= ? AND date_to >= ?""",
                   (s["staff_id"], week_dates[-1], week_dates[0]), fetch=True) or []
        for lv in leaves:
            lv = dict(lv)
            d1 = datetime.strptime(lv["date_from"], "%Y-%m-%d")
            d2 = datetime.strptime(lv["date_to"],   "%Y-%m-%d")
            cur = d1
            while cur <= d2:
                leave_map[(s["staff_id"], cur.strftime("%Y-%m-%d"))] = lv["leave_type"]
                cur += timedelta(days=1)

    # Build shifts from templates
    for s in staff:
        sid = s["staff_id"]
        templates = q("SELECT * FROM rota_templates WHERE staff_id=?", (sid,), fetch=True) or []
        tmpl_map  = {dict(t)["day_of_week"]: dict(t) for t in templates}

        for i, date_str in enumerate(week_dates):
            dow      = i  # 0=Sun
            tmpl     = tmpl_map.get(dow, {})
            leave_type = leave_map.get((sid, date_str))

            if leave_type:
                q("""INSERT OR IGNORE INTO rota_shifts
                        (rota_id, staff_id, shift_date, is_off, absence_type)
                     VALUES(?,?,?,1,?)""",
                  (rota_id, sid, date_str, leave_type))
            elif tmpl.get("is_off", 1):
                q("""INSERT OR IGNORE INTO rota_shifts
                        (rota_id, staff_id, shift_date, is_off)
                     VALUES(?,?,?,1)""",
                  (rota_id, sid, date_str))
            else:
                q("""INSERT OR IGNORE INTO rota_shifts
                        (rota_id, staff_id, shift_date,
                         shift_start, shift_end, hours, is_off)
                     VALUES(?,?,?,?,?,?,0)""",
                  (rota_id, sid, date_str,
                   tmpl.get("shift_start"), tmpl.get("shift_end"),
                   tmpl.get("hours")))

    return get_or_create_rota(store_name, week_start)


# ── Rota page ─────────────────────────────────────────────────────────────────

@app.get("/rota", response_class=HTMLResponse)
def rota_page(
    session:    str | None = Cookie(default=None),
    store:      str = "Uxbridge",
    week_start: str = "",
    msg:        str = ""
):
    redir, user = require_login(session)
    if redir: return redir

    # Store access control
    if user["role"] == "staff" and user.get("store_name"):
        store = user["store_name"]
    elif user["role"] == "manager" and user.get("store_name") and not store:
        store = user["store_name"]

    if not week_start:
        week_start = get_week_start()

    week_dates  = get_week_dates(week_start)
    prev_week   = (datetime.strptime(week_start, "%Y-%m-%d") - timedelta(days=7)).strftime("%Y-%m-%d")
    next_week   = (datetime.strptime(week_start, "%Y-%m-%d") + timedelta(days=7)).strftime("%Y-%m-%d")
    week_end    = week_dates[-1]
    is_mgr      = user["role"] in ("owner","manager")
    flash       = f"<div class='flash-success'>{msg}</div>" if msg else ""

    rota = get_or_create_rota(store, week_start)
    rota_id = rota["rota_id"]
    status  = rota["status"]

    # Build shift lookup: (staff_id, date) → shift
    shift_lookup = {(s["staff_id"], s["shift_date"]): s for s in rota.get("shifts",[])}

    # Get active staff for this store
    staff = q("SELECT * FROM staff_profiles WHERE store_name=? AND is_active=1 ORDER BY first_name",
              (store,), fetch=True) or []

    # Status badge
    status_colours = {
        "draft":     ("#fef3c7","#92400e","Draft"),
        "published": ("#dcfce7","#166534","Published"),
        "locked":    ("#dbeafe","#1e40af","Locked"),
    }
    sc = status_colours.get(status, ("#f1f5f9","#64748b","Unknown"))
    status_badge = f"<span style='background:{sc[0]};color:{sc[1]};font-size:12px;font-weight:700;padding:3px 10px;border-radius:6px'>{sc[2]}</span>"

    # Store switcher (managers/owners only)
    store_switcher = ""
    if user["role"] in ("owner","manager"):
        for sv in ["Uxbridge","Newbury"]:
            cls = "btn-primary" if sv == store else "btn-secondary"
            store_switcher += f"<a href='/rota?store={sv}&week_start={week_start}' class='{cls}' style='padding:5px 14px;font-size:13px'>{sv}</a>"

    # Action buttons
    action_btns = ""
    if is_mgr:
        if status == "draft":
            action_btns += f"<a href='/rota/publish?store={store}&week_start={week_start}' class='btn-primary' style='padding:6px 16px;font-size:13px'>&#128228; Publish Rota</a>"
        elif status == "published":
            action_btns += f"<a href='/rota/unpublish?store={store}&week_start={week_start}' class='btn-secondary' style='padding:6px 16px;font-size:13px'>&#128221; Back to Draft</a>"
        action_btns += f"<a href='/rota/pdf?store={store}&week_start={week_start}' class='btn-secondary' style='padding:6px 16px;font-size:13px'>&#128196; Download PDF</a>"
        action_btns += f"<a href='/rota/templates?store={store}' class='btn-secondary' style='padding:6px 16px;font-size:13px'>&#9881;&#65039; Templates</a>"

    # Build rota grid
    # Header row
    header = "<tr style='background:#0f2942;color:white'>"
    header += "<th style='padding:10px 12px;text-align:left;font-size:12px;min-width:140px'>Staff</th>"
    for i, date_str in enumerate(week_dates):
        d    = datetime.strptime(date_str, "%Y-%m-%d")
        day  = DAYS[i]
        date = d.strftime("%d %b")
        is_today = date_str == datetime.now().strftime("%Y-%m-%d")
        bg = "background:#1e3a5f" if is_today else ""
        header += f"<th style='padding:8px 6px;text-align:center;font-size:11px;min-width:90px;{bg}'>{day}<br><span style='font-size:10px;opacity:.7'>{date}</span></th>"
    header += "<th style='padding:10px 8px;font-size:11px;text-align:center'>Hrs</th></tr>"

    # Staff rows
    grid_rows = ""
    for s in staff:
        sid  = s["staff_id"]
        name = f"{s['first_name']} {s['last_name']}"
        total_hrs = 0

        grid_rows += f"<tr style='border-bottom:1px solid #f1f5f9'>"
        grid_rows += f"<td style='padding:8px 12px;font-weight:700;font-size:13px;color:#0f172a;white-space:nowrap'>{name}</td>"

        for i, date_str in enumerate(week_dates):
            shift = shift_lookup.get((sid, date_str), {})
            is_off      = shift.get("is_off", 1)
            absence     = shift.get("absence_type")
            start       = shift.get("shift_start") or ""
            end         = shift.get("shift_end") or ""
            hrs         = shift.get("hours") or 0
            shift_id    = shift.get("shift_id","")
            if hrs: total_hrs += hrs

            if absence:
                absence_colours = {
                    "H": ("#dcfce7","#166534","Holiday"),
                    "S": ("#fee2e2","#dc2626","Sick"),
                    "B": ("#fef3c7","#92400e","BH"),
                    "AL":("#dbeafe","#1d4ed8","Auth.Leave"),
                }
                ac = absence_colours.get(absence, ("#f1f5f9","#64748b", absence))
                cell = f"<div style='background:{ac[0]};color:{ac[1]};border-radius:6px;padding:4px 6px;font-size:11px;font-weight:700;text-align:center'>{ac[2]}</div>"
            elif is_off:
                cell = "<div style='color:#cbd5e1;font-size:12px;text-align:center'>OFF</div>"
            else:
                cell = f"<div style='background:#eff6ff;border-radius:6px;padding:4px 6px;text-align:center'>"
                cell += f"<div style='font-size:12px;font-weight:700;color:#1e40af'>{start}–{end}</div>"
                if hrs: cell += f"<div style='font-size:10px;color:#64748b'>{hrs}h</div>"
                cell += "</div>"

            # Editable if draft/published and manager
            if is_mgr and status in ("draft","published"):
                cell = f"<a href='/rota/edit-shift?rota_id={rota_id}&staff_id={sid}&date={date_str}' style='text-decoration:none;display:block'>{cell}</a>"

            grid_rows += f"<td style='padding:4px 3px'>{cell}</td>"

        hrs_str = f"<span style='font-size:12px;font-weight:700;color:#0f172a'>{total_hrs:.1f}</span>" if total_hrs else "—"
        grid_rows += f"<td style='padding:4px 8px;text-align:center'>{hrs_str}</td>"
        grid_rows += "</tr>"

    # Weekly totals row — hours and staff count per day
    day_totals_hrs   = []
    day_totals_count = []
    week_total  = 0
    week_staff  = set()

    for date_str in week_dates:
        day_hrs   = 0
        day_count = 0
        for s in staff:
            sid   = dict(s)["staff_id"]
            shift = shift_lookup.get((sid, date_str), {})
            hrs   = shift.get("hours") or 0
            if hrs and not shift.get("is_off") and not shift.get("absence_type"):
                day_hrs   += hrs
                day_count += 1
                week_staff.add(sid)
        week_total += day_hrs

        day_totals_hrs.append(
            f"<td style='padding:4px 3px;text-align:center;background:#f8fafc;border-right:1px solid #e2e8f0'>"
            f"<div style='font-size:12px;font-weight:700;color:#0f2942'>{day_hrs:.1f}h</div>"
            f"<div style='font-size:10px;color:#64748b;margin-top:1px'>{day_count} staff</div>"
            f"</td>"
        )

    totals_row = "<tr style='border-top:2px solid #e2e8f0'>"
    totals_row += "<td style='padding:8px 12px;font-size:11px;font-weight:700;color:#64748b;background:#f8fafc'>TOTALS</td>"
    totals_row += "".join(day_totals_hrs)
    totals_row += (
        f"<td style='padding:8px;text-align:center;background:#f8fafc'>"
        f"<div style='font-size:13px;font-weight:900;color:#0f2942'>{week_total:.1f}h</div>"
        f"<div style='font-size:10px;color:#64748b'>week</div>"
        f"</td>"
    )
    totals_row += "</tr>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <div class='text-2xl font-black text-slate-800'>&#128197; Rota — {store}</div>
        <div style='font-size:13px;color:#64748b;margin-top:2px'>
          Week: {datetime.strptime(week_start,"%Y-%m-%d").strftime("%d %b")} – {datetime.strptime(week_end,"%Y-%m-%d").strftime("%d %b %Y")}
          &nbsp; {status_badge}
        </div>
      </div>
      <div style='display:flex;gap:8px;flex-wrap:wrap;align-items:center'>
        {store_switcher}
        <a href='/rota?store={store}&week_start={prev_week}' class='btn-secondary' style='padding:5px 12px'>&#8592; Prev</a>
        <a href='/rota?store={store}&week_start={get_week_start()}' class='btn-secondary' style='padding:5px 12px'>Today</a>
        <a href='/rota?store={store}&week_start={next_week}' class='btn-secondary' style='padding:5px 12px'>Next &#8594;</a>
        {action_btns}
      </div>
    </div>

    <!-- Legend -->
    <div style='display:flex;gap:12px;flex-wrap:wrap;font-size:12px;font-weight:600'>
      <span><span style='display:inline-block;width:12px;height:12px;background:#eff6ff;border-radius:3px;vertical-align:middle'></span> Working</span>
      <span><span style='display:inline-block;width:12px;height:12px;background:#f1f5f9;border-radius:3px;vertical-align:middle'></span> OFF</span>
      <span><span style='display:inline-block;width:12px;height:12px;background:#dcfce7;border-radius:3px;vertical-align:middle'></span> Holiday</span>
      <span><span style='display:inline-block;width:12px;height:12px;background:#fee2e2;border-radius:3px;vertical-align:middle'></span> Sick</span>
      <span><span style='display:inline-block;width:12px;height:12px;background:#fef3c7;border-radius:3px;vertical-align:middle'></span> Bank Holiday</span>
      {'<span style="color:#64748b;font-size:11px">Click any cell to edit shift</span>' if is_mgr and status in ("draft","published") else ''}
    </div>

    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table style='width:100%;border-collapse:collapse;font-family:DM Sans,sans-serif'>
          <thead>{header}</thead>
          <tbody>{grid_rows}{totals_row}</tbody>
        </table>
      </div>
    </div>"""

    return page("Rota", content, user, "rota")


# ── Edit individual shift ─────────────────────────────────────────────────────

@app.get("/rota/edit-shift", response_class=HTMLResponse)
def edit_shift_form(
    rota_id:  int = 0,
    staff_id: int = 0,
    date:     str = "",
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/rota", status_code=303)

    # Get rota info
    rota = q("SELECT * FROM rotas WHERE rota_id=?", (rota_id,), fetch=True)
    if not rota: return RedirectResponse("/rota", status_code=303)
    rota = dict(rota[0])

    # Get staff info
    staff = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not staff: return RedirectResponse("/rota", status_code=303)
    s    = dict(staff[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get existing shift
    shift = q("SELECT * FROM rota_shifts WHERE rota_id=? AND staff_id=? AND shift_date=?",
              (rota_id, staff_id, date), fetch=True)
    sh = dict(shift[0]) if shift else {}

    d_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%A %d %B %Y")
    store = rota["store_name"]
    back  = f"/rota?store={store}&week_start={rota['week_start']}"

    absence_opts = "".join(
        f"<option value='{k}' {'selected' if sh.get('absence_type')==k else ''}>{v}</option>"
        for k,v in ABSENCE_TYPES.items()
    )

    content = f"""
    <div>
      <a href='{back}' style='color:#1e3a5f;font-size:13px;font-weight:700'>&#8592; Back to Rota</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>Edit Shift — {name}</div>
      <div style='font-size:13px;color:#64748b'>{d_fmt} &middot; {store}</div>
    </div>
    <div class='card' style='max-width:480px'>
      <form action='/rota/save-shift' method='POST' class='space-y-4'>
        <input type='hidden' name='rota_id'  value='{rota_id}'>
        <input type='hidden' name='staff_id' value='{staff_id}'>
        <input type='hidden' name='date'     value='{date}'>
        <input type='hidden' name='store'    value='{store}'>
        <input type='hidden' name='week_start' value='{rota["week_start"]}'>

        <div>
          <label>Shift Type</label>
          <select name='shift_type' id='shift_type' onchange='toggleFields()'>
            <option value='working' {'selected' if not sh.get('is_off') and not sh.get('absence_type') else ''}>Working</option>
            <option value='off' {'selected' if sh.get('is_off') and not sh.get('absence_type') else ''}>OFF</option>
            <option value='absence' {'selected' if sh.get('absence_type') else ''}>Absence / Leave</option>
          </select>
        </div>

        <div id='working_fields' style='display:{"block" if not sh.get("is_off") and not sh.get("absence_type") else "none"}'>
          <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
            <div><label>Start Time</label>
              <input type='time' name='shift_start' value='{sh.get("shift_start") or "09:00"}'></div>
            <div><label>End Time</label>
              <input type='time' name='shift_end' value='{sh.get("shift_end") or "17:00"}'></div>
          </div>
          <div style='margin-top:8px'><label>Total Hours (optional — auto-calculates)</label>
            <input type='number' step='0.25' name='hours' id='hours_field'
                   value='{sh.get("hours") or ""}' placeholder='e.g. 7.5'></div>
        </div>

        <div id='absence_fields' style='display:{"block" if sh.get("absence_type") else "none"}'>
          <label>Absence Type</label>
          <select name='absence_type'>
            <option value=''>-- Select --</option>
            {absence_opts}
          </select>
        </div>

        <div><label>Notes (optional)</label>
          <input type='text' name='notes' value='{sh.get("notes") or ""}' placeholder='e.g. Cover for Jessica'></div>

        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>&#128190; Save Shift</button>
          <a href='{back}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>
    <script>
    function toggleFields() {{
      const type = document.getElementById('shift_type').value;
      document.getElementById('working_fields').style.display = type==='working' ? 'block' : 'none';
      document.getElementById('absence_fields').style.display = type==='absence' ? 'block' : 'none';
    }}
    // Auto-calc hours from times
    document.addEventListener('DOMContentLoaded', function() {{
      const start = document.querySelector('[name="shift_start"]');
      const end   = document.querySelector('[name="shift_end"]');
      const hrs   = document.getElementById('hours_field');
      function calc() {{
        if (start.value && end.value) {{
          const [sh,sm] = start.value.split(':').map(Number);
          const [eh,em] = end.value.split(':').map(Number);
          const diff = (eh*60+em - sh*60-sm) / 60;
          if (diff > 0) hrs.value = diff.toFixed(2);
        }}
      }}
      if (start) start.addEventListener('change', calc);
      if (end)   end.addEventListener('change', calc);
    }});
    </script>"""

    return page("Edit Shift", content, user, "rota")


@app.post("/rota/save-shift")
async def save_shift(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form        = await request.form()
    rota_id     = int(form.get("rota_id", 0))
    staff_id    = int(form.get("staff_id", 0))
    date        = form.get("date","")
    store       = form.get("store","")
    week_start  = form.get("week_start","")
    shift_type  = form.get("shift_type","off")
    shift_start = form.get("shift_start","") or None
    shift_end   = form.get("shift_end","")   or None
    notes       = form.get("notes","")       or None
    absence     = form.get("absence_type","") or None
    try:
        hrs = float(form.get("hours",0) or 0)
    except:
        hrs = 0

    is_off = 1 if shift_type != "working" else 0

    q("""INSERT INTO rota_shifts
            (rota_id, staff_id, shift_date, shift_start, shift_end,
             hours, is_off, absence_type, notes)
         VALUES(?,?,?,?,?,?,?,?,?)
         ON CONFLICT(rota_id, staff_id, shift_date) DO UPDATE SET
            shift_start=excluded.shift_start,
            shift_end=excluded.shift_end,
            hours=excluded.hours,
            is_off=excluded.is_off,
            absence_type=excluded.absence_type,
            notes=excluded.notes""",
      (rota_id, staff_id, date, shift_start, shift_end,
       hrs, is_off, absence if shift_type=="absence" else None, notes))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/rota?store={store}&week_start={week_start}&msg={uq('Shift saved')}",
        status_code=303)


# ── Publish / Unpublish ───────────────────────────────────────────────────────

@app.get("/rota/publish")
def publish_rota(
    store:      str = "",
    week_start: str = "",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/rota", status_code=303)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    q("""UPDATE rotas SET status='published', published_at=?, published_by=?
         WHERE store_name=? AND week_start=?""",
      (now, user.get("username"), store, week_start))
    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/rota?store={store}&week_start={week_start}&msg={uq('Rota published — staff can now see their shifts')}",
        status_code=303)


@app.get("/rota/unpublish")
def unpublish_rota(
    store:      str = "",
    week_start: str = "",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    q("UPDATE rotas SET status='draft' WHERE store_name=? AND week_start=?",
      (store, week_start))
    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/rota?store={store}&week_start={week_start}&msg={uq('Rota moved back to draft')}",
        status_code=303)


# ── WhatsApp text export ──────────────────────────────────────────────────────

@app.get("/rota/whatsapp", response_class=HTMLResponse)
def rota_whatsapp(
    store:      str = "",
    week_start: str = "",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    week_dates = get_week_dates(week_start)
    week_end   = week_dates[-1]
    rota = get_or_create_rota(store, week_start)
    shift_lookup = {(s["staff_id"], s["shift_date"]): s for s in rota.get("shifts",[])}
    staff = q("SELECT * FROM staff_profiles WHERE store_name=? AND is_active=1 ORDER BY first_name",
              (store,), fetch=True) or []

    lines = [
        f"📅 *{store} Rota*",
        f"*Week: {datetime.strptime(week_start,'%Y-%m-%d').strftime('%d %b')} – {datetime.strptime(week_end,'%Y-%m-%d').strftime('%d %b %Y')}*",
        ""
    ]

    for s in staff:
        sid   = s["staff_id"]
        name  = s["first_name"]
        shifts = []
        for i, date_str in enumerate(week_dates):
            sh = shift_lookup.get((sid, date_str), {})
            if sh.get("absence_type"):
                at = ABSENCE_TYPES.get(sh["absence_type"], sh["absence_type"])
                shifts.append(f"{DAYS[i]}: {at}")
            elif not sh.get("is_off", 1):
                start = sh.get("shift_start","")
                end   = sh.get("shift_end","")
                shifts.append(f"{DAYS[i]}: {start}–{end}")
        if shifts:
            lines.append(f"*{name}*")
            lines.append("  " + " | ".join(shifts))
        else:
            lines.append(f"*{name}*: All OFF")
        lines.append("")

    message = "\n".join(lines)

    content = f"""
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>&#128242; WhatsApp Rota — {store}</div>
      <a href='/rota?store={store}&week_start={week_start}' class='btn-secondary'>&#8592; Back to Rota</a>
    </div>
    <div class='card'>
      <pre id='rota_text' style='white-space:pre-wrap;font-family:DM Mono,monospace;
           font-size:13px;line-height:1.6;background:#f8fafc;padding:16px;
           border-radius:10px;border:1px solid #e2e8f0'>{message}</pre>
      <button onclick='copyRota()'
        class='btn-primary' style='margin-top:12px;width:100%;padding:10px'>
        &#128203; Copy to Clipboard — then paste into WhatsApp
      </button>
    </div>
    <script>
    function copyRota() {{
      const text = document.getElementById('rota_text').textContent;
      navigator.clipboard.writeText(text).then(() => alert('&#10003; Copied! Open WhatsApp and paste.'));
    }}
    </script>"""

    return page("WhatsApp Rota", content, user, "rota")


# ── Rota Templates ────────────────────────────────────────────────────────────

@app.get("/rota/templates", response_class=HTMLResponse)
def rota_templates(
    store:   str = "Uxbridge",
    session: str | None = Cookie(default=None),
    msg:     str = ""
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/rota", status_code=303)

    staff = q("SELECT * FROM staff_profiles WHERE store_name=? AND is_active=1 ORDER BY first_name",
              (store,), fetch=True) or []
    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    store_btns = ""
    for sv in ["Uxbridge","Newbury"]:
        cls = "btn-primary" if sv == store else "btn-secondary"
        store_btns += f"<a href='/rota/templates?store={sv}' class='{cls}' style='padding:5px 14px;font-size:13px'>{sv}</a>"

    # Build template grid
    header = "<tr style='background:#0f2942;color:white'>"
    header += "<th style='padding:10px 12px;text-align:left;font-size:12px'>Staff Member</th>"
    for day in FULL_DAYS:
        header += f"<th style='padding:8px 6px;text-align:center;font-size:11px;min-width:110px'>{day}</th>"
    header += "</tr>"

    rows = ""
    for s in staff:
        sid  = s["staff_id"]
        name = f"{s['first_name']} {s['last_name']}"
        tmpls = q("SELECT * FROM rota_templates WHERE staff_id=?", (sid,), fetch=True) or []
        tmpl_map = {dict(t)["day_of_week"]: dict(t) for t in tmpls}

        rows += f"<tr style='border-bottom:1px solid #f1f5f9'>"
        rows += f"<td style='padding:8px 12px;font-weight:700;font-size:13px'>{name}</td>"

        for dow in range(7):
            t = tmpl_map.get(dow, {})
            is_off = t.get("is_off", 1)
            start  = t.get("shift_start","") or ""
            end    = t.get("shift_end","")   or ""
            rows += f"""<td style='padding:3px'>
              <div style='display:flex;flex-direction:column;gap:2px'>
                <input type='time' form='tmpl_{sid}' name='start_{dow}'
                  value='{start}'
                  style='font-size:11px;padding:3px 5px;border:1px solid #e2e8f0;border-radius:6px;{"background:#f1f5f9" if is_off else ""}'>
                <input type='time' form='tmpl_{sid}' name='end_{dow}'
                  value='{end}'
                  style='font-size:11px;padding:3px 5px;border:1px solid #e2e8f0;border-radius:6px;{"background:#f1f5f9" if is_off else ""}'>
                <label style='display:flex;gap:4px;align-items:center;font-size:10px;
                              color:#64748b;text-transform:none;font-weight:400;margin:0'>
                  <input type='checkbox' form='tmpl_{sid}' name='off_{dow}'
                    {'checked' if is_off else ''} style='width:auto'> OFF
                </label>
              </div>
            </td>"""

        rows += f"<td style='padding:4px'>"
        rows += f"<form id='tmpl_{sid}' action='/rota/save-template' method='POST'>"
        rows += f"<input type='hidden' name='staff_id' value='{sid}'>"
        rows += f"<input type='hidden' name='store' value='{store}'>"
        rows += f"<button type='submit' class='btn-primary' style='padding:4px 10px;font-size:11px;white-space:nowrap'>&#128190; Save</button>"
        rows += f"</form></td>"
        rows += "</tr>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/rota?store={store}' style='color:#1e3a5f;font-size:13px;font-weight:700'>&#8592; Back to Rota</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>&#9881;&#65039; Rota Templates — {store}</div>
        <div style='font-size:13px;color:#64748b;margin-top:2px'>
          Set each staff member's standard weekly pattern. New rotas start from these times.
        </div>
      </div>
      <div style='display:flex;gap:8px'>{store_btns}</div>
    </div>
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table style='width:100%;border-collapse:collapse;font-family:DM Sans,sans-serif'>
          <thead>{header}</thead>
          <tbody>{rows}</tbody>
        </table>
      </div>
    </div>"""

    return page("Rota Templates", content, user, "rota")


@app.post("/rota/save-template")
async def save_template(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form     = await request.form()
    staff_id = int(form.get("staff_id", 0))
    store    = form.get("store","")

    # Get store for staff
    s = q("SELECT store_name FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    store_name = dict(s[0])["store_name"] if s else store

    for dow in range(7):
        start  = form.get(f"start_{dow}","") or None
        end    = form.get(f"end_{dow}","")   or None
        is_off = 1 if form.get(f"off_{dow}") else 0
        hrs    = 0
        if start and end and not is_off:
            try:
                sh, sm = map(int, start.split(":"))
                eh, em = map(int, end.split(":"))
                hrs = round((eh*60+em - sh*60-sm)/60, 2)
            except: pass

        q("""INSERT INTO rota_templates
                (staff_id, store_name, day_of_week, shift_start, shift_end, hours, is_off)
             VALUES(?,?,?,?,?,?,?)
             ON CONFLICT(staff_id, day_of_week) DO UPDATE SET
                shift_start=excluded.shift_start,
                shift_end=excluded.shift_end,
                hours=excluded.hours,
                is_off=excluded.is_off""",
          (staff_id, store_name, dow, start, end, hrs, is_off))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/rota/templates?store={store}&msg={uq('Template saved')}",
        status_code=303)



# ── Mobile Clock-In Portal ────────────────────────────────────────────────────

from haversine import haversine, Unit

GEOFENCE_RADIUS_M = 200

@app.get("/mobile-clock", response_class=HTMLResponse)
def mobile_clock_page(msg: str = "", msg_type: str = "success"):
    """Public GPS clock-in portal — no login required (uses staff name selection)."""
    # Get all active staff
    staff = q("SELECT staff_id, first_name, last_name, store_name FROM staff_profiles WHERE is_active=1 ORDER BY store_name, first_name",
              fetch=True) or []

    staff_opts = "<option value=''>-- Select your name --</option>"
    current_store = ""
    for s in staff:
        s = dict(s)
        if s["store_name"] != current_store:
            if current_store: staff_opts += "</optgroup>"
            staff_opts += "<optgroup label='" + s['store_name'] + "'>"
            current_store = s["store_name"]
        staff_opts += "<option value='" + str(s['staff_id']) + "'>" + s['first_name'] + " " + s['last_name'] + "</option>"
    if current_store: staff_opts += "</optgroup>"

    flash = ""
    if msg:
        col = "#dcfce7" if msg_type == "success" else "#fee2e2"
        tcol = "#166534" if msg_type == "success" else "#dc2626"
        flash = f"<div style='background:{col};color:{tcol};border-radius:10px;padding:12px 16px;font-size:14px;font-weight:700;margin-bottom:16px'>{msg}</div>"

    return HTMLResponse(f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1">
  <title>Clock In — Snappy Snaps</title>
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;700;900&display=swap" rel="stylesheet">
  <style>
    * {{ box-sizing:border-box; margin:0; padding:0; }}
    body {{ font-family:'DM Sans',sans-serif; background:#0f2942; min-height:100vh;
            display:flex; align-items:center; justify-content:center; padding:20px; }}
    .card {{ background:white; border-radius:20px; padding:28px; width:100%; max-width:360px;
             box-shadow:0 20px 60px rgba(0,0,0,.3); }}
    h1 {{ font-size:22px; font-weight:900; color:#0f2942; margin-bottom:4px; }}
    .sub {{ font-size:12px; color:#94a3b8; font-weight:700; letter-spacing:.05em;
            text-transform:uppercase; margin-bottom:24px; }}
    label {{ font-size:11px; font-weight:700; color:#64748b; text-transform:uppercase;
             letter-spacing:.05em; display:block; margin-bottom:4px; margin-top:14px; }}
    select {{ width:100%; border:1px solid #e2e8f0; border-radius:10px; padding:12px 14px;
              font-size:15px; font-family:'DM Sans',sans-serif; outline:none; appearance:none;
              background:white url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='12' height='8' viewBox='0 0 12 8'%3E%3Cpath d='M1 1l5 5 5-5' stroke='%2394a3b8' stroke-width='2' fill='none'/%3E%3C/svg%3E") no-repeat right 14px center; }}
    .btn {{ width:100%; padding:14px; border-radius:12px; font-size:16px; font-weight:900;
            border:none; cursor:pointer; font-family:'DM Sans',sans-serif; margin-top:10px;
            transition:all .15s; }}
    .btn-in  {{ background:#16a34a; color:white; }}
    .btn-in:hover  {{ background:#15803d; }}
    .btn-out {{ background:#dc2626; color:white; }}
    .btn-out:hover {{ background:#b91c1c; }}
    .lock {{ font-size:11px; color:#94a3b8; text-align:center; margin-top:14px; }}
  </style>
</head>
<body>
  <div class="card">
    <h1>Staff Clock Portal</h1>
    <div class="sub">&#128274; GPS Verified &middot; Snappy Snaps</div>
    {flash}
    <form id="clockForm" action="/mobile-clock/submit" method="POST">
      <input type="hidden" name="latitude"  id="lat">
      <input type="hidden" name="longitude" id="lon">
      <input type="hidden" name="action"    id="action">
      <label>Your Name</label>
      <select name="staff_id" required>{staff_opts}</select>
      <button type="button" class="btn btn-in"  onclick="punch('clock_in')">&#128994; Clock In</button>
      <button type="button" class="btn btn-out" onclick="punch('clock_out')">&#128308; Clock Out</button>
    </form>
    <div class="lock">&#128205; Location required to verify attendance</div>
  </div>
  <script>
  function punch(action) {{
    const staff = document.querySelector('[name="staff_id"]').value;
    if (!staff) {{ alert('Please select your name first'); return; }}
    if (!navigator.geolocation) {{ alert('GPS not available on this device'); return; }}
    navigator.geolocation.getCurrentPosition(
      function(pos) {{
        document.getElementById('lat').value    = pos.coords.latitude;
        document.getElementById('lon').value    = pos.coords.longitude;
        document.getElementById('action').value = action;
        document.getElementById('clockForm').submit();
      }},
      function(err) {{ alert('Location access required. Please enable GPS and try again.'); }},
      {{ enableHighAccuracy:true, timeout:10000 }}
    );
  }}
  </script>
</body>
</html>""")


@app.post("/mobile-clock/submit", response_class=HTMLResponse)
async def submit_clock(request: Request):
    form      = await request.form()
    staff_id  = int(form.get("staff_id", 0))
    action    = form.get("action","clock_in")
    try:
        lat = float(form.get("latitude",  0))
        lon = float(form.get("longitude", 0))
    except:
        return HTMLResponse("<p>Invalid location data</p>", status_code=400)

    # Get staff details
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        from urllib.parse import quote as uq
        return RedirectResponse(f"/mobile-clock?msg={uq('Staff member not found')}&msg_type=error", status_code=303)
    s          = dict(rows[0])
    store_name = s["store_name"]
    full_name  = f"{s['first_name']} {s['last_name']}"

    # GPS verification
    store_coords = STORE_GPS.get(store_name)
    if not store_coords:
        from urllib.parse import quote as uq
        return RedirectResponse(f"/mobile-clock?msg={uq('Store location not configured')}&msg_type=error", status_code=303)

    distance_m = haversine((lat, lon), store_coords, unit=Unit.METERS)
    on_site    = distance_m <= GEOFENCE_RADIUS_M

    if not on_site:
        from urllib.parse import quote as uq
        msg = f"Location rejected — you are {distance_m:.0f}m from {store_name} (max {GEOFENCE_RADIUS_M}m)"
        return RedirectResponse(f"/mobile-clock?msg={uq(msg)}&msg_type=error", status_code=303)

    # Record punch
    now_time = datetime.now().strftime("%H:%M:%S")
    now_date = datetime.now().strftime("%Y-%m-%d")

    if action == "clock_in":
        q("""INSERT INTO timesheets (staff_name, store_name, work_date, clock_in_time, status_flag)
             VALUES(?,?,?,?,'GPS_VERIFIED')
             ON CONFLICT(staff_name, store_name, work_date) DO UPDATE SET
                clock_in_time=excluded.clock_in_time, status_flag='GPS_VERIFIED'""",
          (full_name, store_name, now_date, now_time))
        msg = f"Clocked IN &#10003; — {full_name} at {store_name} {now_time}"
    else:
        q("""INSERT INTO timesheets (staff_name, store_name, work_date, clock_out_time, status_flag)
             VALUES(?,?,?,?,'GPS_VERIFIED')
             ON CONFLICT(staff_name, store_name, work_date) DO UPDATE SET
                clock_out_time=excluded.clock_out_time""",
          (full_name, store_name, now_date, now_time))
        msg = f"Clocked OUT &#10003; — {full_name} at {store_name} {now_time}"

    from urllib.parse import quote as uq
    return RedirectResponse(f"/mobile-clock?msg={uq(msg)}", status_code=303)


# ── Timesheets page ───────────────────────────────────────────────────────────

@app.get("/timesheets", response_class=HTMLResponse)
def timesheets_page(
    session:    str | None = Cookie(default=None),
    store:      str = "",
    month:      str = "",
    export:     str = ""
):
    redir, user = require_login(session)
    if redir: return redir

    if not month:
        month = datetime.now().strftime("%Y-%m")
    if not store and user.get("store_name"):
        store = user["store_name"]

    year, mon = map(int, month.split("-"))

    # Date range for this month
    from calendar import monthrange
    _, last_day = monthrange(year, mon)
    date_from   = f"{month}-01"
    date_to     = f"{month}-{last_day:02d}"

    # Get records
    conds  = ["work_date BETWEEN ? AND ?"]
    params = [date_from, date_to]
    if store:
        conds.append("store_name=?")
        params.append(store)
    if user["role"] == "staff":
        name = user.get("full_name","")
        if name:
            conds.append("staff_name=?")
            params.append(name)

    records = q(f"""SELECT * FROM timesheets WHERE {' AND '.join(conds)}
                    ORDER BY store_name, staff_name, work_date""",
                params, fetch=True) or []

    # CSV export
    if export == "csv":
        import csv, io
        buf = io.StringIO()
        w   = csv.writer(buf)
        w.writerow(["Staff Name","Store","Date","Clock In","Clock Out","Status","Hours Worked"])
        for r in records:
            r = dict(r)
            # Calculate hours
            hrs = ""
            if r.get("clock_in_time") and r.get("clock_out_time"):
                try:
                    ci = datetime.strptime(r["clock_in_time"],  "%H:%M:%S")
                    co = datetime.strptime(r["clock_out_time"], "%H:%M:%S")
                    hrs = f"{(co-ci).seconds/3600:.2f}"
                except: pass
            w.writerow([r["staff_name"],r["store_name"],r["work_date"],
                        r.get("clock_in_time",""),r.get("clock_out_time",""),
                        r.get("status_flag",""),hrs])
        from fastapi.responses import Response
        return Response(content=buf.getvalue(), media_type="text/csv",
                        headers={"Content-Disposition": f"attachment;filename=timesheets_{month}_{store}.csv"})

    # Month navigation
    prev_d = (datetime(year, mon, 1) - timedelta(days=1))
    next_d = (datetime(year, mon, last_day) + timedelta(days=1))
    prev_m = prev_d.strftime("%Y-%m")
    next_m = next_d.strftime("%Y-%m")

    # Store filter
    store_btns = ""
    if user["role"] in ("owner","manager"):
        for sv, sl in [("","Both"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]:
            cls = "btn-primary" if store==sv else "btn-secondary"
            store_btns += f"<a href='/timesheets?store={sv}&month={month}' class='{cls}' style='padding:5px 14px;font-size:13px'>{sl}</a>"

    # Build table
    rows_html = ""
    for r in records:
        r   = dict(r)
        hrs = ""
        if r.get("clock_in_time") and r.get("clock_out_time"):
            try:
                ci  = datetime.strptime(r["clock_in_time"],  "%H:%M:%S")
                co  = datetime.strptime(r["clock_out_time"], "%H:%M:%S")
                hrs = f"{(co-ci).seconds/3600:.2f}h"
            except: pass
        status_cls = "badge-paid" if r.get("status_flag")=="GPS_VERIFIED" else "badge-unpaid"
        out_val    = r.get("clock_out_time") or "<span style='color:#d97706'>On shift</span>"
        rows_html += f"""<tr>
          <td style='font-weight:700'>{r['staff_name']}</td>
          <td style='font-size:12px;color:#64748b'>{r['store_name']}</td>
          <td class='mono' style='font-size:12px'>{r['work_date']}</td>
          <td class='mono' style='color:#16a34a;font-weight:700'>{r.get('clock_in_time') or '—'}</td>
          <td class='mono' style='color:#dc2626;font-weight:700'>{out_val}</td>
          <td class='mono' style='font-weight:700'>{hrs}</td>
          <td><span class='{status_cls}'>{r.get("status_flag") or "—"}</span></td>
        </tr>"""

    month_label = datetime(year, mon, 1).strftime("%B %Y")

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>&#9200; Timesheets — {month_label}</div>
      <div style='display:flex;gap:8px;flex-wrap:wrap;align-items:center'>
        {store_btns}
        <a href='/timesheets?store={store}&month={prev_m}' class='btn-secondary' style='padding:5px 12px'>&#8592;</a>
        <a href='/timesheets?store={store}&month={next_m}' class='btn-secondary' style='padding:5px 12px'>&#8594;</a>
        <a href='/timesheets?store={store}&month={month}&export=csv'
           class='btn-primary' style='padding:6px 16px;font-size:13px'>
          &#11015;&#65039; Export CSV for Payroll
        </a>
      </div>
    </div>
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead>
            <tr>
              <th>Staff Member</th><th>Store</th><th>Date</th>
              <th>Clock In</th><th>Clock Out</th><th>Hours</th><th>Status</th>
            </tr>
          </thead>
          <tbody>
            {rows_html or '<tr><td colspan="7" style="text-align:center;padding:32px;color:#94a3b8">No records for this period</td></tr>'}
          </tbody>
        </table>
      </div>
    </div>"""

    return page("Timesheets", content, user, "timesheets")



# ── Rota PDF Export ───────────────────────────────────────────────────────────

@app.get("/rota/pdf")
def rota_pdf(
    store:      str = "",
    week_start: str = "",
    session:    str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir

    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from fastapi.responses import Response
    import io

    week_dates = get_week_dates(week_start)
    week_end   = week_dates[-1]
    rota       = get_or_create_rota(store, week_start)
    shifts     = {(s["staff_id"], s["shift_date"]): s for s in rota.get("shifts", [])}
    staff      = q("SELECT * FROM staff_profiles WHERE store_name=? AND is_active=1 ORDER BY first_name",
                   (store,), fetch=True) or []

    # Colours
    COL_NAVY   = colors.HexColor("#0f2942")
    COL_BLUE   = colors.HexColor("#1e3a5f")
    COL_GREEN  = colors.HexColor("#dcfce7")
    COL_GREEN2 = colors.HexColor("#166534")
    COL_RED    = colors.HexColor("#fee2e2")
    COL_RED2   = colors.HexColor("#dc2626")
    COL_AMBER  = colors.HexColor("#fef3c7")
    COL_AMBER2 = colors.HexColor("#92400e")
    COL_LGREY  = colors.HexColor("#f8fafc")
    COL_GREY   = colors.HexColor("#e2e8f0")
    COL_WHITE  = colors.white

    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(buf, pagesize=landscape(A4),
                             leftMargin=10*mm, rightMargin=10*mm,
                             topMargin=12*mm, bottomMargin=12*mm)

    styles  = getSampleStyleSheet()
    title_s = ParagraphStyle("title", fontSize=16, fontName="Helvetica-Bold",
                              textColor=COL_NAVY, alignment=TA_LEFT)
    sub_s   = ParagraphStyle("sub", fontSize=9, fontName="Helvetica",
                              textColor=colors.HexColor("#64748b"), alignment=TA_LEFT)
    cell_s  = ParagraphStyle("cell", fontSize=8, fontName="Helvetica-Bold",
                              alignment=TA_CENTER, leading=10)
    small_s = ParagraphStyle("small", fontSize=7, fontName="Helvetica",
                              alignment=TA_CENTER, textColor=colors.HexColor("#64748b"), leading=8)
    name_s  = ParagraphStyle("name", fontSize=9, fontName="Helvetica-Bold",
                              alignment=TA_LEFT, textColor=COL_NAVY)
    hdr_s   = ParagraphStyle("hdr", fontSize=8, fontName="Helvetica-Bold",
                              alignment=TA_CENTER, textColor=COL_WHITE)

    week_label = f"{datetime.strptime(week_start,'%Y-%m-%d').strftime('%d %b')} – {datetime.strptime(week_end,'%Y-%m-%d').strftime('%d %b %Y')}"
    status     = rota.get("status","draft").upper()

    story = [
        Paragraph(f"Snappy Snaps {store} — Weekly Rota", title_s),
        Spacer(1, 3*mm),
        Paragraph(f"Week: {week_label}  ·  Status: {status}  ·  Generated: {datetime.now().strftime('%d %b %Y %H:%M')}", sub_s),
        Spacer(1, 5*mm),
    ]

    # Build table data
    # Header row
    header = [Paragraph("Staff Member", hdr_s)]
    for i, date_str in enumerate(week_dates):
        d   = datetime.strptime(date_str, "%Y-%m-%d")
        txt = f"{DAYS[i]}\n{d.strftime('%d %b')}"
        header.append(Paragraph(txt, hdr_s))
    header.append(Paragraph("Hrs", hdr_s))

    table_data  = [header]
    table_style = [
        ("BACKGROUND", (0,0), (-1,0), COL_NAVY),
        ("TEXTCOLOR",  (0,0), (-1,0), COL_WHITE),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,0), 8),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("GRID",       (0,0), (-1,-1), 0.3, COL_GREY),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [COL_WHITE, COL_LGREY]),
        ("LEFTPADDING",  (0,0), (-1,-1), 4),
        ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING",   (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0), (-1,-1), 4),
    ]

    # Day totals tracking
    day_hrs   = [0.0] * 7
    day_count = [0]   * 7
    week_hrs  = 0.0

    for row_idx, s in enumerate(staff):
        sid   = s["staff_id"]
        name  = f"{s['first_name']} {s['last_name']}"
        row   = [Paragraph(name, name_s)]
        total = 0.0
        r     = row_idx + 1

        for i, date_str in enumerate(week_dates):
            sh      = shifts.get((sid, date_str), {})
            is_off  = sh.get("is_off", 1)
            absence = sh.get("absence_type")
            start   = sh.get("shift_start") or ""
            end     = sh.get("shift_end")   or ""
            hrs     = sh.get("hours") or 0

            if absence:
                labels = {"H":"Holiday","S":"Sick","B":"Bank Hol","AL":"Auth Leave","L":"Late"}
                lbl    = labels.get(absence, absence)
                cell   = Paragraph(lbl, ParagraphStyle("ab", fontSize=7, fontName="Helvetica-Bold",
                                   alignment=TA_CENTER, textColor=COL_GREEN2))
                table_style.append(("BACKGROUND", (i+1, r), (i+1, r), COL_GREEN))
                if absence == "S":
                    table_style.append(("BACKGROUND", (i+1, r), (i+1, r), COL_RED))
                    cell = Paragraph(lbl, ParagraphStyle("ab", fontSize=7, fontName="Helvetica-Bold",
                                     alignment=TA_CENTER, textColor=COL_RED2))
            elif is_off:
                cell = Paragraph("OFF", ParagraphStyle("off", fontSize=7, fontName="Helvetica",
                                 alignment=TA_CENTER, textColor=colors.HexColor("#cbd5e1")))
            else:
                total += hrs
                day_hrs[i]   += hrs
                day_count[i] += 1
                week_hrs     += hrs
                shift_txt = f"{start}–{end}\n{hrs:.1f}h"
                cell = Paragraph(shift_txt, ParagraphStyle("sh", fontSize=8, fontName="Helvetica-Bold",
                                 alignment=TA_CENTER, textColor=COL_BLUE, leading=10))
                table_style.append(("BACKGROUND", (i+1, r), (i+1, r), colors.HexColor("#eff6ff")))

            row.append(cell)

        hrs_cell = Paragraph(f"{total:.1f}", ParagraphStyle("hrs", fontSize=9,
                             fontName="Helvetica-Bold", alignment=TA_CENTER,
                             textColor=COL_NAVY if total else colors.HexColor("#cbd5e1")))
        row.append(hrs_cell)
        table_data.append(row)

    # Totals row
    totals_row = [Paragraph("TOTALS", ParagraphStyle("tot", fontSize=8, fontName="Helvetica-Bold",
                             alignment=TA_LEFT, textColor=COL_NAVY))]
    for i in range(7):
        txt = f"{day_hrs[i]:.1f}h\n{day_count[i]} staff"
        totals_row.append(Paragraph(txt, ParagraphStyle("dt", fontSize=7, fontName="Helvetica-Bold",
                                    alignment=TA_CENTER, textColor=COL_NAVY, leading=9)))
    totals_row.append(Paragraph(f"{week_hrs:.1f}", ParagraphStyle("wt", fontSize=10,
                                fontName="Helvetica-Bold", alignment=TA_CENTER, textColor=COL_NAVY)))
    table_data.append(totals_row)

    n = len(table_data)
    table_style.append(("BACKGROUND",  (0, n-1), (-1, n-1), COL_LGREY))
    table_style.append(("FONTNAME",    (0, n-1), (-1, n-1), "Helvetica-Bold"))
    table_style.append(("LINEABOVE",   (0, n-1), (-1, n-1), 1.5, COL_GREY))

    # Column widths — name col wider, day cols equal, hrs col narrow
    page_w    = landscape(A4)[0] - 20*mm
    name_w    = 38*mm
    hrs_w     = 14*mm
    day_w     = (page_w - name_w - hrs_w) / 7
    col_widths = [name_w] + [day_w]*7 + [hrs_w]
    row_height = 14*mm

    tbl = Table(table_data, colWidths=col_widths, rowHeights=row_height)
    tbl.setStyle(TableStyle(table_style))
    story.append(tbl)

    # Legend
    story.append(Spacer(1, 4*mm))
    legend_data = [[
        Paragraph("Legend:", ParagraphStyle("lg", fontSize=7, fontName="Helvetica-Bold")),
        Paragraph("■ Working shift", ParagraphStyle("lg2", fontSize=7, fontName="Helvetica",
                  textColor=COL_BLUE)),
        Paragraph("■ Holiday", ParagraphStyle("lg3", fontSize=7, fontName="Helvetica",
                  textColor=COL_GREEN2)),
        Paragraph("■ Sick", ParagraphStyle("lg4", fontSize=7, fontName="Helvetica",
                  textColor=COL_RED2)),
        Paragraph("Hours shown are PAID hours (30 min break deducted for shifts ≥ 4h)",
                  ParagraphStyle("lg5", fontSize=7, fontName="Helvetica",
                  textColor=colors.HexColor("#64748b"))),
    ]]
    legend = Table(legend_data, colWidths=[20*mm, 28*mm, 20*mm, 15*mm, 120*mm])
    legend.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"MIDDLE")]))
    story.append(legend)

    doc.build(story)
    buf.seek(0)

    filename = f"Rota_{store}_{week_start}.pdf"
    return Response(content=buf.read(), media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={filename}"})



# ── My Profile (staff self-service) ──────────────────────────────────────────

@app.get("/my-profile", response_class=HTMLResponse)
def my_profile(session: str | None = Cookie(default=None), msg: str = ""):
    redir, user = require_login(session)
    if redir: return redir

    # Find staff profile by matching full name to username
    full_name = user.get("full_name", "")
    rows = q("""SELECT * FROM staff_profiles
                WHERE first_name || ' ' || last_name = ?
                AND is_active = 1""", (full_name,), fetch=True)

    if not rows:
        content = """
        <div class='text-2xl font-black text-slate-800'>My Profile</div>
        <div class='card'>
          <p style='color:#64748b'>No staff profile linked to your account yet.
          Please contact your manager.</p>
        </div>"""
        return page("My Profile", content, user, "my profile")

    s     = dict(rows[0])
    sid   = s["staff_id"]
    year  = datetime.now().year
    leave = get_leave_summary(sid, year)
    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    content = f"""
    {flash}
    <div class='text-2xl font-black text-slate-800'>My Profile</div>

    <!-- Leave summary -->
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(150px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Leave Entitlement</div>
        <div style='font-size:24px;font-weight:900;color:#0f2942'>{leave.get("entitlement_fmt","—")}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Holiday Taken</div>
        <div style='font-size:24px;font-weight:900;color:#d97706'>{leave.get("taken_days",0)} days</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Balance</div>
        <div style='font-size:24px;font-weight:900;color:#16a34a'>{leave.get("balance_fmt","—")}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Sick Days {year}</div>
        <div style='font-size:24px;font-weight:900;color:{"#dc2626" if leave.get("sick_days",0) else "#0f2942"}'>{leave.get("sick_days",0)}</div>
      </div>
    </div>

    <!-- Editable personal details -->
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:4px'>Personal Details</div>
      <div style='font-size:12px;color:#94a3b8;margin-bottom:16px'>
        You can update your contact details below. Employment details can only be changed by your manager.
      </div>
      <form action='/my-profile' method='POST' class='grid gap-3'
            style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
        <div><label>Phone</label>
          <input type='text' name='phone' value='{s.get("phone") or ""}' placeholder='07700 123456'></div>
        <div><label>Email</label>
          <input type='email' name='email' value='{s.get("email") or ""}' placeholder='your@email.com'></div>
        <div><label>Address Line 1</label>
          <input type='text' name='address_1' value='{s.get("address_1") or ""}'></div>
        <div><label>Address Line 2</label>
          <input type='text' name='address_2' value='{s.get("address_2") or ""}'></div>
        <div><label>Town / City</label>
          <input type='text' name='address_3' value='{s.get("address_3") or ""}'></div>
        <div><label>Postcode</label>
          <input type='text' name='postcode' value='{s.get("postcode") or ""}'></div>
        <div style='grid-column:1/-1'>
          <button type='submit' class='btn-primary'>&#128190; Save Changes</button>
        </div>
      </form>
    </div>

    <!-- Read-only employment info — no pay rates shown to staff -->
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr));font-size:13px'>
        <div><span style='color:#94a3b8;font-weight:700'>Store</span><br>{s.get("store_name") or "—"}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Date Joined</span><br>{s.get("date_joined") or "—"}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Contracted Hours</span><br>{str(s.get("contracted_hrs") or "—") + "h/wk"}</div>
      </div>
    </div>

    <!-- Leave request -->
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>&#128197; Request Leave</div>
      <a href='/staff/{sid}/request-leave' class='btn-primary'>Submit Leave Request</a>
    </div>"""

    return page("My Profile", content, user, "my profile")


@app.post("/my-profile")
async def save_my_profile(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    form      = await request.form()
    full_name = user.get("full_name","")

    q("""UPDATE staff_profiles SET
            phone=?, email=?, address_1=?, address_2=?, address_3=?, postcode=?
         WHERE first_name || ' ' || last_name = ? AND is_active=1""",
      (str(form.get("phone","") or "").strip(),
       str(form.get("email","") or "").strip(),
       str(form.get("address_1","") or "").strip(),
       str(form.get("address_2","") or "").strip(),
       str(form.get("address_3","") or "").strip(),
       str(form.get("postcode","") or "").strip(),
       full_name))

    from urllib.parse import quote as uq
    return RedirectResponse(f"/my-profile?msg={uq('Profile updated successfully')}", status_code=303)



# ══════════════════════════════════════════════════════════════════════════════
# MODULE 5 — SALES & FRANCHISE
# ══════════════════════════════════════════════════════════════════════════════

def ensure_sales_tables():
    conn = db()
    c    = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS daily_cashsheet (
            entry_id       INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name     TEXT NOT NULL,
            sale_date      TEXT NOT NULL,
            z_read_no      INTEGER,
            -- 22 Sales categories
            digital_printing   REAL DEFAULT 0,
            other_dp           REAL DEFAULT 0,
            instant_prints     REAL DEFAULT 0,
            reprint_enlarge    REAL DEFAULT 0,
            internet_orders    REAL DEFAULT 0,
            passport           REAL DEFAULT 0,
            film_media         REAL DEFAULT 0,
            graphic_design     REAL DEFAULT 0,
            large_format       REAL DEFAULT 0,
            toner_laser        REAL DEFAULT 0,
            batteries          REAL DEFAULT 0,
            frames_albums      REAL DEFAULT 0,
            photogifts         REAL DEFAULT 0,
            backup_media       REAL DEFAULT 0,
            dvd_transfer       REAL DEFAULT 0,
            studio             REAL DEFAULT 0,
            sundry             REAL DEFAULT 0,
            promotions         REAL DEFAULT 0,
            rcs_std_vat        REAL DEFAULT 0,
            rcs_zero           REAL DEFAULT 0,
            photobooks         REAL DEFAULT 0,
            type_b_sales       REAL DEFAULT 0,
            discount_amount    REAL DEFAULT 0,
            -- Card breakdown
            card_visa          REAL DEFAULT 0,
            card_visa_debit    REAL DEFAULT 0,
            card_mastercard    REAL DEFAULT 0,
            card_mc_debit      REAL DEFAULT 0,
            card_maestro_dom   REAL DEFAULT 0,
            card_maestro_int   REAL DEFAULT 0,
            card_solo          REAL DEFAULT 0,
            card_electron      REAL DEFAULT 0,
            card_amex          REAL DEFAULT 0,
            card_discover      REAL DEFAULT 0,
            card_other         REAL DEFAULT 0,
            -- Cash
            cash_taken         REAL DEFAULT 0,
            opening_cash_bf    REAL DEFAULT 0,
            -- Paid outs
            paid_out_total     REAL DEFAULT 0,
            paid_out_notes     TEXT,
            -- Till reconciliation
            till_credit_sales  REAL DEFAULT 0,
            till_internet_sales REAL DEFAULT 0,
            total_cash_store   REAL DEFAULT 0,
            -- Meta
            entered_by         TEXT,
            submitted_at       TEXT DEFAULT (datetime('now')),
            is_locked          INTEGER DEFAULT 0,
            notes              TEXT,
            UNIQUE(store_name, sale_date)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS paid_outs (
            paidout_id     INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name     TEXT NOT NULL,
            entry_date     TEXT NOT NULL,
            description    TEXT NOT NULL,
            amount         REAL NOT NULL,
            category       TEXT,
            entered_by     TEXT,
            created_at     TEXT DEFAULT (datetime('now'))
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS sales_targets (
            target_id      INTEGER PRIMARY KEY AUTOINCREMENT,
            store_name     TEXT NOT NULL,
            year           INTEGER NOT NULL,
            month          INTEGER NOT NULL,
            target_amount  REAL DEFAULT 0,
            ly_actual      REAL DEFAULT 0,
            UNIQUE(store_name, year, month)
        )
    """)
    conn.commit()
    conn.close()

ensure_sales_tables()

# Category definitions matching your cashsheet exactly
SALES_CATEGORIES = [
    ("digital_printing",  "1",  "Digital Printing",    "trans_digital_printing"),
    ("other_dp",          "2",  "Other D&P",           "trans_other_dp"),
    ("instant_prints",    "3",  "Instant Prints",      "trans_instant_prints"),
    ("reprint_enlarge",   "4",  "Reprint/Enlarge",     "trans_reprint_enlarge"),
    ("internet_orders",   "5",  "Internet Orders",     "trans_internet_orders"),
    ("passport",          "6",  "Passport",            "trans_passport"),
    ("film_media",        "7",  "Film Media",          "trans_film_media"),
    ("graphic_design",    "8",  "Graphic Design",      "trans_graphic_design"),
    ("large_format",      "9",  "Large Format",        "trans_large_format"),
    ("toner_laser",       "10", "Toner/Laser Output",  "trans_toner_laser"),
    ("batteries",         "11", "Batteries",           "trans_batteries"),
    ("frames_albums",     "12", "Frames & Albums",     "trans_frames_albums"),
    ("photogifts",        "13", "Photogifts",          "trans_photogifts"),
    ("backup_media",      "14", "Backup to Media",     "trans_backup_media"),
    ("dvd_transfer",      "15", "DVD Transfer",        "trans_dvd_transfer"),
    ("studio",            "16", "Studio",              "trans_studio"),
    ("sundry",            "17", "Sundry",              "trans_sundry"),
    ("promotions",        "18", "Promotions",          "trans_promotions"),
    ("rcs_std_vat",       "19", "RCS (STD VAT)",       "trans_rcs_std_vat"),
    ("rcs_zero",          "20", "RCS (ZERO)",          "trans_rcs_zero"),
    ("photobooks",        "21", "Photobooks",          "trans_photobooks"),
    ("type_b_sales",      "22", "TYPE B Sales",        "trans_type_b_sales"),
]

CARD_TYPES = [
    ("card_visa",        "VISA"),
    ("card_visa_debit",  "VISA DEBIT"),
    ("card_mastercard",  "MASTERCARD"),
    ("card_mc_debit",    "MASTERCARD DEBIT"),
    ("card_maestro_dom", "MAESTRO DOM"),
    ("card_maestro_int", "MAESTRO INT"),
    ("card_solo",        "SOLO"),
    ("card_electron",    "ELECTRON"),
    ("card_amex",        "AMEX"),
    ("card_discover",    "DISCOVER"),
    ("card_other",       "OTHER"),
]


# ── Sales overview / home page ────────────────────────────────────────────────

@app.get("/sales", response_class=HTMLResponse)
def sales_page(
    session:    str | None = Cookie(default=None),
    store:      str = "",
    week_start: str = "",
    msg:        str = ""
):
    redir, user = require_login(session)
    if redir: return redir

    if not store and user.get("store_name"):
        store = user["store_name"]
    if not store:
        store = "Uxbridge"
    if not week_start:
        week_start = get_week_start()

    week_dates = get_week_dates(week_start)
    week_end   = week_dates[-1]
    prev_week  = (datetime.strptime(week_start, "%Y-%m-%d") - timedelta(days=7)).strftime("%Y-%m-%d")
    next_week  = (datetime.strptime(week_start, "%Y-%m-%d") + timedelta(days=7)).strftime("%Y-%m-%d")
    flash      = f"<div class='flash-success'>{msg}</div>" if msg else ""
    is_mgr     = user["role"] in ("owner","manager")

    # Get all cashsheet entries for this week
    entries = q("""SELECT * FROM daily_cashsheet
                   WHERE store_name=? AND sale_date BETWEEN ? AND ?
                   ORDER BY sale_date""",
                (store, week_dates[0], week_dates[-1]), fetch=True) or []
    entry_map = {dict(e)["sale_date"]: dict(e) for e in entries}

    # Store switcher
    store_btns = ""
    if user["role"] in ("owner","manager"):
        for sv in ["Uxbridge","Newbury"]:
            cls = "btn-primary" if sv == store else "btn-secondary"
            store_btns += f"<a href='/sales?store={sv}&week_start={week_start}' class='{cls}' style='padding:5px 14px;font-size:13px'>{sv}</a>"

    # Week summary cards
    week_total = sum(
        sum((e.get(col, 0) or 0) for col, _, _, _ in SALES_CATEGORIES) + (e.get("discount_amount", 0) or 0)
        for e in [entry_map.get(d, {}) for d in week_dates]
    )
    days_entered = sum(1 for d in week_dates if d in entry_map)

    # Get this month's target
    today     = datetime.now()
    target_row = q("SELECT target_amount FROM sales_targets WHERE store_name=? AND year=? AND month=?",
                   (store, today.year, today.month), fetch=True)
    monthly_target = dict(target_row[0])["target_amount"] if target_row else 0

    # Month to date
    month_start = today.strftime("%Y-%m-01")
    month_end   = today.strftime("%Y-%m-%d")
    mtd_rows    = q("""SELECT * FROM daily_cashsheet
                       WHERE store_name=? AND sale_date BETWEEN ? AND ?""",
                    (store, month_start, month_end), fetch=True) or []
    mtd_total   = sum(
        sum((dict(e).get(col, 0) or 0) for col, _, _, _ in SALES_CATEGORIES) + (dict(e).get("discount_amount", 0) or 0)
        for e in mtd_rows
    )
    target_pct  = (mtd_total / monthly_target * 100) if monthly_target else 0
    target_col  = "#16a34a" if target_pct >= 100 else ("#d97706" if target_pct >= 75 else "#dc2626")

    summary_cards = f"""
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(160px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>This Week</div>
        <div style='font-size:24px;font-weight:900;color:#0f2942'>£{week_total:,.2f}</div>
        <div style='font-size:11px;color:#94a3b8'>{days_entered}/7 days entered</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Month to Date</div>
        <div style='font-size:24px;font-weight:900;color:#0f2942'>£{mtd_total:,.2f}</div>
        <div style='font-size:11px;color:#94a3b8'>{today.strftime("%B %Y")}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Monthly Target</div>
        <div style='font-size:24px;font-weight:900;color:{target_col}'>
          {'£'+f"{monthly_target:,.0f}" if monthly_target else "Not set"}
        </div>
        <div style='font-size:11px;color:{target_col}'>{target_pct:.0f}% of target</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Remaining</div>
        <div style='font-size:24px;font-weight:900;color:{"#16a34a" if monthly_target-mtd_total<=0 else "#dc2626"}'>
          £{max(0,monthly_target-mtd_total):,.2f}
        </div>
        <div style='font-size:11px;color:#94a3b8'>to reach target</div>
      </div>
    </div>"""

    # Weekly grid
    header = "<tr style='background:#0f2942;color:white'>"
    header += "<th style='padding:10px 12px;text-align:left;font-size:12px;min-width:180px'>Category</th>"
    for i, date_str in enumerate(week_dates):
        d      = datetime.strptime(date_str, "%Y-%m-%d")
        is_today = date_str == today.strftime("%Y-%m-%d")
        has_data = date_str in entry_map
        bg = "background:#1e3a5f" if is_today else ""
        tick = " ✅" if has_data else ""
        header += f"<th style='padding:8px 6px;text-align:right;font-size:11px;min-width:90px;{bg}'>{DAYS[i]}<br><span style='font-size:10px;opacity:.7'>{d.strftime('%d %b')}{tick}</span></th>"
    header += "<th style='padding:8px;text-align:right;font-size:11px'>Week Total</th></tr>"

    rows_html = ""
    cat_totals = {col: 0 for col, _, _, _ in SALES_CATEGORIES}

    for col, num, label, trans_col in SALES_CATEGORIES:
        row_total = sum((entry_map.get(d, {}).get(col, 0) or 0) for d in week_dates)
        cat_totals[col] = row_total
        cells = ""
        for date_str in week_dates:
            val = entry_map.get(date_str, {}).get(col, 0) or 0
            cells += f"<td style='padding:6px 8px;text-align:right;font-size:12px'>{'£'+f'{val:.2f}' if val else '—'}</td>"
        rows_html += f"""<tr style='border-bottom:1px solid #f1f5f9'>
          <td style='padding:6px 12px;font-size:13px;color:#334155'><span style='color:#94a3b8;font-size:11px'>{num}</span> {label}</td>
          {cells}
          <td style='padding:6px 8px;text-align:right;font-size:13px;font-weight:700;color:#0f2942'>{'£'+f'{row_total:.2f}' if row_total else '—'}</td>
        </tr>"""

    # Discount row
    disc_row_total = sum((entry_map.get(d, {}).get("discount_amount", 0) or 0) for d in week_dates)
    disc_cells = ""
    for d in week_dates:
        dval = entry_map.get(d, {}).get('discount_amount', 0) or 0
        disc_cells += "<td style='padding:6px 8px;text-align:right;font-size:12px;color:#dc2626'>" + ('£' + f'{dval:.2f}' if dval else '—') + "</td>"
    rows_html += f"""<tr style='border-bottom:2px solid #e2e8f0;background:#fff5f5'>
      <td style='padding:6px 12px;font-size:13px;color:#dc2626;font-weight:700'>Less: Discounts</td>
      {disc_cells}
      <td style='padding:6px 8px;text-align:right;font-size:13px;font-weight:700;color:#dc2626'>{'£'+f'{disc_row_total:.2f}' if disc_row_total else '—'}</td>
    </tr>"""

    # Total row
    day_totals = []
    for date_str in week_dates:
        e   = entry_map.get(date_str, {})
        tot = sum((e.get(col, 0) or 0) for col, _, _, _ in SALES_CATEGORIES) + (e.get("discount_amount", 0) or 0)
        day_totals.append(f"<td style='padding:8px;text-align:right;font-size:13px;font-weight:900;color:#0f2942;background:#f8fafc'>{'£'+f'{tot:,.2f}' if tot else '—'}</td>")

    rows_html += f"""<tr style='background:#f8fafc;border-top:2px solid #e2e8f0'>
      <td style='padding:8px 12px;font-size:13px;font-weight:900;color:#0f2942'>TOTAL SALES</td>
      {"".join(day_totals)}
      <td style='padding:8px;text-align:right;font-size:14px;font-weight:900;color:#0f2942'>£{week_total:,.2f}</td>
    </tr>"""

    # Action buttons
    action_btns = f"""
    <div style='display:flex;gap:8px;flex-wrap:wrap'>
      <a href='/sales/enter?store={store}&date={today.strftime("%Y-%m-%d")}' class='btn-primary'>
        &#128221; Enter Today's Sales
      </a>
      {'<a href="/sales/targets?store=' + store + '" class="btn-secondary">&#127919; Manage Targets</a>' if is_mgr else ''}
      <a href='/sales/franchise-return?store={store}&week_start={week_start}' class='btn-secondary'>
        &#128196; Franchise Return
      </a>
      <a href='/sales/managers-report?store={store}&week_start={week_start}' class='btn-secondary'>
        &#128200; Manager's Report
      </a>
    </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <div class='text-2xl font-black text-slate-800'>&#128200; Sales — {store}</div>
        <div style='font-size:13px;color:#64748b;margin-top:2px'>
          Week: {datetime.strptime(week_start,"%Y-%m-%d").strftime("%d %b")} –
          {datetime.strptime(week_end,"%Y-%m-%d").strftime("%d %b %Y")}
        </div>
      </div>
      <div style='display:flex;gap:8px;flex-wrap:wrap;align-items:center'>
        {store_btns}
        <a href='/sales?store={store}&week_start={prev_week}' class='btn-secondary' style='padding:5px 12px'>&#8592;</a>
        <a href='/sales?store={store}' class='btn-secondary' style='padding:5px 12px'>This Week</a>
        <a href='/sales?store={store}&week_start={next_week}' class='btn-secondary' style='padding:5px 12px'>&#8594;</a>
      </div>
    </div>
    {summary_cards}
    {action_btns}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table style='width:100%;border-collapse:collapse;font-family:DM Sans,sans-serif'>
          <thead>{header}</thead>
          <tbody>{rows_html}</tbody>
        </table>
      </div>
    </div>"""

    return page("Sales", content, user, "sales")


# ── Daily Cash Entry Form ─────────────────────────────────────────────────────

@app.get("/sales/enter", response_class=HTMLResponse)
def sales_entry_form(
    store:   str = "",
    date:    str = "",
    session: str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if not store and user.get("store_name"):
        store = user["store_name"]
    if not date:
        date = datetime.now().strftime("%Y-%m-%d")

    d_fmt    = datetime.strptime(date, "%Y-%m-%d").strftime("%A %d %B %Y")
    is_thurs = datetime.strptime(date, "%Y-%m-%d").weekday() == 3

    # Get existing entry
    existing = q("SELECT * FROM daily_cashsheet WHERE store_name=? AND sale_date=?",
                 (store, date), fetch=True)
    e = dict(existing[0]) if existing else {}

    # Get previous day for B/F
    prev_date  = (datetime.strptime(date,"%Y-%m-%d") - timedelta(days=1)).strftime("%Y-%m-%d")
    prev_entry = q("SELECT * FROM daily_cashsheet WHERE store_name=? AND sale_date=?",
                   (store, prev_date), fetch=True)
    prev_e     = dict(prev_entry[0]) if prev_entry else {}
    bf_auto    = e.get("opening_cash_bf") or prev_e.get("actual_cash_cf") or 0
    internet_orders_val = e.get("internet_orders") or 0
    prev_z     = prev_e.get("z_read_no") or 0
    prev_z2    = prev_e.get("z2_read_no") or 0

    def fv(col):
        v = e.get(col, 0)
        return ("%.2f" % v) if isinstance(v,(int,float)) and v else ""
    def fvs(col): return str(e.get(col,"") or "")
    def fvi(col): return str(int(e.get(col,0) or 0)) if e.get(col) else ""
    def fvn(col): return str(int(e.get(col,0) or 0)) if e.get(col) else ""

    prev_d = (datetime.strptime(date,"%Y-%m-%d")-timedelta(days=1)).strftime("%Y-%m-%d")
    next_d = (datetime.strptime(date,"%Y-%m-%d")+timedelta(days=1)).strftime("%Y-%m-%d")
    z_cur  = fvi("z_read_no")
    z2_cur = fvi("z2_read_no")

    # Build sales category rows
    cat_rows = ""
    for col, num, label, trans_col in SALES_CATEGORIES:
        v    = fv(col)
        tcnt = fvi(trans_col)
        act  = e.get(col,0) or 0
        tnum = e.get(trans_col,0) or 0
        vpt  = ("%.2f" % (act/tnum)) if tnum else "&mdash;"
        cat_rows += (
            "<tr style='border-bottom:1px solid #f1f5f9'>"
            "<td style='padding:3px 8px;font-size:11px;color:#334155;white-space:nowrap'>"
            "<span style='color:#94a3b8;font-size:10px'>" + num + "</span> - " + label + "</td>"
            "<td style='padding:2px 3px;width:55px'>"
            "<input type='number' name='" + trans_col + "' value='" + tcnt + "'"
            " min='0' step='1' oninput='updVPT(\"" + col + "\",\"" + trans_col + "\")'"
            " onblur='if(this.value)this.value=Math.round(parseFloat(this.value)||0)'"
            " placeholder='0'"
            " style='width:100%;text-align:right;border:1px solid #d1d5db;border-radius:4px;"
            "padding:3px 4px;font-size:12px;background:#fefce8'>"
            "</td>"
            "<td style='padding:3px 5px;width:65px;text-align:right;font-size:11px;"
            "font-family:DM Mono,monospace;background:#f0fdf4;color:#166534'"
            " id='vpt_" + col + "'>" + vpt + "</td>"
            "<td style='padding:2px 3px;width:75px'>"
            "<input type='number' step='0.01' name='" + col + "' value='" + v + "'"
            " onblur='if(this.value&&!isNaN(this.value))this.value=parseFloat(this.value).toFixed(2)'"
            " oninput='updTot();updVPT(\"" + col + "\",\"" + trans_col + "\")'"
            " placeholder='0.00'"
            " style='width:100%;text-align:right;border:1px solid #d1d5db;border-radius:4px;"
            "padding:3px 4px;font-size:12px;background:#fefce8;font-family:DM Mono,monospace'>"
            "</td></tr>"
        )

    # Blank line 24
    cat_rows += (
        "<tr style='border-bottom:1px solid #f1f5f9'>"
        "<td style='padding:3px 8px;font-size:11px;color:#94a3b8'><span style='font-size:10px'>24</span> - </td>"
        "<td style='padding:2px 3px'><input type='number' min='0' step='1' placeholder='0'"
        " style='width:100%;text-align:right;border:1px solid #d1d5db;border-radius:4px;padding:3px 4px;font-size:12px;background:#fefce8'></td>"
        "<td style='background:#f0fdf4'>&mdash;</td>"
        "<td style='padding:2px 3px'><input type='number' step='0.01' placeholder='0.00'"
        " style='width:100%;text-align:right;border:1px solid #d1d5db;border-radius:4px;padding:3px 4px;font-size:12px;background:#fefce8;font-family:DM Mono,monospace'></td>"
        "</tr>"
    )
    # Discount line 25
    disc_amt = fv("discount_amount") or ""
    cat_rows += (
        "<tr style='background:#fff5f5;border-top:1px solid #fca5a5'>"
        "<td style='padding:3px 8px;font-size:11px;font-weight:700;color:#dc2626;white-space:nowrap'>"
        "<span style='font-size:10px;color:#dc2626'>25</span> - % - Discount (Enter as -ve)</td>"
        "<td style='padding:2px 3px'><input type='number' name='discount_trans' min='0' step='1' placeholder='0'"
        " oninput='updDiscVPT()'"
        " style='width:100%;text-align:right;border:1px solid #fca5a5;border-radius:4px;padding:3px 4px;font-size:12px;background:#fef2f2'></td>"
        "<td id='vpt_discount' style='background:#f0fdf4;text-align:right;font-size:11px;font-family:DM Mono,monospace;color:#166534'>&mdash;</td>"
        "<td style='padding:2px 3px'><input type='number' step='0.01' name='discount_amount' value='" + disc_amt + "'"
        " onfocus='highlightRow(this)' onblur='if(this.value&&!isNaN(this.value))this.value=parseFloat(this.value).toFixed(2)'"
        " oninput='updTot();updDiscVPT()' placeholder='e.g. -5.00'"
        " style='width:100%;text-align:right;border:1px solid #fca5a5;border-radius:4px;padding:3px 4px;font-size:12px;color:#dc2626;background:#fef2f2;font-family:DM Mono,monospace'></td>"
        "</tr>"
    )

    # Card rows
    card_rows = ""
    # Internet Orders auto-populated row first
    io_val = fv("internet_orders") or ""
    card_rows += (
        "<tr style='border-bottom:1px solid #f1f5f9;background:#f0f9ff'>"
        "<td style='padding:4px 10px;font-size:12px;color:#334155'>Internet Orders <span style='font-size:10px;color:#94a3b8'>(auto)</span></td>"
        "<td style='padding:2px 6px'>"
        "<input type='number' step='0.01' name='card_internet_orders' id='card_io_field'"
        " value='" + io_val + "' readonly"
        " style='width:100%;text-align:right;border:1px solid #bae6fd;border-radius:5px;"
        "padding:4px 6px;font-size:13px;font-family:DM Mono,monospace;background:#f0f9ff'>"
        "</td></tr>"
    )
    for col, label in CARD_TYPES:
        card_rows += (
            "<tr style='border-bottom:1px solid #f1f5f9'>"
            "<td style='padding:4px 10px;font-size:12px;color:#334155'>" + label + "</td>"
            "<td style='padding:2px 6px'>"
            "<input type='number' step='0.01' name='" + col + "' value='" + fv(col) + "'"
            " onblur='if(this.value)this.value=(parseFloat(this.value)||0).toFixed(2)'"
            " oninput='updCards()' placeholder='0.00'"
            " style='width:100%;text-align:right;border:1px solid #e2e8f0;border-radius:5px;"
            "padding:4px 6px;font-size:13px;font-family:DM Mono,monospace'>"
            "</td></tr>"
        )

    # Denomination rows
    denoms = [
        ("notes_50",50.00,"£50"),("notes_20",20.00,"£20"),("notes_10",10.00,"£10"),
        ("notes_5",5.00,"£5"),("coins_2",2.00,"£2"),("coins_1",1.00,"£1"),
        ("coins_50p",0.50,"50p"),("coins_20p",0.20,"20p"),("coins_10p",0.10,"10p"),
        ("coins_5p",0.05,"5p"),("coins_2p",0.02,"2p"),("coins_1p",0.01,"1p"),
    ]
    denom_rows = ""
    for col, val, label in denoms:
        cnt = fvn(col)
        dv  = ("£" + "%.2f" % (int(e.get(col,0) or 0)*val)) if e.get(col) else "&mdash;"
        denom_rows += (
            "<tr style='border-bottom:1px solid #f1f5f9'>"
            "<td style='padding:3px 8px;font-size:12px;font-weight:700;color:#334155'>" + label + "</td>"
            "<td style='padding:2px 4px'>"
            "<input type='number' name='" + col + "' value='" + cnt + "'"
            " onfocus='highlightRow(this)' oninput='updDenoms()' placeholder='0' min='0' step='1'"
            " style='width:60px;text-align:right;border:1px solid #e2e8f0;border-radius:5px;padding:3px 5px;font-size:12px'>"
            "</td>"
            "<td style='padding:3px 8px;font-size:12px;text-align:right;font-family:DM Mono,monospace'"
            " id='dv_" + col + "'>" + dv + "</td>"
            "</tr>"
        )

    # Checklist items
    chk_items = [
        ("c1","1. Discount entered"),("c2","2. Staff on shift entered"),
        ("c3","3. Person cashing up entered"),("c4","4. Customer number entered"),
        ("c5","5. Print count entered"),("c6","6. CR1 entered"),
        ("c7","7. CR2 entered"),("c8","8. Paid out checked"),
        ("c9","9. Comments if needed"),("c10","10. Z-read number entered"),
        ("c11","11. Actual cash C/F entered"),("c12","12. All checks complete"),
    ]
    chk_html = ""
    for cid, clbl in chk_items:
        chk_html += (
            "<div style='display:flex;gap:8px;align-items:center;padding:6px 0;"
            "border-bottom:1px solid #f1f5f9;font-size:13px;color:#64748b'>"
            "<span id='" + cid + "' style='font-size:16px'>&#9744;</span>"
            "<span>" + clbl + "</span></div>"
        )

    # ZZ row for Thursday
    zz_row = ""
    if is_thurs:
        zz_row = (
            "&nbsp;&nbsp;&nbsp;<span style='font-size:12px;font-weight:700'>Till ZZ' No.:</span>"
            "&nbsp;<span id='z2_p_check' onclick='ztick(this)' style='font-size:16px;cursor:pointer'>&#9744;</span>"
            "&nbsp;<input type='number' name='z2_read_no' id='inp_z2' value='" + z2_cur + "'"
            " min='1' step='1' oninput='chkZ2();chks()' placeholder='ZZ'"
            " style='width:65px;text-align:center;border:2px solid #fefce8;border-radius:6px;"
            "padding:4px 6px;font-size:13px;font-weight:900;background:#fefce8;color:#0f2942'>"
            "&nbsp;<span id='z2_echo' style='font-size:12px;opacity:.7'></span>"
            "&nbsp;<span id='z2_tick' style='font-size:16px'>&#9744;</span>"
            "&nbsp;<span id='z2_status' style='font-size:11px;font-weight:700;color:#dc2626'>Enter ZZ number</span>"
        )

    cats_js   = repr([col for col,_,_,_ in SALES_CATEGORIES])
    cards_js  = repr([col for col,_ in CARD_TYPES])
    denoms_js = repr([(col,val) for col,val,_ in denoms])

    content = (
        "<div class='flex justify-between items-center flex-wrap gap-3'>"
        "<div><a href='/sales?store=" + store + "' style='color:#1e3a5f;font-size:13px;font-weight:700'>&#8592; Back</a>"
        "<div class='text-2xl font-black text-slate-800 mt-1'>&#128221; Daily Cash Sheet &mdash; " + store + "</div>"
        "<div style='font-size:13px;color:#64748b'>" + d_fmt + "</div></div>"
        "<div style='display:flex;gap:8px'>"
        "<a href='/sales/enter?store=" + store + "&date=" + prev_d + "' class='btn-secondary' style='padding:5px 12px'>&#8592; Prev</a>"
        "<a href='/sales/enter?store=" + store + "&date=" + next_d + "' class='btn-secondary' style='padding:5px 12px'>Next &#8594;</a>"
        "</div></div>"

        # Z read bar
        "<div style='display:flex;align-items:center;gap:8px;flex-wrap:wrap;"
        "padding:6px 0 10px 0;border-bottom:2px solid #e2e8f0;width:fit-content;margin-bottom:12px'>"
        "<span style='font-size:12px;font-weight:700'>Till Z' No.:</span>"
        "&nbsp;<span id='z_p_check' onclick='ztick(this)' style='font-size:18px;cursor:pointer' title='Click to confirm Z entered'>&#9744;</span>"
        "&nbsp;<input type='number' name='z_read_no' id='inp_z' value='" + z_cur + "'"
        " min='1' step='1' oninput='chkZ();chks()' placeholder='Z No.'"
        " style='width:70px;text-align:center;border:2px solid #fefce8;border-radius:6px;"
        "padding:5px 6px;font-size:14px;font-weight:900;background:#fefce8;color:#0f2942'>"
        "&nbsp;<span id='z_tick' style='font-size:20px'>&#9744;</span>"
        "&nbsp;<span id='z_status' style='font-size:11px;font-weight:700;color:#dc2626'>Enter Z number to check</span>"
        + zz_row +
        "</div>"

        "<form action='/sales/enter' method='POST' id='salesForm'>"
        "<input type='hidden' name='store' value='" + store + "'>"
        "<input type='hidden' name='date' value='" + date + "'>"
        "<input type='hidden' name='prev_z' value='" + str(prev_z) + "'>"
        "<input type='hidden' name='prev_z2' value='" + str(prev_z2) + "'>"

        # Tab bar
        "<div style='display:flex;gap:0;border-bottom:2px solid #e2e8f0;margin-bottom:4px'>"
        "<button type='button' onclick='showTab(1)' id='tab1' class='tab-btn active'>"
        "&#128200; Sales &amp; Shift <span id='b1' class='tab-badge'>0/25</span></button>"
        "<button type='button' onclick='showTab(2)' id='tab2' class='tab-btn'>"
        "&#128179; Card Breakdown <span id='b2' class='tab-badge'>—</span></button>"
        "<button type='button' onclick='showTab(3)' id='tab3' class='tab-btn'>"
        "&#128181; Cash &amp; Recon <span id='b3' class='tab-badge'>—</span></button>"
        "</div>"
        "<div style='background:#e2e8f0;border-radius:99px;height:4px;margin-bottom:12px'>"
        "<div id='tab-progress' style='background:#0f2942;border-radius:99px;height:4px;width:25%;transition:width .3s'></div>"
        "</div>"

        # Tab 1: Sales
        "<div id='panel1' class='tab-panel'>"
        "<div style='display:flex;gap:20px;align-items:start'>"
        "<div style='display:flex;flex-direction:column;gap:12px'>"
        "<div class='card' style='padding:0;overflow:hidden;width:fit-content'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;font-weight:700;font-size:13px'>Sales by Category</div>"
        "<div>"
        "<table style='border-collapse:collapse;table-layout:fixed'>"
        "<col style='width:215px'>"
        "<col style='width:60px'>"
        "<col style='width:70px'>"
        "<col style='width:85px'>"
        "</colgroup>"
        "<thead><tr style='background:#f8fafc'>"
        "<th style='padding:5px 8px;text-align:left;font-size:10px;color:#64748b'>CATEGORY</th>"
        "<th style='padding:5px 3px;text-align:center;font-size:10px;color:#92400e;background:#fefce8'>TRANS</th>"
        "<th style='padding:5px 5px;text-align:center;font-size:10px;color:#166534;background:#f0fdf4'>PER TRANS</th>"
        "<th style='padding:5px 3px;text-align:right;font-size:10px;color:#92400e;background:#fefce8'>ACTUAL £</th>"
        "</tr></thead>"
        "<tbody>" + cat_rows + "</tbody>"
        "</table></div>"
        "<div style='padding:10px 16px;background:#0f2942;display:flex;justify-content:space-between;align-items:center'>"
        "<span style='font-weight:700;color:white;font-size:13px'>TOTAL SALES</span>"
        "<span id='tot_sales' style='font-weight:900;color:white;font-size:18px;font-family:DM Mono,monospace'>£0.00</span>"
        "</div></div>"
        "<div class='card' style='padding:0;overflow:hidden;max-width:432px'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;font-weight:700;font-size:13px;margin:0 0 0 0;border-radius:8px 8px 0 0'>Shift &amp; Till Details</div>"
        "<table style='width:100%;border-collapse:collapse'><tbody>"
        "<tr style='border-bottom:1px solid #f1f5f9'>"
        "<td style='padding:4px 6px;font-size:12px;font-weight:700;color:#64748b;white-space:nowrap;width:90px'>Staff on Shift</td>"
        "<td style='padding:4px 6px'><input type='text' name='staff_on_shift' value='" + fvs("staff_on_shift") + "' oninput='chks()'"
        " placeholder='e.g. Kaleem / Rhys / Jessica' style='width:100%;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;box-sizing:border-box'></td></tr>"
        "<tr style='border-bottom:1px solid #f1f5f9'>"
        "<td style='padding:4px 8px;font-size:12px;font-weight:700;color:#64748b;white-space:nowrap'>Person Cashing Up</td>"
        "<td style='padding:4px 6px'><input type='text' name='person_cashing_up' value='" + fvs("person_cashing_up") + "' oninput='chks()'"
        " placeholder='One name only' style='width:100%;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;box-sizing:border-box'></td></tr>"
        "<tr style='border-bottom:1px solid #f1f5f9'>"
        "<td style='padding:4px 8px;font-size:12px;font-weight:700;color:#64748b;white-space:nowrap'>Customer Count</td>"
        "<td style='padding:4px 6px'><input type='number' name='customer_count' value='" + fvi("customer_count") + "' oninput='chks()'"
        " placeholder='From till read' style='width:100%;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;box-sizing:border-box'></td></tr>"
        "<tr style='border-bottom:1px solid #f1f5f9'>"
        "<td style='padding:4px 8px;font-size:12px;font-weight:700;color:#64748b;white-space:nowrap'>Print Count (D3000)</td>"
        "<td style='padding:4px 6px'><input type='number' name='print_count' value='" + fvi("print_count") + "' oninput='chks()'"
        " placeholder='From D3000' style='width:100%;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;box-sizing:border-box'></td></tr>"
        "<tr>"
        "<td style='padding:4px 8px;font-size:12px;font-weight:700;color:#64748b;white-space:nowrap'>Apply &amp; Go Count</td>"
        "<td style='padding:4px 6px'><input type='number' name='apply_go_count' value='" + fvi("apply_go_count") + "'"
        " placeholder='e.g. 0' style='width:100%;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;box-sizing:border-box'></td></tr>"
        "</tbody></table>"
        "</div>"
        "</div>"
        "<div style='position:sticky;top:16px'>"
        "<div class='card' style=''>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;font-weight:700;font-size:13px;margin:-14px -14px 12px -14px;border-radius:8px 8px 0 0'>&#9989; Completion Checklist</div>"
        "<div style='font-size:12px;color:#94a3b8;margin-bottom:10px'>Complete all items before saving</div>"
        + chk_html +
        "<div style='margin-top:16px'>"
        "<button type='submit' name='action' value='save' id='sbtn'"
        " class='btn-primary' style='width:100%;padding:12px;font-size:15px'>&#128190; Save Cash Sheet</button>"
        "<div id='swarn' style='display:none;margin-top:8px;background:#fef3c7;border:1px solid #fcd34d;"
        "border-radius:8px;padding:10px;font-size:13px;color:#92400e'>&#9888; Complete all checklist items first</div>"
        "</div></div>"
        "</div>"

        "</div>"
        "</div>"
        "</div>"
        "<button type='button' onclick='showTab(2)' class='btn-primary'>Next: Card Breakdown &#8594;</button>"
        # Tab 2: Till & Cards
        "<div id='panel2' class='tab-panel' style='display:none'>"
        "<div style='max-width:950px'>"
        "<div class='grid gap-4' style='grid-template-columns:1fr 1fr'>"

        # Left: Shift details + Till reads
        "<div style='display:flex;flex-direction:column;gap:10px'>"

        "<div class='card' style='padding:0;overflow:hidden'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;font-weight:700;font-size:13px'>Total Credit Cards as per Till</div><table style='border-collapse:collapse;width:auto'><tbody><tr style='border-bottom:1px solid #f1f5f9'><td style='padding:6px 10px;font-size:13px;color:#334155;white-space:nowrap'>CR1 &mdash; Credit Card Sales</td><td style='padding:4px 8px;text-align:right'><input type='number' step='0.01' name='till_credit_sales' value='" + fv("till_credit_sales") + "' onblur='if(this.value&&!isNaN(this.value))this.value=parseFloat(this.value).toFixed(2)' oninput='updTillTotal();chks()' placeholder='0.00' style='width:80px;text-align:right;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;font-family:DM Mono,monospace'></td></tr><tr style='border-bottom:1px solid #f1f5f9'><td style='padding:6px 10px;font-size:13px;color:#334155;white-space:nowrap'>CR2 &mdash; Internet Sales</td><td style='padding:4px 8px;text-align:right'><input type='number' step='0.01' name='till_internet_sales' id='cr2' value='" + fv("till_internet_sales") + "' onblur='if(this.value&&!isNaN(this.value))this.value=parseFloat(this.value).toFixed(2)' oninput='updTillTotal();chks()' placeholder='0.00' style='width:80px;text-align:right;border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;font-family:DM Mono,monospace'></td></tr><tr style='border-bottom:1px solid #e2e8f0;background:#f8fafc'><td style='padding:6px 10px;font-size:13px;font-weight:700;color:#0f2942'>Total CR1 + CR2</td><td style='padding:6px 10px;text-align:right;font-size:13px;font-weight:900;font-family:DM Mono,monospace;color:#0f2942' id='till_total'>£0.00</td></tr><tr style='border-bottom:1px solid #e2e8f0;background:#f8fafc'><td style='padding:6px 10px;font-size:13px;font-weight:700;color:#0f2942'>Total Cards (PDQ)</td><td style='padding:6px 10px;text-align:right;font-size:13px;font-weight:900;font-family:DM Mono,monospace;color:#0f2942' id='till_cards_ref'>£0.00</td></tr><tr style='background:#f8fafc'><td style='padding:6px 10px;font-size:13px;font-weight:700;color:#0f2942'>Difference</td><td style='padding:6px 10px;text-align:right;font-size:13px;font-weight:900;font-family:DM Mono,monospace' id='till_diff_display'>—</td></tr></tbody></table></div>"
        "<div id='card_chk' style='margin-top:6px;font-size:12px;display:none'></div>"
        "</div>"

        # Right: Card breakdown
        "<div class='card' style='padding:0;overflow:hidden'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;display:flex;justify-content:space-between;font-weight:700;font-size:13px'>Card Breakdown (PDQ)</div>"
        "<table style='width:100%;border-collapse:collapse'><tbody>" + card_rows + "</tbody>"
        "<tfoot><tr style='background:#f8fafc;border-top:2px solid #e2e8f0'>"
        "<td style='padding:6px 10px;font-weight:900;font-size:13px'>Total Cards</td>"
        "<td style='padding:6px 6px;text-align:right;font-weight:900;font-size:14px;font-family:DM Mono,monospace' id='card_tot2'>£0.00</td>"
        "</tr></tfoot></table></div>"
        "</div>"
        "</div>"
        "<div style='display:flex;justify-content:space-between;margin-top:10px'>"
        "<button type='button' onclick='showTab(1)' class='btn-secondary'>&#8592; Back: Sales &amp; Shift</button>"
        "<button type='button' onclick='showTab(3)' class='btn-primary'>Next: Cash &amp; Recon &#8594;</button>"
        "</div></div>"

        # Tab 3: Cash Count
        "<div id='panel3' class='tab-panel' style='display:none'>"
        "<div style='max-width:950px'>"
        "<div class='grid gap-4' style='grid-template-columns:1fr 1fr'>"

        # Left: Denominations
        "<div class='card' style='padding:0;overflow:hidden'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;display:flex;justify-content:space-between;font-weight:700;font-size:13px'>"
        "Cash Count by Denomination<span id='cash_count_total' style='font-family:DM Mono,monospace'>£0.00</span></div>"
        "<table style='width:100%;border-collapse:collapse'>"
        "<thead><tr style='background:#f8fafc'>"
        "<th style='padding:3px 8px;text-align:left;font-size:10px;color:#64748b'>DENOM</th>"
        "<th style='padding:3px 4px;text-align:center;font-size:10px;color:#64748b'>COUNT</th>"
        "<th style='padding:3px 8px;text-align:right;font-size:10px;color:#64748b'>VALUE</th>"
        "</tr></thead>"
        "<tbody>" + denom_rows + "</tbody>"
        "<tfoot>"
        "<tr style='background:#f0f9ff'>"
        "<td colspan='2' style='padding:4px 8px;font-size:11px;font-weight:700;color:#0369a1'>Notes Tin</td>"
        "<td style='padding:4px 4px;text-align:right'>"
        "<input type='number' step='0.01' name='notes_tin' id='notes_tin_inp'"
        " value='" + fv("notes_tin") + "' oninput='updCashStore()' placeholder='0.00'"
        " style='width:80px;text-align:right;border:1px solid #bae6fd;border-radius:5px;padding:3px 5px;font-size:12px;font-family:DM Mono,monospace'>"
        "</td></tr>"
        "<tr style='background:#f0fdf4'>"
        "<td colspan='2' style='padding:4px 8px;font-size:11px;font-weight:700;color:#166534'>Change Tin</td>"
        "<td style='padding:4px 4px;text-align:right'>"
        "<input type='number' step='0.01' name='change_tin' id='change_tin_inp'"
        " value='" + fv("change_tin") + "' oninput='updCashStore()' placeholder='0.00'"
        " style='width:80px;text-align:right;border:1px solid #bbf7d0;border-radius:5px;padding:3px 5px;font-size:12px;font-family:DM Mono,monospace'>"
        "</td></tr>"
        "<tr style='border-top:2px solid #e2e8f0;background:#f8fafc'>"
        "<td colspan='2' style='padding:5px 8px;font-size:12px;font-weight:900;color:#0f2942'>Total Cash in Store</td>"
        "<td style='padding:5px 8px;text-align:right;font-size:14px;font-weight:900;color:#0f2942;font-family:DM Mono,monospace' id='total_cash_store'>£0.00</td>"
        "</tr></tfoot></table>"
        "<input type='hidden' name='notes_tin' id='hid_notes_tin'>"
        "<input type='hidden' name='change_tin' id='hid_change_tin'>"
        "<input type='hidden' name='total_cash_store' id='hid_total_cash'>"
        "</div>"

        # Right: Cash reconciliation
        "<div style='display:flex;flex-direction:column;gap:10px'>"
        "<div class='card'>"
        "<div style='padding:8px 14px;background:#0f2942;color:white;font-weight:700;font-size:13px;margin:-14px -14px 12px -14px;border-radius:8px 8px 0 0'>Cash Reconciliation</div>"
        "<div class='grid gap-2' style='grid-template-columns:1fr 1fr'>"
        "<div><label>Opening Cash B/F <span style='font-weight:400;color:#94a3b8'>(auto)</span></label>"
        "<input type='number' step='0.01' name='opening_cash_bf'"
        " value='" + ("%.2f" % bf_auto if bf_auto else "") + "'"
        " oninput='updCash()' placeholder='From yesterday'"
        " style='border:1px solid #bae6fd;border-radius:6px;padding:6px 8px;font-size:13px;width:100%;text-align:right;font-family:DM Mono,monospace;background:#f0f9ff'></div>"
        "<div><label>Paid Out Total</label>"
        "<input type='number' step='0.01' name='paid_out_total' value='" + fv("paid_out_total") + "'"
        " id='inp_paidout' oninput='chks();updCash()' placeholder='0.00'"
        " style='border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;width:100%;text-align:right;font-family:DM Mono,monospace'></div>"
        "<div style='grid-column:1/-1'><label>Paid Out Details</label>"
        "<textarea name='paid_out_notes' rows='2' placeholder='e.g. Cleaning £5.00'"
        " style='border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:12px;width:100%'>" + fvs("paid_out_notes") + "</textarea></div>"
        "</div>"
        "<div style='margin-top:8px;background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:8px'>"
        "<label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:600;color:#166534'>"
        "<input type='checkbox' name='paid_out_checked' id='po_chk' " + ("checked" if e.get("paid_out_checked") else "") + " oninput='chks()'"
        " style='width:18px;height:18px'>"
        "I confirm paid out sheet has been checked and total is correct"
        "</label></div>"
        "<div style='margin-top:8px;padding:8px;background:#f8fafc;border-radius:8px;font-size:12px;line-height:1.9'>"
        "<div style='display:flex;justify-content:space-between'><span>Total Sales</span><span id='r_sales' style='font-family:DM Mono,monospace'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between'><span>+ Opening B/F</span><span id='r_bf' style='font-family:DM Mono,monospace'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between;border-top:1px solid #e2e8f0;padding-top:2px;font-weight:700'><span>Sub Total</span><span id='r_sub' style='font-family:DM Mono,monospace'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between'><span>- Paid Out</span><span id='r_po' style='font-family:DM Mono,monospace;color:#dc2626'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between'><span>- Total Cards</span><span id='r_cards' style='font-family:DM Mono,monospace;color:#dc2626'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between;border-top:1px solid #e2e8f0;padding-top:2px;font-weight:900'><span>Theoretical Cash</span><span id='r_theo' style='font-family:DM Mono,monospace'>£0.00</span></div>"
        "<div style='display:flex;justify-content:space-between;font-weight:700'><span>Difference</span><span id='r_diff' style='font-family:DM Mono,monospace'>—</span></div>"
        "</div>"
        "<div class='grid gap-2' style='grid-template-columns:1fr 1fr;margin-top:8px'>"
        "<div><label>Paid Into Bank</label>"
        "<input type='number' step='0.01' name='total_paid_bank' value='" + fv("total_paid_bank") + "'"
        " oninput='updCash()' placeholder='0.00'"
        " style='border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;width:100%;text-align:right;font-family:DM Mono,monospace'></div>"
        "<div><label>Actual Cash C/F Tomorrow</label>"
        "<input type='number' step='0.01' name='actual_cash_cf' value='" + fv("actual_cash_cf") + "'"
        " id='inp_cashcf' oninput='chks();updCash()' placeholder='0.00'"
        " style='border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;width:100%;text-align:right;font-family:DM Mono,monospace'></div>"
        "</div>"
        "<div id='diff_box' style='margin-top:6px;display:none'>"
        "<label style='color:#dc2626'>Reason for Difference (required)</label>"
        "<textarea name='till_diff_reason' rows='2' placeholder='Please explain'"
        " style='border:1px solid #fca5a5;border-radius:6px;padding:6px 8px;font-size:12px;width:100%;margin-top:3px'>" + fvs("till_diff_reason") + "</textarea></div>"
        "</div>"
        "<div class='card'><label style='font-size:13px;font-weight:700;color:#0f2942'>Comments</label>"
        "<textarea name='notes' rows='3' oninput='chks()'"
        " placeholder='Any notes, discrepancies or issues to report'"
        " style='border:1px solid #e2e8f0;border-radius:6px;padding:6px 8px;font-size:13px;width:100%;margin-top:6px'>" + fvs("notes") + "</textarea></div>"
        "</div>"
        "</div>"
        "<div style='display:flex;justify-content:space-between;margin-top:10px'>"
        "<button type='button' onclick='showTab(2)' class='btn-secondary'>&#8592; Back: Card Breakdown</button>"
        "</div></div>"

        "</div></form>"
    )

    # Add tab styles and JS
    tab_css = """<style>
tr.sales-row-active td {  background:#dbeafe !important;}.sales-row-active td input {  background:#bfdbfe !important;}.sales-table tr.discount-rowtr.sales-row-active td {  background:#fee2e2 !important;}.sales-table tr.discount-row.sales-row-active td input {  background:#fecaca !important;}.tab-btn{padding:9px 18px;font-size:13px;font-weight:700;color:#64748b;background:none;
  border:none;cursor:pointer;border-bottom:3px solid transparent;margin-bottom:-2px;
  font-family:'DM Sans',sans-serif;white-space:nowrap;}
.tab-btn:hover{color:#0f2942;}
.tab-btn.active{color:#0f2942;border-bottom-color:#0f2942;}
.tab-badge{background:#e2e8f0;color:#64748b;border-radius:99px;padding:1px 7px;
  font-size:11px;margin-left:6px;}
.tab-btn.active .tab-badge{background:#0f2942;color:white;}
.tab-btn.done .tab-badge{background:#16a34a;color:white;}
</style>"""

    js_code = """
function highlightRow(el){  var tbl=el.closest('table');  if(tbl)tbl.querySelectorAll('.sales-row-active').forEach(function(r){r.classList.remove('sales-row-active');});  var row=el.closest('tr');  if(row)row.classList.add('sales-row-active');}function showTab(n){
  for(var i=1;i<=3;i++){
    document.getElementById('panel'+i).style.display=i===n?'block':'none';
    document.getElementById('tab'+i).classList.toggle('active',i===n);
  }
  document.getElementById('tab-progress').style.width=(n*33.3)+'%';
}
function ztick(t){
  var on=t.textContent==='\\u2610';
  t.innerHTML=on?'\\u2713':'\\u2610';
  t.style.color=on?'#16a34a':'';
}
function gv(n){var es=document.getElementsByName(n);return es.length?(parseFloat(es[0].value||0)||0):0;}
function gs(n){var es=document.getElementsByName(n);return es.length?es[0].value.trim():'';}
function fm(n){return'\\xa3'+n.toFixed(2);}
function tk(id,ok){var e=document.getElementById(id);if(!e)return;
  e.innerHTML=ok?'&#10003;':'&#9744;';e.style.color=ok?'#16a34a':'#94a3b8';e.style.fontSize=ok?'18px':'15px';}

var CATS=CAT_PH;
var CARDS=CARD_PH;var ALL_CARDS=CARDS.concat(['card_internet_orders']);
var DENOMS=DENOM_PH;
var PREV_Z=PREVZ_PH;
var PREV_Z2=PREVZ2_PH;

function updVPT(col,tcol){
  var a=parseFloat(document.getElementsByName(col)[0]?.value||0)||0;
  var t=parseInt(document.getElementsByName(tcol)[0]?.value||0)||0;
  var el=document.getElementById('vpt_'+col);
  if(el)el.textContent=t>0?(a/t).toFixed(2):'\\u2014';
}
function updDiscVPT(){
  var a=parseFloat(document.getElementsByName('discount_amount')[0]?.value||0)||0;
  var t=parseInt(document.getElementsByName('discount_trans')[0]?.value||0)||0;
  var el=document.getElementById('vpt_discount');
  if(el)el.textContent=t>0?(a/t).toFixed(2):'\\u2014';
}
function updTot(){var io=parseFloat(document.getElementsByName('internet_orders')[0]?.value||0)||0;var ioF=document.getElementById('card_io_field');if(ioF){ioF.value=io>0?io.toFixed(2):'';updCards();}  updTillTotal();
  var t=0;
  CATS.forEach(function(c){var es=document.getElementsByName(c);if(es.length)t+=parseFloat(es[0].value||0)||0;});
  document.getElementById('tot_sales').textContent=fm(t);
  var filled=0;
  CATS.forEach(function(c){var es=document.getElementsByName(c);if(es.length&&es[0].value)filled++;});
  document.getElementById('b1').textContent=filled+'/25';
  updCash();chks();
}
function updTillTotal(){var cr1=parseFloat(document.getElementsByName('till_credit_sales')[0]?.value||0)||0;var cr2=parseFloat(document.getElementsByName('till_internet_sales')[0]?.value||0)||0;var tot=cr1+cr2;var tt=document.getElementById('till_total');if(tt)tt.textContent='\xa3'+tot.toFixed(2);var cards=0;ALL_CARDS.forEach(function(c){var es=document.getElementsByName(c);if(es.length)cards+=parseFloat(es[0].value||0)||0;});var tr=document.getElementById('till_cards_ref');if(tr)tr.textContent='\xa3'+cards.toFixed(2);var diff=tot-cards;var de=document.getElementById('till_diff_display');if(tot>0||cards>0){if(Math.abs(diff)<0.01){de.textContent='\u2713 Balanced';de.style.color='#16a34a';}else{de.textContent=(diff>0?'\u26a0 CR1+CR2 exceeds PDQ by \xa3':'\u26a0 PDQ exceeds CR1+CR2 by \xa3')+Math.abs(diff).toFixed(2);de.style.color='#dc2626';}}else{de.textContent='\u2014';de.style.color='';}}function updCards(){
  var t=0;
  ALL_CARDS.forEach(function(c){var es=document.getElementsByName(c);if(es.length)t+=parseFloat(es[0].value||0)||0;});
  var ct=document.getElementById('card_tot');if(ct)ct.textContent=fm(t);
  var ct2=document.getElementById('card_tot2');if(ct2)ct2.textContent=fm(t);
  var tcr=document.getElementById('till_cards_ref');if(tcr)tcr.textContent=fm(t);
  updTillTotal();
  updCash();
}
function updDenoms(){
  var notes=0,coins=0;
  DENOMS.forEach(function(d){
    var col=d[0],val=d[1];
    var inps=document.getElementsByName(col);
    var cnt=inps.length?(parseInt(inps[0].value)||0):0;
    var amt=cnt*val;
    var dv=document.getElementById('dv_'+col);
    if(dv)dv.textContent=cnt?fm(amt):'\\u2014';
    if(val>=5)notes+=amt;else coins+=amt;
  });
  document.getElementById('cash_count_total').textContent=fm(notes+coins);
  updCashStore();
}
function updCashStore(){
  var notes=0,coins=0;
  DENOMS.forEach(function(d){
    var inps=document.getElementsByName(d[0]);
    var cnt=inps.length?(parseInt(inps[0].value)||0):0;
    if(d[1]>=5)notes+=cnt*d[1];else coins+=cnt*d[1];
  });
  var ntin=parseFloat(document.getElementById('notes_tin_inp')?.value||0)||0;
  var ctin=parseFloat(document.getElementById('change_tin_inp')?.value||0)||0;
  var tot=(notes+coins)+ntin+ctin;
  document.getElementById('total_cash_store').textContent=fm(tot);
  document.getElementById('hid_notes_tin').value=ntin.toFixed(2);
  document.getElementById('hid_change_tin').value=ctin.toFixed(2);
  document.getElementById('hid_total_cash').value=tot.toFixed(2);
}
function updCash(){
  var sales=0;
  CATS.forEach(function(c){var es=document.getElementsByName(c);if(es.length)sales+=parseFloat(es[0].value||0)||0;});
  var cards=0;
  CARDS.forEach(function(c){var es=document.getElementsByName(c);if(es.length)cards+=parseFloat(es[0].value||0)||0;});
  var bf=gv('opening_cash_bf'),po=gv('paid_out_total');
  var sub=sales+bf,theo=sub-po-cards;
  var cf=gv('actual_cash_cf'),diff=cf-theo;
  document.getElementById('r_sales').textContent=fm(sales);
  document.getElementById('r_bf').textContent=fm(bf);
  document.getElementById('r_sub').textContent=fm(sub);
  document.getElementById('r_po').textContent=fm(po);
  document.getElementById('r_cards').textContent=fm(cards);
  document.getElementById('r_theo').textContent=fm(theo);
  var de=document.getElementById('r_diff');
  if(cf>0){de.textContent=(diff>=0?'+':'')+fm(diff);de.style.color=Math.abs(diff)<0.01?'#16a34a':'#dc2626';
    document.getElementById('diff_box').style.display=Math.abs(diff)>0.01?'block':'none';}
  else{de.textContent='\\u2014';de.style.color='';document.getElementById('diff_box').style.display='none';}
}
function chkZ(){
  var zel=document.getElementsByName('z_read_no');
  var z=zel.length?(parseInt(zel[0].value)||0):0;
  var el=document.getElementById('z_status');if(!el)return;
  if(!z){el.textContent='Enter Z number to check';el.style.color='#dc2626';return;}
  if(PREV_Z===0){el.textContent='No previous Z to compare';el.style.color='#d97706';return;}
  if(z===PREV_Z+1){el.style.color='#16a34a';el.textContent='\\u2713 Z No. OK ('+PREV_Z+'\\u2192'+z+')';}
  else{el.style.color='#dc2626';el.textContent='\\u26a0 Expected Z '+(PREV_Z+1)+' got '+z+' \\u2014 comment required';}
}
function chkZ2(){
  var zel=document.getElementsByName('z2_read_no');
  var z2=zel.length?(parseInt(zel[0].value)||0):0;
  var el=document.getElementById('z2_status');if(!el)return;
  if(!z2){el.textContent='Enter ZZ number';el.style.color='#dc2626';return;}
  if(PREV_Z2===0){el.textContent='No previous ZZ';el.style.color='#d97706';return;}
  if(z2===PREV_Z2+1){el.style.color='#16a34a';el.textContent='\\u2713 ZZ No. OK';}
  else{el.style.color='#dc2626';el.textContent='\\u26a0 Expected ZZ '+(PREV_Z2+1);}
}
function chks(){
  var de=document.getElementsByName('discount_amount');
  tk('c1',de.length&&de[0].value!=='');
  tk('c2',gs('staff_on_shift').length>1);
  tk('c3',gs('person_cashing_up').length>1);
  tk('c4',gv('customer_count')>0);
  var pe=document.getElementsByName('print_count');tk('c5',pe.length&&pe[0].value!=='');
  var r1=document.getElementsByName('till_credit_sales');tk('c6',r1.length&&r1[0].value!=='');
  var r2=document.getElementsByName('till_internet_sales');tk('c7',r2.length&&r2[0].value!=='');
  var poc=document.getElementById('po_chk');tk('c8',poc&&poc.checked);
  tk('c9',true);
  var zel=document.getElementsByName('z_read_no');tk('c10',zel.length&&parseInt(zel[0].value||0)>0);
  var cfe=document.getElementsByName('actual_cash_cf');var c11ok=cfe.length&&cfe[0].value!=='';tk('c11',c11ok);
  var r1v=r1.length&&r1[0].value!=='';
  var allOk=gs('staff_on_shift').length>1&&gs('person_cashing_up').length>1&&
    gv('customer_count')>0&&zel.length&&parseInt(zel[0].value||0)>0&&c11ok&&r1v;
  tk('c12',allOk);
  var done=0;for(var i=1;i<=12;i++){var e=document.getElementById('c'+i);if(e&&e.textContent==='\\u2713')done++;}
  document.getElementById('b4').textContent=done+'/12';
  if(done>=10)document.getElementById('tab3').classList.add('done');
  var btn=document.getElementById('sbtn'),wrn=document.getElementById('swarn');
  if(btn){btn.style.opacity=allOk?'1':'0.6';
    btn.onclick=function(ev){if(!allOk){ev.preventDefault();wrn.style.display='block';return false;}wrn.style.display='none';};}
}
document.addEventListener('DOMContentLoaded',function(){
  updTot();updCards();updDenoms();chks();chkZ();
  ALL_CARDS.forEach(function(c){var es=document.getElementsByName(c);if(es.length&&es[0]&&!es[0].readOnly)es[0].addEventListener('input',updCards);});
  var cr2=document.getElementsByName('till_internet_sales');
  if(cr2.length)cr2[0].addEventListener('input',function(){this.dataset.manual='1';});
});
"""

    js_code = (js_code
        .replace("CAT_PH", cats_js)
        .replace("CARD_PH", cards_js)
        .replace("DENOM_PH", denoms_js)
        .replace("PREVZ_PH", str(prev_z))
        .replace("PREVZ2_PH", str(prev_z2))
    )

    full_content = tab_css + content + "<script>" + js_code + "</script>"
    return page("Daily Cash Entry", full_content, user, "sales")




@app.post("/sales/enter")
async def save_sales_entry(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    form  = await request.form()
    store = form.get("store","")
    date  = form.get("date","")

    def fn(k):
        try: return float(form.get(k, 0) or 0)
        except: return 0.0
    def fi(k):
        try: return int(form.get(k, 0) or 0)
        except: return 0

    # All numeric columns
    num_cols = (
        [col for col,_,_,_ in SALES_CATEGORIES] +
        ["discount_amount"] +
        [col for col,_ in CARD_TYPES] +
        ["opening_cash_bf","paid_out_total","till_credit_sales","till_internet_sales",
         "total_cash_store","notes_tin","change_tin","total_paid_bank","actual_cash_cf"]
    )
    int_cols = ["notes_50","notes_20","notes_10","notes_5","coins_2","coins_1",
                "coins_50p","coins_20p","coins_10p","coins_5p","coins_2p","coins_1p"] + ["trans_"+col for col,_,_,_ in SALES_CATEGORIES]

    all_cols   = num_cols + int_cols
    num_values = [fn(c) for c in num_cols]
    int_values = [fi(c) for c in int_cols]
    values     = num_values + int_values

    z_read     = fi("z_read_no") or None
    z2_read    = fi("z2_read_no") or None
    notes      = str(form.get("notes","") or "").strip() or None
    po_notes   = str(form.get("paid_out_notes","") or "").strip() or None
    diff_reason= str(form.get("till_diff_reason","") or "").strip() or None
    z2_comment = str(form.get("z2_diff_comment","") or "").strip() or None
    po_checked = 1 if form.get("paid_out_checked") else 0
    staff_shift= str(form.get("staff_on_shift","") or "").strip() or None
    cashup_by  = str(form.get("person_cashing_up","") or "").strip() or None
    cust_count = fi("customer_count") or None
    print_count= fi("print_count") or None
    apply_go   = fi("apply_go_count") or None
    entered_by = user.get("username","")

    set_clause = ", ".join(f"{c}=?" for c in all_cols)
    col_list   = ", ".join(all_cols)
    ph         = ", ".join("?" for _ in all_cols)

    q(f"""INSERT INTO daily_cashsheet
            (store_name, sale_date, z_read_no, z2_read_no, {col_list},
             paid_out_notes, notes, till_diff_reason, z2_diff_comment,
             paid_out_checked, staff_on_shift, person_cashing_up,
             customer_count, print_count, apply_go_count, entered_by)
         VALUES(?,?,?,?,{ph},?,?,?,?,?,?,?,?,?,?,?)
         ON CONFLICT(store_name, sale_date) DO UPDATE SET
            z_read_no=excluded.z_read_no,
            z2_read_no=excluded.z2_read_no,
            {set_clause},
            paid_out_notes=excluded.paid_out_notes,
            notes=excluded.notes,
            till_diff_reason=excluded.till_diff_reason,
            z2_diff_comment=excluded.z2_diff_comment,
            paid_out_checked=excluded.paid_out_checked,
            staff_on_shift=excluded.staff_on_shift,
            person_cashing_up=excluded.person_cashing_up,
            customer_count=excluded.customer_count,
            print_count=excluded.print_count,
            apply_go_count=excluded.apply_go_count,
            entered_by=excluded.entered_by""",
      [store, date, z_read, z2_read] + values +
      [po_notes, notes, diff_reason, z2_comment, po_checked,
       staff_shift, cashup_by, cust_count, print_count, apply_go, entered_by] +
      values)

    from urllib.parse import quote as uq
    week_start = get_week_start(date)
    return RedirectResponse(
        f"/sales?store={store}&week_start={week_start}&msg={uq('Cash sheet saved for ' + date)}",
        status_code=303)


# ── Sales Targets ─────────────────────────────────────────────────────────────

@app.get("/sales/targets", response_class=HTMLResponse)
def sales_targets(
    store:   str = "Uxbridge",
    year:    int = 0,
    session: str | None = Cookie(default=None),
    msg:     str = ""
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/sales", status_code=303)

    if not year: year = datetime.now().year

    targets = q("SELECT * FROM sales_targets WHERE store_name=? AND year=? ORDER BY month",
                (store, year), fetch=True) or []
    tmap    = {dict(t)["month"]: dict(t) for t in targets}

    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]

    rows = ""
    for m in range(1, 13):
        t       = tmap.get(m, {})
        target  = t.get("target_amount", 0)
        ly      = t.get("ly_actual", 0)
        rows += f"""<tr style='border-bottom:1px solid #f1f5f9'>
          <td style='padding:8px 12px;font-weight:700'>{months[m-1]}</td>
          <td style='padding:4px 8px'>
            <input type='number' step='0.01' form='targets_form' name='target_{m}'
              value='{"%.2f" % target if target else ""}'
              placeholder='0.00'
              style='width:100%;text-align:right;border:1px solid #e2e8f0;border-radius:6px;
                     padding:6px 8px;font-size:13px;font-family:DM Mono,monospace'>
          </td>
          <td style='padding:4px 8px'>
            <input type='number' step='0.01' form='targets_form' name='ly_{m}'
              value='{"%.2f" % ly if ly else ""}'
              placeholder='0.00'
              style='width:100%;text-align:right;border:1px solid #e2e8f0;border-radius:6px;
                     padding:6px 8px;font-size:13px;font-family:DM Mono,monospace'>
          </td>
        </tr>"""

    store_btns = ""
    for sv in ["Uxbridge","Newbury"]:
        cls = "btn-primary" if sv == store else "btn-secondary"
        store_btns += f"<a href='/sales/targets?store={sv}&year={year}' class='{cls}' style='padding:5px 14px;font-size:13px'>{sv}</a>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/sales?store={store}' style='color:#1e3a5f;font-size:13px;font-weight:700'>&#8592; Back to Sales</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>&#127919; Sales Targets — {store} {year}</div>
      </div>
      <div style='display:flex;gap:8px'>
        {store_btns}
        <a href='/sales/targets?store={store}&year={year-1}' class='btn-secondary' style='padding:5px 12px'>&#8592; {year-1}</a>
        <a href='/sales/targets?store={store}&year={year+1}' class='btn-secondary' style='padding:5px 12px'>{year+1} &#8594;</a>
      </div>
    </div>
    <form id='targets_form' action='/sales/targets' method='POST'>
      <input type='hidden' name='store' value='{store}'>
      <input type='hidden' name='year'  value='{year}'>
      <div class='card' style='padding:0;overflow:hidden'>
        <table style='width:100%;border-collapse:collapse;font-family:DM Sans,sans-serif'>
          <thead><tr style='background:#0f2942;color:white'>
            <th style='padding:10px 12px;text-align:left;font-size:12px'>Month</th>
            <th style='padding:10px 8px;text-align:right;font-size:12px'>Target (£)</th>
            <th style='padding:10px 8px;text-align:right;font-size:12px'>Last Year Actual (£)</th>
          </tr></thead>
          <tbody>{rows}</tbody>
        </table>
      </div>
      <div style='margin-top:12px'>
        <button type='submit' class='btn-primary'>&#128190; Save Targets</button>
      </div>
    </form>"""

    return page("Sales Targets", content, user, "sales")


@app.post("/sales/targets")
async def save_targets(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    form  = await request.form()
    store = form.get("store","")
    year  = int(form.get("year", datetime.now().year))

    for m in range(1, 13):
        try: target = float(form.get(f"target_{m}", 0) or 0)
        except: target = 0
        try: ly = float(form.get(f"ly_{m}", 0) or 0)
        except: ly = 0
        q("""INSERT INTO sales_targets (store_name,year,month,target_amount,ly_actual)
             VALUES(?,?,?,?,?)
             ON CONFLICT(store_name,year,month) DO UPDATE SET
                target_amount=excluded.target_amount,
                ly_actual=excluded.ly_actual""",
          (store, year, m, target, ly))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/sales/targets?store={store}&year={year}&msg={uq('Targets saved')}",
        status_code=303)


# ── Franchise Return & Manager's Report (placeholders for now) ────────────────

@app.get("/sales/franchise-return", response_class=HTMLResponse)
def franchise_return(store: str="", week_start: str="", session: str|None=Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    content = f"""
    <div class='text-2xl font-black text-slate-800'>&#128196; Franchise Return — {store}</div>
    <div class='card text-center' style='padding:40px;color:#94a3b8'>
      <div style='font-size:40px;margin-bottom:12px'>&#128196;</div>
      <div style='font-weight:700;font-size:16px;color:#334155'>Coming Soon</div>
      <div style='font-size:13px;margin-top:8px'>
        The Franchise Return PDF will be generated here once daily sales data is entered for the full week.
      </div>
      <a href='/sales?store={store}&week_start={week_start}' class='btn-secondary' style='margin-top:16px;display:inline-block'>
        &#8592; Back to Sales
      </a>
    </div>"""
    return page("Franchise Return", content, user, "sales")

@app.get("/sales/managers-report", response_class=HTMLResponse)
def managers_report(store: str="", week_start: str="", session: str|None=Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    content = f"""
    <div class='text-2xl font-black text-slate-800'>&#128200; Manager's Report — {store}</div>
    <div class='card text-center' style='padding:40px;color:#94a3b8'>
      <div style='font-size:40px;margin-bottom:12px'>&#128200;</div>
      <div style='font-weight:700;font-size:16px;color:#334155'>Coming Soon</div>
      <div style='font-size:13px;margin-top:8px'>
        The Manager's Report will be generated here showing weekly performance vs targets and last year.
      </div>
      <a href='/sales?store={store}&week_start={week_start}' class='btn-secondary' style='margin-top:16px;display:inline-block'>
        &#8592; Back to Sales
      </a>
    </div>"""
    return page("Manager's Report", content, user, "sales")

