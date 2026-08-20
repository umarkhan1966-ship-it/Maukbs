"""staff routes."""
import os, io, re, uuid, math, shutil, secrets, hashlib, html
from datetime import datetime, timedelta, date
from fastapi import APIRouter, Request, Form, Cookie, UploadFile, File
from fastapi.responses import (HTMLResponse, RedirectResponse, FileResponse,
                               JSONResponse, StreamingResponse, Response,
                               PlainTextResponse)
from core.db import DB_FILE, db, q
from core.constants import *
from core.security import (hash_password, verify_password,
                           get_session, require_login, user_staff_id)
from core.layout import page
from core.rota_utils import (calc_paid_hours, parse_hours,
                             get_week_start, get_week_dates)
from docx import Document as DocxDocument
from docx.shared import Pt

router = APIRouter()


# ── Access-control helpers ──────────────────────────────────────────────────
# The Staff module holds sensitive personal/pay/tax data. Management routes are
# owner/manager only; self-service {staff_id} routes let a staff member reach
# ONLY their own record (owner/manager may reach anyone).
def _is_mgr(user) -> bool:
    return user.get("role") in ("owner", "manager")


def _require_mgr(user):
    """Bail (RedirectResponse) if the user isn't owner/manager, else None."""
    if not _is_mgr(user):
        return RedirectResponse("/?msg=That+area+is+for+managers+only&msg_type=error", status_code=303)
    return None


def _require_owner(user):
    """Bail (RedirectResponse) if the user isn't the owner, else None. Used for
    the sensitive areas managers must never reach (pay, edit, entitlement)."""
    if user.get("role") != "owner":
        return RedirectResponse("/?msg=That+area+is+owner+only&msg_type=error", status_code=303)
    return None


def _own_staff_id(user):
    """staff_id of the logged-in user's own profile — via the robust users.staff_id
    link (falls back to full-name match for any unlinked legacy account), or None."""
    return user_staff_id(user)


def _gen_username(first, last):
    """Auto 'firstname.lastname' login name (lowercased, de-punctuated), made
    unique — appends a number on the rare identical-full-name clash."""
    base = (re.sub(r"[^a-z0-9]", "", (first or "").lower()) + "." +
            re.sub(r"[^a-z0-9]", "", (last or "").lower())).strip(".") or "staff"
    uname, n = base, 1
    while q("SELECT 1 FROM users WHERE username=?", (uname,), fetch=True):
        n += 1
        uname = f"{base}{n}"
    return uname


def _gen_temp_password(length=10):
    """Readable random temp password (letters+digits), handed over once."""
    import string
    return "".join(secrets.choice(string.ascii_letters + string.digits) for _ in range(length))


def _staff_access_guard(user, staff_id, allow_staff=True):
    """Staff-scoped routes: the owner may access anyone; a MANAGER only staff at
    their OWN store; a staff user only their own record. Bail (RedirectResponse)
    otherwise, else None.

    allow_staff=False marks areas that are NOT staff self-service (documents,
    onboarding) — owner + own-store manager only; a staff user is turned away
    even from their own record."""
    if user.get("role") == "owner":
        return None
    if user.get("role") == "manager":
        row = q("SELECT store_name FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
        if row and row[0]["store_name"] and row[0]["store_name"] == user.get("store_name"):
            return None
        return RedirectResponse("/staff?msg=You+can+only+access+your+own+store&msg_type=error",
                                status_code=303)
    # staff role
    if not allow_staff:
        return RedirectResponse("/my-profile?msg=That+area+is+not+available+in+your+view&msg_type=error",
                                status_code=303)
    if _own_staff_id(user) == staff_id:
        return None
    return RedirectResponse("/my-profile?msg=You+can+only+access+your+own+record&msg_type=error",
                            status_code=303)


def _safe_part(s):
    """Filename-safe token — blocks path traversal via user-supplied name parts."""
    return (re.sub(r"[^A-Za-z0-9_-]", "_", str(s or ""))[:40]) or "x"


def _safe_ext(ext):
    """Whitelist upload extensions; unknown types become .dat (never executable)."""
    e = re.sub(r"[^a-z0-9.]", "", str(ext or "").lower())[:6]
    return e if e in (".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".webp") else ".dat"


def esc(x):
    """HTML-escape a value (guards stored XSS from user-entered fields like
    names/addresses that render into pages other users view)."""
    return html.escape(str(x), quote=True) if x is not None else ""


# One-off / special bank holidays (e.g. Jubilees, Coronations) that no formula
# predicts — add them here by "YYYY-MM-DD" as and when the government announces them.
EXTRA_BANK_HOLIDAYS = set()

_BH_CACHE = {}

def uk_bank_holidays(year: int) -> set:
    """UK (England & Wales) bank holidays for any year, computed — so this never
    needs a manual yearly update. Covers New Year, Good Friday, Easter Monday,
    the early-May/spring/summer Mondays, Christmas and Boxing Day, applying the
    'substitute Monday/Tuesday' rule when a fixed date lands on a weekend."""
    if year in _BH_CACHE:
        return _BH_CACHE[year]
    from datetime import date as _date
    # Easter Sunday via the Anonymous Gregorian algorithm (Computus)
    a = year % 19; b = year // 100; c = year % 100; d = b // 4; e = b % 4
    f = (b + 8) // 25; g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30; i = c // 4; k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7; m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31; day = ((h + l - 7 * m + 114) % 31) + 1
    easter = _date(year, month, day)
    fixed = [
        easter - timedelta(days=2),                       # Good Friday
        easter + timedelta(days=1),                       # Easter Monday
        _date(year, 5, 1)  + timedelta(days=(7 - _date(year, 5, 1).weekday()) % 7),   # first Mon of May
        _date(year, 5, 31) - timedelta(days=_date(year, 5, 31).weekday()),            # last Mon of May
        _date(year, 8, 31) - timedelta(days=_date(year, 8, 31).weekday()),            # last Mon of Aug
    ]
    taken = set(fixed)
    def _subst(d0):
        d0 = d0
        while d0.weekday() >= 5 or d0 in taken:          # roll weekends/clashes to next weekday
            d0 += timedelta(days=1)
        taken.add(d0); return d0
    for fixed_date in (_date(year, 1, 1), _date(year, 12, 25), _date(year, 12, 26)):
        _subst(fixed_date)
    result = {d.strftime("%Y-%m-%d") for d in taken} | {
        x for x in EXTRA_BANK_HOLIDAYS if x.startswith(f"{year}-")}
    _BH_CACHE[year] = result
    return result


def calc_entitlement(contracted_hrs: float) -> float:
    """5.6 weeks × contracted hours per week, including bank holidays."""
    if not contracted_hrs:
        return 0.0
    return round(5.6 * contracted_hrs, 1)  # entitlement in hours


def hrs_to_days(hrs: float, contracted_hrs: float) -> float:
    """Convert hours to days based on contracted daily hours."""
    if not contracted_hrs or contracted_hrs == 0:
        return 0.0
    daily = contracted_hrs / 5
    return round(hrs / daily, 1) if daily > 0 else 0.0


def is_full_time(contracted_hrs: float) -> bool:
    return contracted_hrs is not None and contracted_hrs >= 30.0


def fmt_entitlement(hrs: float, contracted_hrs: float) -> str:
    """Format entitlement as days for FT (>=30h), hours for PT (<30h)."""
    if not contracted_hrs:
        return "—"
    if is_full_time(contracted_hrs):
        days = hrs_to_days(hrs, contracted_hrs)
        return f"{days} days"
    else:
        return f"{hrs} hrs ({hrs_to_days(hrs, contracted_hrs)} days)"


def get_leave_summary(staff_id: int, year: int = None) -> dict:
    """Return entitlement, taken, balance for a staff member."""
    if year is None:
        year = datetime.now().year
    staff = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not staff:
        return {}
    s = dict(staff[0])
    contracted = s.get("contracted_hrs") or 0

    # Check for custom entitlement first
    custom = q("""SELECT effective_days, statutory_days FROM leave_entitlements
                  WHERE staff_id=? AND year=?""", (staff_id, year), fetch=True)
    if custom:
        c = dict(custom[0])
        effective_days   = c["effective_days"] or c["statutory_days"] or 0
        entitlement_hrs  = effective_days * (contracted/5) if contracted else effective_days
    else:
        entitlement_hrs = calc_entitlement(contracted)

    # Days taken this year by type
    daily = contracted / 5 if contracted else 7.5

    # Sum the DAYS in each approved request (days_taken) — NOT the number of
    # requests. A single 5-day holiday is one row with days_taken=5, so counting
    # rows would record it as 1 day. COALESCE handles missing rows/values.
    def _days_taken(ltype):
        r = q("""SELECT COALESCE(SUM(COALESCE(days_taken,1)),0) AS d FROM leave_requests
                 WHERE staff_id=? AND status='approved'
                 AND leave_type=? AND strftime('%Y',date_from)=?""",
              (staff_id, ltype, str(year)), fetch=True)
        d = (r[0]["d"] if r else 0) or 0
        return int(d) if float(d).is_integer() else round(d, 1)   # 5.0 -> 5, keep 0.5

    holiday_days = _days_taken('H')
    taken_hrs    = holiday_days * daily
    bh_days      = _days_taken('B')
    bh_hrs       = bh_days * daily
    sick_days    = _days_taken('S')

    balance_hrs = entitlement_hrs - taken_hrs - bh_hrs
    # Sick days tracked separately — don't affect holiday balance
    daily = contracted / 5 if contracted else 7.5

    ft = is_full_time(contracted)
    return {
        "entitlement_hrs":  entitlement_hrs,
        "entitlement_days": hrs_to_days(entitlement_hrs, contracted),
        "entitlement_fmt":  fmt_entitlement(entitlement_hrs, contracted),
        "taken_hrs":        round(taken_hrs, 1),
        "taken_days":       holiday_days,
        "taken_fmt":        f"{holiday_days:g} days",
        "bh_days":          bh_days,
        "bh_hrs":           round(bh_hrs, 1),
        "sick_days":        sick_days,
        "balance_hrs":      round(balance_hrs, 1),
        "balance_days":     hrs_to_days(balance_hrs, contracted),
        "balance_fmt":      fmt_entitlement(round(balance_hrs,1), contracted),
        "daily_hrs":        round(daily, 2),
        "contracted_hrs":   contracted,
        "is_full_time":     ft,
    }


def ensure_staff_tables():
    conn = db()
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS leave_requests (
            request_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            leave_type    TEXT NOT NULL DEFAULT 'H',
            date_from     TEXT NOT NULL,
            date_to       TEXT NOT NULL,
            days_taken    REAL DEFAULT 1,
            status        TEXT DEFAULT 'pending',
            requested_by  TEXT,
            approved_by   TEXT,
            approved_at   TEXT,
            notes         TEXT,
            created_at    TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS leave_entitlements (
            entitlement_id   INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id         INTEGER NOT NULL,
            year             INTEGER NOT NULL,
            statutory_days   REAL,
            custom_days      REAL,
            effective_days   REAL,
            notes            TEXT,
            UNIQUE(staff_id, year),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS document_templates (
            template_id    INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_type       TEXT NOT NULL,
            version        INTEGER DEFAULT 1,
            file_path      TEXT NOT NULL,
            file_name      TEXT,
            is_current     INTEGER DEFAULT 1,
            uploaded_by    TEXT,
            uploaded_at    TEXT DEFAULT (datetime('now')),
            notes          TEXT,
            UNIQUE(doc_type, version)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_documents (
            doc_id         INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id       INTEGER NOT NULL,
            doc_type       TEXT NOT NULL,
            version        INTEGER DEFAULT 1,
            file_path      TEXT NOT NULL,
            file_name      TEXT,
            is_current     INTEGER DEFAULT 1,
            generated      INTEGER DEFAULT 0,
            uploaded_by    TEXT,
            uploaded_at    TEXT DEFAULT (datetime('now')),
            notes          TEXT,
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS pay_history (
            pay_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            effective_date TEXT NOT NULL,
            pay_basis     TEXT,
            hourly_rate   REAL,
            previous_rate REAL,
            salary_amount REAL,
            previous_salary REAL,
            contracted_hrs REAL,
            previous_hrs  REAL,
            change_reason TEXT,
            recorded_by   TEXT,
            created_at    TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_attendance (
            att_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id      INTEGER NOT NULL,
            work_date     TEXT NOT NULL,
            day           TEXT,
            a_type        TEXT,
            status        TEXT,            -- Worked / Holiday / Sick / Maternity / Bank Holiday
            sched_start   TEXT,
            sched_finish  TEXT,
            clock_in      TEXT,
            clock_out     TEXT,
            hours_worked  REAL,            -- = the sheet's D.Hours (decimal)
            paid_hours    REAL,            -- = the sheet's P.Hours (breaks deducted)
            comments      TEXT,
            source        TEXT,
            UNIQUE(staff_id, work_date),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS staff_rtw_checks (
            rtw_id            INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id          INTEGER NOT NULL,
            id_type           TEXT,            -- e.g. British Passport, Driving Licence, Share code
            rtw_confirmed     INTEGER DEFAULT 0,-- 1 = right to work verified
            check_date        TEXT,            -- date the check was carried out
            expiry_date       TEXT,            -- for time-limited RTW (visa/BRP); blank = no limit
            evidence_location TEXT,            -- where the actual copy is held (e.g. offline, owner's laptop)
            notes             TEXT,
            checked_by        TEXT,
            created_at        TEXT DEFAULT (datetime('now')),
            UNIQUE(staff_id),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS nmw_rates (
            nmw_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            effective_date TEXT NOT NULL,
            rate_21_plus  REAL,
            rate_18_20    REAL,
            rate_16_17    REAL,
            rate_apprentice REAL,
            UNIQUE(effective_date)
        )
    """)
    # Seed NMW rates from historical data
    nmw_data = [
        ("2026-04-01", 12.71, 10.85, 8.00,  7.55),
        ("2025-04-01", 12.21, 10.00, 7.55,  7.55),
        ("2024-04-01", 11.44,  8.60, 6.40,  6.40),
        ("2023-04-01", 10.42,  7.49, 5.28,  5.28),
        ("2022-04-01",  9.50,  6.83, 4.81,  4.81),
        ("2021-04-01",  8.91,  6.56, 4.62,  4.30),
        ("2020-04-01",  8.72,  6.45, 4.55,  4.15),
        ("2019-04-01",  8.21,  6.15, 4.35,  3.90),
        ("2018-04-01",  7.83,  5.90, 4.20,  3.70),
        ("2017-04-01",  7.50,  5.60, 4.05,  3.50),
    ]
    for row in nmw_data:
        try: c.execute("INSERT OR IGNORE INTO nmw_rates (effective_date,rate_21_plus,rate_18_20,rate_16_17,rate_apprentice) VALUES(?,?,?,?,?)", row)
        except: pass

    # ── Migration: two-stage leave approval (manager recommends → owner signs off) ──
    _lr_cols = {r[1] for r in c.execute("PRAGMA table_info(leave_requests)")}
    for _name, _ddl in [("mgr_approved_by", "mgr_approved_by TEXT"),
                        ("mgr_approved_at", "mgr_approved_at TEXT")]:
        if _name not in _lr_cols:
            c.execute(f"ALTER TABLE leave_requests ADD COLUMN {_ddl}")

    conn.commit()
    conn.close()


@router.get("/staff", response_class=HTMLResponse)
def staff_page(
    session:  str | None = Cookie(default=None),
    store:    str = "",
    show:     str = "active",
    msg:      str = "",
    msg_type: str = "success"
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r
    is_owner = user["role"] == "owner"

    # Build filter
    conds  = []
    params = []
    if show == "active":
        conds.append("is_active = 1")
    elif show == "leavers":
        if not is_owner:
            return RedirectResponse("/staff", status_code=303)
        conds.append("is_active = 0")
    # Store filter
    if store:
        conds.append("store_name = ?")
        params.append(store)
    elif user["role"] == "manager" and user.get("store_name"):
        conds.append("store_name = ?")
        params.append(user["store_name"])

    where = ("WHERE " + " AND ".join(conds)) if conds else ""
    staff = q(f"SELECT * FROM staff_profiles {where} ORDER BY store_name, first_name",
              params, fetch=True) or []

    # Leave-request badge — count only what THIS user needs to action:
    #   owner   → awaiting final sign-off (mgr_approved) + any unreviewed pending
    #   manager → own-store requests still pending their approval
    if is_owner:
        _cnt = q("""SELECT COUNT(*) n FROM leave_requests
                    WHERE status IN ('pending','mgr_approved')""", fetch=True)
    else:
        _cnt = q("""SELECT COUNT(*) n FROM leave_requests lr
                    JOIN staff_profiles sp ON lr.staff_id=sp.staff_id
                    WHERE lr.status='pending' AND sp.store_name=?""",
                 (user.get("store_name") or "",), fetch=True)
    pending_n = _cnt[0]["n"] if _cnt else 0

    flash = f"<div class='flash-{'success' if msg_type=='success' else 'error'}'>{msg}</div>" if msg else ""

    # ── Tab bar ──
    tabs = ""
    for val, label in [("active","Active Staff"),("leavers","Former Staff (Leavers)")]:
        if val == "leavers" and not is_owner:
            continue
        active_cls = "border-b-2 border-blue-900 font-black text-blue-900" if show == val else "text-slate-500 hover:text-slate-700"
        tabs += f"<a href='/staff?show={val}' class='px-4 py-2 text-sm {active_cls} transition'>{label}</a>"

    # ── Store filter buttons (owner only — a manager has just their one store) ──
    store_btns = ""
    if is_owner:
        for sv, sl in [("","Both Stores"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]:
            cls = "btn-primary" if store == sv else "btn-secondary"
            store_btns += f"<a href='/staff?show={show}&store={sv}' class='{cls}' style='padding:6px 14px;font-size:13px'>{sl}</a>"

    # ── Staff cards ──
    cards_html = ""
    year = datetime.now().year
    # Show a store header before each store's block when viewing both stores,
    # so the store grouping is obvious (cards flow across a multi-column grid).
    group_by_store = not store
    store_counts = {}
    for x in staff:
        k = (dict(x).get("store_name") or "—")
        store_counts[k] = store_counts.get(k, 0) + 1
    last_store = None
    for s in staff:
        s = dict(s)
        sid   = s["staff_id"]
        this_store = s.get("store_name") or "—"
        if group_by_store and this_store != last_store:
            cards_html += (
                f"<div style='grid-column:1/-1;display:flex;align-items:center;gap:8px;"
                f"margin-top:6px;padding:6px 4px;border-bottom:2px solid #cbd5e1'>"
                f"<span style='font-size:15px'>&#128205;</span>"
                f"<span style='font-weight:900;font-size:14px;color:#0f172a'>{esc(this_store)}</span>"
                f"<span style='font-size:12px;color:#94a3b8;font-weight:700'>({store_counts[this_store]})</span></div>")
            last_store = this_store
        name  = esc(f"{s['first_name']} {s['last_name']}")
        store_badge = f"<span style='background:#e0f2fe;color:#0369a1;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>{s.get('store_name','')}</span>"
        status_badge = "<span class='badge-paid'>Active</span>" if s["is_active"] else "<span class='badge-overdue'>Left</span>"
        # Day-one statutory: flag active staff with no current Employment Contract on file.
        contract_badge = ""
        if s["is_active"] and not q("SELECT 1 FROM staff_documents WHERE staff_id=? AND doc_type='Employment Contract' AND is_current=1 AND deleted_at IS NULL LIMIT 1", (sid,), fetch=True):
            contract_badge = "<span style='background:#fef2f2;color:#b91c1c;border:1px solid #fecaca;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>&#9888;&#65039; No contract</span>"
        # Legal check: flag active staff whose Right to Work hasn't been verified yet.
        rtw_badge = ""
        if s["is_active"] and not q("SELECT 1 FROM staff_rtw_checks WHERE staff_id=? AND rtw_confirmed=1 LIMIT 1", (sid,), fetch=True):
            rtw_badge = "<span style='background:#fffbeb;color:#b45309;border:1px solid #fde68a;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>&#9888;&#65039; RTW not verified</span>"

        # Quick leave summary — from ATTENDANCE (same source as the profile cards),
        # so the list and the profile always agree.
        leave = attendance_leave(sid, year)
        bal      = leave.get("balance", 0)
        bal_fmt  = leave.get("balance_fmt", "—")
        tak_fmt  = leave.get("taken_fmt", "—")
        ent_fmt  = leave.get("entitlement_fmt", "—")
        bal_col  = "#16a34a" if bal > 5 else ("#d97706" if bal >= 0 else "#dc2626")

        rate = f"£{s['hourly_rate']:.2f}/hr" if s.get("hourly_rate") else "—"
        hrs  = f"{s['contracted_hrs']}h/wk" if s.get("contracted_hrs") else "—"
        joined = s.get("date_joined") or "—"
        # Pay rate + edit are owner-only on the list card.
        rate_cell = (f"<div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Rate</div>"
                     f"<div style='font-size:13px;font-weight:600;color:#334155'>{rate}</div></div>") if is_owner else ""
        edit_btn = f"<a href='/staff/{sid}/edit' class='btn-secondary' style='padding:5px 12px;font-size:12px'>&#9999;&#65039; Edit</a>" if is_owner else ""

        cards_html += f"""
        <div class='card' style='padding:0;overflow:hidden'>
          <div style='background:#f8fafc;padding:12px 16px;display:flex;justify-content:space-between;align-items:center;border-bottom:1px solid #e2e8f0'>
            <div>
              <div style='font-weight:900;font-size:15px;color:#0f172a'>{name}</div>
              <div style='display:flex;gap:6px;margin-top:4px;flex-wrap:wrap'>{store_badge} {status_badge} {contract_badge} {rtw_badge}</div>
            </div>
            <div style='display:flex;gap:8px'>
              <a href='/staff/{sid}' class='btn-primary' style='padding:5px 12px;font-size:12px'>👁 View</a>
              {edit_btn}
            </div>
          </div>
          <div style='padding:12px 16px;display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:8px'>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Joined</div>
                 <div style='font-size:13px;font-weight:600;color:#334155'>{joined}</div></div>
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Hours</div>
                 <div style='font-size:13px;font-weight:600;color:#334155'>{hrs}</div></div>
            {rate_cell}
            <div><div style='font-size:11px;color:#94a3b8;font-weight:700;text-transform:uppercase'>Leave Balance {year}</div>
                 <div style='font-size:13px;font-weight:700;color:{bal_col}'>{bal_fmt} left <span style='color:#94a3b8;font-weight:400'>({tak_fmt} of {ent_fmt})</span></div></div>
          </div>
        </div>"""

    if not cards_html:
        cards_html = "<div class='card text-center' style='padding:40px;color:#94a3b8'>No staff found</div>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>👤 Staff</div>
      <div style='display:flex;gap:8px;flex-wrap:wrap'>
        {'<a href="/staff/leave-requests" class="btn-secondary" style="position:relative">📋 Leave Requests' + (f'<span style="position:absolute;top:-6px;right:-6px;background:#dc2626;color:white;border-radius:50%;width:18px;height:18px;font-size:10px;font-weight:900;display:flex;align-items:center;justify-content:center">{pending_n}</span>' if pending_n > 0 else '') + '</a>' if is_owner or user["role"]=="manager" else ''}
        <a href='/staff/leave-planner' class='btn-secondary'>📅 Leave Planner</a>
        {'<a href="/staff/pay-overview" class="btn-secondary">💰 Pay Overview</a>' if is_owner else ''}
        {'<a href="/staff/document-templates" class="btn-secondary">📋 Doc Templates</a>' if is_owner else ''}
        {'<a href="/staff/new" class="btn-primary">➕ Add Staff Member</a>' if is_owner or user["role"]=="manager" else ''}
      </div>
    </div>
    <div style='display:flex;gap:0;border-bottom:1px solid #e2e8f0'>{tabs}</div>
    <div style='display:flex;gap:8px;flex-wrap:wrap'>{store_btns}</div>
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(420px,1fr))'>
      {cards_html}
    </div>"""

    return page("Staff", content, user, "staff")


@router.get("/staff/document-templates", response_class=HTMLResponse)
def document_templates(session: str | None = Cookie(default=None), msg: str = ""):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    templates = q("SELECT * FROM document_templates ORDER BY doc_type, version DESC",
                  fetch=True) or []

    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    from collections import defaultdict
    by_type = defaultdict(list)
    for t in templates:
        by_type[dict(t)["doc_type"]].append(dict(t))

    tmpl_html = ""
    for dtype in DOC_TYPES:
        type_tmpls = by_type.get(dtype, [])
        current    = next((t for t in type_tmpls if t["is_current"]), None)
        older      = [t for t in type_tmpls if not t["is_current"]]

        current_html = ""
        if current:
            current_html = f"""
            <div style='background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:12px 14px'>
              <div style='font-size:13px;font-weight:700;color:#166534;margin-bottom:4px'>
                ✅ Current — v{current['version']} uploaded {current['uploaded_at'][:10]}
              </div>
              <div style='font-size:11px;color:#64748b;margin-bottom:10px'>{current.get('notes') or ''}</div>
              <div style='display:flex;gap:12px;align-items:center'>
                <a href='/staff/document-templates/{current["template_id"]}/download'
                   style='color:#64748b;font-size:12px;text-decoration:underline'>⬇️ download template</a>
                <a href='/staff/document-templates/{current["template_id"]}/delete'
                   onclick='return confirm("Delete this version? The previous version will become current.")'
                   class='btn-danger' style='padding:5px 14px;font-size:12px'>🗑️ Delete</a>
              </div>
            </div>"""

        older_html = "".join(
            f"<div style='font-size:12px;color:#94a3b8;padding:4px 10px'>v{t['version']} — {t['uploaded_at'][:10]} (superseded)</div>"
            for t in older
        )

        tmpl_html += f"""
        <div class='card'>
          <div style='font-weight:900;color:#0f2942;margin-bottom:8px'>{dtype}</div>
          {current_html or "<div style='color:#94a3b8;font-size:13px;padding:8px 0'>No template uploaded yet</div>"}
          {older_html}
          <form action='/staff/document-templates/upload' method='POST'
                enctype='multipart/form-data'
                style='margin-top:12px;padding-top:12px;border-top:1px solid #f1f5f9'
                onsubmit='showUploading(this)'>
            <input type='hidden' name='doc_type' value='{dtype}'>
            <div style='margin-bottom:8px'>
              <label style='font-size:11px;font-weight:700;color:#64748b;
                            text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px'>
                Select Template File (.docx or .dotx)
              </label>
              <input type='file' name='template_file' accept='.docx,.dotx' required
                     style='width:100%;border:1px solid #e2e8f0;border-radius:8px;
                            padding:8px 10px;font-size:13px;background:white;cursor:pointer'
                     onchange='previewFile(this)'>
              <div id='preview_{dtype.replace(" ","_")}' style='font-size:12px;color:#16a34a;
                   font-weight:700;margin-top:4px;display:none'>
                ✅ Selected: <span class='filename'></span>
              </div>
            </div>
            <div style='margin-bottom:8px'>
              <label style='font-size:11px;font-weight:700;color:#64748b;
                            text-transform:uppercase;letter-spacing:.05em;display:block;margin-bottom:4px'>
                Version Notes
              </label>
              <input type='text' name='notes' id='notes_{dtype.replace(" ","_")}'
                     placeholder='e.g. Updated Nov 2024 — new holiday clause'
                     style='width:100%;border:1px solid #e2e8f0;border-radius:8px;
                            padding:8px 10px;font-size:13px'>
            </div>
            <button type='submit' class='btn-primary' style='width:100%;padding:8px;font-size:13px'>
              ⬆️ Upload New Version
            </button>
          </form>
        </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>📋 Document Templates</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>
    <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
      <div style='font-size:13px;font-weight:700;color:#92400e'>📝 How to set up templates</div>
      <div style='font-size:13px;color:#78350f;margin-top:4px'>
        Create a Word (.docx) document with your letter/contract content.
        Use these placeholders where you want staff details inserted:
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{FULL_NAME}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{STORE}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{DATE_JOINED}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{HOURLY_RATE}}}}</code>
        <code style='background:#fff;padding:2px 6px;border-radius:4px;margin:0 4px'>{{{{TODAY}}}}</code>
        and more. See the full list when generating a document.
      </div>
    </div>
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(380px,1fr))'>
      {tmpl_html}
    </div>
    <script>
    function previewFile(input) {{
      if (!input.files.length) return;
      const fname = input.files[0].name;
      const form  = input.closest('form');
      // Show selected filename
      const preview = form.querySelector('[id^="preview_"]');
      if (preview) {{
        preview.querySelector('.filename').textContent = fname;
        preview.style.display = 'block';
      }}
      // Auto-fill notes with filename + date if empty
      const notes = form.querySelector('[id^="notes_"]');
      if (notes && !notes.value) {{
        const today = new Date().toLocaleDateString('en-GB', {{day:'2-digit',month:'short',year:'numeric'}});
        notes.value = fname + ' — uploaded ' + today;
      }}
    }}
    function showUploading(form) {{
      const btn = form.querySelector('button[type="submit"]');
      if (btn) {{ btn.textContent = '⏳ Uploading...'; btn.disabled = true; }}
    }}
    </script>"""

    return page("Document Templates", content, user, "staff")


@router.post("/staff/document-templates/upload")
async def upload_template(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff/document-templates", status_code=303)

    form     = await request.form()
    doc_type = form.get("doc_type","")
    tmpl     = form.get("template_file")
    notes    = str(form.get("notes","") or "").strip()

    if not tmpl or not hasattr(tmpl,"filename") or not tmpl.filename:
        return RedirectResponse("/staff/document-templates?msg=No+file+selected", status_code=303)

    existing = q("SELECT MAX(version) as v FROM document_templates WHERE doc_type=?",
                 (doc_type,), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1

    q("UPDATE document_templates SET is_current=0 WHERE doc_type=?", (doc_type,))

    filename = f"template_{_safe_part(doc_type)}_v{next_ver}.docx"
    filepath = os.path.join(TEMPLATES_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(await tmpl.read())

    q("""INSERT INTO document_templates
            (doc_type, version, file_path, file_name, is_current, uploaded_by, notes)
         VALUES(?,?,?,?,1,?,?)""",
      (doc_type, next_ver, filepath, tmpl.filename, user.get("username"), notes or None))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/document-templates?msg={uq(doc_type + ' template uploaded (v' + str(next_ver) + ')')}",
        status_code=303)


@router.get("/staff/document-templates/{template_id}/delete")
def delete_template(template_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff/document-templates", status_code=303)
    # Get the doc_type before deleting
    rows = q("SELECT * FROM document_templates WHERE template_id=?", (template_id,), fetch=True)
    if rows:
        t = dict(rows[0])
        # Delete the file from disk
        if os.path.exists(t["file_path"]):
            os.remove(t["file_path"])
        # Delete from database
        q("DELETE FROM document_templates WHERE template_id=?", (template_id,))
        # If this was current, make the previous version current
        q("""UPDATE document_templates SET is_current=1
             WHERE doc_type=? AND template_id=(
                SELECT MAX(template_id) FROM document_templates WHERE doc_type=?
             )""", (t["doc_type"], t["doc_type"]))
    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/document-templates?msg={uq('Template version deleted')}",
        status_code=303)


@router.get("/staff/document-templates/{template_id}/download")
def download_template(template_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r
    rows = q("SELECT * FROM document_templates WHERE template_id=?", (template_id,), fetch=True)
    if not rows: return HTMLResponse("<p>Not found</p>", status_code=404)
    t = dict(rows[0])
    if not os.path.exists(t["file_path"]):
        return HTMLResponse("<p>File not found</p>", status_code=404)
    return FileResponse(t["file_path"], filename=t["file_name"] or os.path.basename(t["file_path"]),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/staff/new", response_class=HTMLResponse)
def new_staff_form(session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse("/staff", status_code=303)
    return render_staff_form(user, None)


@router.get("/staff/leave-requests", response_class=HTMLResponse)
def leave_requests(session: str | None = Cookie(default=None), msg: str = "", msg_type: str = "success"):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse("/staff", status_code=303)

    is_owner = user["role"] == "owner"
    # Managers only ever see their own store; the owner sees every store.
    scope, sp = ("", [])
    if not is_owner:
        scope = "AND sp.store_name = ?"
        sp    = [user.get("store_name") or ""]

    def _fetch(where, params, order):
        return q(f"""SELECT lr.*, sp.first_name, sp.last_name, sp.store_name
                     FROM leave_requests lr
                     JOIN staff_profiles sp ON lr.staff_id=sp.staff_id
                     WHERE {where} {scope} ORDER BY {order}""",
                 params + sp, fetch=True) or []

    def req_row(lr, action="none", cols=9):
        lr    = dict(lr)
        name  = f"{lr['first_name']} {lr['last_name']}"
        ltype = ABSENCE_TYPES.get(lr['leave_type'], lr['leave_type'])
        badge = {"approved":"<span class='badge-paid'>Approved</span>",
                 "pending": "<span class='badge-partial'>Pending manager</span>",
                 "mgr_approved":"<span class='badge-partial' style='background:#e0e7ff;color:#3730a3'>Manager OK · awaiting owner</span>",
                 "declined":"<span class='badge-overdue'>Declined</span>"}.get(lr["status"],"")
        # action: "approve" shows Approve/Decline; "signoff" is the owner's final
        # sign-off wording; "none" is read-only.
        actions = ""
        if action in ("approve", "signoff"):
            label = "🖊️ Sign off" if action == "signoff" else "✅ Approve"
            actions = f"""
            <form method='POST' action='/staff/leave-requests/{lr['request_id']}/approve' style='display:inline'>
              <button type='submit' class='btn-success' style='padding:4px 10px;font-size:11px'>{label}</button></form>
            <form method='POST' action='/staff/leave-requests/{lr['request_id']}/decline' style='display:inline'
                  onsubmit="return confirm('Decline this leave request?');">
              <button type='submit' class='btn-danger' style='padding:4px 10px;font-size:11px'>❌ Decline</button></form>"""
        action_cell = f"<td><div style='display:flex;gap:4px'>{actions}</div></td>" if action != "none" else "<td></td>"
        return f"""<tr>
          <td style='font-weight:700'>{name}</td>
          <td style='font-size:12px;color:#64748b'>{lr.get('store_name','')}</td>
          <td>{ltype}</td>
          <td class='mono'>{lr['date_from']}</td>
          <td class='mono'>{lr['date_to']}</td>
          <td class='mono'>{lr['days_taken']}</td>
          <td>{badge}</td>
          <td style='font-size:12px;color:#64748b'>{lr.get('notes') or '—'}</td>
          {action_cell}
        </tr>"""

    def _table(title, colour, rows_html, empty):
        return f"""
        <div class='card' style='padding:0;overflow:hidden'>
          <div style='padding:12px 16px;background:{colour};color:white;font-weight:700;font-size:14px'>{title}</div>
          <div style='overflow-x:auto'>
            <table class='tbl'>
              <thead><tr><th>Staff</th><th>Store</th><th>Type</th><th>From</th><th>To</th><th>Days</th><th>Status</th><th>Notes</th><th>Action</th></tr></thead>
              <tbody>{rows_html or f'<tr><td colspan="9" style="text-align:center;padding:24px;color:#94a3b8">{empty}</td></tr>'}</tbody>
            </table>
          </div>
        </div>"""

    recent = _fetch("lr.status IN ('approved','declined')", [], "lr.created_at DESC LIMIT 20")
    recent_html = "".join(req_row(lr, "none") for lr in recent)

    if is_owner:
        # The owner's sign-off queue first, then anything staff submitted that no
        # manager has looked at yet (the owner can act on those directly too).
        signoff = _fetch("lr.status='mgr_approved'", [], "lr.date_from ASC")
        newpend = _fetch("lr.status='pending'", [], "lr.date_from ASC")
        blocks  = (
            _table(f"🖊️ Awaiting your sign-off ({len(signoff)})", "#7c3aed",
                   "".join(req_row(lr, "signoff") for lr in signoff),
                   "Nothing awaiting sign-off") +
            _table(f"⏳ Not yet seen by a manager ({len(newpend)})", "#d97706",
                   "".join(req_row(lr, "signoff") for lr in newpend),
                   "No unreviewed requests") +
            _table("Recent Decisions", "#0f2942", recent_html, "No recent decisions"))
    else:
        # Manager: their approval queue, then what they've passed up to the owner.
        needs = _fetch("lr.status='pending'", [], "lr.date_from ASC")
        upped = _fetch("lr.status='mgr_approved'", [], "lr.date_from ASC")
        blocks = (
            _table(f"⏳ Needs your approval ({len(needs)})", "#d97706",
                   "".join(req_row(lr, "approve") for lr in needs),
                   "No requests awaiting you") +
            _table(f"🖊️ Awaiting owner sign-off ({len(upped)})", "#7c3aed",
                   "".join(req_row(lr, "none") for lr in upped),
                   "Nothing awaiting the owner") +
            _table("Recent Decisions", "#0f2942", recent_html, "No recent decisions"))

    flash = ""
    if msg:
        cls = "flash-success" if msg_type == "success" else "flash-error"
        flash = f"<div class='{cls}'>{esc(msg)}</div>"

    content = f"""
    {flash}
    <div class='flex justify-between items-center'>
      <div class='text-2xl font-black text-slate-800'>📋 Leave Requests</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>
    {blocks}"""
    return page("Leave Requests", content, user, "staff")


@router.get("/staff/leave-planner", response_class=HTMLResponse)
def leave_planner(
    session: str | None = Cookie(default=None),
    year:    int = 0,
    store:   str = ""
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r
    if not year: year = datetime.now().year

    # Get store filter — a manager is FORCED to their own store (ignore any
    # ?store= override), the owner may pick any.
    if user["role"] == "manager":
        store = user.get("store_name") or ""

    # Get active staff
    conds  = ["is_active=1"]
    params = []
    if store:
        conds.append("store_name=?")
        params.append(store)
    staff = q(f"SELECT * FROM staff_profiles WHERE {' AND '.join(conds)} ORDER BY first_name",
              params, fetch=True) or []

    # Get all approved leave for this year
    leave_data = q("""
        SELECT lr.staff_id, lr.date_from, lr.date_to, lr.leave_type
        FROM leave_requests lr
        WHERE lr.status='approved'
          AND strftime('%Y',lr.date_from)=?
    """, (str(year),), fetch=True) or []

    # Build a set of (staff_id, date) → leave_type
    leave_map = {}
    for lr in leave_data:
        lr = dict(lr)
        try:
            d1 = datetime.strptime(lr["date_from"], "%Y-%m-%d")
            d2 = datetime.strptime(lr["date_to"],   "%Y-%m-%d")
            cur = d1
            while cur <= d2:
                leave_map[(lr["staff_id"], cur.strftime("%Y-%m-%d"))] = lr["leave_type"]
                cur = cur + timedelta(days=1)
        except Exception: pass

    # BH set (computed for the year being viewed)
    bh_set = uk_bank_holidays(year)

    months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    import calendar

    # Build the planner grid
    grid_html = ""
    for mi, mname in enumerate(months, 1):
        _, days_in_month = calendar.monthrange(year, mi)
        # Header row
        grid_html += f"<tr><td style='background:#0f2942;color:white;font-weight:900;font-size:12px;padding:6px 10px;white-space:nowrap'>{mname}</td>"
        for d in range(1, 32):
            if d > days_in_month:
                grid_html += "<td style='background:#f8fafc'></td>"
                continue
            date_str = f"{year}-{mi:02d}-{d:02d}"
            dow = datetime(year, mi, d).weekday()
            is_weekend = dow >= 5
            is_bh = date_str in bh_set
            if is_bh:
                bg = "#fef3c7"; txt = "<span style='font-size:9px;color:#92400e;font-weight:700'>BH</span>"
            elif is_weekend:
                bg = "#f1f5f9"; txt = ""
            else:
                bg = "white"; txt = ""
            grid_html += f"<td style='background:{bg};border:1px solid #e2e8f0;text-align:center;padding:2px;min-width:28px;font-size:10px'>{txt}</td>"
        grid_html += "</tr>"

        # Staff rows for this month
        for s in staff:
            s = dict(s)
            sid   = s["staff_id"]
            # First name or nickname (first 3 chars)
            initials = s["first_name"][:3]
            grid_html += f"<tr><td style='font-size:11px;font-weight:700;color:#334155;padding:3px 10px;white-space:nowrap;border-bottom:1px solid #f1f5f9'>{initials}</td>"
            for d in range(1, 32):
                if d > days_in_month:
                    grid_html += "<td style='background:#f8fafc'></td>"
                    continue
                date_str = f"{year}-{mi:02d}-{d:02d}"
                dow = datetime(year, mi, d).weekday()
                is_weekend = dow >= 5
                ltype = leave_map.get((sid, date_str))
                if ltype == "H":
                    bg = "#dcfce7"; cell = f"<span style='font-size:9px;font-weight:900;color:#166534'>{initials}</span>"
                elif ltype == "S":
                    bg = "#fee2e2"; cell = f"<span style='font-size:9px;font-weight:900;color:#991b1b'>S</span>"
                elif ltype == "B":
                    bg = "#fef3c7"; cell = ""
                elif is_weekend:
                    bg = "#f1f5f9"; cell = ""
                else:
                    bg = "white"; cell = ""
                grid_html += f"<td style='background:{bg};border:1px solid #f1f5f9;text-align:center;padding:1px;font-size:9px'>{cell}</td>"
            grid_html += "</tr>"
        # Spacer between months
        grid_html += f"<tr><td colspan='32' style='height:4px;background:#f8fafc'></td></tr>"

    # Store filter buttons (owner only — a manager has just their one store)
    store_btns = ""
    if user["role"] == "owner":
        for sv,sl in [("","Both"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]:
            cls = "btn-primary" if store==sv else "btn-secondary"
            store_btns += f"<a href='/staff/leave-planner?year={year}&store={sv}' class='{cls}' style='padding:5px 12px;font-size:12px'>{sl}</a>"

    # Legend
    legend = """
    <div style='display:flex;gap:12px;flex-wrap:wrap;font-size:12px;font-weight:600'>
      <span><span style='display:inline-block;width:14px;height:14px;background:#dcfce7;border:1px solid #86efac;border-radius:3px;vertical-align:middle'></span> Holiday</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#fee2e2;border:1px solid #fca5a5;border-radius:3px;vertical-align:middle'></span> Sick</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#fef3c7;border:1px solid #fcd34d;border-radius:3px;vertical-align:middle'></span> Bank Holiday</span>
      <span><span style='display:inline-block;width:14px;height:14px;background:#f1f5f9;border:1px solid #e2e8f0;border-radius:3px;vertical-align:middle'></span> Weekend</span>
    </div>"""

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>📅 Leave Planner {year}</div>
      <div style='display:flex;gap:8px;align-items:center;flex-wrap:wrap'>
        {store_btns}
        <a href='/staff/leave-planner?year={year-1}&store={store}' class='btn-secondary' style='padding:5px 12px;font-size:12px'>← {year-1}</a>
        <a href='/staff/leave-planner?year={year+1}&store={store}' class='btn-secondary' style='padding:5px 12px;font-size:12px'>{year+1} →</a>
        <a href='/staff' class='btn-secondary' style='padding:5px 12px;font-size:12px'>← Staff</a>
      </div>
    </div>
    {legend}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table style='border-collapse:collapse;font-family:DM Mono,monospace;width:100%'>
          <!-- Day numbers header -->
          <tr>
            <td style='background:#0f2942;color:white;font-size:11px;font-weight:700;padding:6px 10px;white-space:nowrap'>Month / Day</td>
            {"".join(f"<td style='background:#0f2942;color:white;font-size:10px;font-weight:700;text-align:center;padding:3px;min-width:28px'>{d}</td>" for d in range(1,32))}
          </tr>
          {grid_html}
        </table>
      </div>
    </div>"""
    return page("Leave Planner", content, user, "staff")


def get_nmw_for_person(dob_str: str, check_date: str = None) -> float:
    """Return current NMW rate for a person based on their age."""
    if not dob_str: return 0.0
    if not check_date: check_date = datetime.now().strftime("%Y-%m-%d")
    try:
        dob  = datetime.strptime(dob_str, "%Y-%m-%d")
        chk  = datetime.strptime(check_date, "%Y-%m-%d")
        # exact age: subtract a year if this year's birthday hasn't happened yet
        # (days//365 drifts with leap years and can tip someone into the next
        #  NMW bracket a few days early — the 21 boundary is a big rate jump)
        age  = chk.year - dob.year - ((chk.month, chk.day) < (dob.month, dob.day))
        # Get most recent NMW rates effective on or before check_date
        rates = q("""SELECT * FROM nmw_rates WHERE effective_date <= ?
                     ORDER BY effective_date DESC LIMIT 1""",
                  (check_date,), fetch=True)
        if not rates: return 0.0
        r = dict(rates[0])
        if age >= 21: return r["rate_21_plus"]
        elif age >= 18: return r["rate_18_20"]
        else: return r["rate_16_17"]
    except: return 0.0


@router.get("/staff/pay-overview", response_class=HTMLResponse)
def pay_overview(session: str | None = Cookie(default=None)):
    """Owner-only overview of all staff pay vs NMW."""
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    today = datetime.now().strftime("%Y-%m-%d")
    staff = q("SELECT * FROM staff_profiles WHERE is_active=1 ORDER BY store_name, first_name",
              fetch=True) or []

    rows_html = ""
    warnings  = 0
    for s in staff:
        s       = dict(s)
        current = s.get("hourly_rate") or 0
        nmw     = get_nmw_for_person(s.get("date_of_birth",""), today)
        diff    = round(current - nmw, 2) if nmw else 0
        annual  = current * (s.get("contracted_hrs") or 0) * 52

        if nmw == 0:
            status = "<span class='badge-unpaid'>No DOB</span>"
        elif diff < 0:
            status = f"<span class='badge-overdue'>⚠️ £{abs(diff):.2f} BELOW</span>"
            warnings += 1
        elif diff < 0.50:
            status = f"<span class='badge-partial'>Near NMW +£{diff:.2f}</span>"
        else:
            status = f"<span class='badge-paid'>✅ +£{diff:.2f}</span>"

        dob = s.get("date_of_birth","")
        if dob:
            _d = datetime.strptime(dob, "%Y-%m-%d"); _t = datetime.now()
            age = _t.year - _d.year - ((_t.month, _t.day) < (_d.month, _d.day))
        else:
            age = "?"
        rows_html += f"""<tr>
          <td style='font-weight:700'>{esc(s['first_name'])} {esc(s['last_name'])}</td>
          <td style='font-size:12px;color:#64748b'>{esc(s.get('store_name',''))}</td>
          <td>{age}</td>
          <td class='mono' style='font-weight:700'>£{current:.2f}</td>
          <td class='mono' style='color:#64748b'>£{nmw:.2f}</td>
          <td>{status}</td>
          <td class='mono'>£{annual:,.0f}</td>
          <td><a href='/staff/{s["staff_id"]}/pay-history' class='btn-secondary' style='padding:3px 10px;font-size:11px'>History</a></td>
        </tr>"""

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div class='text-2xl font-black text-slate-800'>💰 Pay Overview — All Staff</div>
      <a href='/staff' class='btn-secondary'>← Back to Staff</a>
    </div>
    {'<div class="flash-error">⚠️ ' + str(warnings) + ' staff member(s) may be below National Minimum Wage — please review</div>' if warnings else ''}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Name</th><th>Store</th><th>Age</th><th>Current Rate</th><th>NMW</th><th>Status</th><th>Annual Equiv.</th><th></th></tr></thead>
          <tbody>{rows_html}</tbody>
        </table>
      </div>
    </div>"""
    return page("Pay Overview", content, user, "staff")


from core.paths import data_path
DOCS_DIR      = data_path("staff_docs")   # generated docs → persistent volume


TEMPLATES_DIR = "doc_templates"           # shipped assets → stay next to the code


DOC_TYPES = [
    "Offer Letter",
    "Employment Contract",
    "Right to Work",
    "P45/P46",
    "New Employee Notification",
    "Application / CV",
    "Pay & Role Changes",
    "Return-to-Work Interview",
    "Self-Certification",
    "Disciplinary & Warnings",
    "DBS Check",
    "Other",
]

# Categories too sensitive for a staff member's own view — owner/manager only.
# (Staff logins aren't live yet; this makes the app correct for when they are.)
RESTRICTED_DOC_TYPES = {"Pay & Role Changes", "Disciplinary & Warnings"}

# Managers may CAPTURE (upload) only these — the sick-return + onboarding forms
# they complete WITH the staff member. They can never open/download any document
# (that's owner-only); after upload they just see a "✓ on file" tick.
MANAGER_CAPTURE_DOC_TYPES = {
    "Return-to-Work Interview", "Self-Certification",
    "Application / CV", "Right to Work", "New Employee Notification", "P45/P46",
}


def get_store_entity(store_name: str) -> dict:
    """The legal entity a store trades as, read from the company_entities table
    (never hardcoded) so a future change — e.g. the Uxbridge LLP being replaced
    by its Ltd partner — is a one-row edit that every generated document picks up."""
    rows = q("SELECT * FROM company_entities WHERE store_name=?", (store_name or "",), fetch=True)
    if rows:
        return dict(rows[0])
    # Fallback keeps generation working even if a store has no entity row yet
    return {"store_name": store_name, "legal_name": store_name or "the Company",
            "trading_name": "", "addr_line1": "", "addr_line2": "",
            "addr_line3": "", "addr_line4": ""}


def employer_options(selected_val: str = "") -> str:
    """Employer <option>s for onboarding dropdowns, built from company_entities
    (retail stores) so the names match everywhere and a future company appears
    automatically. Values/labels are HTML-escaped."""
    rows = q("""SELECT legal_name, trading_name, store_name,
                       addr_line1, addr_line2, addr_line3, addr_line4
                FROM company_entities WHERE kind='retail' ORDER BY store_name""", fetch=True) or []
    out = []
    for r in rows:
        r = dict(r)
        addr    = ", ".join(x for x in [r["addr_line1"], r["addr_line2"], r["addr_line3"], r["addr_line4"]] if x)
        trading = r["trading_name"] or ""
        value   = f"{r['legal_name']} T/A {trading}, {addr}" if trading else f"{r['legal_name']}, {addr}"
        label   = f"{trading or r['legal_name']} — {r['addr_line1']}, {r['store_name']} {r['addr_line4']}"
        # match a previously-saved value by store or by the first word of the legal name
        sel = "selected" if selected_val and (
                r["store_name"] in selected_val
                or r["legal_name"].split()[0] in selected_val) else ""
        out.append(f"<option value='{esc(value)}' {sel}>{esc(label)}</option>")
    return "".join(out)


def get_merge_fields(staff: dict) -> dict:
    """Return all merge fields for Word template substitution.
    Supports both <<field>> (your existing format) and {{FIELD}} formats.
    """
    today    = datetime.now().strftime("%d %B %Y")

    def _fmt(d):   # ISO date -> "11 September 2023" for generated documents
        if not d:
            return ''
        try:
            return datetime.strptime(str(d)[:10], "%Y-%m-%d").strftime("%d %B %Y")
        except Exception:
            return str(d)

    name     = f"{staff.get('first_name','')} {staff.get('last_name','')}".strip()
    store    = staff.get('store_name','')
    ent      = get_store_entity(store)
    employer      = (f"{ent['legal_name']} T/A {ent['trading_name']}"
                     if ent.get('trading_name') else ent['legal_name'])
    store_lines   = [ent.get('addr_line1'), ent.get('addr_line2'),
                     ent.get('addr_line3'), ent.get('addr_line4')]
    store_addr    = ", ".join(x for x in store_lines if x)   # postal address, one line
    employer_addr = f"{employer}, {store_addr}" if store_addr else employer
    hrs      = staff.get('contracted_hrs') or 0
    try:                                   # show "20 hours" not "20.0 hours"
        hrs_txt = int(hrs) if float(hrs).is_integer() else hrs
    except Exception:
        hrs_txt = hrs
    rate     = staff.get('hourly_rate') or 0
    emp_type = staff.get('employment_type') or ('Full-time' if hrs >= 30 else 'Part-time')
    pay_type = 'salary' if staff.get('is_salaried') == 'Y' else 'hourly rate of'
    wages    = f"£{staff.get('salary_amount',0):,.2f} per annum" if staff.get('is_salaried') == 'Y' else f"£{rate:.2f} per hour"
    job_title  = staff.get('job_title') or 'Sales Assistant'
    reports_to = staff.get('reports_to') or 'Store Manager'

    # Support both << >> and {{ }} formats
    fields = {}

    # Your existing << >> format
    angle = {
        "<<employee name>>":           name,
        "<<employee first name>>":     staff.get('first_name',''),
        "<<address line 1>>":          staff.get('address_1','') or '',
        "<<address line 2>>":          staff.get('address_2','') or '',
        "<<address line 3>>":          staff.get('address_3','') or '',
        "<<address line 4>>":          staff.get('address_4','') or '',
        "<<post code>>":               staff.get('postcode','') or '',
        "<<today's date>>":            today,
        "<<today’s date>>":       today,   # templates use a curly apostrophe
        "<<position>>":                job_title,
        "<<Position>>":                job_title,   # offer-letter template capitalises it
        "<<FT or PT>>":                emp_type,
        "<<salary or hourly>>":        pay_type,
        "<<wages>>":                   wages,
        "<<employer>>":                employer,
        "<<employer and store address>>": employer_addr,
        "<<store address>>":           store_addr,
        "<<s tore address >>":         store_addr,
        "<<store address line 1>>":    ent.get('addr_line1','') or '',
        "<<store address line 2>>":    ent.get('addr_line2','') or '',
        "<<store address line 3>>":    ent.get('addr_line3','') or '',
        "<<store address line 4>>":    ent.get('addr_line4','') or '',
        "<<reporting to>>":            reports_to,
        "<<notice period>>":           staff.get('notice_period','') or '',
        "<<contracted hours>>":        f"{hrs_txt} hours per week",
        "<<hours of work>>":           f"{hrs_txt} hours per week",   # contract template token
        "<<hourly rate>>":             f"£{rate:.2f}",
        "<<date of joining>>":         _fmt(staff.get('date_joined')),
        "<<DOJ>>":                     _fmt(staff.get('date_joined')),   # contract template token
        "<<date of birth>>":           _fmt(staff.get('date_of_birth')),
        "<<p osition>>":               job_title,
        "<<e mployer>>":               employer,
    }
    fields.update(angle)

    # Also {{ }} format for new templates
    curly = {
        "{{FULL_NAME}}":        name,
        "{{FIRST_NAME}}":       staff.get('first_name',''),
        "{{LAST_NAME}}":        staff.get('last_name',''),
        "{{ADDRESS_1}}":        staff.get('address_1','') or '',
        "{{ADDRESS_2}}":        staff.get('address_2','') or '',
        "{{ADDRESS_3}}":        staff.get('address_3','') or '',
        "{{POSTCODE}}":         staff.get('postcode','') or '',
        "{{EMAIL}}":            staff.get('email','') or '',
        "{{PHONE}}":            staff.get('phone','') or '',
        "{{STORE}}":            store,
        "{{STORE_ADDRESS}}":    store_addr,
        "{{DATE_JOINED}}":      _fmt(staff.get('date_joined')),
        "{{DATE_OF_BIRTH}}":    _fmt(staff.get('date_of_birth')),
        "{{CONTRACTED_HOURS}}": str(hrs),
        "{{HOURLY_RATE}}":      f"£{rate:.2f}" if rate else '',
        "{{JOB_TITLE}}":        job_title,
        "{{REPORTS_TO}}":       reports_to,
        "{{EMP_TYPE}}":         emp_type,
        "{{TODAY}}":            today,
        "{{YEAR}}":             str(datetime.now().year),
        "{{EMPLOYER}}":         employer,
    }
    fields.update(curly)
    return fields


def fill_word_template(template_path: str, fields: dict) -> bytes:
    """Fill a Word .docx/.dotx template with merge fields and return as bytes.
    Handles split runs where a merge field spans multiple runs in the same paragraph.
    """
    # Open the template. Word *template* files (.dotx, or .docx that were saved
    # from a .dotx) are marked internally as a "template" content-type, which the
    # docx library refuses to open. Detect that by what's INSIDE the file (not the
    # filename) and rewrite the content-type to a normal document on the fly.
    import zipfile
    with open(template_path, "rb") as fh:
        raw = fh.read()
    try:
        doc = DocxDocument(io.BytesIO(raw))
    except ValueError:
        buf_out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(raw)) as zin, \
             zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == "[Content_Types].xml":
                    data = data.replace(
                        b"wordprocessingml.template.main+xml",
                        b"wordprocessingml.document.main+xml")
                zout.writestr(item, data)
        buf_out.seek(0)
        doc = DocxDocument(buf_out)

    from docx.shared import RGBColor
    FILL_RGB = RGBColor(0x00, 0x00, 0x80)   # navy — house style for merged-in values (matches Umar's macro)

    def _style_value_run(run):
        run.font.color.rgb = FILL_RGB
        run.font.bold = True

    def replace_para_text(para):
        """Replace merge fields, even where a field is split across several runs,
        and colour ONLY the filled-in value navy+bold (leaving all other text —
        labels, body, the red signing 'X' — exactly as the template has it)."""
        runs = para.runs
        if not runs:
            return
        guard = 0
        while guard < 500:                      # each pass replaces one field; guard against loops
            guard += 1
            texts = [r.text for r in runs]
            full  = "".join(texts)
            # earliest-appearing placeholder in this paragraph
            hit = None
            for key, val in fields.items():
                i = full.find(key)
                if i != -1 and (hit is None or i < hit[1]):
                    hit = (key, i, val)
            if hit is None:
                break
            key, start, val = hit
            end = start + len(key)
            val = "" if val is None else str(val)

            # locate the run(s) the placeholder spans
            pos = start_run = start_off = end_run = end_off = None
            pos = 0
            for ri, t in enumerate(texts):
                rlen = len(t)
                if start_run is None and start < pos + rlen:
                    start_run, start_off = ri, start - pos
                if end <= pos + rlen:
                    end_run, end_off = ri, end - pos
                    break
                pos += rlen
            if start_run is None or end_run is None:
                break

            before = texts[start_run][:start_off]
            after  = texts[end_run][end_off:]
            if start_run == end_run:
                runs[start_run].text = before + val + after
                # style unless real (non-space) text shares the run — a stray
                # leading/trailing space must NOT block styling (it's invisible)
                can_style = (before.strip() == "" and after.strip() == "")
            else:
                runs[start_run].text = before + val
                for ri in range(start_run + 1, end_run):
                    runs[ri].text = ""
                runs[end_run].text = after          # keeps the end run's own formatting
                can_style = (before.strip() == "")
            if can_style and val:
                _style_value_run(runs[start_run])

    for para in doc.paragraphs:
        replace_para_text(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para_text(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/staff/{staff_id}", response_class=HTMLResponse)
def staff_profile(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])

    # Staff never use the full /staff profile — they get their own limited self-view.
    if user["role"] == "staff":
        return RedirectResponse("/my-profile", status_code=303)
    is_owner = user["role"] == "owner"   # managers get a deliberately trimmed view

    # Managers only see staff at their OWN store (the list already scopes this, but
    # a hand-typed /staff/<id> for another store must be blocked too).
    if user["role"] == "manager" and s.get("store_name") != user.get("store_name"):
        return RedirectResponse("/staff?msg=You+can+only+access+your+own+store&msg_type=error",
                                status_code=303)

    is_leaver = not s["is_active"]
    if is_leaver and user["role"] != "owner":
        return RedirectResponse("/staff", status_code=303)

    year  = datetime.now().year
    # Leave figure for the cards comes from ATTENDANCE (same helper as the
    # Attendance page), so the profile and Attendance page always agree.
    att_leave = attendance_leave(staff_id, year)

    # Staff-login status card — owner only.
    login_card = ""
    if is_owner:
        _lg = q("SELECT username, is_active, must_change_pw FROM users WHERE staff_id=?", (staff_id,), fetch=True)
        if _lg:
            _lg = dict(_lg[0])
            _lstat = ("<span style='color:#16a34a;font-weight:700'>active</span>" if _lg["is_active"]
                      else "<span style='color:#dc2626;font-weight:700'>disabled</span>")
            _pw = " &middot; <span style='color:#d97706'>awaiting first-login password change</span>" if _lg.get("must_change_pw") else ""
            login_card = (f"<div class='card' style='border-left:4px solid #0ea5e9'>"
                          f"<div style='font-weight:900;color:#0f2942;margin-bottom:4px'>&#128273; Staff Login</div>"
                          f"<div style='font-size:13px;color:#334155'>Username: <strong class='mono'>{esc(_lg['username'])}</strong> &middot; {_lstat}{_pw}</div>"
                          f"<div style='font-size:11px;color:#94a3b8;margin-top:6px'>Reset password / disable via <a href='/manage-users' style='color:#0369a1'>Manage Users</a>.</div></div>")
        elif s["is_active"] and user["role"] == "owner":
            login_card = (f"<div class='card' style='border-left:4px solid #0ea5e9'>"
                          f"<div style='font-weight:900;color:#0f2942;margin-bottom:6px'>&#128273; Staff Login</div>"
                          f"<div style='font-size:13px;color:#64748b;margin-bottom:10px'>No login yet &mdash; create one so {esc(s['first_name'])} can sign in to their own details, leave and documents.</div>"
                          f"<form method='POST' action='/staff/{staff_id}/create-login'><button type='submit' class='btn-primary'>&#128273; Create login</button></form></div>")
    # Personal Details card — OWNER ONLY (holds DOB / address / phone / pay / RTW).
    personal_card = ""
    if is_owner:
        personal_card = f"""<div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Personal Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr));font-size:13px'>
        <div><span style='color:#94a3b8;font-weight:700'>Date of Birth</span><br>{s.get('date_of_birth') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Phone</span><br>{esc(s.get('phone')) or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Email</span><br>{esc(s.get('email')) or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Address</span><br>{esc(', '.join(filter(None,[s.get('address_1'),s.get('address_2'),s.get('address_3'),s.get('postcode')]))) or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Date Joined</span><br>{s.get('date_joined') or '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Hourly Rate</span><br>{'£'+str(s['hourly_rate'])+'/hr' if s.get('hourly_rate') else '—'}</div>
        <div><span style='color:#94a3b8;font-weight:700'>Right to Work</span><br><a href='/staff/{staff_id}/documents' style='text-decoration:none'>{_rtw_status_summary(get_rtw_check(staff_id))}</a></div>
      </div>
    </div>"""
    _plieu = att_leave.get("lieu", 0)
    plieu_note = (f"<div style='display:flex;align-items:center;gap:5px;font-size:11px;color:#64748b;margin-top:8px'>"
                  f"<span style='font-size:13px'>&#8505;&#65039;</span>Holiday Taken includes {_plieu} bank holiday{'s' if _plieu != 1 else ''} "
                  f"worked, kept in your balance as day{'s' if _plieu != 1 else ''} in lieu.</div>") if _plieu else ""
    _ppb = att_leave.get("proj_balance")
    pproj_note = (f"<div style='display:flex;align-items:center;gap:5px;font-size:11px;color:#94a3b8;"
                  f"margin-top:6px;padding:6px 10px;background:#f8fafc;border-radius:8px'>"
                  f"<span style='font-size:13px'>&#128200;</span>Projected balance at year-end: "
                  f"<strong style='color:#475569'>{_ppb:g} hrs</strong>"
                  f"<span>&middot; estimate at current pace, before further leave</span></div>") if _ppb is not None else ""

    # Leave taken this year — derived from ATTENDANCE (single source of truth), so it
    # agrees with the summary cards. (The old leave_requests list was stale seed data.)
    leave_hist = q("""
        SELECT work_date, status, paid_hours, comments FROM staff_attendance
        WHERE staff_id=? AND substr(work_date,1,4)=?
          AND status IN('Holiday','Bank Holiday','Sick','Maternity','Unpaid Leave')
        ORDER BY work_date DESC
    """, (staff_id, str(year)), fetch=True) or []


    name = esc(f"{s['first_name']} {s['last_name']}")

    # ── Leave-taken table (from attendance) ──
    leave_rows = ""
    _lv_colour = {"Holiday": "#0369a1", "Bank Holiday": "#7c3aed", "Sick": "#dc2626",
                  "Maternity": "#db2777", "Unpaid Leave": "#b45309"}
    for lr in leave_hist:
        lr = dict(lr)
        st = lr["status"]
        badge = f"<span style='font-weight:700;color:{_lv_colour.get(st, '#334155')}'>{esc(st)}</span>"
        hrs = f"{lr['paid_hours']:g}h" if lr.get('paid_hours') else "—"
        leave_rows += f"""
        <tr>
          <td class='mono'>{lr['work_date']}</td>
          <td>{badge}</td>
          <td class='mono'>{hrs}</td>
          <td style='font-size:12px;color:#64748b'>{esc(lr.get('comments') or '—')}</td>
        </tr>"""

    can_edit = user["role"] in ("owner", "manager")

    content = f"""
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Staff</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>{name}</div>
        <div style='color:#64748b;font-size:13px'>{s.get('store_name') or ''} {'· Left ' + s['date_left'] if is_leaver and s.get('date_left') else ''}</div>
      </div>
      <div style='display:flex;gap:8px;flex-wrap:wrap'>
        {'<a href="/staff/' + str(staff_id) + '/edit" class="btn-primary">✏️ Edit Profile</a>' if is_owner else ''}
        <a href='/staff/{staff_id}/request-leave' class='btn-secondary'>📅 Request Leave</a>
        {'<a href="/staff/' + str(staff_id) + '/pay-history" class="btn-secondary">💰 Pay History</a>' if is_owner else ''}
        {'<a href="/staff/' + str(staff_id) + '/set-entitlement" class="btn-secondary">⚙️ Set Entitlement</a>' if is_owner else ''}
        <a href='/staff/{staff_id}/documents' class='btn-secondary'>&#128193; Documents</a>
        {'<a href="/staff/' + str(staff_id) + '/attendance" class="btn-secondary">🕒 Attendance</a>' if can_edit else ''}
        {'<a href="/staff/' + str(staff_id) + '/onboarding" class="btn-secondary">&#128203; Onboarding</a>' if is_owner else ''}
      </div>
    </div>

    <!-- Summary cards -->
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(150px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Entitlement {year}</div>
        <div style='font-size:20px;font-weight:900;color:#0f2942'>{att_leave.get("entitlement_fmt","—")}</div>
        <div style='font-size:10px;color:#94a3b8'>inc. bank holidays</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Holiday Taken</div>
        <div style='font-size:20px;font-weight:900;color:#d97706'>{att_leave.get("taken_fmt","—")}</div>
        <div style='font-size:10px;color:#94a3b8'>{att_leave.get("bh_days",0)} bank hols incl.</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Holiday Balance</div>
        <div style='font-size:20px;font-weight:900;color:{"#16a34a" if att_leave.get("balance",0)>=0 else "#dc2626"}'>{att_leave.get("balance_fmt","—")}</div>
        <div style='font-size:10px;color:#94a3b8'>from attendance</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Sick Days {year}</div>
        <div style='font-size:20px;font-weight:900;color:{"#dc2626" if att_leave.get("sick_days",0)>0 else "#0f2942"}'>{att_leave.get("sick_days",0)}</div>
        <div style='font-size:10px;color:#94a3b8'>does not affect holiday</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Contract</div>
        <div style='font-size:24px;font-weight:900;color:#0f2942'>{s.get('contracted_hrs') or '—'}</div>
        <div style='font-size:11px;color:#94a3b8'>hrs/week</div>
      </div>
    </div>
    {plieu_note}
    {pproj_note}

    <!-- Personal details (owner only) -->
    {personal_card}
    {(f"<div class='card' style='border-left:4px solid #dc2626'><div style='font-weight:900;color:#991b1b;margin-bottom:6px'>&#128682; Left" + (f" &middot; {s.get('date_left')}" if s.get('date_left') else "") + "</div><div style='font-size:13px;color:#334155;white-space:pre-wrap'>" + esc(s.get('leaving_reason') or '—') + "</div></div>") if is_leaver else ''}
    {(f"<div class='card' style='border-left:4px solid #f59e0b'><div style='font-weight:900;color:#92400e;margin-bottom:6px'>&#128221; Notes</div><div style='font-size:13px;color:#334155;white-space:pre-wrap'>" + esc(s.get('notes')) + "</div></div>") if is_owner and s.get('notes') else ''}
    {login_card}


        <!-- Leave history -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px;display:flex;justify-content:space-between;align-items:center'>
        <span>📅 Leave Taken {year}</span>
        <a href='/staff/{staff_id}/request-leave' style='background:rgba(255,255,255,.15);color:white;font-size:12px;font-weight:700;padding:4px 12px;border-radius:6px;text-decoration:none'>+ Request Leave</a>
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Date</th><th>Type</th><th>Paid hrs</th><th>Notes</th></tr></thead>
          <tbody>{leave_rows or '<tr><td colspan="4" style="text-align:center;padding:24px;color:#94a3b8">No leave taken this year</td></tr>'}</tbody>
        </table>
      </div>
    </div>

    <div style='text-align:right;margin-top:4px'>
      {'<a href="/staff/' + str(staff_id) + '/attendance" style="font-size:12px;color:#64748b">Full attendance record →</a>' if can_edit else ''}
    </div>"""

    return page(name, content, user, "staff")


@router.get("/staff/{staff_id}/edit", response_class=HTMLResponse)
def edit_staff_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # managers don't edit staff records; staff use /my-profile
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])
    return render_staff_form(user, s)


def render_staff_form(user: dict, s: dict | None) -> HTMLResponse:
    is_edit   = s is not None
    is_owner  = user["role"] == "owner"
    is_mgr    = user["role"] in ("owner", "manager")
    is_self   = user["role"] == "staff"
    title     = f"✏️ Edit — {s['first_name']} {s['last_name']}" if is_edit else "➕ New Staff Member"
    action    = f"/staff/{s['staff_id']}/save" if is_edit else "/staff/save"
    back_url  = f"/staff/{s['staff_id']}" if is_edit else "/staff"
    sv        = s or {}

    # Staff number is SYSTEM-GENERATED (next free number in sequence) and read-only,
    # so it's always unique and can't be mistyped into a clash.
    import json as _json
    _next_no = (q("SELECT MAX(staff_number) m FROM staff_profiles", fetch=True)[0]["m"] or 0) + 1
    if is_edit:
        _staff_number_field = (
            "<div><label>Staff Number</label>"
            f"<input type='text' name='staff_number' value='{esc(sv.get('staff_number') or '')}' readonly "
            "style='background:#f8fafc;color:#64748b' title='Permanent staff number — not editable'></div>")
    else:
        _staff_number_field = (
            "<div><label>Staff Number</label>"
            f"<input type='text' name='staff_number' value='{_next_no}' readonly "
            "style='background:#f8fafc;color:#64748b' title='Assigned automatically'>"
            "<div style='font-size:11px;color:#94a3b8;margin-top:2px'>Assigned automatically (next in sequence)</div></div>")

    # Soft duplicate-NAME warning on ADD (never a hard block — same/similar names
    # happen with rejoiners and coincidences; the owner decides).
    _exist = {}
    for _r in (q("SELECT first_name,last_name,store_name,date_joined,is_active FROM staff_profiles", fetch=True) or []):
        _d = dict(_r)
        _key = f"{(_d['first_name'] or '').strip().lower()}|{(_d['last_name'] or '').strip().lower()}"
        _exist[_key] = f"{_d.get('store_name') or '?'}, joined {_d.get('date_joined') or '?'}, {'Active' if _d.get('is_active') else 'Left'}"
    _exist_json = _json.dumps(_exist)

    def fi(name, label, ftype="text", val=None, req=False, opts=None, disabled=False, placeholder="", full=False):
        safe = val if val is not None else ""
        req_a = "required" if req else ""
        dis_a = "disabled style='background:#f8fafc;color:#94a3b8'" if disabled else ""
        step  = "step='0.01'" if ftype=="number" else ""
        ph    = f"placeholder='{placeholder}'" if placeholder else ""
        wrap  = " style='grid-column:1/-1'" if full else ""
        if ftype == "textarea":
            return (f"<div{wrap}><label>{label}</label><textarea name='{name}' rows='2' {ph} "
                    f"style='width:100%;font-family:inherit;font-size:13px'>{esc(safe)}</textarea></div>")
        if opts is not None:
            o = "".join(f"<option value='{ov}' {'selected' if str(safe)==str(ov) else ''}>{ol}</option>"
                        for ov,ol in opts)
            return f"<div{wrap}><label>{label}</label><select name='{name}' {req_a} {dis_a}>{o}</select></div>"
        return f"<div{wrap}><label>{label}</label><input type='{ftype}' name='{name}' value='{esc(safe)}' {req_a} {dis_a} {step} {ph}></div>"

    store_opts = [("","-- Select --"),("Uxbridge","Uxbridge"),("Newbury","Newbury")]
    sex_opts   = [("","--"),("M","Male"),("F","Female"),("O","Other")]
    active_opts= [("1","Active"),("0","Left / Leaver")]

    # Personal details — staff can edit these themselves
    personal = f"""
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Personal Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
        {fi('first_name','First Name','text',sv.get('first_name'),req=True,disabled=is_self)}
        {fi('last_name', 'Last Name', 'text',sv.get('last_name'), req=True,disabled=is_self)}
        {fi('date_of_birth','Date of Birth','date',sv.get('date_of_birth'))}
        {fi('sex','Gender',opts=sex_opts,val=sv.get('sex'))}
        {fi('phone','Phone','text',sv.get('phone'),placeholder='07700 000000')}
        {fi('email','Email','email',sv.get('email'))}
        {fi('address_1','Address Line 1','text',sv.get('address_1'))}
        {fi('address_2','Address Line 2','text',sv.get('address_2'))}
        {fi('address_3','Town/City','text',sv.get('address_3'))}
        {fi('postcode','Postcode','text',sv.get('postcode'))}
      </div>
    </div>"""

    # Employment details — manager/owner only
    employment = ""
    if is_mgr:
        employment = f"""
    <div class='card'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment Details</div>
      <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
        {_staff_number_field}
        {fi('store_name','Store',opts=store_opts,val=sv.get('store_name'),req=True)}
        {fi('job_title','Job Title','text',sv.get('job_title',''),placeholder='e.g. Sales Assistant')}
        {fi('employment_type','Employment Type',opts=[('Full-time','Full-time'),('Part-time','Part-time')],val=sv.get('employment_type','Part-time'))}
        {fi('reports_to','Reports To','text',sv.get('reports_to',''),placeholder='e.g. Store Manager')}
        {fi('notice_period','Notice Period','text',sv.get('notice_period',''),placeholder='e.g. 1 week, 12 weeks')}
        {fi('date_joined','Date Joined','date',sv.get('date_joined'))}
        {fi('contracted_hrs','Contracted Hours/Week','number',sv.get('contracted_hrs'),placeholder='e.g. 37.5')}
        {fi('hourly_rate','Hourly Rate (£)','number',sv.get('hourly_rate'),placeholder='e.g. 11.44')}
        {fi('is_salaried','Salaried?',opts=[('N','No'),('Y','Yes')],val=sv.get('is_salaried','N'))}
        {fi('salary_amount','Salary Amount (£/yr)','number',sv.get('salary_amount'))}
        {fi('is_active','Status',opts=active_opts,val=str(sv.get('is_active',1)))}
        {fi('date_left','Date Left','date',sv.get('date_left')) if is_edit else ''}
        {fi('leaving_reason','Reason for Leaving','textarea',sv.get('leaving_reason'),full=True) if is_edit else ''}
        {fi('notes','Notes / Comments','textarea',sv.get('notes'),full=True,placeholder='Internal notes — e.g. resigned via WhatsApp 5 Aug (effective 1 Aug); on reduced hours; etc.')}
      </div>
    </div>"""

    content = f"""
    <div class='flex justify-between items-center'>
      <div>
        <a href='{back_url}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>{title}</div>
      </div>
    </div>
    <form action='{action}' method='POST' enctype='multipart/form-data'>
      {personal}
      {employment}
      <div class='card'>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>{'💾 Save Changes' if is_edit else '➕ Add Staff Member'}</button>
          <a href='{back_url}' class='btn-secondary'>Cancel</a>
          {'<a href="/staff/' + str(s["staff_id"]) + '/delete" class="btn-danger" onclick="return confirm(\'Are you sure?\')">🗑️ Delete</a>' if is_edit and is_owner else ''}
        </div>
      </div>
    </form>
    {'' if is_edit else '''
    <script>
    (function(){
      var existing = ''' + _exist_json + ''';
      var form = document.querySelector('form[action="/staff/save"]');
      if(!form) return;
      form.addEventListener('submit', function(e){
        var fn=(form.first_name.value||'').trim().toLowerCase();
        var ln=(form.last_name.value||'').trim().toLowerCase();
        var m=existing[fn+'|'+ln];
        if(m){
          if(!confirm('\\u26a0\\ufe0f A staff member named '+form.first_name.value+' '+form.last_name.value+' already exists ('+m+').\\n\\nIs this a genuinely NEW starter?\\n\\nOK = add anyway     Cancel = stop')){
            e.preventDefault();
          }
        }
      });
    })();
    </script>'''}"""

    return HTMLResponse(page(title, content, user, "staff"))


@router.post("/staff/save")
async def save_new_staff(request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r
    form = await request.form()
    fv = lambda k, d="": str(form.get(k, d) or d).strip()
    fn = lambda k: float(form.get(k, 0) or 0) if form.get(k) else None
    # Staff number is system-assigned: next free number in sequence, always unique.
    next_no = (q("SELECT MAX(staff_number) m FROM staff_profiles", fetch=True)[0]["m"] or 0) + 1
    q("""INSERT INTO staff_profiles
        (staff_number,first_name,last_name,store_name,sex,phone,email,
         address_1,address_2,address_3,postcode,date_joined,date_of_birth,
         contracted_hrs,hourly_rate,is_salaried,salary_amount,is_active,
         job_title,employment_type,reports_to,notice_period,notes)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
      (next_no, fv("first_name"), fv("last_name"),
       fv("store_name"), fv("sex"), fv("phone"), fv("email"),
       fv("address_1"), fv("address_2"), fv("address_3"), fv("postcode"),
       fv("date_joined") or None, fv("date_of_birth") or None,
       fn("contracted_hrs"), fn("hourly_rate"),
       fv("is_salaried","N"), fn("salary_amount"),
       int(form.get("is_active", 1)),
       fv("job_title") or None, fv("employment_type") or None,
       fv("reports_to") or None, fv("notice_period") or None,
       fv("notes") or None))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff?msg={uq('Staff member added successfully')}", status_code=303)


@router.post("/staff/{staff_id}/save")
async def save_staff(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # only the owner edits staff records
    form = await request.form()
    fv = lambda k, d="": str(form.get(k, d) or d).strip()
    fn = lambda k: float(form.get(k, 0) or 0) if form.get(k) else None
    is_mgr = user["role"] in ("owner","manager")

    if is_mgr:
        q("""UPDATE staff_profiles SET
            staff_number=?,first_name=?,last_name=?,store_name=?,sex=?,
            phone=?,email=?,address_1=?,address_2=?,address_3=?,postcode=?,
            date_joined=?,date_of_birth=?,contracted_hrs=?,hourly_rate=?,
            is_salaried=?,salary_amount=?,is_active=?,date_left=?,leaving_reason=?,
            job_title=?,employment_type=?,reports_to=?,notice_period=?,notes=?
            WHERE staff_id=?""",
          (form.get("staff_number") or None, fv("first_name"), fv("last_name"),
           fv("store_name"), fv("sex"), fv("phone"), fv("email"),
           fv("address_1"), fv("address_2"), fv("address_3"), fv("postcode"),
           fv("date_joined") or None, fv("date_of_birth") or None,
           fn("contracted_hrs"), fn("hourly_rate"),
           fv("is_salaried","N"), fn("salary_amount"),
           int(form.get("is_active",1)),
           fv("date_left") or None, fv("leaving_reason") or None,
           fv("job_title") or None, fv("employment_type") or None,
           fv("reports_to") or None, fv("notice_period") or None,
           fv("notes") or None,
           staff_id))
    else:
        # Staff can only update personal contact details
        q("""UPDATE staff_profiles SET
            phone=?,email=?,address_1=?,address_2=?,address_3=?,postcode=?,date_of_birth=?
            WHERE staff_id=?""",
          (fv("phone"),fv("email"),fv("address_1"),fv("address_2"),
           fv("address_3"),fv("postcode"),fv("date_of_birth") or None, staff_id))

    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}?msg={uq('Profile updated')}", status_code=303)


@router.post("/staff/{staff_id}/create-login")
def create_staff_login(staff_id: int, session: str | None = Cookie(default=None)):
    """Owner-only: create a staff login linked by staff_id, with an auto username
    and a one-time temp password (forced change on first login)."""
    from urllib.parse import quote as uq
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}?msg={uq('Only the owner can create staff logins')}&msg_type=error",
                                status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows:
        return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])
    ex = q("SELECT username FROM users WHERE staff_id=?", (staff_id,), fetch=True)
    if ex:
        return RedirectResponse(f"/staff/{staff_id}?msg={uq('Already has a login: ' + ex[0]['username'])}&msg_type=error",
                                status_code=303)
    uname = _gen_username(s["first_name"], s["last_name"])
    temp  = _gen_temp_password()
    q("""INSERT INTO users (username, password, full_name, role, store_name, is_active, staff_id, must_change_pw)
         VALUES(?,?,?,?,?,1,?,1)""",
      (uname, hash_password(temp), f"{s['first_name']} {s['last_name']}",
       "staff", s.get("store_name"), staff_id))
    name = f"{s['first_name']} {s['last_name']}"
    content = f"""
    <div class='text-2xl font-black text-slate-800'>&#128273; Login created &mdash; {esc(name)}</div>
    <div class='card' style='max-width:540px;border-left:5px solid #16a34a'>
      <p style='color:#334155;margin-bottom:12px'>Give these to {esc(s['first_name'])}. They'll be asked to set their own password the first time they sign in.</p>
      <div style='background:#f0fdf4;border:1px solid #86efac;border-radius:10px;padding:14px 16px;font-size:15px'>
        <div>Username: <strong class='mono'>{esc(uname)}</strong></div>
        <div style='margin-top:6px'>Temporary password: <strong class='mono'>{esc(temp)}</strong></div>
      </div>
      <p style='font-size:12px;color:#94a3b8;margin-top:10px'>&#9888;&#65039; Shown once &mdash; note it now. You can reset it later from Manage Users.</p>
      <a href='/staff/{staff_id}' class='btn-primary' style='margin-top:14px;display:inline-block'>&larr; Back to profile</a>
    </div>"""
    return page("Login created", content, user, "staff")


@router.get("/staff/{staff_id}/request-leave", response_class=HTMLResponse)
def request_leave_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id)): return r
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"
    leave = attendance_leave(staff_id)   # single source of truth = imported attendance
    bal   = leave.get("balance", 0)

    type_opts = "".join(f"<option value='{k}'>{v}</option>" for k,v in ABSENCE_TYPES.items())

    content = f"""
    <div>
      <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>📅 Request Leave — {name}</div>
    </div>
    <div class='card' style='max-width:520px'>
      <div style='background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;padding:12px 16px;margin-bottom:16px;font-size:13px'>
        Current balance: <strong style='color:#0369a1'>{leave.get('balance_fmt','—')} remaining</strong>
      </div>
      <form action='/staff/{staff_id}/request-leave' method='POST' class='space-y-4'>
        <div><label>Leave Type</label>
          <select name='leave_type'>{type_opts}</select></div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          <div><label>From Date</label>
            <input type='date' name='date_from' required></div>
          <div><label>To Date</label>
            <input type='date' name='date_to' required></div>
        </div>
        <div><label>Notes (optional)</label>
          <textarea name='notes' rows='2' placeholder='Any additional details...'></textarea></div>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>📤 Submit Request</button>
          <a href='/staff/{staff_id}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>
    <script>
    // Auto-calculate working days when dates change
    document.addEventListener('DOMContentLoaded', function() {{
      const from = document.querySelector('[name="date_from"]');
      const to   = document.querySelector('[name="date_to"]');
      if (from) from.addEventListener('change', function() {{
        if (to && !to.value) to.value = from.value;
      }});
    }});
    </script>"""
    return page("Request Leave", content, user, "staff")


@router.post("/staff/{staff_id}/request-leave")
async def submit_leave(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id)): return r
    form       = await request.form()
    leave_type = form.get("leave_type","H")
    date_from  = form.get("date_from","")
    date_to    = form.get("date_to","")
    notes      = str(form.get("notes","") or "").strip()

    # Calculate working days
    try:
        d1 = datetime.strptime(date_from, "%Y-%m-%d")
        d2 = datetime.strptime(date_to,   "%Y-%m-%d")
        bh   = uk_bank_holidays(d1.year) | uk_bank_holidays(d2.year)
        days = 0
        cur  = d1
        while cur <= d2:
            if cur.weekday() < 5 and cur.strftime("%Y-%m-%d") not in bh:
                days += 1
            cur = cur + timedelta(days=1)
    except Exception:
        days = 1

    # Two-stage approval: the owner's own entry is final; a manager entering leave
    # counts as the manager stage (still needs owner sign-off); a staff request
    # starts at the very beginning and waits for the manager.
    mgr_by = mgr_at = None
    if user["role"] == "owner":
        status = "approved"
    elif user["role"] == "manager":
        status = "mgr_approved"
        mgr_by = user.get("username")
        mgr_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    else:
        status = "pending"

    q("""INSERT INTO leave_requests
        (staff_id, leave_type, date_from, date_to, days_taken,
         status, requested_by, notes, mgr_approved_by, mgr_approved_at)
        VALUES(?,?,?,?,?,?,?,?,?,?)""",
      (staff_id, leave_type, date_from, date_to, days,
       status, user.get("username"), notes or None, mgr_by, mgr_at))

    from urllib.parse import quote as uq
    msg = {"approved": "Leave approved ✅",
           "mgr_approved": "Recorded — awaiting owner sign-off ⏳",
           "pending": "Leave request submitted — awaiting approval ⏳"}[status]
    return RedirectResponse(f"/staff/{staff_id}?msg={uq(msg)}", status_code=303)


def _leave_req_with_store(req_id: int):
    """The leave request joined to its staff member's store, or None."""
    rows = q("""SELECT lr.*, sp.store_name FROM leave_requests lr
                JOIN staff_profiles sp ON lr.staff_id = sp.staff_id
                WHERE lr.request_id = ?""", (req_id,), fetch=True)
    return dict(rows[0]) if rows else None


def _mgr_can_act(user, lr) -> bool:
    """A manager may only act on requests for their OWN store."""
    if user["role"] == "owner":
        return True
    return bool(user.get("store_name")) and lr.get("store_name") == user["store_name"]


@router.post("/staff/leave-requests/{req_id}/approve")
def approve_leave(req_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r

    from urllib.parse import quote as uq
    lr = _leave_req_with_store(req_id)
    if not lr or not _mgr_can_act(user, lr):
        return RedirectResponse(f"/staff/leave-requests?msg={uq('Not permitted')}&msg_type=error", status_code=303)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    if user["role"] == "owner":
        # Owner gives the final sign-off (works from pending or mgr_approved).
        q("""UPDATE leave_requests SET status='approved', approved_by=?, approved_at=?
             WHERE request_id=?""", (user.get("username"), now, req_id))
        msg = "Leave approved ✅"
    else:
        # Manager stage: recommend → moves into the owner's sign-off queue.
        if lr["status"] != "pending":
            return RedirectResponse(f"/staff/leave-requests?msg={uq('Already actioned')}&msg_type=error", status_code=303)
        q("""UPDATE leave_requests SET status='mgr_approved', mgr_approved_by=?, mgr_approved_at=?
             WHERE request_id=?""", (user.get("username"), now, req_id))
        msg = "Approved — sent to owner for sign-off ⏳"
    return RedirectResponse(f"/staff/leave-requests?msg={uq(msg)}", status_code=303)


@router.post("/staff/leave-requests/{req_id}/decline")
def decline_leave(req_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r

    from urllib.parse import quote as uq
    lr = _leave_req_with_store(req_id)
    if not lr or not _mgr_can_act(user, lr):
        return RedirectResponse(f"/staff/leave-requests?msg={uq('Not permitted')}&msg_type=error", status_code=303)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    # A manager records who declined in the manager field; the owner in the final field.
    if user["role"] == "owner":
        q("""UPDATE leave_requests SET status='declined', approved_by=?, approved_at=?
             WHERE request_id=?""", (user.get("username"), now, req_id))
    else:
        q("""UPDATE leave_requests SET status='declined', mgr_approved_by=?, mgr_approved_at=?
             WHERE request_id=?""", (user.get("username"), now, req_id))
    return RedirectResponse(f"/staff/leave-requests?msg={uq('Leave declined')}", status_code=303)


@router.get("/staff/{staff_id}/pay-history", response_class=HTMLResponse)
def pay_history_page(staff_id: int, session: str | None = Cookie(default=None), msg: str = ""):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # pay is owner-only
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse("/staff", status_code=303)

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Current NMW
    today   = datetime.now().strftime("%Y-%m-%d")
    nmw     = get_nmw_for_person(s.get("date_of_birth",""), today)
    current = s.get("hourly_rate") or 0
    _is_sal = s.get("is_salaried") == 'Y'
    _sal    = s.get("salary_amount") or 0
    _hrs    = s.get("contracted_hrs") or 0
    # For the NMW check a salaried person's effective hourly = salary / (hrs*52).
    eff_hourly = (_sal / (_hrs * 52)) if (_is_sal and _hrs) else current
    diff    = round(eff_hourly - nmw, 2)
    dob_str = s.get("date_of_birth","")
    age     = ((datetime.now() - datetime.strptime(dob_str,"%Y-%m-%d")).days // 365) if dob_str else 0

    # NMW status badge
    if nmw == 0:
        nmw_badge = "<span class='badge-unpaid'>DOB not set</span>"
    elif diff < 0:
        nmw_badge = f"<span class='badge-overdue'>⚠️ £{abs(diff):.2f} BELOW NMW</span>"
    elif diff < 0.50:
        nmw_badge = f"<span class='badge-partial'>⚠️ Only £{diff:.2f} above NMW</span>"
    else:
        nmw_badge = f"<span class='badge-paid'>✅ £{diff:.2f} above NMW</span>"

    # Pay history
    history = q("""SELECT * FROM pay_history WHERE staff_id=?
                   ORDER BY effective_date DESC""", (staff_id,), fetch=True) or []

    hist_rows = ""
    for h in history:
        h = dict(h)
        _b = h.get("pay_basis") or "hourly"
        if _b == "salary":
            pay_now  = f"£{(h.get('salary_amount') or 0):,.0f}/yr <span style='color:#94a3b8;font-weight:400'>salary</span>"
            if h.get("previous_salary"):   prev_txt = f"was £{h['previous_salary']:,.0f}/yr"
            elif h.get("previous_rate"):   prev_txt = f"was £{h['previous_rate']:.2f}/hr"
            else:                          prev_txt = ""
        else:
            pay_now  = f"£{(h.get('hourly_rate') or 0):.2f}/hr"
            if h.get("previous_rate"):     prev_txt = f"was £{h['previous_rate']:.2f}/hr"
            elif h.get("previous_salary"): prev_txt = f"was £{h['previous_salary']:,.0f}/yr salary"
            else:                          prev_txt = ""
        hrs_txt = "—"
        if h.get("contracted_hrs") is not None:
            hrs_txt = f"{h['contracted_hrs']:g}h/wk"
            if h.get("previous_hrs") is not None and h["previous_hrs"] != h["contracted_hrs"]:
                hrs_txt += f" <span style='color:#94a3b8'>(was {h['previous_hrs']:g})</span>"
        hist_rows += f"""<tr>
          <td class='mono'>{h['effective_date']}</td>
          <td style='font-weight:700'>{pay_now}<div style='font-size:11px;color:#94a3b8;font-weight:400'>{prev_txt}</div></td>
          <td class='mono' style='font-size:12px'>{hrs_txt}</td>
          <td style='font-size:12px;color:#64748b'>{h.get('change_reason') or '—'}</td>
          <td style='font-size:12px;color:#94a3b8'>{h.get('recorded_by') or '—'}</td>
        </tr>"""

    # NMW history table
    nmw_rows = ""
    all_nmw = q("SELECT * FROM nmw_rates ORDER BY effective_date DESC", fetch=True) or []
    for n in all_nmw:
        n   = dict(n)
        age_rate = n["rate_21_plus"] if age >= 21 else (n["rate_18_20"] if age >= 18 else n["rate_16_17"])
        nmw_rows += f"""<tr>
          <td class='mono'>{n['effective_date']}</td>
          <td class='mono'>£{n['rate_21_plus']:.2f}</td>
          <td class='mono'>£{n['rate_18_20']:.2f}</td>
          <td class='mono'>£{n['rate_16_17']:.2f}</td>
          <td class='mono' style='font-weight:700;color:#0369a1'>£{age_rate:.2f}</td>
        </tr>"""

    flash = f"<div class='flash-success'>{msg}</div>" if msg else ""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>💰 Pay History — {name}</div>
      </div>
    </div>

    <!-- Current pay status -->
    <div class='grid gap-4' style='grid-template-columns:repeat(auto-fit,minmax(160px,1fr))'>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Current {'Salary' if _is_sal else 'Rate'}</div>
        <div style='font-size:28px;font-weight:900;color:#0f2942'>{f'£{_sal:,.0f}' if _is_sal else f'£{current:.2f}'}</div>
        <div style='font-size:11px;color:#94a3b8'>{'per year' + (f' (≈£{eff_hourly:.2f}/hr)' if _hrs else '') if _is_sal else 'per hour'}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>NMW (Age {age})</div>
        <div style='font-size:28px;font-weight:900;color:#0f2942'>£{nmw:.2f}</div>
        <div style='font-size:11px;color:#94a3b8'>minimum wage</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Status</div>
        <div style='margin-top:8px'>{nmw_badge}</div>
      </div>
      <div class='card py-3 text-center'>
        <div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Annual {'Salary' if _is_sal else 'Equiv.'}</div>
        <div style='font-size:22px;font-weight:900;color:#0f2942'>£{(_sal if _is_sal else current * _hrs * 52):,.0f}</div>
        <div style='font-size:11px;color:#94a3b8'>{'fixed salary' if _is_sal else f'based on {_hrs}h/wk'}</div>
      </div>
    </div>

    <!-- Record new pay change -->
    <div class='card' style='max-width:560px'>
      <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>➕ Record Pay Change</div>
      <form action='/staff/{staff_id}/pay-history' method='POST' class='grid gap-3' style='grid-template-columns:1fr 1fr'>
        <div><label>Effective Date</label>
          <input type='date' name='effective_date' value='{today}' required></div>
        <div><label>Pay Basis</label>
          <select name='pay_basis' id='payBasis' onchange='payBasisToggle()'>
            <option value='hourly'>Hourly</option>
            <option value='salary'>Salaried</option>
          </select></div>
        <div id='rateBox'><label>New Hourly Rate (£)</label>
          <input type='number' step='0.01' name='hourly_rate' placeholder='e.g. 12.71'></div>
        <div id='salaryBox' style='display:none'><label>New Salary (£/yr)</label>
          <input type='number' step='0.01' name='salary_amount' placeholder='e.g. 24000'></div>
        <div><label>Contracted Hours/Week</label>
          <input type='number' step='0.01' name='contracted_hrs' placeholder='e.g. 32.5' value='{s.get('contracted_hrs') or ''}'></div>
        <div style='grid-column:1/-1'><label>Reason</label>
          <input type='text' name='change_reason' placeholder='e.g. Annual review, NMW increase, Promotion to Store Manager'></div>
        <div style='grid-column:1/-1'>
          <button type='submit' class='btn-primary'>💾 Save Pay Change</button>
        </div>
      </form>
    </div>
    <script>
    function payBasisToggle(){{
      var b=document.getElementById('payBasis').value;
      document.getElementById('rateBox').style.display   = (b==='salary')?'none':'';
      document.getElementById('salaryBox').style.display = (b==='salary')?'':'none';
    }}
    payBasisToggle();
    </script>

    <!-- Pay history -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>Pay History</div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Effective Date</th><th>Pay</th><th>Hours</th><th>Reason</th><th>Recorded By</th></tr></thead>
          <tbody>{hist_rows or '<tr><td colspan="5" style="text-align:center;padding:24px;color:#94a3b8">No pay history recorded yet</td></tr>'}</tbody>
        </table>
      </div>
    </div>

    <!-- NMW reference table -->
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>
        National Minimum Wage Reference
        <span style='font-size:12px;font-weight:400;color:#93c5fd;margin-left:8px'>Highlighted column = this employee's applicable rate</span>
      </div>
      <div style='overflow-x:auto'>
        <table class='tbl'>
          <thead><tr><th>Effective</th><th>21+</th><th>18-20</th><th>16-17</th><th style='background:#1e3a5f'>Applicable Rate</th></tr></thead>
          <tbody>{nmw_rows}</tbody>
        </table>
      </div>
    </div>"""
    return page(f"Pay History — {name}", content, user, "staff")


@router.post("/staff/{staff_id}/pay-history")
async def save_pay_change(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # pay is owner-only
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}", status_code=303)
    form = await request.form()
    _num = lambda k: (float(form.get(k)) if str(form.get(k) or "").strip() not in ("", "None") else None)
    basis     = (form.get("pay_basis") or "hourly").strip()   # 'hourly' or 'salary'
    eff_date  = form.get("effective_date","")
    reason    = str(form.get("change_reason","") or "").strip()
    new_hrs   = _num("contracted_hrs")
    if basis == "salary":
        new_rate, new_salary = None, _num("salary_amount")
    else:
        new_rate, new_salary = _num("hourly_rate"), None

    # Snapshot the CURRENT terms as the "previous" side of this change.
    cur = dict(q("SELECT hourly_rate,salary_amount,contracted_hrs FROM staff_profiles WHERE staff_id=?",
                 (staff_id,), fetch=True)[0])
    prev_rate, prev_salary, prev_hrs = cur["hourly_rate"], cur["salary_amount"], cur["contracted_hrs"]
    # If hours weren't given on the change, keep the existing contracted hours.
    if new_hrs is None:
        new_hrs = prev_hrs

    q("""INSERT INTO pay_history
           (staff_id,effective_date,pay_basis,hourly_rate,previous_rate,
            salary_amount,previous_salary,contracted_hrs,previous_hrs,change_reason,recorded_by)
         VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
      (staff_id, eff_date, basis, new_rate, prev_rate,
       new_salary, prev_salary, new_hrs, prev_hrs, reason or None, user.get("username")))

    # Update the current terms on the profile to match this change.
    q("""UPDATE staff_profiles SET is_salaried=?, hourly_rate=?, salary_amount=?, contracted_hrs=?
         WHERE staff_id=?""",
      ('Y' if basis == "salary" else 'N', new_rate, new_salary, new_hrs, staff_id))

    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/pay-history?msg={uq('Pay change recorded')}", status_code=303)


def attendance_leave(staff_id: int, year=None) -> dict:
    """Leave (holiday + bank holidays, inclusive) for a year, DERIVED FROM ATTENDANCE.
    Salaried -> days (5.6 wks × days/week). Hourly -> hours (12.07% of hours worked).
    A worked bank holiday is NOT counted (it's a worked day / day in lieu)."""
    year = str(year or datetime.now().year)
    rows = q("SELECT is_salaried, days_per_week, contracted_hrs FROM staff_profiles WHERE staff_id=?",
             (staff_id,), fetch=True)
    if not rows:
        return {}
    s = dict(rows[0])
    # Holiday accrual AND deduction use PAID hours (breaks deducted) — Umar's basis.
    lv = dict(q("""SELECT
          COALESCE(SUM(CASE WHEN status IN('Holiday','Bank Holiday') THEN 1 ELSE 0 END),0) tdays,
          COALESCE(SUM(CASE WHEN status IN('Holiday','Bank Holiday') THEN paid_hours ELSE 0 END),0) thrs,
          COALESCE(SUM(CASE WHEN status='Bank Holiday' THEN 1 ELSE 0 END),0) bh,
          COALESCE(SUM(CASE WHEN status='Sick' THEN 1 ELSE 0 END),0) sick,
          COALESCE(SUM(CASE WHEN status='Worked' THEN paid_hours ELSE 0 END),0) whrs,
          COALESCE(SUM(CASE WHEN status='Worked' THEN 1 ELSE 0 END),0) wdays,
          COALESCE(SUM(CASE WHEN status='Worked' AND a_type='B' THEN 1 ELSE 0 END),0) lieu
        FROM staff_attendance WHERE staff_id=? AND substr(work_date,1,4)=?""",
        (staff_id, year), fetch=True)[0])
    salaried = s.get("is_salaried") == "Y"
    if salaried:
        dpw = s.get("days_per_week") or 5
        ent = round(5.6 * dpw, 1); taken = round(lv["tdays"], 1); unit = "days"
        basis = f"5.6 weeks × {dpw:g} days/week (salaried)"
    else:
        ent = round(lv["whrs"] * 0.1207, 1); taken = round(lv["thrs"], 1); unit = "hrs"
        avg = (lv["whrs"] / lv["wdays"]) if lv["wdays"] else 0
        basis = (f"12.07% of {lv['whrs']:,.0f} paid hrs"
                 + (f" · ≈ {ent/avg:.1f} days at {avg:.1f}h/day" if avg else ""))
    bal = round(ent - taken, 1)
    # Projected year-end balance — hourly, current year only. Kept SEPARATE from the
    # real figures above: entitlement/taken/balance are untouched. Projects the
    # person's current weekly pace of PAID hours forward to 31 Dec, then applies
    # 12.07% and subtracts leave already taken. A "before further leave" estimate.
    proj = {}
    if (not salaried) and int(year) == datetime.now().year and lv["whrs"] > 0:
        # Span = first to last attendance day of the year (≈ Jan 1, or their start
        # date if they joined mid-year) — the elapsed calendar period, so the weekly
        # pace matches the remaining calendar weeks it's projected over.
        span = q("""SELECT MIN(work_date) a, MAX(work_date) b FROM staff_attendance
                    WHERE staff_id=? AND substr(work_date,1,4)=?""",
                 (staff_id, year), fetch=True)
        if span and span[0]["a"]:
            a = datetime.fromisoformat(span[0]["a"]); b = datetime.fromisoformat(span[0]["b"])
            span_wks = max((b - a).days / 7, 1)
            avg_wk = lv["whrs"] / span_wks
            rem_wks = max((datetime(int(year), 12, 31) - b).days / 7, 0)
            proj_hrs = lv["whrs"] + avg_wk * rem_wks
            proj_ent = round(proj_hrs * 0.1207, 1)
            proj_bal = round(proj_ent - taken, 1)
            proj = {"proj_balance": proj_bal, "proj_balance_fmt": f"{proj_bal:g} hrs",
                    "proj_entitlement": proj_ent}
    return {"unit": unit, "salaried": salaried, "basis": basis,
            "entitlement": ent, "taken": taken, "balance": bal,
            "entitlement_fmt": f"{ent:g} {unit}", "taken_fmt": f"{taken:g} {unit}",
            "balance_fmt": f"{bal:g} {unit}", "bh_days": lv["bh"], "sick_days": lv["sick"],
            "lieu": lv["lieu"], **proj}


@router.get("/staff/{staff_id}/attendance", response_class=HTMLResponse)
def attendance_page(staff_id: int, session: str | None = Cookie(default=None),
                    year: str = "", view: str = "all"):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner", "manager"):
        return RedirectResponse("/staff", status_code=303)
    # Managers get a deliberately trimmed view: absence FLAGS only (who's in /
    # holiday / sick / absent) — no worked hours, no daily hours/comments detail,
    # no lifetime history. The owner sees everything.
    is_owner = user["role"] == "owner"
    if not is_owner:   # managers: own store only
        _st = q("SELECT store_name FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
        if not _st or _st[0]["store_name"] != user.get("store_name"):
            return RedirectResponse("/staff?msg=You+can+only+view+your+own+store&msg_type=error",
                                    status_code=303)
    if not is_owner:
        view = "absences"
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0]); name = f"{s['first_name']} {s['last_name']}"

    tot = dict(q("SELECT COUNT(*) c, MIN(work_date) mn, MAX(work_date) mx FROM staff_attendance WHERE staff_id=?",
                 (staff_id,), fetch=True)[0])
    if not tot["c"]:
        body = ("<div class='card' style='margin-top:14px;color:#64748b'>No attendance imported for "
                f"{esc(name)} yet.</div>")
    else:
        from datetime import datetime as _dt
        cur_year = str(_dt.now().year)
        def _uk(dd):
            p = (dd or "").split("-"); return f"{p[2]}/{p[1]}/{p[0]}" if len(p) == 3 else (dd or "")

        def _summary(extra_where, extra_params):
            return {dict(r)["status"]: dict(r) for r in q(
                "SELECT status, COUNT(*) days, COALESCE(SUM(hours_worked),0) hw "
                f"FROM staff_attendance WHERE staff_id=? {extra_where} GROUP BY status",
                [staff_id] + extra_params, fetch=True)}
        def _cards(label, sm):
            def dd(k): return sm.get(k, {}).get("days", 0)
            w = sm.get("Worked", {})
            worked_sub = f"days · {w.get('hw',0):,.0f}h" if is_owner else "days"
            spec = [("Worked", str(w.get("days", 0)), worked_sub, "#16a34a"),
                    ("Holiday", str(dd("Holiday")), "days", "#0369a1"),
                    ("Sick", str(dd("Sick")), "days", "#dc2626"),
                    ("Maternity", str(dd("Maternity")), "days", "#7c3aed"),
                    ("Bank Hol", str(dd("Bank Holiday")), "days", "#475569")]
            inner = "".join(
                f"<div class='card py-3 text-center'><div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>{t}</div>"
                f"<div style='font-size:24px;font-weight:900;color:{col}'>{v}</div><div style='font-size:11px;color:#94a3b8'>{sub}</div></div>"
                for t, v, sub, col in spec)
            return (f"<div style='font-size:12px;font-weight:800;color:#334155;margin:16px 0 6px'>{label}</div>"
                    f"<div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(120px,1fr))'>{inner}</div>")
        # Year selector — drives BOTH the summary and the daily list below.
        years = [r["yr"] for r in q("SELECT DISTINCT substr(work_date,1,4) yr FROM staff_attendance "
                                    "WHERE staff_id=? ORDER BY yr DESC", (staff_id,), fetch=True) or []]
        if is_owner:
            sel = year if year in years else (cur_year if cur_year in years else (years[0] if years else cur_year))
        else:
            sel = cur_year   # managers are locked to the current year — no history
        def _chip(y):
            on = (y == sel)
            st = "background:#0f2942;color:#fff" if on else "background:#fff;color:#0f2942;border:1px solid #cbd5e1"
            return (f"<a href='/staff/{staff_id}/attendance?year={y}&view={view}' style='{st};padding:5px 14px;"
                    f"border-radius:999px;font-weight:700;font-size:13px;text-decoration:none'>{y}</a>")
        chips = " ".join(_chip(y) for y in years) if is_owner else ""

        # Summary reflects the SELECTED year.
        cards = _cards(f"Summary — {sel}", _summary("AND substr(work_date,1,4)=?", [sel]))

        # Leave (holiday + bank holidays, inclusive) for the selected year — via the
        # SAME shared helper the profile cards use, so the two can never disagree.
        _L = attendance_leave(staff_id, sel)
        _ent, _taken, _bal = _L["entitlement"], _L["taken"], _L["balance"]
        _unit, _basis = _L["unit"], _L["basis"]
        _balcol = "#16a34a" if _bal >= 0 else "#dc2626"
        _lieu = _L.get("lieu", 0)
        _lieu_note = (f"<div style='display:flex;align-items:center;gap:5px;font-size:11px;color:#64748b;margin-top:6px'>"
                      f"<span style='font-size:13px'>&#8505;&#65039;</span>Includes {_lieu} bank holiday{'s' if _lieu != 1 else ''} worked, "
                      f"kept in balance as day{'s' if _lieu != 1 else ''} in lieu.</div>") if _lieu else ""
        _pb = _L.get("proj_balance")
        _proj_note = (f"<div style='display:flex;align-items:center;gap:5px;font-size:11px;color:#94a3b8;"
                      f"margin-top:6px;padding:6px 10px;background:#f8fafc;border-radius:8px'>"
                      f"<span style='font-size:13px'>&#128200;</span>Projected balance at year-end: "
                      f"<strong style='color:#475569'>{_pb:g} hrs</strong>"
                      f"<span>&middot; estimate at current pace, before further leave</span></div>") if _pb is not None else ""
        leave_block = f"""
        <div style='font-size:12px;font-weight:800;color:#334155;margin:18px 0 6px'>Leave — {sel}
          <span style='font-weight:400;color:#94a3b8'>· holiday + bank holidays (inclusive)</span></div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(150px,1fr))'>
          <div class='card py-3 text-center'><div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Entitlement</div>
            <div style='font-size:24px;font-weight:900;color:#0f2942'>{_ent:g}</div><div style='font-size:11px;color:#94a3b8'>{_unit}</div></div>
          <div class='card py-3 text-center'><div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Taken</div>
            <div style='font-size:24px;font-weight:900;color:#0369a1'>{_taken:g}</div><div style='font-size:11px;color:#94a3b8'>{_unit}</div></div>
          <div class='card py-3 text-center'><div style='font-size:11px;font-weight:700;color:#94a3b8;text-transform:uppercase'>Balance</div>
            <div style='font-size:24px;font-weight:900;color:{_balcol}'>{_bal:g}</div><div style='font-size:11px;color:#94a3b8'>{_unit}</div></div>
        </div>
        {_lieu_note}
        {_proj_note}
        <div style='font-size:11px;color:#94a3b8;margin-top:3px'>Basis: {_basis}. Derived from attendance.</div>"""

        # Lifetime totals, one row per year.
        yr = q("""SELECT substr(work_date,1,4) yr,
                    SUM(CASE WHEN status='Worked' THEN 1 ELSE 0 END) w,
                    COALESCE(SUM(CASE WHEN status='Worked' THEN hours_worked ELSE 0 END),0) hw,
                    SUM(CASE WHEN status='Holiday' THEN 1 ELSE 0 END) hol,
                    SUM(CASE WHEN status='Sick' THEN 1 ELSE 0 END) sick,
                    SUM(CASE WHEN status='Maternity' THEN 1 ELSE 0 END) mat,
                    SUM(CASE WHEN status='Bank Holiday' THEN 1 ELSE 0 END) bh
                  FROM staff_attendance WHERE staff_id=? GROUP BY yr ORDER BY yr DESC""",
               (staff_id,), fetch=True) or []
        yr_html = "".join(
            f"<tr style='{'background:#eff6ff' if r['yr']==sel else ''}'><td class='mono' style='font-weight:700'>{r['yr']}</td>"
            f"<td class='mono' style='text-align:right'>{r['w']}</td>"
            f"<td class='mono' style='text-align:right'>{r['hw']:,.1f}</td>"
            f"<td class='mono' style='text-align:right'>{r['hol']}</td>"
            f"<td class='mono' style='text-align:right'>{r['sick']}</td>"
            f"<td class='mono' style='text-align:right'>{r['mat']}</td>"
            f"<td class='mono' style='text-align:right'>{r['bh']}</td></tr>" for r in yr)

        def _vt(v, label):
            on = (v == view)
            st = "background:#0f2942;color:#fff" if on else "background:#fff;color:#0f2942;border:1px solid #cbd5e1"
            return (f"<a href='/staff/{staff_id}/attendance?year={sel}&view={v}' style='{st};padding:4px 12px;"
                    f"border-radius:8px;font-weight:700;font-size:12px;text-decoration:none'>{label}</a>")
        toggle = (_vt("all", "All days") + " " + _vt("absences", "Absences only")) if is_owner else ""

        _col = {"Worked":"#16a34a","Holiday":"#0369a1","Sick":"#dc2626","Maternity":"#7c3aed","Bank Holiday":"#475569"}
        _bg  = {"Holiday":"#eff6ff","Sick":"#fef2f2","Maternity":"#faf5ff","Bank Holiday":"#f1f5f9"}
        vf = "AND status!='Worked'" if view == "absences" else ""
        det = q(f"""SELECT work_date,day,status,hours_worked,paid_hours,comments
                    FROM staff_attendance WHERE staff_id=? AND substr(work_date,1,4)=? {vf}
                    ORDER BY work_date DESC""", (staff_id, sel), fetch=True) or []
        if is_owner:
            det_head = ("<tr><th>Date</th><th>Day</th><th>Status</th><th style='text-align:right'>Hours</th>"
                        "<th style='text-align:right'>Paid</th><th>Comments</th></tr>")
            det_cols = 6
            rec_html = "".join(
                f"<tr style='background:{_bg.get(r['status'],'')}'>"
                f"<td class='mono' style='font-size:12px'>{_uk(r['work_date'])}</td>"
                f"<td style='font-size:12px'>{r['day'] or ''}</td>"
                f"<td><span style='color:{_col.get(r['status'],'#64748b')};font-weight:{'800' if r['status']!='Worked' else '600'};font-size:12px'>{r['status']}</span></td>"
                f"<td class='mono' style='text-align:right;font-size:12px'>{('%.2f'%r['hours_worked']) if r['hours_worked'] is not None else '—'}</td>"
                f"<td class='mono' style='text-align:right;font-size:12px'>{('%.2f'%r['paid_hours']) if r['paid_hours'] is not None else '—'}</td>"
                f"<td style='font-size:12px;color:#64748b'>{esc(r['comments'] or '')}</td></tr>" for r in det)
        else:
            # Managers: absence flags only — no hours, no comments.
            det_head = "<tr><th>Date</th><th>Day</th><th>Status</th></tr>"
            det_cols = 3
            rec_html = "".join(
                f"<tr style='background:{_bg.get(r['status'],'')}'>"
                f"<td class='mono' style='font-size:12px'>{_uk(r['work_date'])}</td>"
                f"<td style='font-size:12px'>{r['day'] or ''}</td>"
                f"<td><span style='color:{_col.get(r['status'],'#64748b')};font-weight:800;font-size:12px'>{r['status']}</span></td></tr>"
                for r in det)

        lifetime_block = f"""
        <div class='card' style='padding:0;overflow:hidden;margin-top:14px'>
          <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>By year (lifetime)</div>
          <div style='overflow-x:auto'><table class='tbl'>
            <thead><tr><th>Year</th><th style='text-align:right'>Worked days</th><th style='text-align:right'>Worked hrs</th>
              <th style='text-align:right'>Holiday</th><th style='text-align:right'>Sick</th>
              <th style='text-align:right'>Maternity</th><th style='text-align:right'>Bank Hol</th></tr></thead>
            <tbody>{yr_html}</tbody>
          </table></div>
        </div>""" if is_owner else ""

        _daily_title = f"Daily record — {sel}" if is_owner else f"Absences — {sel}"
        # Owner gets the year selector; managers are locked to the current year.
        year_bar = (f"""
        <div class='card' style='margin-top:14px;padding:12px 16px;display:flex;gap:8px;flex-wrap:wrap;align-items:center'>
          <span style='font-size:13px;font-weight:800;color:#334155;margin-right:4px'>Year:</span>{chips}
        </div>""" if is_owner else "")
        body = f"""
        {year_bar}
        {cards}
        {leave_block}
        {lifetime_block}
        <div class='card' style='padding:0;overflow:hidden;margin-top:14px'>
          <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
            <span>{_daily_title}</span><span style='display:flex;gap:6px'>{toggle}</span>
          </div>
          <div style='overflow-x:auto'><table class='tbl'>
            <thead>{det_head}</thead>
            <tbody>{rec_html or f"<tr><td colspan='{det_cols}' style='text-align:center;padding:20px;color:#94a3b8'>No entries for this selection</td></tr>"}</tbody>
          </table></div>
        </div>"""

    content = f"""
    <div>
      <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {esc(name)}</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>🕒 Attendance — {esc(name)}</div>
      <div style='color:#64748b;font-size:13px'>{tot['c']:,} days on record{(' · ' + str(tot['mn']) + ' → ' + str(tot['mx'])) if tot['c'] else ''}</div>
    </div>
    {body}"""
    return page(f"Attendance — {name}", content, user, "staff")


@router.get("/staff/{staff_id}/set-entitlement", response_class=HTMLResponse)
def set_entitlement_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # entitlement overrides are owner-only
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}", status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"
    year = datetime.now().year
    contracted = s.get("contracted_hrs") or 0
    # Statutory minimum is always 5.6 weeks x 5 days = 28 days
    # (regardless of hours — part-timers get pro-rated days but
    #  the day count is the same; the difference is day length)
    statutory  = 28.0
    # Pro-rata = 5.6 weeks × contracted hours ÷ 5 (actual days entitlement)
    pro_rata   = round(5.6 * contracted / 5, 1) if contracted else 28.0
    existing   = q("SELECT * FROM leave_entitlements WHERE staff_id=? AND year=?",
                   (staff_id, year), fetch=True)
    current    = dict(existing[0]) if existing else {}
    cur_val    = current.get("custom_days") or pro_rata

    html  = f"""
    <div>
      <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>
        &larr; Back to {name}</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>
        &#9881;&#65039; Leave Entitlement &mdash; {name} ({year})</div>
    </div>
    <div class='card' style='max-width:520px'>
      <div style='background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;
                  padding:12px 16px;margin-bottom:16px;font-size:13px'>
        <div><strong>Contracted hours:</strong> {contracted}h/week</div>
        <div><strong>Statutory minimum:</strong> 28 days (5.6 weeks)</div>
        <div style='font-size:12px;color:#64748b;margin-top:4px'>
          For part-time staff you can set below 28 days — pro-rata
          is acceptable (e.g. 3 days/wk = 5.6 × 3 = 16.8 days minimum).
          Set whatever you have agreed contractually.
        </div>
      </div>
      <form action='/staff/{staff_id}/set-entitlement' method='POST' class='space-y-4'>
        <input type='hidden' name='year' value='{year}'>
        <div>
          <label>Custom Entitlement (days) for {year}</label>
          <div style='display:flex;gap:8px;align-items:center'>
            <input type='number' step='0.5' name='custom_days' id='custom_days_input'
                   value='{cur_val}' min='1' style='flex:1'>
            <button type='button'
              onclick="document.getElementById('custom_days_input').value='{pro_rata}';
                       document.getElementById('notes_input').value='Reset to statutory pro-rata';"
              class='btn-secondary' style='white-space:nowrap;padding:8px 14px;font-size:12px'>
              🔄 Reset to Statutory ({pro_rata} days)
            </button>
          </div>
          <div style='font-size:11px;color:#94a3b8;margin-top:4px'>
            Statutory pro-rata for {contracted}h/wk =
            5.6 weeks × {contracted}h ÷ 5 days/wk = {pro_rata} days
          </div>
        </div>
        <div>
          <label>Notes</label>
          <input type='text' name='notes' id='notes_input'
                 value='{current.get("notes") or ""}'
                 placeholder='e.g. Pro-rated, contractual agreement'>
        </div>
        <div style='display:flex;gap:8px'>
          <button type='submit' class='btn-primary'>💾 Save Entitlement</button>
          <a href='/staff/{staff_id}' class='btn-secondary'>Cancel</a>
        </div>
      </form>
    </div>"""
    return page(f"Entitlement", html, user, "staff")


@router.post("/staff/{staff_id}/set-entitlement")
async def save_entitlement(staff_id: int, request: Request,
                           session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # entitlement overrides are owner-only
    form        = await request.form()
    year        = int(form.get("year", datetime.now().year))
    custom_days = float(form.get("custom_days", 0) or 0)
    notes       = str(form.get("notes","") or "").strip()
    s           = q("SELECT contracted_hrs FROM staff_profiles WHERE staff_id=?",
                    (staff_id,), fetch=True)
    contracted  = dict(s[0])["contracted_hrs"] if s else 0
    statutory   = 28.0  # Always 28 days statutory minimum
    q("""INSERT INTO leave_entitlements
            (staff_id,year,statutory_days,custom_days,effective_days,notes)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,year) DO UPDATE SET
            custom_days=excluded.custom_days,
            effective_days=excluded.effective_days,
            notes=excluded.notes""",
      (staff_id, year, statutory, custom_days, custom_days, notes or None))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}?msg={uq('Entitlement updated')}",
                            status_code=303)


def _doc_kind(path: str) -> str:
    """Human label for a document's file type, from its extension."""
    ext = os.path.splitext(path or "")[1].lower()
    return {".pdf": "PDF", ".jpg": "Image", ".jpeg": "Image", ".png": "Image",
            ".docx": "Word", ".doc": "Word", ".dotx": "Word"}.get(ext, (ext.lstrip(".").upper() or "File"))


def _kind_chip(path: str) -> str:
    kind = _doc_kind(path)
    bg, fg = {"PDF": ("#fee2e2", "#991b1b"), "Word": ("#dbeafe", "#1d4ed8"),
              "Image": ("#dcfce7", "#166534")}.get(kind, ("#f1f5f9", "#64748b"))
    return (f"<span style='background:{bg};color:{fg};font-size:10px;font-weight:700;"
            f"padding:2px 6px;border-radius:4px'>{kind}</span>")


def _doc_action_buttons(staff_id: int, d: dict, is_owner: bool, small: bool = False) -> str:
    """View (PDFs/images only — Word can't preview in a browser) + Download + owner-only Delete."""
    pad  = "3px 8px" if small else "4px 10px"
    kind = _doc_kind(d["file_path"])
    btns = ""
    if kind in ("PDF", "Image"):
        btns += (f"<a href='/staff/{staff_id}/documents/{d['doc_id']}/view' target='_blank' "
                 f"class='btn-secondary' style='padding:{pad};font-size:11px'>👁 View</a>")
    btns += (f"<a href='/staff/{staff_id}/documents/{d['doc_id']}/download' "
             f"class='btn-secondary' style='padding:{pad};font-size:11px'>⬇️ Download</a>")
    if is_owner:
        btns += (f"<form method='POST' action='/staff/{staff_id}/documents/{d['doc_id']}/delete' "
                 f"style='display:inline' onsubmit=\"return confirm('Delete v{d['version']} of "
                 f"{d['doc_type']}? This cannot be undone.');\">"
                 f"<button type='submit' class='btn-danger' style='padding:{pad};font-size:11px'>🗑 Delete</button></form>")
    return btns


RTW_ID_TYPES = [
    "British/Irish Passport",
    "Passport + Visa/BRP",
    "Biometric Residence Permit (BRP)",
    "Driving Licence",
    "Birth Certificate + NI proof",
    "Online share-code check",
    "Other",
]


def get_rtw_check(staff_id: int) -> dict:
    """Current Right-to-Work / ID check record for a staff member (or {})."""
    r = q("SELECT * FROM staff_rtw_checks WHERE staff_id=?", (staff_id,), fetch=True)
    return dict(r[0]) if r else {}


def _rtw_status_summary(rtw: dict) -> str:
    """Short chip describing RTW status — reused on profile + documents page."""
    if not rtw:
        return "<span style='background:#fee2e2;color:#991b1b;font-weight:800;font-size:11px;padding:2px 8px;border-radius:6px'>RTW not recorded</span>"
    if rtw.get("rtw_confirmed"):
        base = "<span style='background:#dcfce7;color:#166534;font-weight:800;font-size:11px;padding:2px 8px;border-radius:6px'>✓ Right to work confirmed</span>"
    else:
        base = "<span style='background:#fef9c3;color:#854d0e;font-weight:800;font-size:11px;padding:2px 8px;border-radius:6px'>RTW recorded — not confirmed</span>"
    exp = rtw.get("expiry_date")
    if exp:
        try:
            days = (date.fromisoformat(exp) - date.today()).days
            if days < 0:
                base += f" <span style='color:#dc2626;font-weight:800;font-size:11px'>⚠ expired {exp}</span>"
            elif days <= 60:
                base += f" <span style='color:#d97706;font-weight:800;font-size:11px'>⚠ expires {exp}</span>"
        except ValueError:
            pass
    return base


def _rtw_panel_html(staff_id: int, rtw: dict, is_mgr: bool) -> str:
    """The Right to Work / ID Check panel on the Documents page."""
    if rtw:
        exp = rtw.get("expiry_date")
        exp_html = rtw.get("expiry_date") or "—"
        if exp:
            try:
                days = (date.fromisoformat(exp) - date.today()).days
                if days < 0:
                    exp_html = f"<span style='color:#dc2626;font-weight:800'>⚠ {exp} (expired)</span>"
                elif days <= 60:
                    exp_html = f"<span style='color:#d97706;font-weight:800'>⚠ {exp} ({days}d left)</span>"
            except ValueError:
                pass
        details = f"""
          <div style='display:grid;grid-template-columns:repeat(auto-fit,minmax(150px,1fr));gap:8px;font-size:13px;margin-top:12px'>
            <div><span style='color:#94a3b8;font-weight:700;font-size:11px;text-transform:uppercase'>ID seen</span><br>{esc(rtw.get('id_type') or '—')}</div>
            <div><span style='color:#94a3b8;font-weight:700;font-size:11px;text-transform:uppercase'>Date checked</span><br>{rtw.get('check_date') or '—'}</div>
            <div><span style='color:#94a3b8;font-weight:700;font-size:11px;text-transform:uppercase'>Expiry</span><br>{exp_html}</div>
          </div>
          {f"<div style='font-size:12px;color:#64748b;margin-top:8px'>{esc(rtw.get('notes'))}</div>" if rtw.get('notes') else ''}
          <div style='font-size:11px;color:#94a3b8;margin-top:6px'>Recorded by {esc(rtw.get('checked_by') or '—')}</div>"""
    else:
        details = "<div style='color:#94a3b8;font-size:13px;margin-top:8px'>No right-to-work / ID check recorded yet.</div>"

    form_html = ""
    if is_mgr:
        opts = "".join(f"<option {'selected' if rtw.get('id_type') == t else ''}>{t}</option>" for t in RTW_ID_TYPES)
        form_html = f"""
        <details style='margin-top:12px'>
          <summary style='cursor:pointer;font-weight:700;font-size:13px;color:#1e3a5f'>{'Update' if rtw else 'Record'} right-to-work / ID check</summary>
          <form method='POST' action='/staff/{staff_id}/rtw-check' style='margin-top:10px;display:grid;gap:10px;grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
            <label style='font-size:12px;color:#475569'>ID document seen
              <select name='id_type' style='width:100%'>{opts}</select></label>
            <label style='font-size:12px;color:#475569'>Date checked
              <input type='date' name='check_date' value="{rtw.get('check_date') or ''}" style='width:100%'></label>
            <label style='font-size:12px;color:#475569'>Expiry (if time-limited)
              <input type='date' name='expiry_date' value="{rtw.get('expiry_date') or ''}" style='width:100%'></label>
            <label style='font-size:12px;color:#475569;grid-column:1/-1'>What was seen / notes
              <input type='text' name='notes' value="{esc(rtw.get('notes') or '')}" placeholder="e.g. original British passport seen in person; photo + name/DOB matched" style='width:100%'></label>
            <label style='font-size:13px;font-weight:700;color:#166534;display:flex;align-items:center;gap:6px'>
              <input type='checkbox' name='rtw_confirmed' value='1' {'checked' if rtw.get('rtw_confirmed') else ''}> Right to work confirmed</label>
            <div style='grid-column:1/-1'><button type='submit' class='btn-primary' style='padding:6px 16px;font-size:13px'>💾 Save check</button></div>
          </form>
        </details>"""

    return f"""
    <div class='card' style='border-left:4px solid #1e3a5f'>
      <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
        <div style='font-weight:900;color:#0f2942;font-size:14px'>🛡️ Right to Work / ID Check</div>
        {_rtw_status_summary(rtw)}
      </div>
      {details}
      {form_html}
      <div style='font-size:11px;color:#94a3b8;margin-top:10px'>🔒 Photo-ID originals are kept offline — this panel is the record that the check was done.</div>
    </div>"""


@router.get("/staff/{staff_id}/documents", response_class=HTMLResponse)
def staff_documents(
    staff_id: int,
    session:  str | None = Cookie(default=None),
    msg:      str = "",
    msg_type: str = "success"
):
    redir, user = require_login(session)
    if redir: return redir
    # Documents are not a staff self-service area — owner + own-store manager only.
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"
    is_owner = user["role"] == "owner"

    # Get all documents for this staff member
    docs = q("""SELECT * FROM staff_documents WHERE staff_id=? AND deleted_at IS NULL
                ORDER BY doc_type, version DESC""",
             (staff_id,), fetch=True) or []

    # Get available templates
    templates = q("SELECT * FROM document_templates WHERE is_current=1 ORDER BY doc_type",
                  fetch=True) or []
    template_types = {dict(t)["doc_type"] for t in templates}

    flash = f"<div class='flash-{'success' if msg_type=='success' else 'error'}'>{msg}</div>" if msg else ""

    # Group docs by type
    from collections import defaultdict
    by_type = defaultdict(list)
    for d in docs:
        by_type[dict(d)["doc_type"]].append(dict(d))

    # Build document cards
    doc_cards = ""
    viewer_is_mgr = user["role"] in ("owner", "manager")
    for dtype in DOC_TYPES:
        # MANAGER: capture-only — only the allowed types, and NO view/download, no
        # versions, no auto-fill; just an upload box + a "✓ on file" tick.
        if not is_owner:
            if dtype not in MANAGER_CAPTURE_DOC_TYPES:
                continue
            _cur = next((d for d in by_type.get(dtype, []) if d["is_current"]), None)
            _tick = ("<span style='color:#16a34a;font-weight:700;font-size:13px'>&#10003; on file</span>"
                     if _cur else "<span style='color:#94a3b8;font-size:13px'>&mdash; not captured yet</span>")
            doc_cards += f"""
        <div class='card'>
          <div style='font-weight:900;color:#0f2942;margin-bottom:6px;font-size:14px'>{dtype}</div>
          <div style='margin-bottom:8px'>{_tick}</div>
          <form action='/staff/{staff_id}/documents/upload' method='POST' enctype='multipart/form-data' style='display:flex;gap:6px'>
            <input type='hidden' name='doc_type' value='{dtype}'>
            <input type='file' name='doc_file' accept='.pdf,.doc,.docx' style='flex:1;font-size:12px;padding:4px 8px'>
            <button type='submit' class='btn-primary' style='padding:4px 12px;font-size:11px;white-space:nowrap'>&#8593;&#65039; Upload</button>
          </form>
        </div>"""
            continue
        # ---- OWNER: full document card (view/download, versions, auto-fill) ----
        type_docs = by_type.get(dtype, [])
        has_template = dtype in template_types

        # Current version
        current = next((d for d in type_docs if d["is_current"]), None)
        older   = [d for d in type_docs if not d["is_current"]]

        current_html = ""
        if current:
            gen_badge = "<span style='background:#dbeafe;color:#1d4ed8;font-size:10px;font-weight:700;padding:2px 6px;border-radius:4px'>AUTO-GENERATED</span>" if current["generated"] else ""
            current_html = f"""
            <div style='background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:10px 14px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
              <div style='min-width:0'>
                <div style='font-size:13px;font-weight:700;color:#166534'>
                  ✅ {esc(current.get('file_name') or 'Document')} {_kind_chip(current['file_path'])} {gen_badge}
                </div>
                <div style='font-size:11px;color:#64748b'>v{current['version']} · added {current['uploaded_at'][:10]}{(' · ' + esc(current['notes'])) if current.get('notes') else ''}</div>
              </div>
              <div style='display:flex;gap:6px;flex-wrap:wrap'>
                {_doc_action_buttons(staff_id, current, is_owner)}
              </div>
            </div>"""

        older_html = ""
        if older:
            older_html = "<div style='margin-top:6px'>"
            for od in older:
                older_html += f"""
                <div style='display:flex;justify-content:space-between;align-items:center;padding:6px 10px;font-size:12px;color:#64748b;border-bottom:1px solid #f1f5f9;gap:8px;flex-wrap:wrap'>
                  <span style='min-width:0'>{esc(od.get('file_name') or 'Document')} <span style='color:#94a3b8'>· v{od['version']} · {od['uploaded_at'][:10]}</span> {_kind_chip(od['file_path'])}</span>
                  <span style='display:flex;gap:6px;flex-wrap:wrap'>{_doc_action_buttons(staff_id, od, is_owner, small=True)}</span>
                </div>"""
            older_html += "</div>"

        # Upload / generate form
        action_html = f"""
        <div style='margin-top:10px;padding-top:10px;border-top:1px solid #f1f5f9'>
          <div style='display:flex;gap:8px;flex-wrap:wrap'>
            <form action='/staff/{staff_id}/documents/upload' method='POST'
                  enctype='multipart/form-data' style='display:flex;gap:6px;flex:1;min-width:200px'>
              <input type='hidden' name='doc_type' value='{dtype}'>
              <input type='file' name='doc_file' accept='.pdf,.doc,.docx,.dotx'
                     style='flex:1;font-size:12px;padding:4px 8px'>
              <button type='submit' class='btn-primary' style='padding:4px 12px;font-size:11px;white-space:nowrap'>
                ⬆️ Upload
              </button>
            </form>
            {"<a href='/staff/" + str(staff_id) + "/documents/generate?doc_type=" + dtype + "' class='btn-secondary' style='padding:4px 12px;font-size:11px;white-space:nowrap'>⚡ Auto-fill</a>" if has_template else ""}
          </div>
        </div>"""

        doc_cards += f"""
        <div class='card'>
          <div style='font-weight:900;color:#0f2942;margin-bottom:8px;font-size:14px'>{dtype}</div>
          {current_html or "<div style='color:#94a3b8;font-size:13px'>No document uploaded yet</div>"}
          {older_html}
          {action_html}
        </div>"""

    bin_n = q("SELECT COUNT(*) c FROM staff_documents WHERE staff_id=? AND deleted_at IS NOT NULL",
              (staff_id,), fetch=True)
    bin_n = bin_n[0]["c"] if bin_n else 0
    bin_btn = (f"<a href='/staff/{staff_id}/documents/bin' class='btn-secondary'>&#128465; Recycle Bin ({bin_n})</a>"
               if user["role"] == "owner" and bin_n else "")
    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {name}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>📁 Documents — {name}</div>
      </div>
      <div style='display:flex;gap:8px'>
        {bin_btn}
        {'<a href="/staff/document-templates" class="btn-secondary">📋 Manage Templates</a>' if user["role"] == "owner" else ''}
      </div>
    </div>
    {_rtw_panel_html(staff_id, get_rtw_check(staff_id), True) if is_owner else ''}
    <div style='display:grid;gap:12px;grid-template-columns:repeat(auto-fill,minmax(400px,1fr))'>
      {doc_cards}
    </div>"""

    return page(f"Documents — {name}", content, user, "staff")


@router.post("/staff/{staff_id}/rtw-check")
async def save_rtw_check(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r

    form = await request.form()
    id_type   = str(form.get("id_type", "") or "").strip()
    confirmed = 1 if form.get("rtw_confirmed") else 0
    check_date = str(form.get("check_date", "") or "").strip() or None
    expiry     = str(form.get("expiry_date", "") or "").strip() or None
    evidence   = str(form.get("evidence_location", "") or "").strip() or None
    notes      = str(form.get("notes", "") or "").strip() or None

    q("""INSERT INTO staff_rtw_checks
            (staff_id, id_type, rtw_confirmed, check_date, expiry_date,
             evidence_location, notes, checked_by)
         VALUES(?,?,?,?,?,?,?,?)
         ON CONFLICT(staff_id) DO UPDATE SET
            id_type=excluded.id_type, rtw_confirmed=excluded.rtw_confirmed,
            check_date=excluded.check_date, expiry_date=excluded.expiry_date,
            evidence_location=excluded.evidence_location, notes=excluded.notes,
            checked_by=excluded.checked_by""",
      (staff_id, id_type, confirmed, check_date, expiry, evidence, notes,
       user.get("username")))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/documents?msg={uq('Right-to-work / ID check saved')}",
        status_code=303)


@router.post("/staff/{staff_id}/documents/upload")
async def upload_staff_doc(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r

    form     = await request.form()
    doc_type = form.get("doc_type","Other")
    # Managers may only CAPTURE the allowed onboarding/absence types.
    if user["role"] != "owner" and doc_type not in MANAGER_CAPTURE_DOC_TYPES:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/documents?msg={uq('You can only add sick-return and onboarding documents')}&msg_type=error",
            status_code=303)
    doc_file = form.get("doc_file")
    notes    = str(form.get("notes","") or "").strip()

    if not doc_file or not hasattr(doc_file, "filename") or not doc_file.filename:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/documents?msg={uq('No file selected')}&msg_type=error",
            status_code=303)

    # Get next version number
    existing = q("""SELECT MAX(version) as v FROM staff_documents
                    WHERE staff_id=? AND doc_type=?""",
                 (staff_id, doc_type), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1

    # Mark previous versions as not current
    q("UPDATE staff_documents SET is_current=0 WHERE staff_id=? AND doc_type=?",
      (staff_id, doc_type))

    # Save file (sanitise name parts to prevent path traversal; whitelist ext; cap size)
    ext      = os.path.splitext(doc_file.filename)[1].lower()
    filename = f"staff_{staff_id}_{_safe_part(doc_type)}_v{next_ver}{_safe_ext(ext)}"
    filepath = os.path.join(DOCS_DIR, filename)
    data = await doc_file.read()
    if len(data) > 25 * 1024 * 1024:
        from urllib.parse import quote as uq
        return RedirectResponse(f"/staff/{staff_id}/documents?msg={uq('File too large (max 25 MB)')}&msg_type=error",
                                status_code=303)
    os.makedirs(DOCS_DIR, exist_ok=True)
    with open(filepath, "wb") as f:
        f.write(data)

    q("""INSERT INTO staff_documents
            (staff_id, doc_type, version, file_path, file_name,
             is_current, generated, uploaded_by, notes)
         VALUES(?,?,?,?,?,1,0,?,?)""",
      (staff_id, doc_type, next_ver, filepath,
       doc_file.filename, user.get("username"), notes or None))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/documents?msg={uq(doc_type + ' uploaded successfully')}",
        status_code=303)


@router.get("/staff/{staff_id}/documents/generate", response_class=HTMLResponse)
def generate_doc_form(
    staff_id: int,
    doc_type: str = "",
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] not in ("owner","manager"):
        return RedirectResponse(f"/staff/{staff_id}/documents", status_code=303)

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get available templates
    templates = q("SELECT * FROM document_templates WHERE is_current=1 ORDER BY doc_type",
                  fetch=True) or []

    # Show merge fields preview
    fields     = get_merge_fields(s)
    fields_html = ""
    for k, v in fields.items():
        fields_html += f"""
        <div style='display:flex;gap:12px;padding:4px 0;border-bottom:1px solid #f1f5f9;font-size:12px'>
          <code style='color:#7c3aed;min-width:180px'>{k}</code>
          <span style='color:#334155'>{v or '—'}</span>
        </div>"""

    type_opts = ""
    for t in templates:
        td = dict(t)
        sel = "selected" if td["doc_type"] == doc_type else ""
        type_opts += f'<option value="{td["doc_type"]}" {sel}>{td["doc_type"]} (v{td["version"]})</option>'

    content = f"""
    <div>
      <a href='/staff/{staff_id}/documents' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Documents</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>⚡ Auto-Generate Document — {name}</div>
    </div>
    <div class='grid gap-6' style='grid-template-columns:1fr 1fr'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Generate Document</div>
        <form action='/staff/{staff_id}/documents/generate' method='POST' class='space-y-4'>
          <div>
            <label>Document Type</label>
            <select name='doc_type' required>
              <option value=''>-- Select template --</option>
              {type_opts}
            </select>
          </div>
          <div>
            <label>Notes (optional)</label>
            <input type='text' name='notes' placeholder='e.g. Initial offer, Updated contract'>
          </div>
          <button type='submit' class='btn-primary'>⚡ Generate & Download</button>
        </form>
        {'<div class="flash-error" style="margin-top:12px">No templates uploaded yet. <a href=\'  /staff/document-templates\' style=\'color:#1e3a5f;font-weight:700\'>Upload templates here →</a></div>' if not templates else ''}
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Available Merge Fields</div>
        <div style='font-size:11px;color:#64748b;margin-bottom:8px'>
          Use these placeholders in your Word template — they will be replaced with this staff member's details.
        </div>
        <div style='max-height:400px;overflow-y:auto'>
          {fields_html}
        </div>
      </div>
    </div>"""

    return page("Generate Document", content, user, "staff")


@router.post("/staff/{staff_id}/documents/generate")
async def generate_doc(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_owner(user)): return r   # auto-fill templates are owner-only

    form     = await request.form()
    doc_type = form.get("doc_type","")
    notes    = str(form.get("notes","") or "").strip()

    # Get staff details
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0])

    # Get template
    tmpl = q("SELECT * FROM document_templates WHERE doc_type=? AND is_current=1",
             (doc_type,), fetch=True)
    if not tmpl:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/documents?msg={uq('No template found for ' + doc_type)}&msg_type=error",
            status_code=303)
    tmpl = dict(tmpl[0])

    # Fill template
    fields   = get_merge_fields(s)
    doc_bytes = fill_word_template(tmpl["file_path"], fields)

    # Save generated file
    existing = q("""SELECT MAX(version) as v FROM staff_documents
                    WHERE staff_id=? AND doc_type=?""",
                 (staff_id, doc_type), fetch=True)
    next_ver = (existing[0]["v"] or 0) + 1 if existing else 1
    q("UPDATE staff_documents SET is_current=0 WHERE staff_id=? AND doc_type=?",
      (staff_id, doc_type))

    filename = f"staff_{staff_id}_{_safe_part(doc_type)}_v{next_ver}.docx"
    filepath = os.path.join(DOCS_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(doc_bytes)

    q("""INSERT INTO staff_documents
            (staff_id, doc_type, version, file_path, file_name,
             is_current, generated, uploaded_by, notes)
         VALUES(?,?,?,?,?,1,1,?,?)""",
      (staff_id, doc_type, next_ver, filepath, filename,
       user.get("username"), notes or None))

    # Return the file for download
    name = f"{s['first_name']} {s['last_name']}"
    download_name = f"{doc_type} - {name}.docx"
    return FileResponse(filepath, filename=download_name,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/staff/{staff_id}/documents/{doc_id}/download")
def download_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id)): return r
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=?",
             (doc_id, staff_id), fetch=True)
    if not rows: return HTMLResponse("<p>Document not found</p>", status_code=404)
    d = dict(rows[0])
    if user["role"] != "owner":   # only the owner may open/download a document
        return HTMLResponse("<p>Not available</p>", status_code=403)
    if not os.path.exists(d["file_path"]):
        return HTMLResponse("<p>File not found on disk</p>", status_code=404)
    ext = os.path.splitext(d["file_path"])[1].lower()
    media = "application/pdf" if ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(d["file_path"], filename=d["file_name"] or os.path.basename(d["file_path"]),
                        media_type=media)


@router.get("/staff/{staff_id}/documents/{doc_id}/view")
def view_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id)): return r
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=?",
             (doc_id, staff_id), fetch=True)
    if not rows: return HTMLResponse("<p>Document not found</p>", status_code=404)
    d = dict(rows[0])
    if user["role"] != "owner":   # only the owner may open/download a document
        return HTMLResponse("<p>Not available</p>", status_code=403)
    if not os.path.exists(d["file_path"]):
        return HTMLResponse("<p>File not found on disk</p>", status_code=404)
    # Serve with the file's REAL type: PDFs/images preview inline in the browser;
    # Word docs (.docx) can't preview inline, so they open/download in Word. Forcing
    # application/pdf on a .docx is what made the viewer fail with "Failed to load PDF".
    ext   = os.path.splitext(d["file_path"])[1].lower()
    media = {".pdf": "application/pdf", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
             ".png": "image/png",
             ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             ".doc":  "application/msword"}.get(ext, "application/octet-stream")
    return FileResponse(d["file_path"], media_type=media)


@router.post("/staff/{staff_id}/documents/{doc_id}/delete")
def delete_staff_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    """SOFT-delete one filed document version to the Recycle Bin. Owner-only +
    version-aware: the file is KEPT (recoverable), the row is flagged deleted, and
    if it was the current version the newest remaining LIVE version is promoted to
    current so the record is never left broken."""
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/documents?msg=Only+the+owner+can+delete+documents&msg_type=error",
                                status_code=303)
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=? AND deleted_at IS NULL",
             (doc_id, staff_id), fetch=True)
    if not rows:
        return RedirectResponse(f"/staff/{staff_id}/documents", status_code=303)
    d = dict(rows[0])
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # Move to bin (keep the file), clear its current flag.
    q("UPDATE staff_documents SET deleted_at=?, deleted_by=?, is_current=0 WHERE doc_id=?",
      (now, user.get("username", ""), doc_id))
    # If it was the current version, promote the newest remaining LIVE version.
    if d["is_current"]:
        rem = q("""SELECT doc_id FROM staff_documents
                   WHERE staff_id=? AND doc_type=? AND deleted_at IS NULL
                   ORDER BY version DESC LIMIT 1""", (staff_id, d["doc_type"]), fetch=True)
        if rem:
            q("UPDATE staff_documents SET is_current=1 WHERE doc_id=?", (dict(rem[0])["doc_id"],))
    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/documents?msg={uq('Moved to Recycle Bin — you can restore it from there')}", status_code=303)


@router.post("/staff/{staff_id}/documents/{doc_id}/restore")
def restore_staff_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    """Owner-only: bring a document back from the Recycle Bin. It returns as a
    (non-current) version unless nothing of its type is current, in which case it
    becomes current again."""
    redir, user = require_login(session)
    if redir: return redir
    from urllib.parse import quote as uq
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/documents?msg=Owner+only&msg_type=error", status_code=303)
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=? AND deleted_at IS NOT NULL",
             (doc_id, staff_id), fetch=True)
    if not rows:
        return RedirectResponse(f"/staff/{staff_id}/documents/bin", status_code=303)
    d = dict(rows[0])
    q("UPDATE staff_documents SET deleted_at=NULL, deleted_by=NULL WHERE doc_id=?", (doc_id,))
    # If no live version of this type is current, make the restored one current.
    has_current = q("""SELECT 1 FROM staff_documents WHERE staff_id=? AND doc_type=?
                       AND is_current=1 AND deleted_at IS NULL LIMIT 1""",
                    (staff_id, d["doc_type"]), fetch=True)
    if not has_current:
        q("UPDATE staff_documents SET is_current=1 WHERE doc_id=?", (doc_id,))
    return RedirectResponse(
        f"/staff/{staff_id}/documents/bin?msg={uq('Document restored')}", status_code=303)


@router.post("/staff/{staff_id}/documents/{doc_id}/purge")
def purge_staff_doc(staff_id: int, doc_id: int, session: str | None = Cookie(default=None)):
    """Owner-only: permanently delete a document already in the Recycle Bin
    (removes the file from disk + the DB row). Only reachable for binned docs."""
    redir, user = require_login(session)
    if redir: return redir
    from urllib.parse import quote as uq
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/documents?msg=Owner+only&msg_type=error", status_code=303)
    rows = q("SELECT * FROM staff_documents WHERE doc_id=? AND staff_id=? AND deleted_at IS NOT NULL",
             (doc_id, staff_id), fetch=True)
    if not rows:   # never hard-delete something that isn't already in the bin
        return RedirectResponse(f"/staff/{staff_id}/documents/bin", status_code=303)
    d = dict(rows[0])
    try:
        if d["file_path"] and os.path.exists(d["file_path"]):
            os.remove(d["file_path"])
    except Exception:
        pass
    q("DELETE FROM staff_documents WHERE doc_id=?", (doc_id,))
    return RedirectResponse(
        f"/staff/{staff_id}/documents/bin?msg={uq('Document permanently deleted')}", status_code=303)


@router.get("/staff/{staff_id}/documents/bin", response_class=HTMLResponse)
def documents_bin(staff_id: int, session: str | None = Cookie(default=None),
                  msg: str = "", msg_type: str = "success"):
    """Owner-only Recycle Bin — restore a deleted document, or delete it forever."""
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/documents?msg=Owner+only&msg_type=error", status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s = dict(rows[0]); name = f"{s['first_name']} {s['last_name']}"
    binned = q("""SELECT * FROM staff_documents WHERE staff_id=? AND deleted_at IS NOT NULL
                  ORDER BY deleted_at DESC""", (staff_id,), fetch=True) or []
    flash = f"<div class='flash-{'success' if msg_type=='success' else 'error'}'>{esc(msg)}</div>" if msg else ""
    rows_html = ""
    for d in (dict(x) for x in binned):
        view = (f"<a href='/staff/{staff_id}/documents/{d['doc_id']}/view' target='_blank' "
                f"class='btn-secondary' style='padding:4px 10px;font-size:12px'>View</a>")
        restore = (f"<form method='POST' action='/staff/{staff_id}/documents/{d['doc_id']}/restore' style='display:inline'>"
                   f"<button class='btn-secondary' style='padding:4px 10px;font-size:12px;color:#16a34a'>&#8617; Restore</button></form>")
        purge = (f"<form method='POST' action='/staff/{staff_id}/documents/{d['doc_id']}/purge' style='display:inline' "
                 f"onsubmit=\"return confirm('Permanently delete this document? This CANNOT be undone.');\">"
                 f"<button class='btn-danger' style='padding:4px 10px;font-size:12px'>&#128465; Delete forever</button></form>")
        rows_html += (f"<tr><td style='font-weight:700'>{esc(d['doc_type'])}</td>"
                      f"<td class='mono' style='font-size:12px'>v{d.get('version','')}</td>"
                      f"<td style='font-size:12px;color:#64748b'>{esc(d.get('file_name') or '')}</td>"
                      f"<td class='mono' style='font-size:12px;color:#64748b'>{esc(d.get('deleted_at') or '')}</td>"
                      f"<td style='font-size:12px;color:#64748b'>{esc(d.get('deleted_by') or '')}</td>"
                      f"<td><div style='display:flex;gap:6px'>{view}{restore}{purge}</div></td></tr>")
    content = f"""
    {flash}
    <div>
      <a href='/staff/{staff_id}/documents' style='color:#1e3a5f;font-size:13px;font-weight:700'>&larr; Back to Documents</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>&#128465; Recycle Bin &mdash; {esc(name)}</div>
      <div style='font-size:13px;color:#64748b;margin-top:2px'>Deleted documents are kept here so an accidental delete can be undone.
      <strong>Restore</strong> brings one back; <strong>Delete forever</strong> removes it permanently.</div>
    </div>
    <div class='card' style='padding:0;overflow:hidden;margin-top:12px'>
      <div style='overflow-x:auto'><table class='tbl'>
        <thead><tr><th>Type</th><th>Version</th><th>File</th><th>Deleted</th><th>By</th><th>Action</th></tr></thead>
        <tbody>{rows_html or "<tr><td colspan='6' style='text-align:center;padding:24px;color:#94a3b8'>Recycle bin is empty</td></tr>"}</tbody>
      </table></div>
    </div>"""
    return page(f"Recycle Bin — {name}", content, user, "staff")


def ensure_onboarding_tables():
    conn = db()
    c    = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS onboarding_forms (
            form_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id       INTEGER NOT NULL,
            form_type      TEXT NOT NULL,
            status         TEXT DEFAULT 'not_started',
            started_at     TEXT,
            completed_at   TEXT,
            form_data      TEXT,
            pdf_path       TEXT,
            UNIQUE(staff_id, form_type),
            FOREIGN KEY (staff_id) REFERENCES staff_profiles(staff_id)
        )
    """)
    conn.commit()
    conn.close()


ONBOARD_FORMS = [
    ("employment_application", "Employment Application",  "staff"),
    ("p46",                    "P46 Tax Form",            "staff"),
    ("new_employee_notify",    "New Employee Notification","owner"),
    ("offer_letter",           "Offer Letter",            "owner"),
    ("employment_contract",    "Employment Contract",     "owner"),
    ("right_to_work",          "Right to Work Checked",  "owner"),
]


DIGITAL_FORMS = {"employment_application", "p46", "new_employee_notify"}


def get_onboarding_status(staff_id: int) -> dict:
    """Return completion status for each onboarding form and document."""
    rows = q("SELECT form_type, status FROM onboarding_forms WHERE staff_id=?",
             (staff_id,), fetch=True) or []
    status_map = {dict(r)["form_type"]: dict(r)["status"] for r in rows}

    # Check staff_documents for document-based items
    doc_rows = q("SELECT doc_type, is_current FROM staff_documents WHERE staff_id=? AND is_current=1 AND deleted_at IS NULL",
                 (staff_id,), fetch=True) or []
    doc_types = {dict(d)["doc_type"] for d in doc_rows}

    # Map document types to onboarding form types
    doc_type_map = {
        "offer_letter":        "Offer Letter",
        "employment_contract": "Employment Contract",
        "right_to_work":       "Right to Work",
    }

    result = {}
    for ftype, flabel, fwho in ONBOARD_FORMS:
        # Check if it's a document-based item
        if ftype in doc_type_map:
            doc_label = doc_type_map[ftype]
            status = "completed" if doc_label in doc_types else status_map.get(ftype, "not_started")
        else:
            status = status_map.get(ftype, "not_started")
        result[ftype] = {
            "label":   flabel,
            "who":     fwho,
            "status":  status,
            "is_doc":  ftype not in DIGITAL_FORMS,
        }
    return result


def onboard_status_badge(status: str) -> str:
    return {
        "not_started": "<span style='background:#f1f5f9;color:#64748b;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>Not Started</span>",
        "in_progress": "<span style='background:#fef3c7;color:#d97706;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>In Progress</span>",
        "completed":   "<span style='background:#dcfce7;color:#16a34a;font-size:11px;font-weight:700;padding:2px 8px;border-radius:6px'>✅ Complete</span>",
    }.get(status, status)


@router.get("/staff/{staff_id}/onboarding", response_class=HTMLResponse)
def onboarding_overview(
    staff_id: int,
    session:  str | None = Cookie(default=None),
    msg:      str = ""
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    ob_status = get_onboarding_status(staff_id)
    is_owner  = user["role"] == "owner"
    flash     = f"<div class='flash-success'>{msg}</div>" if msg else ""

    # Build checklist
    checklist = ""
    all_done  = all(v["status"] == "completed" for v in ob_status.values())

    for ftype, info in ob_status.items():
        # Staff only see their own forms
        if info["who"] == "owner" and not is_owner:
            continue

        status  = info["status"]
        badge   = onboard_status_badge(status)
        is_doc  = info.get("is_doc", False)

        if is_doc:
            # Document-based — link to documents page
            btn_url = f"/staff/{staff_id}/documents"
            btn_lbl = "Go to Documents →" if status != "completed" else "View Documents →"
            btn_cls = "btn-primary" if status != "completed" else "btn-secondary"
        else:
            btn_url = f"/staff/{staff_id}/onboarding/{ftype}"
            btn_lbl = "Start →" if status == "not_started" else ("Continue →" if status == "in_progress" else "View →")
            btn_cls = "btn-primary" if status != "completed" else "btn-secondary"

        # Upload signed copy option (for forms only, not doc-based items)
        upload_html = ""
        if not is_doc and status != "completed":
            upload_html = f"""
        <form action='/staff/{staff_id}/onboarding/{ftype}/upload-paper' method='POST'
              enctype='multipart/form-data'
              style='display:inline-flex;gap:6px;align-items:center;margin-left:8px;margin-top:6px'>
          <input type='file' name='paper_form' accept='.pdf,.jpg,.jpeg,.png'
                 style='font-size:11px;max-width:160px;padding:3px'>
          <button type='submit' class='btn-secondary' style='padding:3px 8px;font-size:11px;white-space:nowrap'>
            &#128196; Upload Signed Copy (PDF/Scan)
          </button>
        </form>"""

        # PDF download if completed
        pdf_link = ""
        if status == "completed":
            pdf_row = q("SELECT pdf_path FROM onboarding_forms WHERE staff_id=? AND form_type=?",
                        (staff_id, ftype), fetch=True)
            if pdf_row and dict(pdf_row[0]).get("pdf_path"):
                pdf_link = f"<a href='/staff/{staff_id}/onboarding/{ftype}/pdf' target='_blank' style='color:#1e3a5f;font-size:12px;font-weight:700;margin-left:8px'>📄 PDF</a>"

        owner_tag = "<span style='font-size:10px;color:#94a3b8;margin-left:6px'>(owner only)</span>" if info["who"] == "owner" else ""

        checklist += f"""
        <div style='padding:14px 16px;border-bottom:1px solid #f1f5f9'>
          <div style='display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px'>
            <div>
              <div style='font-weight:700;color:#0f172a;font-size:14px'>{info['label']}{owner_tag}</div>
              <div style='margin-top:4px'>{badge}{pdf_link}</div>
            </div>
            <a href='{btn_url}' class='{btn_cls}' style='padding:6px 16px;font-size:13px'>{btn_lbl}</a>
          </div>
          {upload_html}
        </div>"""

    completion_bar = ""
    completed_n = sum(1 for v in ob_status.values() if v["status"] == "completed")
    total_n     = len(ob_status)
    pct         = int(completed_n / total_n * 100)
    bar_col     = "#16a34a" if pct == 100 else ("#d97706" if pct > 0 else "#e2e8f0")
    completion_bar = f"""
    <div class='card'>
      <div style='display:flex;justify-content:space-between;margin-bottom:6px'>
        <span style='font-size:13px;font-weight:700;color:#0f172a'>Onboarding Progress</span>
        <span style='font-size:13px;font-weight:700;color:{bar_col}'>{completed_n}/{total_n} complete</span>
      </div>
      <div style='background:#f1f5f9;border-radius:99px;height:8px'>
        <div style='background:{bar_col};border-radius:99px;height:8px;width:{pct}%;transition:width .3s'></div>
      </div>
      {'<div style="font-size:12px;color:#16a34a;font-weight:700;margin-top:6px">🎉 All onboarding forms complete!</div>' if pct==100 else ''}
    </div>"""

    content = f"""
    {flash}
    <div class='flex justify-between items-center flex-wrap gap-3'>
      <div>
        <a href='/staff/{staff_id}' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to {esc(name)}</a>
        <div class='text-2xl font-black text-slate-800 mt-1'>📋 Onboarding — {esc(name)}</div>
      </div>
    </div>
    {completion_bar}
    <div class='card' style='padding:0;overflow:hidden'>
      <div style='padding:12px 16px;background:#0f2942;color:white;font-weight:700;font-size:14px'>
        Onboarding Checklist
      </div>
      {checklist}
    </div>"""

    return page("Onboarding", content, user, "staff")


@router.get("/staff/{staff_id}/onboarding/employment_application", response_class=HTMLResponse)
def employment_application_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r

    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    # Get any saved data
    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='employment_application'",
              (staff_id,), fetch=True)
    import json
    data = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}

    # Only pre-fill from profile if form has been started before
    has_data = bool(data)
    def fv(k, default=""): return data.get(k, default) if has_data else ""
    def fi(name, label, ftype="text", val=None, req=False, placeholder=""):
        v    = val if val is not None else fv(name)
        req_a = "required" if req else ""
        ph    = f"placeholder='{placeholder}'" if placeholder else ""
        return f"<div><label>{label}</label><input type='{ftype}' name='{name}' value='{esc(v)}' {req_a} {ph}></div>"

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>Employment Application — {esc(name)}</div>
      <div style='font-size:13px;color:#64748b;margin-top:2px'>Snappy Snaps — Equal Opportunity Employer</div>
    </div>

    <form action='/staff/{staff_id}/onboarding/employment_application' method='POST' class='space-y-6'>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Position & Personal Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('position_applied', 'Position Applied For', req=True)}
          {fi('full_name', 'Full Name', val=fv('full_name') or f"{s.get('first_name','')} {s.get('last_name','')}", req=True)}
          {fi('address',         'Address',               val=fv('address') or ', '.join(filter(None,[s.get('address_1',''),s.get('address_2',''),s.get('address_3',''),s.get('postcode','')])))}
          <div><label>Telephone No.</label>
            <input type='text' name='phone' value='{esc(fv("phone") or s.get("phone",""))}'
              placeholder='01234 567890'>
          </div>
          <div><label>Mobile No. <span style="font-size:10px;color:#94a3b8;font-weight:400">(preferred format: 07700 123456)</span></label>
            <input type='text' name='mobile' value='{esc(fv("mobile"))}'
              placeholder='07700 123456'>
          </div>
          {fi('ni_number',       'National Insurance No.',placeholder='AB 12 34 56 C')}
          <div><label>Driving Licence</label>
            <select name='driving_licence'>
              <option value=''>-- Select --</option>
              <option {'selected' if fv('driving_licence')=='Full' else ''}>Full</option>
              <option {'selected' if fv('driving_licence')=='Provisional' else ''}>Provisional</option>
              <option {'selected' if fv('driving_licence')=='None' else ''}>None</option>
            </select></div>
          <div><label>Work Permit Required?</label>
            <select name='work_permit'>
              <option value='No' {'selected' if fv('work_permit','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('work_permit')=='Yes' else ''}>Yes</option>
            </select></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Health Information</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('health_state', 'Current State of Health', placeholder='e.g. Good')}
          <div><label>Respiratory Problems?</label>
            <select name='respiratory'><option value='No'>No</option><option value='Yes' {'selected' if fv('respiratory')=='Yes' else ''}>Yes</option></select></div>
          <div><label>Skin Irritation?</label>
            <select name='skin_irritation'><option value='No'>No</option><option value='Yes' {'selected' if fv('skin_irritation')=='Yes' else ''}>Yes</option></select></div>
          <div style='grid-column:1/-1'>
            <label>Absence from work through illness in past 12 months</label>
            <textarea name='illness_absence' rows='2' placeholder='Please give details if any'>{esc(fv('illness_absence'))}</textarea>
          </div>
          <div><label>Do you smoke?</label>
            <select name='smoking'>
              <option value='Never' {'selected' if fv('smoking','Never')=='Never' else ''}>Never</option>
              <option value='Socially' {'selected' if fv('smoking')=='Socially' else ''}>Socially</option>
              <option value='Sometimes' {'selected' if fv('smoking')=='Sometimes' else ''}>Sometimes</option>
              <option value='Over 20/day' {'selected' if fv('smoking')=='Over 20/day' else ''}>Over 20/day</option>
            </select></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Educational History</div>
        <div class='grid gap-3' style='grid-template-columns:1fr'>
          <div><label>O Levels / GCSEs (subjects and grades)</label>
            <textarea name='gcse' rows='2'>{esc(fv('gcse'))}</textarea></div>
          <div><label>A Levels</label>
            <textarea name='a_levels' rows='2'>{esc(fv('a_levels'))}</textarea></div>
          <div><label>University / Degree</label>
            <input type='text' name='university' value='{esc(fv('university'))}'></div>
          <div><label>Other Qualifications or Skills</label>
            <textarea name='other_quals' rows='2'>{esc(fv('other_quals'))}</textarea></div>
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>General Information</div>
        <div class='grid gap-3' style='grid-template-columns:1fr'>
          <div><label>What do you seek most from this position?</label>
            <textarea name='seeks' rows='2'>{esc(fv('seeks'))}</textarea></div>
          <div><label>Where do you see yourself in 5 years?</label>
            <textarea name='five_years' rows='2'>{esc(fv('five_years'))}</textarea></div>
          <div><label>Interests and hobbies</label>
            <textarea name='hobbies' rows='2'>{esc(fv('hobbies'))}</textarea></div>
          <div><label>Greatest strengths</label>
            <textarea name='strengths' rows='2'>{esc(fv('strengths'))}</textarea></div>
          <div><label>Greatest weaknesses</label>
            <textarea name='weaknesses' rows='2'>{esc(fv('weaknesses'))}</textarea></div>
          <div><label>Any court convictions or outstanding hearings?</label>
            <textarea name='convictions' rows='2' placeholder='Please declare if any'>{esc(fv('convictions'))}</textarea></div>
          <div><label>Have you previously applied to or worked at Snappy Snaps?</label>
            <select name='prev_snappy'>
              <option value='No' {'selected' if fv('prev_snappy','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('prev_snappy')=='Yes' else ''}>Yes</option>
            </select></div>
          {fi('prev_snappy_details', 'If yes, please give details', placeholder='Store, position, dates')}
        </div>
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment History (most recent first)</div>"""

    for i in range(1, 4):
        content += f"""
        <div style='border:1px solid #e2e8f0;border-radius:10px;padding:14px;margin-bottom:10px'>
          <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Employer {i}</div>
          <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(200px,1fr))'>
            <div><label>Employer Name</label><input type='text' name='emp{i}_name' value='{esc(fv(f"emp{i}_name"))}'></div>
            <div><label>Address</label><input type='text' name='emp{i}_address' value='{esc(fv(f"emp{i}_address"))}'></div>
            <div><label>Date Commenced</label><input type='date' name='emp{i}_start' value='{esc(fv(f"emp{i}_start"))}'></div>
            <div><label>Date Left</label><input type='date' name='emp{i}_end' value='{esc(fv(f"emp{i}_end"))}'></div>
            <div><label>Position</label><input type='text' name='emp{i}_position' value='{esc(fv(f"emp{i}_position"))}'></div>
            <div><label>Salary</label><input type='text' name='emp{i}_salary' value='{esc(fv(f"emp{i}_salary"))}'></div>
            <div style='grid-column:1/-1'><label>Reason for leaving</label>
              <input type='text' name='emp{i}_reason' value='{esc(fv(f"emp{i}_reason"))}'></div>
          </div>
        </div>"""

    content += f"""
      </div>

      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>References</div>
        <div class='grid gap-6' style='grid-template-columns:1fr 1fr'>"""

    for i in range(1, 3):
        content += f"""
          <div>
            <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Reference {i}</div>
            <div class='space-y-2'>
              <div><label>Name</label><input type='text' name='ref{i}_name' value='{esc(fv(f"ref{i}_name"))}'></div>
              <div><label>Address</label><textarea name='ref{i}_address' rows='2'>{esc(fv(f"ref{i}_address"))}</textarea></div>
            </div>
          </div>"""

    content += f"""
        </div>
      </div>

      <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
        <div style='font-size:13px;color:#92400e;font-weight:600;margin-bottom:12px'>
          Declaration: I warrant that the information given is complete, true and accurate.
          I understand that any false statement may disqualify me from employment.
        </div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          {fi('declaration_name', 'Printed Name', val=f"{s.get('first_name','')} {s.get('last_name','')}", req=True)}
          {fi('declaration_date', 'Date', 'date', req=True)}
        </div>
      </div>

      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Submit</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("Employment Application", content, user, "staff")


@router.post("/staff/{staff_id}/onboarding/employment_application")
async def save_employment_application(
    staff_id: int,
    request:  Request,
    session:  str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    q("""INSERT INTO onboarding_forms (staff_id, form_type, status, started_at, completed_at, form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status,
            completed_at=excluded.completed_at,
            form_data=excluded.form_data""",
      (staff_id, "employment_application", status, now, now if status=="completed" else None,
       json.dumps(data)))

    # Update staff profile with key fields
    q("""UPDATE staff_profiles SET phone=?, address_1=?
         WHERE staff_id=?""",
      (data.get("mobile") or data.get("phone",""),
       data.get("address",""), staff_id))

    from urllib.parse import quote as uq
    msg = "Application submitted ✅" if status=="completed" else "Progress saved"
    return RedirectResponse(
        f"/staff/{staff_id}/onboarding?msg={uq(msg)}", status_code=303)


@router.get("/staff/{staff_id}/onboarding/p46", response_class=HTMLResponse)
def p46_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='p46'",
              (staff_id,), fetch=True)
    import json
    data  = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}
    has_data = bool(data)
    def fv(k, d=""): return data.get(k, d) if has_data else d

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>P46 — Employee without a P45</div>
      <div style='font-size:13px;color:#64748b;margin-top:2px'>Section one — to be completed by the employee</div>
    </div>
    <form action='/staff/{staff_id}/onboarding/p46' method='POST' class='space-y-6'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Your Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Title</label>
            <select name='title'>
              <option>Mr</option><option>Mrs</option><option>Miss</option><option>Ms</option><option>Dr</option>
            </select></div>
          <div><label>Surname</label><input type='text' name='surname' value='{esc(fv("surname", s.get("last_name","")))}' required></div>
          <div><label>First Name(s)</label><input type='text' name='first_name' value='{esc(fv("first_name", s.get("first_name","")))}' required></div>
          <div><label>Gender</label>
            <select name='gender'>
              <option value='Male' {'selected' if fv('gender','Male')=='Male' else ''}>Male</option>
              <option value='Female' {'selected' if fv('gender')=='Female' else ''}>Female</option>
            </select></div>
          <div><label>Date of Birth</label><input type='date' name='dob' value='{esc(fv("dob", s.get("date_of_birth","")))}' required></div>
          <div><label>National Insurance Number</label><input type='text' name='nino' value='{esc(fv("nino"))}' placeholder='AB 12 34 56 C' required></div>
          <div style='grid-column:1/-1'><label>Address</label>
            <input type='text' name='address' value='{esc(fv("address", ", ".join(filter(None,[s.get("address_1",""),s.get("address_2",""),s.get("address_3",""),s.get("postcode","")]))))}'></div>
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Your Present Circumstances</div>
        <div style='font-size:13px;color:#64748b;margin-bottom:12px'>Please select the statement that applies to you:</div>
        <div class='space-y-3'>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='A' {'checked' if fv('circumstance')=='A' else ''} style='width:auto;margin-top:3px'>
            <span><strong>A</strong> — This is my first job since last 6 April and I have not been receiving taxable Jobseeker's Allowance, Employment and Support Allowance or a state/occupational pension.</span>
          </label>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='B' {'checked' if fv('circumstance')=='B' else ''} style='width:auto;margin-top:3px'>
            <span><strong>B</strong> — This is now my only job, but since last 6 April I have had another job or received taxable Jobseeker's Allowance or Employment Support Allowance.</span>
          </label>
          <label style='display:flex;gap:10px;align-items:flex-start;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='radio' name='circumstance' value='C' {'checked' if fv('circumstance')=='C' else ''} style='width:auto;margin-top:3px'>
            <span><strong>C</strong> — I have another job or receive a state or occupational pension.</span>
          </label>
        </div>
        <div style='margin-top:12px'>
          <label style='display:flex;gap:10px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='student_loan' value='D' {'checked' if fv('student_loan')=='D' else ''} style='width:auto'>
            <span><strong>D</strong> — I have a Student Loan to repay (box D on P46)</span>
          </label>
        </div>
      </div>
      <div class='card' style='background:#fef3c7;border-color:#fcd34d'>
        <div style='font-size:13px;color:#92400e;font-weight:600;margin-bottom:10px'>
          Declaration: I confirm that this information is correct.
        </div>
        <div class='grid gap-3' style='grid-template-columns:1fr 1fr'>
          <div><label>Date</label><input type='date' name='sign_date' value='{esc(fv("sign_date"))}' required></div>
        </div>
      </div>

      {'"""' if user["role"] != "owner" else f"""
      <div class='card' style='border:2px solid #0f2942'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:4px'>Section 2 — To be completed by the Employer</div>
        <div style='font-size:12px;color:#94a3b8;margin-bottom:12px'>Owner only — not visible to staff</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Employer Name &amp; Address</label>
            <select name='s2_employer'>{employer_options(fv('s2_employer',''))}</select>
          </div>
          <div><label>Date Employment Started</label>
            <input type='date' name='s2_start_date' value='{esc(fv("s2_start_date") or s.get("date_joined",""))}'></div>
          <div><label>Job Title</label>
            <input type='text' name='s2_job_title' value='{esc(fv("s2_job_title") or s.get("job_title",""))}'
              placeholder='e.g. Sales Assistant'></div>
          <div><label>Works/Payroll Number</label>
            <input type='text' name='s2_payroll_no' value='{esc(fv("s2_payroll_no"))}' placeholder='e.g. P001'></div>
          <div><label>Employer PAYE Reference</label>
            <input type='text' name='s2_paye_ref' value='{esc(fv("s2_paye_ref"))}' placeholder='e.g. 123/AB456'></div>
          <div><label>Tax Code Used</label>
            <input type='text' name='s2_tax_code' value='{esc(fv("s2_tax_code"))}' placeholder='e.g. 1257L'></div>
        </div>
        <div style='margin-top:12px'>
          <div style='font-size:12px;font-weight:700;color:#64748b;margin-bottom:8px;text-transform:uppercase'>Tax Code Basis</div>
          <div class='space-y-2'>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='A_cumulative' {'checked' if fv("s2_tax_basis")=="A_cumulative" else ''} style='width:auto'>
              Box A — Emergency code on a cumulative basis
            </label>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='B_week1' {'checked' if fv("s2_tax_basis")=="B_week1" else ''} style='width:auto'>
              Box B — Emergency code on a non-cumulative Week 1/Month 1 basis
            </label>
            <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
              <input type='radio' name='s2_tax_basis' value='C_BR' {'checked' if fv("s2_tax_basis")=="C_BR" else ''} style='width:auto'>
              Box C — Code BR (or 0T if employee fails to complete Section 1) Week 1/Month 1 basis
            </label>
          </div>
        </div>
      </div>""" if user["role"] == "owner" else ""}

      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Submit</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("P46", content, user, "staff")


@router.post("/staff/{staff_id}/onboarding/p46")
async def save_p46(staff_id: int, request: Request, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    q("""INSERT INTO onboarding_forms (staff_id,form_type,status,started_at,completed_at,form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status, completed_at=excluded.completed_at, form_data=excluded.form_data""",
      (staff_id,"p46",status,now,now if status=="completed" else None, json.dumps(data)))
    # Update DOB on profile
    if data.get("dob"):
        q("UPDATE staff_profiles SET date_of_birth=? WHERE staff_id=?", (data["dob"], staff_id))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/onboarding?msg={uq('P46 saved')}", status_code=303)


@router.get("/staff/{staff_id}/onboarding/new_employee_notify", response_class=HTMLResponse)
def new_employee_notify_form(staff_id: int, session: str | None = Cookie(default=None)):
    redir, user = require_login(session)
    if redir: return redir
    if user["role"] != "owner":
        return RedirectResponse(f"/staff/{staff_id}/onboarding", status_code=303)
    rows = q("SELECT * FROM staff_profiles WHERE staff_id=?", (staff_id,), fetch=True)
    if not rows: return RedirectResponse("/staff", status_code=303)
    s    = dict(rows[0])
    name = f"{s['first_name']} {s['last_name']}"

    saved = q("SELECT form_data FROM onboarding_forms WHERE staff_id=? AND form_type='new_employee_notify'",
              (staff_id,), fetch=True)
    import json
    data     = json.loads(dict(saved[0])["form_data"]) if saved and dict(saved[0])["form_data"] else {}
    has_data = bool(data)
    def fv(k, d=""): return data.get(k, d) if has_data else d

    def fi(nm, lbl, ft="text", val=None, ph=""):
        v  = val if val is not None else fv(nm)
        ph = f"placeholder='{ph}'" if ph else ""
        return f"<div><label>{lbl}</label><input type='{ft}' name='{nm}' value='{esc(v)}' {ph}></div>"

    content = f"""
    <div>
      <a href='/staff/{staff_id}/onboarding' style='color:#1e3a5f;font-size:13px;font-weight:700'>← Back to Onboarding</a>
      <div class='text-2xl font-black text-slate-800 mt-1'>New Employee Notification — {esc(name)}</div>
      <div style='font-size:12px;color:#94a3b8;margin-top:2px'>Owner only — not visible to staff</div>
    </div>
    <form action='/staff/{staff_id}/onboarding/new_employee_notify' method='POST' class='space-y-6'>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employee Details</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          {fi('surname',      'Surname',       val=fv('surname') or s.get('last_name',''))}
          {fi('first_name',   'First Name(s)', val=fv('first_name') or s.get('first_name',''))}
          <div><label>Title</label><select name='title'><option>Mr</option><option>Mrs</option><option>Miss</option><option>Ms</option></select></div>
          <div><label>Gender</label><select name='gender'><option>Male</option><option>Female</option></select></div>
          <div><label>Married</label><select name='married'><option value='No'>No</option><option value='Yes' {'selected' if fv('married')=='Yes' else ''}>Yes</option></select></div>
          {fi('dob',          'Date of Birth', 'date', s.get('date_of_birth',''))}
          {fi('nino',         'NI Number',     ph='AB 12 34 56 C')}
          {fi('start_date',   'Start Date',    'date', s.get('date_joined',''))}
          {fi('address',      'Employee Address', val=fv('address') or ', '.join(filter(None,[s.get('address_1',''),s.get('address_2',''),s.get('address_3',''),s.get('postcode','')])))}
          {fi('postcode',     'Post Code',     val=fv('postcode') or s.get('postcode',''))}
          {fi('phone',        'Phone',         val=fv('phone') or s.get('phone',''))}
          {fi('emergency',    'Emergency Contact')}
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Employment Details (Employer)</div>
        <div class='grid gap-3' style='grid-template-columns:repeat(auto-fit,minmax(220px,1fr))'>
          <div><label>Employer Name & Address</label>
            <select name='employer_name'>{employer_options(fv('employer_name',''))}</select>
          </div>
          <div><label>Pay Frequency</label>
            <select name='pay_frequency'>
              <option {'selected' if fv('pay_frequency','Monthly')=='Monthly' else ''}>Monthly</option>
              <option {'selected' if fv('pay_frequency')=='Weekly' else ''}>Weekly</option>
              <option {'selected' if fv('pay_frequency')=='4 Weekly' else ''}>4 Weekly</option>
            </select></div>
          {fi('pay_day',       'Pay Day & Date', ph='e.g. Last Friday of month')}
          <div><label>Pay Method</label>
            <select name='pay_method'>
              <option {'selected' if fv('pay_method','BACS')=='BACS' else ''}>BACS</option>
              <option {'selected' if fv('pay_method')=='Cash' else ''}>Cash</option>
              <option {'selected' if fv('pay_method')=='Cheque' else ''}>Cheque</option>
            </select></div>
          {fi('payroll_no',    'Payroll No.',    ph='e.g. P001')}
          {fi('tax_code',      'Tax Code',       ph='e.g. 1257L')}
          {fi('nic_letter',    'NIC Letter',     ph='e.g. A')}
          {fi('contracted_hrs','Contracted Hours/Week', val=fv('contracted_hrs') or str(s.get('contracted_hrs','')))}
          {fi('wage',          'Wage/Salary',    ph='e.g. £12.71 per hour')}
          {fi('holiday_start', 'Holiday Year Start', 'date', fv('holiday_start','2026-01-01'))}
          {fi('holiday_end',   'Holiday Year End',   'date', fv('holiday_end','2026-12-31'))}
          {fi('holiday_days',  'Holiday Entitlement (days)', ph='e.g. 19')}
          <div><label>Employment Type</label>
            <select name='emp_type'>
              <option {'selected' if fv('emp_type','Permanent')=='Permanent' else ''}>Permanent</option>
              <option {'selected' if fv('emp_type')=='Temporary' else ''}>Temporary</option>
            </select></div>
          <div><label>Student?</label>
            <select name='is_student'>
              <option value='No' {'selected' if fv('is_student','No')=='No' else ''}>No</option>
              <option value='Yes' {'selected' if fv('is_student')=='Yes' else ''}>Yes</option>
            </select></div>
          <div><label>Only Employment?</label>
            <select name='only_employment'>
              <option value='Yes' {'selected' if fv('only_employment','Yes')=='Yes' else ''}>Yes</option>
              <option value='No' {'selected' if fv('only_employment')=='No' else ''}>No</option>
            </select></div>
        </div>
      </div>
      <div class='card'>
        <div style='font-weight:900;color:#0f2942;margin-bottom:12px'>Right to Work Documents Checked</div>
        <div class='space-y-2'>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_passport' value='1' {'checked' if fv('rtw_passport') else ''} style='width:auto'>
            UK or EEA Passport
          </label>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_birth_cert' value='1' {'checked' if fv('rtw_birth_cert') else ''} style='width:auto'>
            Full British Birth Certificate
          </label>
          <label style='display:flex;gap:8px;align-items:center;cursor:pointer;text-transform:none;font-size:13px;font-weight:400'>
            <input type='checkbox' name='rtw_work_permit' value='1' {'checked' if fv('rtw_work_permit') else ''} style='width:auto'>
            Work Permit with Passport
          </label>
        </div>
      </div>
      <div style='display:flex;gap:8px'>
        <button type='submit' name='action' value='save' class='btn-secondary'>💾 Save Progress</button>
        <button type='submit' name='action' value='complete' class='btn-primary'>✅ Mark Complete</button>
        <a href='/staff/{staff_id}/onboarding' class='btn-secondary'>Cancel</a>
      </div>
    </form>"""

    return page("New Employee Notification", content, user, "staff")


@router.post("/staff/{staff_id}/onboarding/new_employee_notify")
async def save_new_employee_notify(
    staff_id: int, request: Request, session: str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _require_mgr(user)): return r
    import json
    form   = await request.form()
    action = form.get("action","save")
    data   = {k: str(v) for k, v in form.items() if k != "action"}
    status = "completed" if action == "complete" else "in_progress"
    now    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    q("""INSERT INTO onboarding_forms (staff_id,form_type,status,started_at,completed_at,form_data)
         VALUES(?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status=excluded.status, completed_at=excluded.completed_at, form_data=excluded.form_data""",
      (staff_id,"new_employee_notify",status,now,now if status=="completed" else None,json.dumps(data)))
    from urllib.parse import quote as uq
    return RedirectResponse(f"/staff/{staff_id}/onboarding?msg={uq('Notification saved')}", status_code=303)


@router.post("/staff/{staff_id}/onboarding/{form_type}/upload-paper")
async def upload_paper_form(
    staff_id:  int,
    form_type: str,
    request:   Request,
    session:   str | None = Cookie(default=None)
):
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r

    form      = await request.form()
    paper     = form.get("paper_form")

    if not paper or not hasattr(paper, "filename") or not paper.filename:
        from urllib.parse import quote as uq
        return RedirectResponse(
            f"/staff/{staff_id}/onboarding?msg={uq('No file selected')}&msg_type=error",
            status_code=303)

    # Save the file (sanitise name parts to prevent path traversal; whitelist ext; cap size)
    ext      = os.path.splitext(paper.filename)[1].lower()
    filename = f"onboard_{staff_id}_{_safe_part(form_type)}_paper{_safe_ext(ext)}"
    filepath = os.path.join(DOCS_DIR, filename)
    data = await paper.read()
    if len(data) > 25 * 1024 * 1024:
        from urllib.parse import quote as uq
        return RedirectResponse(f"/staff/{staff_id}/onboarding?msg={uq('File too large (max 25 MB)')}&msg_type=error",
                                status_code=303)
    os.makedirs(DOCS_DIR, exist_ok=True)
    with open(filepath, "wb") as f:
        f.write(data)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Mark form as completed with paper upload
    q("""INSERT INTO onboarding_forms
            (staff_id, form_type, status, started_at, completed_at, form_data, pdf_path)
         VALUES(?,?,?,?,?,?,?)
         ON CONFLICT(staff_id,form_type) DO UPDATE SET
            status='completed',
            completed_at=excluded.completed_at,
            pdf_path=excluded.pdf_path""",
      (staff_id, form_type, "completed", now, now,
       '{"source":"paper_upload"}', filepath))

    from urllib.parse import quote as uq
    return RedirectResponse(
        f"/staff/{staff_id}/onboarding?msg={uq('Paper form uploaded and marked complete')}",
        status_code=303)


@router.get("/staff/{staff_id}/onboarding/{form_type}/pdf")
def onboarding_form_file(staff_id: int, form_type: str, session: str | None = Cookie(default=None)):
    """Serve the uploaded signed copy of an onboarding form so it can be reviewed.
    (The checklist links here whenever a form has an uploaded file.)"""
    redir, user = require_login(session)
    if redir: return redir
    if (r := _staff_access_guard(user, staff_id, allow_staff=False)): return r
    rows = q("SELECT pdf_path FROM onboarding_forms WHERE staff_id=? AND form_type=?",
             (staff_id, form_type), fetch=True)
    path = dict(rows[0])["pdf_path"] if rows else None
    if not path or not os.path.exists(path):
        return HTMLResponse("<p>No uploaded file for this form.</p>", status_code=404)
    ext   = os.path.splitext(path)[1].lower()
    media = {".pdf": "application/pdf", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
             ".png": "image/png"}.get(ext, "application/octet-stream")
    return FileResponse(path, media_type=media)


ensure_staff_tables()
ensure_onboarding_tables()
